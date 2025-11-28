# api.py
"""Directly exposed API functions and classes, :func:`Document` for now.

Provides a syntactically more convenient API for interacting with the OpcPackage graph.
"""

from __future__ import annotations

import os
from typing import IO, TYPE_CHECKING, cast

from docx.opc.constants import CONTENT_TYPE as CT
from docx.package import Package

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject
    from docx.parts.document import DocumentPart


def Document(docx: str | IO[bytes] | None = None) -> DocumentObject:
    """Return a |Document| object loaded from `docx`, where `docx` can be either a path
    to a ``.docx`` file (a string) or a file-like object.

    If `docx` is missing or ``None``, the built-in default document "template" is
    loaded.
    """
    docx = _default_docx_path() if docx is None else docx
    document_part = cast("DocumentPart", Package.open(docx).main_document_part)
    if document_part.content_type != CT.WML_DOCUMENT_MAIN:
        tmpl = "file '%s' is not a Word file, content type is '%s'"
        raise ValueError(tmpl % (docx, document_part.content_type))
    return document_part.document


def _default_docx_path():
    """Return the path to the built-in default .docx package."""
    _thisdir = os.path.split(__file__)[0]
    return os.path.join(_thisdir, "templates", "default.docx")


# blkcntnr.py
# pyright: reportImportCycles=false

"""Block item container, used by body, cell, header, etc.

Block level items are things like paragraph and table, although there are a few other
specialized ones like structured document tags.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Iterator

from typing_extensions import TypeAlias

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import StoryChild
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    import docx.types as t
    from docx.oxml.document import CT_Body
    from docx.oxml.section import CT_HdrFtr
    from docx.oxml.table import CT_Tc
    from docx.shared import Length
    from docx.styles.style import ParagraphStyle
    from docx.table import Table

BlockItemElement: TypeAlias = "CT_Body | CT_HdrFtr | CT_Tc"


class BlockItemContainer(StoryChild):
    """Base class for proxy objects that can contain block items.

    These containers include _Body, _Cell, header, footer, footnote, endnote, comment,
    and text box objects. Provides the shared functionality to add a block item like a
    paragraph or table.
    """

    def __init__(self, element: BlockItemElement, parent: t.ProvidesStoryPart):
        super(BlockItemContainer, self).__init__(parent)
        self._element = element

    def add_paragraph(self, text: str = "", style: str | ParagraphStyle | None = None) -> Paragraph:
        """Return paragraph newly added to the end of the content in this container.

        The paragraph has `text` in a single run if present, and is given paragraph
        style `style`.

        If `style` is |None|, no paragraph style is applied, which has the same effect
        as applying the 'Normal' style.
        """
        paragraph = self._add_paragraph()
        if text:
            paragraph.add_run(text)
        if style is not None:
            paragraph.style = style
        return paragraph

    def add_table(self, rows: int, cols: int, width: Length) -> Table:
        """Return table of `width` having `rows` rows and `cols` columns.

        The table is appended appended at the end of the content in this container.

        `width` is evenly distributed between the table columns.
        """
        from docx.table import Table

        tbl = CT_Tbl.new_tbl(rows, cols, width)
        self._element._insert_tbl(tbl)  #  # pyright: ignore[reportPrivateUsage]
        return Table(tbl, self)

    def iter_inner_content(self) -> Iterator[Paragraph | Table]:
        """Generate each `Paragraph` or `Table` in this container in document order."""
        from docx.table import Table

        for element in self._element.inner_content_elements:
            yield (Paragraph(element, self) if isinstance(element, CT_P) else Table(element, self))

    @property
    def paragraphs(self):
        """A list containing the paragraphs in this container, in document order.

        Read-only.
        """
        return [Paragraph(p, self) for p in self._element.p_lst]

    @property
    def tables(self):
        """A list containing the tables in this container, in document order.

        Read-only.
        """
        from docx.table import Table

        return [Table(tbl, self) for tbl in self._element.tbl_lst]

    def _add_paragraph(self):
        """Return paragraph newly added to the end of the content in this container."""
        return Paragraph(self._element.add_p(), self)


# color.py
"""DrawingML objects related to color, ColorFormat being the most prominent."""

from ..enum.dml import MSO_COLOR_TYPE
from ..oxml.simpletypes import ST_HexColorAuto
from ..shared import ElementProxy


class ColorFormat(ElementProxy):
    """Provides access to color settings such as RGB color, theme color, and luminance
    adjustments."""

    def __init__(self, rPr_parent):
        super(ColorFormat, self).__init__(rPr_parent)

    @property
    def rgb(self):
        """An |RGBColor| value or |None| if no RGB color is specified.

        When :attr:`type` is `MSO_COLOR_TYPE.RGB`, the value of this property will
        always be an |RGBColor| value. It may also be an |RGBColor| value if
        :attr:`type` is `MSO_COLOR_TYPE.THEME`, as Word writes the current value of a
        theme color when one is assigned. In that case, the RGB value should be
        interpreted as no more than a good guess however, as the theme color takes
        precedence at rendering time. Its value is |None| whenever :attr:`type` is
        either |None| or `MSO_COLOR_TYPE.AUTO`.

        Assigning an |RGBColor| value causes :attr:`type` to become `MSO_COLOR_TYPE.RGB`
        and any theme color is removed. Assigning |None| causes any color to be removed
        such that the effective color is inherited from the style hierarchy.
        """
        color = self._color
        if color is None:
            return None
        if color.val == ST_HexColorAuto.AUTO:
            return None
        return color.val

    @rgb.setter
    def rgb(self, value):
        if value is None and self._color is None:
            return
        rPr = self._element.get_or_add_rPr()
        rPr._remove_color()
        if value is not None:
            rPr.get_or_add_color().val = value

    @property
    def theme_color(self):
        """Member of :ref:`MsoThemeColorIndex` or |None| if no theme color is specified.

        When :attr:`type` is `MSO_COLOR_TYPE.THEME`, the value of this property will
        always be a member of :ref:`MsoThemeColorIndex`. When :attr:`type` has any other
        value, the value of this property is |None|.

        Assigning a member of :ref:`MsoThemeColorIndex` causes :attr:`type` to become
        `MSO_COLOR_TYPE.THEME`. Any existing RGB value is retained but ignored by Word.
        Assigning |None| causes any color specification to be removed such that the
        effective color is inherited from the style hierarchy.
        """
        color = self._color
        if color is None or color.themeColor is None:
            return None
        return color.themeColor

    @theme_color.setter
    def theme_color(self, value):
        if value is None:
            if self._color is not None:
                self._element.rPr._remove_color()
            return
        self._element.get_or_add_rPr().get_or_add_color().themeColor = value

    @property
    def type(self) -> MSO_COLOR_TYPE:
        """Read-only.

        A member of :ref:`MsoColorType`, one of RGB, THEME, or AUTO, corresponding to
        the way this color is defined. Its value is |None| if no color is applied at
        this level, which causes the effective color to be inherited from the style
        hierarchy.
        """
        color = self._color
        if color is None:
            return None
        if color.themeColor is not None:
            return MSO_COLOR_TYPE.THEME
        if color.val == ST_HexColorAuto.AUTO:
            return MSO_COLOR_TYPE.AUTO
        return MSO_COLOR_TYPE.RGB

    @property
    def _color(self):
        """Return `w:rPr/w:color` or |None| if not present.

        Helper to factor out repetitive element access.
        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.color


# __init__.py


# document.py
# pyright: reportImportCycles=false
# pyright: reportPrivateUsage=false

"""|Document| and closely related objects."""

from __future__ import annotations

from typing import IO, TYPE_CHECKING, Iterator, List

from docx.blkcntnr import BlockItemContainer
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.section import Section, Sections
from docx.shared import ElementProxy, Emu

if TYPE_CHECKING:
    import docx.types as t
    from docx.oxml.document import CT_Body, CT_Document
    from docx.parts.document import DocumentPart
    from docx.settings import Settings
    from docx.shared import Length
    from docx.styles.style import ParagraphStyle, _TableStyle
    from docx.table import Table
    from docx.text.paragraph import Paragraph


class Document(ElementProxy):
    """WordprocessingML (WML) document.

    Not intended to be constructed directly. Use :func:`docx.Document` to open or create
    a document.
    """

    def __init__(self, element: CT_Document, part: DocumentPart):
        super(Document, self).__init__(element)
        self._element = element
        self._part = part
        self.__body = None

    def add_heading(self, text: str = "", level: int = 1):
        """Return a heading paragraph newly added to the end of the document.

        The heading paragraph will contain `text` and have its paragraph style
        determined by `level`. If `level` is 0, the style is set to `Title`. If `level`
        is 1 (or omitted), `Heading 1` is used. Otherwise the style is set to `Heading
        {level}`. Raises |ValueError| if `level` is outside the range 0-9.
        """
        if not 0 <= level <= 9:
            raise ValueError("level must be in range 0-9, got %d" % level)
        style = "Title" if level == 0 else "Heading %d" % level
        return self.add_paragraph(text, style)

    def add_page_break(self):
        """Return newly |Paragraph| object containing only a page break."""
        paragraph = self.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return paragraph

    def add_paragraph(self, text: str = "", style: str | ParagraphStyle | None = None) -> Paragraph:
        """Return paragraph newly added to the end of the document.

        The paragraph is populated with `text` and having paragraph style `style`.

        `text` can contain tab (``\\t``) characters, which are converted to the
        appropriate XML form for a tab. `text` can also include newline (``\\n``) or
        carriage return (``\\r``) characters, each of which is converted to a line
        break.
        """
        return self._body.add_paragraph(text, style)

    def add_picture(
        self,
        image_path_or_stream: str | IO[bytes],
        width: int | Length | None = None,
        height: int | Length | None = None,
    ):
        """Return new picture shape added in its own paragraph at end of the document.

        The picture contains the image at `image_path_or_stream`, scaled based on
        `width` and `height`. If neither width nor height is specified, the picture
        appears at its native size. If only one is specified, it is used to compute a
        scaling factor that is then applied to the unspecified dimension, preserving the
        aspect ratio of the image. The native size of the picture is calculated using
        the dots-per-inch (dpi) value specified in the image file, defaulting to 72 dpi
        if no value is specified, as is often the case.
        """
        run = self.add_paragraph().add_run()
        return run.add_picture(image_path_or_stream, width, height)

    def add_section(self, start_type: WD_SECTION = WD_SECTION.NEW_PAGE):
        """Return a |Section| object newly added at the end of the document.

        The optional `start_type` argument must be a member of the :ref:`WdSectionStart`
        enumeration, and defaults to ``WD_SECTION.NEW_PAGE`` if not provided.
        """
        new_sectPr = self._element.body.add_section_break()
        new_sectPr.start_type = start_type
        return Section(new_sectPr, self._part)

    def add_table(self, rows: int, cols: int, style: str | _TableStyle | None = None):
        """Add a table having row and column counts of `rows` and `cols` respectively.

        `style` may be a table style object or a table style name. If `style` is |None|,
        the table inherits the default table style of the document.
        """
        table = self._body.add_table(rows, cols, self._block_width)
        table.style = style
        return table

    @property
    def core_properties(self):
        """A |CoreProperties| object providing Dublin Core properties of document."""
        return self._part.core_properties

    @property
    def inline_shapes(self):
        """The |InlineShapes| collection for this document.

        An inline shape is a graphical object, such as a picture, contained in a run of
        text and behaving like a character glyph, being flowed like other text in a
        paragraph.
        """
        return self._part.inline_shapes

    def iter_inner_content(self) -> Iterator[Paragraph | Table]:
        """Generate each `Paragraph` or `Table` in this document in document order."""
        return self._body.iter_inner_content()

    @property
    def paragraphs(self) -> List[Paragraph]:
        """The |Paragraph| instances in the document, in document order.

        Note that paragraphs within revision marks such as ``<w:ins>`` or ``<w:del>`` do
        not appear in this list.
        """
        return self._body.paragraphs

    @property
    def part(self) -> DocumentPart:
        """The |DocumentPart| object of this document."""
        return self._part

    def save(self, path_or_stream: str | IO[bytes]):
        """Save this document to `path_or_stream`.

        `path_or_stream` can be either a path to a filesystem location (a string) or a
        file-like object.
        """
        self._part.save(path_or_stream)

    @property
    def sections(self) -> Sections:
        """|Sections| object providing access to each section in this document."""
        return Sections(self._element, self._part)

    @property
    def settings(self) -> Settings:
        """A |Settings| object providing access to the document-level settings."""
        return self._part.settings

    @property
    def styles(self):
        """A |Styles| object providing access to the styles in this document."""
        return self._part.styles

    @property
    def tables(self) -> List[Table]:
        """All |Table| instances in the document, in document order.

        Note that only tables appearing at the top level of the document appear in this
        list; a table nested inside a table cell does not appear. A table within
        revision marks such as ``<w:ins>`` or ``<w:del>`` will also not appear in the
        list.
        """
        return self._body.tables

    @property
    def _block_width(self) -> Length:
        """A |Length| object specifying the space between margins in last section."""
        section = self.sections[-1]
        return Emu(section.page_width - section.left_margin - section.right_margin)

    @property
    def _body(self) -> _Body:
        """The |_Body| instance containing the content for this document."""
        if self.__body is None:
            self.__body = _Body(self._element.body, self)
        return self.__body


class _Body(BlockItemContainer):
    """Proxy for `<w:body>` element in this document.

    It's primary role is a container for document content.
    """

    def __init__(self, body_elm: CT_Body, parent: t.ProvidesStoryPart):
        super(_Body, self).__init__(body_elm, parent)
        self._body = body_elm

    def clear_content(self):
        """Return this |_Body| instance after clearing it of all content.

        Section properties for the main document story, if present, are preserved.
        """
        self._body.clear_content()
        return self


# __init__.py
"""DrawingML-related objects are in this subpackage."""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.oxml.drawing import CT_Drawing
from docx.shared import Parented

if TYPE_CHECKING:
    import docx.types as t


class Drawing(Parented):
    """Container for a DrawingML object."""

    def __init__(self, drawing: CT_Drawing, parent: t.ProvidesStoryPart):
        super().__init__(parent)
        self._parent = parent
        self._drawing = self._element = drawing


# base.py
"""Base classes and other objects used by enumerations."""

from __future__ import annotations

import enum
import textwrap
from typing import TYPE_CHECKING, Any, Dict, Type, TypeVar

if TYPE_CHECKING:
    from typing_extensions import Self

_T = TypeVar("_T", bound="BaseXmlEnum")


class BaseEnum(int, enum.Enum):
    """Base class for Enums that do not map XML attr values.

    The enum's value will be an integer, corresponding to the integer assigned the
    corresponding member in the MS API enum of the same name.
    """

    def __new__(cls, ms_api_value: int, docstr: str):
        self = int.__new__(cls, ms_api_value)
        self._value_ = ms_api_value
        self.__doc__ = docstr.strip()
        return self

    def __str__(self):
        """The symbolic name and string value of this member, e.g. 'MIDDLE (3)'."""
        return f"{self.name} ({self.value})"


class BaseXmlEnum(int, enum.Enum):
    """Base class for Enums that also map XML attr values.

    The enum's value will be an integer, corresponding to the integer assigned the
    corresponding member in the MS API enum of the same name.
    """

    xml_value: str

    def __new__(cls, ms_api_value: int, xml_value: str, docstr: str):
        self = int.__new__(cls, ms_api_value)
        self._value_ = ms_api_value
        self.xml_value = xml_value
        self.__doc__ = docstr.strip()
        return self

    def __str__(self):
        """The symbolic name and string value of this member, e.g. 'MIDDLE (3)'."""
        return f"{self.name} ({self.value})"

    @classmethod
    def from_xml(cls, xml_value: str | None) -> Self:
        """Enumeration member corresponding to XML attribute value `xml_value`.

        Example::

            >>> WD_PARAGRAPH_ALIGNMENT.from_xml("center")
            WD_PARAGRAPH_ALIGNMENT.CENTER

        """
        member = next((member for member in cls if member.xml_value == xml_value), None)
        if member is None:
            raise ValueError(f"{cls.__name__} has no XML mapping for '{xml_value}'")
        return member

    @classmethod
    def to_xml(cls: Type[_T], value: int | _T | None) -> str | None:
        """XML value of this enum member, generally an XML attribute value."""
        # -- presence of multi-arg `__new__()` method fools type-checker, but getting a
        # -- member by its value using EnumCls(val) works as usual.
        return cls(value).xml_value


class DocsPageFormatter:
    """Generate an .rst doc page for an enumeration.

    Formats a RestructuredText documention page (string) for the enumeration class parts
    passed to the constructor. An immutable one-shot service object.
    """

    def __init__(self, clsname: str, clsdict: Dict[str, Any]):
        self._clsname = clsname
        self._clsdict = clsdict

    @property
    def page_str(self):
        """The RestructuredText documentation page for the enumeration.

        This is the only API member for the class.
        """
        tmpl = ".. _%s:\n\n%s\n\n%s\n\n----\n\n%s"
        components = (
            self._ms_name,
            self._page_title,
            self._intro_text,
            self._member_defs,
        )
        return tmpl % components

    @property
    def _intro_text(self):
        """Docstring of the enumeration, formatted for documentation page."""
        try:
            cls_docstring = self._clsdict["__doc__"]
        except KeyError:
            cls_docstring = ""

        if cls_docstring is None:
            return ""

        return textwrap.dedent(cls_docstring).strip()

    def _member_def(self, member: BaseEnum | BaseXmlEnum):
        """Return an individual member definition formatted as an RST glossary entry,
        wrapped to fit within 78 columns."""
        assert member.__doc__ is not None
        member_docstring = textwrap.dedent(member.__doc__).strip()
        member_docstring = textwrap.fill(
            member_docstring,
            width=78,
            initial_indent=" " * 4,
            subsequent_indent=" " * 4,
        )
        return "%s\n%s\n" % (member.name, member_docstring)

    @property
    def _member_defs(self):
        """A single string containing the aggregated member definitions section of the
        documentation page."""
        members = self._clsdict["__members__"]
        member_defs = [self._member_def(member) for member in members if member.name is not None]
        return "\n".join(member_defs)

    @property
    def _ms_name(self):
        """The Microsoft API name for this enumeration."""
        return self._clsdict["__ms_name__"]

    @property
    def _page_title(self):
        """The title for the documentation page, formatted as code (surrounded in
        double-backtics) and underlined with '=' characters."""
        title_underscore = "=" * (len(self._clsname) + 4)
        return "``%s``\n%s" % (self._clsname, title_underscore)


# dml.py
"""Enumerations used by DrawingML objects."""

from .base import BaseEnum, BaseXmlEnum


class MSO_COLOR_TYPE(BaseEnum):
    """Specifies the color specification scheme.

    Example::

        from docx.enum.dml import MSO_COLOR_TYPE

        assert font.color.type == MSO_COLOR_TYPE.SCHEME

    MS API name: `MsoColorType`

    http://msdn.microsoft.com/en-us/library/office/ff864912(v=office.15).aspx
    """

    RGB = (1, "Color is specified by an |RGBColor| value.")
    """Color is specified by an |RGBColor| value."""

    THEME = (2, "Color is one of the preset theme colors.")
    """Color is one of the preset theme colors."""

    AUTO = (101, "Color is determined automatically by the application.")
    """Color is determined automatically by the application."""


class MSO_THEME_COLOR_INDEX(BaseXmlEnum):
    """Indicates the Office theme color, one of those shown in the color gallery on the
    formatting ribbon.

    Alias: ``MSO_THEME_COLOR``

    Example::

        from docx.enum.dml import MSO_THEME_COLOR

        font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

    MS API name: `MsoThemeColorIndex`

    http://msdn.microsoft.com/en-us/library/office/ff860782(v=office.15).aspx
    """

    NOT_THEME_COLOR = (0, "UNMAPPED", "Indicates the color is not a theme color.")
    """Indicates the color is not a theme color."""

    ACCENT_1 = (5, "accent1", "Specifies the Accent 1 theme color.")
    """Specifies the Accent 1 theme color."""

    ACCENT_2 = (6, "accent2", "Specifies the Accent 2 theme color.")
    """Specifies the Accent 2 theme color."""

    ACCENT_3 = (7, "accent3", "Specifies the Accent 3 theme color.")
    """Specifies the Accent 3 theme color."""

    ACCENT_4 = (8, "accent4", "Specifies the Accent 4 theme color.")
    """Specifies the Accent 4 theme color."""

    ACCENT_5 = (9, "accent5", "Specifies the Accent 5 theme color.")
    """Specifies the Accent 5 theme color."""

    ACCENT_6 = (10, "accent6", "Specifies the Accent 6 theme color.")
    """Specifies the Accent 6 theme color."""

    BACKGROUND_1 = (14, "background1", "Specifies the Background 1 theme color.")
    """Specifies the Background 1 theme color."""

    BACKGROUND_2 = (16, "background2", "Specifies the Background 2 theme color.")
    """Specifies the Background 2 theme color."""

    DARK_1 = (1, "dark1", "Specifies the Dark 1 theme color.")
    """Specifies the Dark 1 theme color."""

    DARK_2 = (3, "dark2", "Specifies the Dark 2 theme color.")
    """Specifies the Dark 2 theme color."""

    FOLLOWED_HYPERLINK = (
        12,
        "followedHyperlink",
        "Specifies the theme color for a clicked hyperlink.",
    )
    """Specifies the theme color for a clicked hyperlink."""

    HYPERLINK = (11, "hyperlink", "Specifies the theme color for a hyperlink.")
    """Specifies the theme color for a hyperlink."""

    LIGHT_1 = (2, "light1", "Specifies the Light 1 theme color.")
    """Specifies the Light 1 theme color."""

    LIGHT_2 = (4, "light2", "Specifies the Light 2 theme color.")
    """Specifies the Light 2 theme color."""

    TEXT_1 = (13, "text1", "Specifies the Text 1 theme color.")
    """Specifies the Text 1 theme color."""

    TEXT_2 = (15, "text2", "Specifies the Text 2 theme color.")
    """Specifies the Text 2 theme color."""


MSO_THEME_COLOR = MSO_THEME_COLOR_INDEX


# section.py
"""Enumerations related to the main document in WordprocessingML files."""

from .base import BaseXmlEnum


class WD_HEADER_FOOTER_INDEX(BaseXmlEnum):
    """Alias: **WD_HEADER_FOOTER**

    Specifies one of the three possible header/footer definitions for a section.

    For internal use only; not part of the python-docx API.

    MS API name: `WdHeaderFooterIndex`
    URL: https://docs.microsoft.com/en-us/office/vba/api/word.wdheaderfooterindex
    """

    PRIMARY = (1, "default", "Header for odd pages or all if no even header.")
    """Header for odd pages or all if no even header."""

    FIRST_PAGE = (2, "first", "Header for first page of section.")
    """Header for first page of section."""

    EVEN_PAGE = (3, "even", "Header for even pages of recto/verso section.")
    """Header for even pages of recto/verso section."""


WD_HEADER_FOOTER = WD_HEADER_FOOTER_INDEX


class WD_ORIENTATION(BaseXmlEnum):
    """Alias: **WD_ORIENT**

    Specifies the page layout orientation.

    Example::

        from docx.enum.section import WD_ORIENT

        section = document.sections[-1] section.orientation = WD_ORIENT.LANDSCAPE

    MS API name: `WdOrientation`
    MS API URL: http://msdn.microsoft.com/en-us/library/office/ff837902.aspx
    """

    PORTRAIT = (0, "portrait", "Portrait orientation.")
    """Portrait orientation."""

    LANDSCAPE = (1, "landscape", "Landscape orientation.")
    """Landscape orientation."""


WD_ORIENT = WD_ORIENTATION


class WD_SECTION_START(BaseXmlEnum):
    """Alias: **WD_SECTION**

    Specifies the start type of a section break.

    Example::

        from docx.enum.section import WD_SECTION

        section = document.sections[0] section.start_type = WD_SECTION.NEW_PAGE

    MS API name: `WdSectionStart`
    MS API URL: http://msdn.microsoft.com/en-us/library/office/ff840975.aspx
    """

    CONTINUOUS = (0, "continuous", "Continuous section break.")
    """Continuous section break."""

    NEW_COLUMN = (1, "nextColumn", "New column section break.")
    """New column section break."""

    NEW_PAGE = (2, "nextPage", "New page section break.")
    """New page section break."""

    EVEN_PAGE = (3, "evenPage", "Even pages section break.")
    """Even pages section break."""

    ODD_PAGE = (4, "oddPage", "Section begins on next odd page.")
    """Section begins on next odd page."""


WD_SECTION = WD_SECTION_START


# shape.py
"""Enumerations related to DrawingML shapes in WordprocessingML files."""

import enum


class WD_INLINE_SHAPE_TYPE(enum.Enum):
    """Corresponds to WdInlineShapeType enumeration.

    http://msdn.microsoft.com/en-us/library/office/ff192587.aspx.
    """

    CHART = 12
    LINKED_PICTURE = 4
    PICTURE = 3
    SMART_ART = 15
    NOT_IMPLEMENTED = -6


WD_INLINE_SHAPE = WD_INLINE_SHAPE_TYPE


# style.py
"""Enumerations related to styles."""

from .base import BaseEnum, BaseXmlEnum


class WD_BUILTIN_STYLE(BaseEnum):
    """Alias: **WD_STYLE**

    Specifies a built-in Microsoft Word style.

    Example::

        from docx import Document
        from docx.enum.style import WD_STYLE

        document = Document()
        styles = document.styles
        style = styles[WD_STYLE.BODY_TEXT]


    MS API name: `WdBuiltinStyle`

    http://msdn.microsoft.com/en-us/library/office/ff835210.aspx
    """

    BLOCK_QUOTATION = (-85, "Block Text.")
    """Block Text."""

    BODY_TEXT = (-67, "Body Text.")
    """Body Text."""

    BODY_TEXT_2 = (-81, "Body Text 2.")
    """Body Text 2."""

    BODY_TEXT_3 = (-82, "Body Text 3.")
    """Body Text 3."""

    BODY_TEXT_FIRST_INDENT = (-78, "Body Text First Indent.")
    """Body Text First Indent."""

    BODY_TEXT_FIRST_INDENT_2 = (-79, "Body Text First Indent 2.")
    """Body Text First Indent 2."""

    BODY_TEXT_INDENT = (-68, "Body Text Indent.")
    """Body Text Indent."""

    BODY_TEXT_INDENT_2 = (-83, "Body Text Indent 2.")
    """Body Text Indent 2."""

    BODY_TEXT_INDENT_3 = (-84, "Body Text Indent 3.")
    """Body Text Indent 3."""

    BOOK_TITLE = (-265, "Book Title.")
    """Book Title."""

    CAPTION = (-35, "Caption.")
    """Caption."""

    CLOSING = (-64, "Closing.")
    """Closing."""

    COMMENT_REFERENCE = (-40, "Comment Reference.")
    """Comment Reference."""

    COMMENT_TEXT = (-31, "Comment Text.")
    """Comment Text."""

    DATE = (-77, "Date.")
    """Date."""

    DEFAULT_PARAGRAPH_FONT = (-66, "Default Paragraph Font.")
    """Default Paragraph Font."""

    EMPHASIS = (-89, "Emphasis.")
    """Emphasis."""

    ENDNOTE_REFERENCE = (-43, "Endnote Reference.")
    """Endnote Reference."""

    ENDNOTE_TEXT = (-44, "Endnote Text.")
    """Endnote Text."""

    ENVELOPE_ADDRESS = (-37, "Envelope Address.")
    """Envelope Address."""

    ENVELOPE_RETURN = (-38, "Envelope Return.")
    """Envelope Return."""

    FOOTER = (-33, "Footer.")
    """Footer."""

    FOOTNOTE_REFERENCE = (-39, "Footnote Reference.")
    """Footnote Reference."""

    FOOTNOTE_TEXT = (-30, "Footnote Text.")
    """Footnote Text."""

    HEADER = (-32, "Header.")
    """Header."""

    HEADING_1 = (-2, "Heading 1.")
    """Heading 1."""

    HEADING_2 = (-3, "Heading 2.")
    """Heading 2."""

    HEADING_3 = (-4, "Heading 3.")
    """Heading 3."""

    HEADING_4 = (-5, "Heading 4.")
    """Heading 4."""

    HEADING_5 = (-6, "Heading 5.")
    """Heading 5."""

    HEADING_6 = (-7, "Heading 6.")
    """Heading 6."""

    HEADING_7 = (-8, "Heading 7.")
    """Heading 7."""

    HEADING_8 = (-9, "Heading 8.")
    """Heading 8."""

    HEADING_9 = (-10, "Heading 9.")
    """Heading 9."""

    HTML_ACRONYM = (-96, "HTML Acronym.")
    """HTML Acronym."""

    HTML_ADDRESS = (-97, "HTML Address.")
    """HTML Address."""

    HTML_CITE = (-98, "HTML Cite.")
    """HTML Cite."""

    HTML_CODE = (-99, "HTML Code.")
    """HTML Code."""

    HTML_DFN = (-100, "HTML Definition.")
    """HTML Definition."""

    HTML_KBD = (-101, "HTML Keyboard.")
    """HTML Keyboard."""

    HTML_NORMAL = (-95, "Normal (Web).")
    """Normal (Web)."""

    HTML_PRE = (-102, "HTML Preformatted.")
    """HTML Preformatted."""

    HTML_SAMP = (-103, "HTML Sample.")
    """HTML Sample."""

    HTML_TT = (-104, "HTML Typewriter.")
    """HTML Typewriter."""

    HTML_VAR = (-105, "HTML Variable.")
    """HTML Variable."""

    HYPERLINK = (-86, "Hyperlink.")
    """Hyperlink."""

    HYPERLINK_FOLLOWED = (-87, "Followed Hyperlink.")
    """Followed Hyperlink."""

    INDEX_1 = (-11, "Index 1.")
    """Index 1."""

    INDEX_2 = (-12, "Index 2.")
    """Index 2."""

    INDEX_3 = (-13, "Index 3.")
    """Index 3."""

    INDEX_4 = (-14, "Index 4.")
    """Index 4."""

    INDEX_5 = (-15, "Index 5.")
    """Index 5."""

    INDEX_6 = (-16, "Index 6.")
    """Index 6."""

    INDEX_7 = (-17, "Index 7.")
    """Index 7."""

    INDEX_8 = (-18, "Index 8.")
    """Index 8."""

    INDEX_9 = (-19, "Index 9.")
    """Index 9."""

    INDEX_HEADING = (-34, "Index Heading")
    """Index Heading"""

    INTENSE_EMPHASIS = (-262, "Intense Emphasis.")
    """Intense Emphasis."""

    INTENSE_QUOTE = (-182, "Intense Quote.")
    """Intense Quote."""

    INTENSE_REFERENCE = (-264, "Intense Reference.")
    """Intense Reference."""

    LINE_NUMBER = (-41, "Line Number.")
    """Line Number."""

    LIST = (-48, "List.")
    """List."""

    LIST_2 = (-51, "List 2.")
    """List 2."""

    LIST_3 = (-52, "List 3.")
    """List 3."""

    LIST_4 = (-53, "List 4.")
    """List 4."""

    LIST_5 = (-54, "List 5.")
    """List 5."""

    LIST_BULLET = (-49, "List Bullet.")
    """List Bullet."""

    LIST_BULLET_2 = (-55, "List Bullet 2.")
    """List Bullet 2."""

    LIST_BULLET_3 = (-56, "List Bullet 3.")
    """List Bullet 3."""

    LIST_BULLET_4 = (-57, "List Bullet 4.")
    """List Bullet 4."""

    LIST_BULLET_5 = (-58, "List Bullet 5.")
    """List Bullet 5."""

    LIST_CONTINUE = (-69, "List Continue.")
    """List Continue."""

    LIST_CONTINUE_2 = (-70, "List Continue 2.")
    """List Continue 2."""

    LIST_CONTINUE_3 = (-71, "List Continue 3.")
    """List Continue 3."""

    LIST_CONTINUE_4 = (-72, "List Continue 4.")
    """List Continue 4."""

    LIST_CONTINUE_5 = (-73, "List Continue 5.")
    """List Continue 5."""

    LIST_NUMBER = (-50, "List Number.")
    """List Number."""

    LIST_NUMBER_2 = (-59, "List Number 2.")
    """List Number 2."""

    LIST_NUMBER_3 = (-60, "List Number 3.")
    """List Number 3."""

    LIST_NUMBER_4 = (-61, "List Number 4.")
    """List Number 4."""

    LIST_NUMBER_5 = (-62, "List Number 5.")
    """List Number 5."""

    LIST_PARAGRAPH = (-180, "List Paragraph.")
    """List Paragraph."""

    MACRO_TEXT = (-46, "Macro Text.")
    """Macro Text."""

    MESSAGE_HEADER = (-74, "Message Header.")
    """Message Header."""

    NAV_PANE = (-90, "Document Map.")
    """Document Map."""

    NORMAL = (-1, "Normal.")
    """Normal."""

    NORMAL_INDENT = (-29, "Normal Indent.")
    """Normal Indent."""

    NORMAL_OBJECT = (-158, "Normal (applied to an object).")
    """Normal (applied to an object)."""

    NORMAL_TABLE = (-106, "Normal (applied within a table).")
    """Normal (applied within a table)."""

    NOTE_HEADING = (-80, "Note Heading.")
    """Note Heading."""

    PAGE_NUMBER = (-42, "Page Number.")
    """Page Number."""

    PLAIN_TEXT = (-91, "Plain Text.")
    """Plain Text."""

    QUOTE = (-181, "Quote.")
    """Quote."""

    SALUTATION = (-76, "Salutation.")
    """Salutation."""

    SIGNATURE = (-65, "Signature.")
    """Signature."""

    STRONG = (-88, "Strong.")
    """Strong."""

    SUBTITLE = (-75, "Subtitle.")
    """Subtitle."""

    SUBTLE_EMPHASIS = (-261, "Subtle Emphasis.")
    """Subtle Emphasis."""

    SUBTLE_REFERENCE = (-263, "Subtle Reference.")
    """Subtle Reference."""

    TABLE_COLORFUL_GRID = (-172, "Colorful Grid.")
    """Colorful Grid."""

    TABLE_COLORFUL_LIST = (-171, "Colorful List.")
    """Colorful List."""

    TABLE_COLORFUL_SHADING = (-170, "Colorful Shading.")
    """Colorful Shading."""

    TABLE_DARK_LIST = (-169, "Dark List.")
    """Dark List."""

    TABLE_LIGHT_GRID = (-161, "Light Grid.")
    """Light Grid."""

    TABLE_LIGHT_GRID_ACCENT_1 = (-175, "Light Grid Accent 1.")
    """Light Grid Accent 1."""

    TABLE_LIGHT_LIST = (-160, "Light List.")
    """Light List."""

    TABLE_LIGHT_LIST_ACCENT_1 = (-174, "Light List Accent 1.")
    """Light List Accent 1."""

    TABLE_LIGHT_SHADING = (-159, "Light Shading.")
    """Light Shading."""

    TABLE_LIGHT_SHADING_ACCENT_1 = (-173, "Light Shading Accent 1.")
    """Light Shading Accent 1."""

    TABLE_MEDIUM_GRID_1 = (-166, "Medium Grid 1.")
    """Medium Grid 1."""

    TABLE_MEDIUM_GRID_2 = (-167, "Medium Grid 2.")
    """Medium Grid 2."""

    TABLE_MEDIUM_GRID_3 = (-168, "Medium Grid 3.")
    """Medium Grid 3."""

    TABLE_MEDIUM_LIST_1 = (-164, "Medium List 1.")
    """Medium List 1."""

    TABLE_MEDIUM_LIST_1_ACCENT_1 = (-178, "Medium List 1 Accent 1.")
    """Medium List 1 Accent 1."""

    TABLE_MEDIUM_LIST_2 = (-165, "Medium List 2.")
    """Medium List 2."""

    TABLE_MEDIUM_SHADING_1 = (-162, "Medium Shading 1.")
    """Medium Shading 1."""

    TABLE_MEDIUM_SHADING_1_ACCENT_1 = (-176, "Medium Shading 1 Accent 1.")
    """Medium Shading 1 Accent 1."""

    TABLE_MEDIUM_SHADING_2 = (-163, "Medium Shading 2.")
    """Medium Shading 2."""

    TABLE_MEDIUM_SHADING_2_ACCENT_1 = (-177, "Medium Shading 2 Accent 1.")
    """Medium Shading 2 Accent 1."""

    TABLE_OF_AUTHORITIES = (-45, "Table of Authorities.")
    """Table of Authorities."""

    TABLE_OF_FIGURES = (-36, "Table of Figures.")
    """Table of Figures."""

    TITLE = (-63, "Title.")
    """Title."""

    TOAHEADING = (-47, "TOA Heading.")
    """TOA Heading."""

    TOC_1 = (-20, "TOC 1.")
    """TOC 1."""

    TOC_2 = (-21, "TOC 2.")
    """TOC 2."""

    TOC_3 = (-22, "TOC 3.")
    """TOC 3."""

    TOC_4 = (-23, "TOC 4.")
    """TOC 4."""

    TOC_5 = (-24, "TOC 5.")
    """TOC 5."""

    TOC_6 = (-25, "TOC 6.")
    """TOC 6."""

    TOC_7 = (-26, "TOC 7.")
    """TOC 7."""

    TOC_8 = (-27, "TOC 8.")
    """TOC 8."""

    TOC_9 = (-28, "TOC 9.")
    """TOC 9."""


WD_STYLE = WD_BUILTIN_STYLE


class WD_STYLE_TYPE(BaseXmlEnum):
    """Specifies one of the four style types: paragraph, character, list, or table.

    Example::

        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE

        styles = Document().styles
        assert styles[0].type == WD_STYLE_TYPE.PARAGRAPH

    MS API name: `WdStyleType`

    http://msdn.microsoft.com/en-us/library/office/ff196870.aspx
    """

    CHARACTER = (2, "character", "Character style.")
    """Character style."""

    LIST = (4, "numbering", "List style.")
    """List style."""

    PARAGRAPH = (1, "paragraph", "Paragraph style.")
    """Paragraph style."""

    TABLE = (3, "table", "Table style.")
    """Table style."""


# table.py
"""Enumerations related to tables in WordprocessingML files."""

from docx.enum.base import BaseEnum, BaseXmlEnum


class WD_CELL_VERTICAL_ALIGNMENT(BaseXmlEnum):
    """Alias: **WD_ALIGN_VERTICAL**

    Specifies the vertical alignment of text in one or more cells of a table.

    Example::

        from docx.enum.table import WD_ALIGN_VERTICAL

        table = document.add_table(3, 3)
        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    MS API name: `WdCellVerticalAlignment`

    https://msdn.microsoft.com/en-us/library/office/ff193345.aspx
    """

    TOP = (0, "top", "Text is aligned to the top border of the cell.")
    """Text is aligned to the top border of the cell."""

    CENTER = (1, "center", "Text is aligned to the center of the cell.")
    """Text is aligned to the center of the cell."""

    BOTTOM = (3, "bottom", "Text is aligned to the bottom border of the cell.")
    """Text is aligned to the bottom border of the cell."""

    BOTH = (
        101,
        "both",
        "This is an option in the OpenXml spec, but not in Word itself. It's not"
        " clear what Word behavior this setting produces. If you find out please"
        " let us know and we'll update this documentation. Otherwise, probably best"
        " to avoid this option.",
    )
    """This is an option in the OpenXml spec, but not in Word itself.

    It's not clear what Word behavior this setting produces. If you find out please let
    us know and we'll update this documentation. Otherwise, probably best to avoid this
    option.
    """


WD_ALIGN_VERTICAL = WD_CELL_VERTICAL_ALIGNMENT


class WD_ROW_HEIGHT_RULE(BaseXmlEnum):
    """Alias: **WD_ROW_HEIGHT**

    Specifies the rule for determining the height of a table row

    Example::

        from docx.enum.table import WD_ROW_HEIGHT_RULE

        table = document.add_table(3, 3)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    MS API name: `WdRowHeightRule`

    https://msdn.microsoft.com/en-us/library/office/ff193620.aspx
    """

    AUTO = (
        0,
        "auto",
        "The row height is adjusted to accommodate the tallest value in the row.",
    )
    """The row height is adjusted to accommodate the tallest value in the row."""

    AT_LEAST = (1, "atLeast", "The row height is at least a minimum specified value.")
    """The row height is at least a minimum specified value."""

    EXACTLY = (2, "exact", "The row height is an exact value.")
    """The row height is an exact value."""


WD_ROW_HEIGHT = WD_ROW_HEIGHT_RULE


class WD_TABLE_ALIGNMENT(BaseXmlEnum):
    """Specifies table justification type.

    Example::

        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = document.add_table(3, 3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    MS API name: `WdRowAlignment`

    http://office.microsoft.com/en-us/word-help/HV080607259.aspx
    """

    LEFT = (0, "left", "Left-aligned")
    """Left-aligned"""

    CENTER = (1, "center", "Center-aligned.")
    """Center-aligned."""

    RIGHT = (2, "right", "Right-aligned.")
    """Right-aligned."""


class WD_TABLE_DIRECTION(BaseEnum):
    """Specifies the direction in which an application orders cells in the specified
    table or row.

    Example::

        from docx.enum.table import WD_TABLE_DIRECTION

        table = document.add_table(3, 3)
        table.direction = WD_TABLE_DIRECTION.RTL

    MS API name: `WdTableDirection`

    http://msdn.microsoft.com/en-us/library/ff835141.aspx
    """

    LTR = (
        0,
        "The table or row is arranged with the first column in the leftmost position.",
    )
    """The table or row is arranged with the first column in the leftmost position."""

    RTL = (
        1,
        "The table or row is arranged with the first column in the rightmost position.",
    )
    """The table or row is arranged with the first column in the rightmost position."""


# text.py
"""Enumerations related to text in WordprocessingML files."""

from __future__ import annotations

import enum

from docx.enum.base import BaseXmlEnum


class WD_PARAGRAPH_ALIGNMENT(BaseXmlEnum):
    """Alias: **WD_ALIGN_PARAGRAPH**

    Specifies paragraph justification type.

    Example::

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    """

    LEFT = (0, "left", "Left-aligned")
    """Left-aligned"""

    CENTER = (1, "center", "Center-aligned.")
    """Center-aligned."""

    RIGHT = (2, "right", "Right-aligned.")
    """Right-aligned."""

    JUSTIFY = (3, "both", "Fully justified.")
    """Fully justified."""

    DISTRIBUTE = (
        4,
        "distribute",
        "Paragraph characters are distributed to fill entire width of paragraph.",
    )
    """Paragraph characters are distributed to fill entire width of paragraph."""

    JUSTIFY_MED = (
        5,
        "mediumKashida",
        "Justified with a medium character compression ratio.",
    )
    """Justified with a medium character compression ratio."""

    JUSTIFY_HI = (
        7,
        "highKashida",
        "Justified with a high character compression ratio.",
    )
    """Justified with a high character compression ratio."""

    JUSTIFY_LOW = (8, "lowKashida", "Justified with a low character compression ratio.")
    """Justified with a low character compression ratio."""

    THAI_JUSTIFY = (
        9,
        "thaiDistribute",
        "Justified according to Thai formatting layout.",
    )
    """Justified according to Thai formatting layout."""


WD_ALIGN_PARAGRAPH = WD_PARAGRAPH_ALIGNMENT


class WD_BREAK_TYPE(enum.Enum):
    """Corresponds to WdBreakType enumeration.

    http://msdn.microsoft.com/en-us/library/office/ff195905.aspx.
    """

    COLUMN = 8
    LINE = 6
    LINE_CLEAR_LEFT = 9
    LINE_CLEAR_RIGHT = 10
    LINE_CLEAR_ALL = 11  # -- added for consistency, not in MS version --
    PAGE = 7
    SECTION_CONTINUOUS = 3
    SECTION_EVEN_PAGE = 4
    SECTION_NEXT_PAGE = 2
    SECTION_ODD_PAGE = 5
    TEXT_WRAPPING = 11


WD_BREAK = WD_BREAK_TYPE


class WD_COLOR_INDEX(BaseXmlEnum):
    """Specifies a standard preset color to apply.

    Used for font highlighting and perhaps other applications.

    * MS API name: `WdColorIndex`
    * URL: https://msdn.microsoft.com/EN-US/library/office/ff195343.aspx
    """

    INHERITED = (-1, None, "Color is inherited from the style hierarchy.")
    """Color is inherited from the style hierarchy."""

    AUTO = (0, "default", "Automatic color. Default; usually black.")
    """Automatic color. Default; usually black."""

    BLACK = (1, "black", "Black color.")
    """Black color."""

    BLUE = (2, "blue", "Blue color")
    """Blue color"""

    BRIGHT_GREEN = (4, "green", "Bright green color.")
    """Bright green color."""

    DARK_BLUE = (9, "darkBlue", "Dark blue color.")
    """Dark blue color."""

    DARK_RED = (13, "darkRed", "Dark red color.")
    """Dark red color."""

    DARK_YELLOW = (14, "darkYellow", "Dark yellow color.")
    """Dark yellow color."""

    GRAY_25 = (16, "lightGray", "25% shade of gray color.")
    """25% shade of gray color."""

    GRAY_50 = (15, "darkGray", "50% shade of gray color.")
    """50% shade of gray color."""

    GREEN = (11, "darkGreen", "Green color.")
    """Green color."""

    PINK = (5, "magenta", "Pink color.")
    """Pink color."""

    RED = (6, "red", "Red color.")
    """Red color."""

    TEAL = (10, "darkCyan", "Teal color.")
    """Teal color."""

    TURQUOISE = (3, "cyan", "Turquoise color.")
    """Turquoise color."""

    VIOLET = (12, "darkMagenta", "Violet color.")
    """Violet color."""

    WHITE = (8, "white", "White color.")
    """White color."""

    YELLOW = (7, "yellow", "Yellow color.")
    """Yellow color."""


WD_COLOR = WD_COLOR_INDEX


class WD_LINE_SPACING(BaseXmlEnum):
    """Specifies a line spacing format to be applied to a paragraph.

    Example::

        from docx.enum.text import WD_LINE_SPACING

        paragraph = document.add_paragraph()
        paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY


    MS API name: `WdLineSpacing`

    URL: http://msdn.microsoft.com/en-us/library/office/ff844910.aspx
    """

    SINGLE = (0, "UNMAPPED", "Single spaced (default).")
    """Single spaced (default)."""

    ONE_POINT_FIVE = (1, "UNMAPPED", "Space-and-a-half line spacing.")
    """Space-and-a-half line spacing."""

    DOUBLE = (2, "UNMAPPED", "Double spaced.")
    """Double spaced."""

    AT_LEAST = (
        3,
        "atLeast",
        "Minimum line spacing is specified amount. Amount is specified separately.",
    )
    """Minimum line spacing is specified amount. Amount is specified separately."""

    EXACTLY = (
        4,
        "exact",
        "Line spacing is exactly specified amount. Amount is specified separately.",
    )
    """Line spacing is exactly specified amount. Amount is specified separately."""

    MULTIPLE = (
        5,
        "auto",
        "Line spacing is specified as multiple of line heights. Changing font size"
        " will change line spacing proportionately.",
    )
    """Line spacing is specified as multiple of line heights. Changing font size will
       change the line spacing proportionately."""


class WD_TAB_ALIGNMENT(BaseXmlEnum):
    """Specifies the tab stop alignment to apply.

    MS API name: `WdTabAlignment`

    URL: https://msdn.microsoft.com/EN-US/library/office/ff195609.aspx
    """

    LEFT = (0, "left", "Left-aligned.")
    """Left-aligned."""

    CENTER = (1, "center", "Center-aligned.")
    """Center-aligned."""

    RIGHT = (2, "right", "Right-aligned.")
    """Right-aligned."""

    DECIMAL = (3, "decimal", "Decimal-aligned.")
    """Decimal-aligned."""

    BAR = (4, "bar", "Bar-aligned.")
    """Bar-aligned."""

    LIST = (6, "list", "List-aligned. (deprecated)")
    """List-aligned. (deprecated)"""

    CLEAR = (101, "clear", "Clear an inherited tab stop.")
    """Clear an inherited tab stop."""

    END = (102, "end", "Right-aligned.  (deprecated)")
    """Right-aligned.  (deprecated)"""

    NUM = (103, "num", "Left-aligned.  (deprecated)")
    """Left-aligned.  (deprecated)"""

    START = (104, "start", "Left-aligned.  (deprecated)")
    """Left-aligned.  (deprecated)"""


class WD_TAB_LEADER(BaseXmlEnum):
    """Specifies the character to use as the leader with formatted tabs.

    MS API name: `WdTabLeader`

    URL: https://msdn.microsoft.com/en-us/library/office/ff845050.aspx
    """

    SPACES = (0, "none", "Spaces. Default.")
    """Spaces. Default."""

    DOTS = (1, "dot", "Dots.")
    """Dots."""

    DASHES = (2, "hyphen", "Dashes.")
    """Dashes."""

    LINES = (3, "underscore", "Double lines.")
    """Double lines."""

    HEAVY = (4, "heavy", "A heavy line.")
    """A heavy line."""

    MIDDLE_DOT = (5, "middleDot", "A vertically-centered dot.")
    """A vertically-centered dot."""


class WD_UNDERLINE(BaseXmlEnum):
    """Specifies the style of underline applied to a run of characters.

    MS API name: `WdUnderline`

    URL: http://msdn.microsoft.com/en-us/library/office/ff822388.aspx
    """

    INHERITED = (-1, None, "Inherit underline setting from containing paragraph.")
    """Inherit underline setting from containing paragraph."""

    NONE = (
        0,
        "none",
        "No underline.\n\nThis setting overrides any inherited underline value, so can"
        " be used to remove underline from a run that inherits underlining from its"
        " containing paragraph. Note this is not the same as assigning |None| to"
        " Run.underline. |None| is a valid assignment value, but causes the run to"
        " inherit its underline value. Assigning `WD_UNDERLINE.NONE` causes"
        " underlining to be unconditionally turned off.",
    )
    """No underline.

    This setting overrides any inherited underline value, so can be used to remove
    underline from a run that inherits underlining from its containing paragraph. Note
    this is not the same as assigning |None| to Run.underline. |None| is a valid
    assignment value, but causes the run to inherit its underline value. Assigning
    ``WD_UNDERLINE.NONE`` causes underlining to be unconditionally turned off.
    """

    SINGLE = (
        1,
        "single",
        "A single line.\n\nNote that this setting is write-only in the sense that"
        " |True| (rather than `WD_UNDERLINE.SINGLE`) is returned for a run having"
        " this setting.",
    )
    """A single line.

    Note that this setting is write-only in the sense that |True|
    (rather than ``WD_UNDERLINE.SINGLE``) is returned for a run having this setting.
    """

    WORDS = (2, "words", "Underline individual words only.")
    """Underline individual words only."""

    DOUBLE = (3, "double", "A double line.")
    """A double line."""

    DOTTED = (4, "dotted", "Dots.")
    """Dots."""

    THICK = (6, "thick", "A single thick line.")
    """A single thick line."""

    DASH = (7, "dash", "Dashes.")
    """Dashes."""

    DOT_DASH = (9, "dotDash", "Alternating dots and dashes.")
    """Alternating dots and dashes."""

    DOT_DOT_DASH = (10, "dotDotDash", "An alternating dot-dot-dash pattern.")
    """An alternating dot-dot-dash pattern."""

    WAVY = (11, "wave", "A single wavy line.")
    """A single wavy line."""

    DOTTED_HEAVY = (20, "dottedHeavy", "Heavy dots.")
    """Heavy dots."""

    DASH_HEAVY = (23, "dashedHeavy", "Heavy dashes.")
    """Heavy dashes."""

    DOT_DASH_HEAVY = (25, "dashDotHeavy", "Alternating heavy dots and heavy dashes.")
    """Alternating heavy dots and heavy dashes."""

    DOT_DOT_DASH_HEAVY = (
        26,
        "dashDotDotHeavy",
        "An alternating heavy dot-dot-dash pattern.",
    )
    """An alternating heavy dot-dot-dash pattern."""

    WAVY_HEAVY = (27, "wavyHeavy", "A heavy wavy line.")
    """A heavy wavy line."""

    DASH_LONG = (39, "dashLong", "Long dashes.")
    """Long dashes."""

    WAVY_DOUBLE = (43, "wavyDouble", "A double wavy line.")
    """A double wavy line."""

    DASH_LONG_HEAVY = (55, "dashLongHeavy", "Long heavy dashes.")
    """Long heavy dashes."""


# __init__.py


# exceptions.py
"""Exceptions used with python-docx.

The base exception class is PythonDocxError.
"""


class PythonDocxError(Exception):
    """Generic error class."""


class InvalidSpanError(PythonDocxError):
    """Raised when an invalid merge region is specified in a request to merge table
    cells."""


class InvalidXmlError(PythonDocxError):
    """Raised when invalid XML is encountered, such as on attempt to access a missing
    required child element."""


# bmp.py
from .constants import MIME_TYPE
from .helpers import LITTLE_ENDIAN, StreamReader
from .image import BaseImageHeader


class Bmp(BaseImageHeader):
    """Image header parser for BMP images."""

    @classmethod
    def from_stream(cls, stream):
        """Return |Bmp| instance having header properties parsed from the BMP image in
        `stream`."""
        stream_rdr = StreamReader(stream, LITTLE_ENDIAN)

        px_width = stream_rdr.read_long(0x12)
        px_height = stream_rdr.read_long(0x16)

        horz_px_per_meter = stream_rdr.read_long(0x26)
        vert_px_per_meter = stream_rdr.read_long(0x2A)

        horz_dpi = cls._dpi(horz_px_per_meter)
        vert_dpi = cls._dpi(vert_px_per_meter)

        return cls(px_width, px_height, horz_dpi, vert_dpi)

    @property
    def content_type(self):
        """MIME content type for this image, unconditionally `image/bmp` for BMP
        images."""
        return MIME_TYPE.BMP

    @property
    def default_ext(self):
        """Default filename extension, always 'bmp' for BMP images."""
        return "bmp"

    @staticmethod
    def _dpi(px_per_meter):
        """Return the integer pixels per inch from `px_per_meter`, defaulting to 96 if
        `px_per_meter` is zero."""
        if px_per_meter == 0:
            return 96
        return int(round(px_per_meter * 0.0254))


# constants.py
"""Constants specific the the image sub-package."""


class JPEG_MARKER_CODE:
    """JPEG marker codes."""

    TEM = b"\x01"
    DHT = b"\xC4"
    DAC = b"\xCC"
    JPG = b"\xC8"

    SOF0 = b"\xC0"
    SOF1 = b"\xC1"
    SOF2 = b"\xC2"
    SOF3 = b"\xC3"
    SOF5 = b"\xC5"
    SOF6 = b"\xC6"
    SOF7 = b"\xC7"
    SOF9 = b"\xC9"
    SOFA = b"\xCA"
    SOFB = b"\xCB"
    SOFD = b"\xCD"
    SOFE = b"\xCE"
    SOFF = b"\xCF"

    RST0 = b"\xD0"
    RST1 = b"\xD1"
    RST2 = b"\xD2"
    RST3 = b"\xD3"
    RST4 = b"\xD4"
    RST5 = b"\xD5"
    RST6 = b"\xD6"
    RST7 = b"\xD7"

    SOI = b"\xD8"
    EOI = b"\xD9"
    SOS = b"\xDA"
    DQT = b"\xDB"  # Define Quantization Table(s)
    DNL = b"\xDC"
    DRI = b"\xDD"
    DHP = b"\xDE"
    EXP = b"\xDF"

    APP0 = b"\xE0"
    APP1 = b"\xE1"
    APP2 = b"\xE2"
    APP3 = b"\xE3"
    APP4 = b"\xE4"
    APP5 = b"\xE5"
    APP6 = b"\xE6"
    APP7 = b"\xE7"
    APP8 = b"\xE8"
    APP9 = b"\xE9"
    APPA = b"\xEA"
    APPB = b"\xEB"
    APPC = b"\xEC"
    APPD = b"\xED"
    APPE = b"\xEE"
    APPF = b"\xEF"

    STANDALONE_MARKERS = (TEM, SOI, EOI, RST0, RST1, RST2, RST3, RST4, RST5, RST6, RST7)

    SOF_MARKER_CODES = (
        SOF0,
        SOF1,
        SOF2,
        SOF3,
        SOF5,
        SOF6,
        SOF7,
        SOF9,
        SOFA,
        SOFB,
        SOFD,
        SOFE,
        SOFF,
    )

    marker_names = {
        b"\x00": "UNKNOWN",
        b"\xC0": "SOF0",
        b"\xC2": "SOF2",
        b"\xC4": "DHT",
        b"\xDA": "SOS",  # start of scan
        b"\xD8": "SOI",  # start of image
        b"\xD9": "EOI",  # end of image
        b"\xDB": "DQT",
        b"\xE0": "APP0",
        b"\xE1": "APP1",
        b"\xE2": "APP2",
        b"\xED": "APP13",
        b"\xEE": "APP14",
    }

    @classmethod
    def is_standalone(cls, marker_code):
        return marker_code in cls.STANDALONE_MARKERS


class MIME_TYPE:
    """Image content types."""

    BMP = "image/bmp"
    GIF = "image/gif"
    JPEG = "image/jpeg"
    PNG = "image/png"
    TIFF = "image/tiff"


class PNG_CHUNK_TYPE:
    """PNG chunk type names."""

    IHDR = "IHDR"
    pHYs = "pHYs"
    IEND = "IEND"


class TIFF_FLD_TYPE:
    """Tag codes for TIFF Image File Directory (IFD) entries."""

    BYTE = 1
    ASCII = 2
    SHORT = 3
    LONG = 4
    RATIONAL = 5

    field_type_names = {
        1: "BYTE",
        2: "ASCII char",
        3: "SHORT",
        4: "LONG",
        5: "RATIONAL",
    }


TIFF_FLD = TIFF_FLD_TYPE


class TIFF_TAG:
    """Tag codes for TIFF Image File Directory (IFD) entries."""

    IMAGE_WIDTH = 0x0100
    IMAGE_LENGTH = 0x0101
    X_RESOLUTION = 0x011A
    Y_RESOLUTION = 0x011B
    RESOLUTION_UNIT = 0x0128

    tag_names = {
        0x00FE: "NewSubfileType",
        0x0100: "ImageWidth",
        0x0101: "ImageLength",
        0x0102: "BitsPerSample",
        0x0103: "Compression",
        0x0106: "PhotometricInterpretation",
        0x010E: "ImageDescription",
        0x010F: "Make",
        0x0110: "Model",
        0x0111: "StripOffsets",
        0x0112: "Orientation",
        0x0115: "SamplesPerPixel",
        0x0117: "StripByteCounts",
        0x011A: "XResolution",
        0x011B: "YResolution",
        0x011C: "PlanarConfiguration",
        0x0128: "ResolutionUnit",
        0x0131: "Software",
        0x0132: "DateTime",
        0x0213: "YCbCrPositioning",
        0x8769: "ExifTag",
        0x8825: "GPS IFD",
        0xC4A5: "PrintImageMatching",
    }


# exceptions.py
"""Exceptions specific the the image sub-package."""


class InvalidImageStreamError(Exception):
    """The recognized image stream appears to be corrupted."""


class UnexpectedEndOfFileError(Exception):
    """EOF was unexpectedly encountered while reading an image stream."""


class UnrecognizedImageError(Exception):
    """The provided image stream could not be recognized."""


# gif.py
from struct import Struct

from .constants import MIME_TYPE
from .image import BaseImageHeader


class Gif(BaseImageHeader):
    """Image header parser for GIF images.

    Note that the GIF format does not support resolution (DPI) information. Both
    horizontal and vertical DPI default to 72.
    """

    @classmethod
    def from_stream(cls, stream):
        """Return |Gif| instance having header properties parsed from GIF image in
        `stream`."""
        px_width, px_height = cls._dimensions_from_stream(stream)
        return cls(px_width, px_height, 72, 72)

    @property
    def content_type(self):
        """MIME content type for this image, unconditionally `image/gif` for GIF
        images."""
        return MIME_TYPE.GIF

    @property
    def default_ext(self):
        """Default filename extension, always 'gif' for GIF images."""
        return "gif"

    @classmethod
    def _dimensions_from_stream(cls, stream):
        stream.seek(6)
        bytes_ = stream.read(4)
        struct = Struct("<HH")
        px_width, px_height = struct.unpack(bytes_)
        return px_width, px_height


# helpers.py
from struct import Struct

from .exceptions import UnexpectedEndOfFileError

BIG_ENDIAN = ">"
LITTLE_ENDIAN = "<"


class StreamReader:
    """Wraps a file-like object to provide access to structured data from a binary file.

    Byte-order is configurable. `base_offset` is added to any base value provided to
    calculate actual location for reads.
    """

    def __init__(self, stream, byte_order, base_offset=0):
        super(StreamReader, self).__init__()
        self._stream = stream
        self._byte_order = LITTLE_ENDIAN if byte_order == LITTLE_ENDIAN else BIG_ENDIAN
        self._base_offset = base_offset

    def read(self, count):
        """Allow pass-through read() call."""
        return self._stream.read(count)

    def read_byte(self, base, offset=0):
        """Return the int value of the byte at the file position defined by
        self._base_offset + `base` + `offset`.

        If `base` is None, the byte is read from the current position in the stream.
        """
        fmt = "B"
        return self._read_int(fmt, base, offset)

    def read_long(self, base, offset=0):
        """Return the int value of the four bytes at the file position defined by
        self._base_offset + `base` + `offset`.

        If `base` is None, the long is read from the current position in the stream. The
        endian setting of this instance is used to interpret the byte layout of the
        long.
        """
        fmt = "<L" if self._byte_order is LITTLE_ENDIAN else ">L"
        return self._read_int(fmt, base, offset)

    def read_short(self, base, offset=0):
        """Return the int value of the two bytes at the file position determined by
        `base` and `offset`, similarly to ``read_long()`` above."""
        fmt = b"<H" if self._byte_order is LITTLE_ENDIAN else b">H"
        return self._read_int(fmt, base, offset)

    def read_str(self, char_count, base, offset=0):
        """Return a string containing the `char_count` bytes at the file position
        determined by self._base_offset + `base` + `offset`."""

        def str_struct(char_count):
            format_ = "%ds" % char_count
            return Struct(format_)

        struct = str_struct(char_count)
        chars = self._unpack_item(struct, base, offset)
        unicode_str = chars.decode("UTF-8")
        return unicode_str

    def seek(self, base, offset=0):
        location = self._base_offset + base + offset
        self._stream.seek(location)

    def tell(self):
        """Allow pass-through tell() call."""
        return self._stream.tell()

    def _read_bytes(self, byte_count, base, offset):
        self.seek(base, offset)
        bytes_ = self._stream.read(byte_count)
        if len(bytes_) < byte_count:
            raise UnexpectedEndOfFileError
        return bytes_

    def _read_int(self, fmt, base, offset):
        struct = Struct(fmt)
        return self._unpack_item(struct, base, offset)

    def _unpack_item(self, struct, base, offset):
        bytes_ = self._read_bytes(struct.size, base, offset)
        return struct.unpack(bytes_)[0]


# image.py
"""Provides objects that can characterize image streams.

That characterization is as to content type and size, as a required step in including
them in a document.
"""

from __future__ import annotations

import hashlib
import io
import os
from typing import IO, Tuple

from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Emu, Inches, Length, lazyproperty


class Image:
    """Graphical image stream such as JPEG, PNG, or GIF with properties and methods
    required by ImagePart."""

    def __init__(self, blob: bytes, filename: str, image_header: BaseImageHeader):
        super(Image, self).__init__()
        self._blob = blob
        self._filename = filename
        self._image_header = image_header

    @classmethod
    def from_blob(cls, blob: bytes) -> Image:
        """Return a new |Image| subclass instance parsed from the image binary contained
        in `blob`."""
        stream = io.BytesIO(blob)
        return cls._from_stream(stream, blob)

    @classmethod
    def from_file(cls, image_descriptor: str | IO[bytes]):
        """Return a new |Image| subclass instance loaded from the image file identified
        by `image_descriptor`, a path or file-like object."""
        if isinstance(image_descriptor, str):
            path = image_descriptor
            with open(path, "rb") as f:
                blob = f.read()
                stream = io.BytesIO(blob)
            filename = os.path.basename(path)
        else:
            stream = image_descriptor
            stream.seek(0)
            blob = stream.read()
            filename = None
        return cls._from_stream(stream, blob, filename)

    @property
    def blob(self):
        """The bytes of the image 'file'."""
        return self._blob

    @property
    def content_type(self) -> str:
        """MIME content type for this image, e.g. ``'image/jpeg'`` for a JPEG image."""
        return self._image_header.content_type

    @lazyproperty
    def ext(self):
        """The file extension for the image.

        If an actual one is available from a load filename it is used. Otherwise a
        canonical extension is assigned based on the content type. Does not contain the
        leading period, e.g. 'jpg', not '.jpg'.
        """
        return os.path.splitext(self._filename)[1][1:]

    @property
    def filename(self):
        """Original image file name, if loaded from disk, or a generic filename if
        loaded from an anonymous stream."""
        return self._filename

    @property
    def px_width(self) -> int:
        """The horizontal pixel dimension of the image."""
        return self._image_header.px_width

    @property
    def px_height(self) -> int:
        """The vertical pixel dimension of the image."""
        return self._image_header.px_height

    @property
    def horz_dpi(self) -> int:
        """Integer dots per inch for the width of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        return self._image_header.horz_dpi

    @property
    def vert_dpi(self) -> int:
        """Integer dots per inch for the height of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        return self._image_header.vert_dpi

    @property
    def width(self) -> Inches:
        """A |Length| value representing the native width of the image, calculated from
        the values of `px_width` and `horz_dpi`."""
        return Inches(self.px_width / self.horz_dpi)

    @property
    def height(self) -> Inches:
        """A |Length| value representing the native height of the image, calculated from
        the values of `px_height` and `vert_dpi`."""
        return Inches(self.px_height / self.vert_dpi)

    def scaled_dimensions(
        self, width: int | Length | None = None, height: int | Length | None = None
    ) -> Tuple[Length, Length]:
        """(cx, cy) pair representing scaled dimensions of this image.

        The native dimensions of the image are scaled by applying the following rules to
        the `width` and `height` arguments.

        * If both `width` and `height` are specified, the return value is (`width`,
        `height`); no scaling is performed.
        * If only one is specified, it is used to compute a scaling factor that is then
        applied to the unspecified dimension, preserving the aspect ratio of the image.
        * If both `width` and `height` are |None|, the native dimensions are returned.

        The native dimensions are calculated using the dots-per-inch (dpi) value
        embedded in the image, defaulting to 72 dpi if no value is specified, as is
        often the case. The returned values are both |Length| objects.
        """
        if width is None and height is None:
            return self.width, self.height

        if width is None:
            assert height is not None
            scaling_factor = float(height) / float(self.height)
            width = round(self.width * scaling_factor)

        if height is None:
            scaling_factor = float(width) / float(self.width)
            height = round(self.height * scaling_factor)

        return Emu(width), Emu(height)

    @lazyproperty
    def sha1(self):
        """SHA1 hash digest of the image blob."""
        return hashlib.sha1(self._blob).hexdigest()

    @classmethod
    def _from_stream(
        cls,
        stream: IO[bytes],
        blob: bytes,
        filename: str | None = None,
    ) -> Image:
        """Return an instance of the |Image| subclass corresponding to the format of the
        image in `stream`."""
        image_header = _ImageHeaderFactory(stream)
        if filename is None:
            filename = "image.%s" % image_header.default_ext
        return cls(blob, filename, image_header)


def _ImageHeaderFactory(stream: IO[bytes]):
    """A |BaseImageHeader| subclass instance that can parse headers of image in `stream`."""
    from docx.image import SIGNATURES

    def read_32(stream: IO[bytes]):
        stream.seek(0)
        return stream.read(32)

    header = read_32(stream)
    for cls, offset, signature_bytes in SIGNATURES:
        end = offset + len(signature_bytes)
        found_bytes = header[offset:end]
        if found_bytes == signature_bytes:
            return cls.from_stream(stream)
    raise UnrecognizedImageError


class BaseImageHeader:
    """Base class for image header subclasses like |Jpeg| and |Tiff|."""

    def __init__(self, px_width: int, px_height: int, horz_dpi: int, vert_dpi: int):
        self._px_width = px_width
        self._px_height = px_height
        self._horz_dpi = horz_dpi
        self._vert_dpi = vert_dpi

    @property
    def content_type(self) -> str:
        """Abstract property definition, must be implemented by all subclasses."""
        msg = "content_type property must be implemented by all subclasses of " "BaseImageHeader"
        raise NotImplementedError(msg)

    @property
    def default_ext(self) -> str:
        """Default filename extension for images of this type.

        An abstract property definition, must be implemented by all subclasses.
        """
        raise NotImplementedError(
            "default_ext property must be implemented by all subclasses of " "BaseImageHeader"
        )

    @property
    def px_width(self):
        """The horizontal pixel dimension of the image."""
        return self._px_width

    @property
    def px_height(self):
        """The vertical pixel dimension of the image."""
        return self._px_height

    @property
    def horz_dpi(self):
        """Integer dots per inch for the width of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        return self._horz_dpi

    @property
    def vert_dpi(self):
        """Integer dots per inch for the height of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        return self._vert_dpi


# jpeg.py
"""Objects related to parsing headers of JPEG image streams.

Includes both JFIF and Exif sub-formats.
"""

import io

from docx.image.constants import JPEG_MARKER_CODE, MIME_TYPE
from docx.image.helpers import BIG_ENDIAN, StreamReader
from docx.image.image import BaseImageHeader
from docx.image.tiff import Tiff


class Jpeg(BaseImageHeader):
    """Base class for JFIF and EXIF subclasses."""

    @property
    def content_type(self):
        """MIME content type for this image, unconditionally `image/jpeg` for JPEG
        images."""
        return MIME_TYPE.JPEG

    @property
    def default_ext(self):
        """Default filename extension, always 'jpg' for JPG images."""
        return "jpg"


class Exif(Jpeg):
    """Image header parser for Exif image format."""

    @classmethod
    def from_stream(cls, stream):
        """Return |Exif| instance having header properties parsed from Exif image in
        `stream`."""
        markers = _JfifMarkers.from_stream(stream)
        # print('\n%s' % markers)

        px_width = markers.sof.px_width
        px_height = markers.sof.px_height
        horz_dpi = markers.app1.horz_dpi
        vert_dpi = markers.app1.vert_dpi

        return cls(px_width, px_height, horz_dpi, vert_dpi)


class Jfif(Jpeg):
    """Image header parser for JFIF image format."""

    @classmethod
    def from_stream(cls, stream):
        """Return a |Jfif| instance having header properties parsed from image in
        `stream`."""
        markers = _JfifMarkers.from_stream(stream)

        px_width = markers.sof.px_width
        px_height = markers.sof.px_height
        horz_dpi = markers.app0.horz_dpi
        vert_dpi = markers.app0.vert_dpi

        return cls(px_width, px_height, horz_dpi, vert_dpi)


class _JfifMarkers:
    """Sequence of markers in a JPEG file, perhaps truncated at first SOS marker for
    performance reasons."""

    def __init__(self, markers):
        super(_JfifMarkers, self).__init__()
        self._markers = list(markers)

    def __str__(self):  # pragma: no cover
        """Returns a tabular listing of the markers in this instance, which can be handy
        for debugging and perhaps other uses."""
        header = " offset  seglen  mc  name\n=======  ======  ==  ====="
        tmpl = "%7d  %6d  %02X  %s"
        rows = []
        for marker in self._markers:
            rows.append(
                tmpl
                % (
                    marker.offset,
                    marker.segment_length,
                    ord(marker.marker_code),
                    marker.name,
                )
            )
        lines = [header] + rows
        return "\n".join(lines)

    @classmethod
    def from_stream(cls, stream):
        """Return a |_JfifMarkers| instance containing a |_JfifMarker| subclass instance
        for each marker in `stream`."""
        marker_parser = _MarkerParser.from_stream(stream)
        markers = []
        for marker in marker_parser.iter_markers():
            markers.append(marker)
            if marker.marker_code == JPEG_MARKER_CODE.SOS:
                break
        return cls(markers)

    @property
    def app0(self):
        """First APP0 marker in image markers."""
        for m in self._markers:
            if m.marker_code == JPEG_MARKER_CODE.APP0:
                return m
        raise KeyError("no APP0 marker in image")

    @property
    def app1(self):
        """First APP1 marker in image markers."""
        for m in self._markers:
            if m.marker_code == JPEG_MARKER_CODE.APP1:
                return m
        raise KeyError("no APP1 marker in image")

    @property
    def sof(self):
        """First start of frame (SOFn) marker in this sequence."""
        for m in self._markers:
            if m.marker_code in JPEG_MARKER_CODE.SOF_MARKER_CODES:
                return m
        raise KeyError("no start of frame (SOFn) marker in image")


class _MarkerParser:
    """Service class that knows how to parse a JFIF stream and iterate over its
    markers."""

    def __init__(self, stream_reader):
        super(_MarkerParser, self).__init__()
        self._stream = stream_reader

    @classmethod
    def from_stream(cls, stream):
        """Return a |_MarkerParser| instance to parse JFIF markers from `stream`."""
        stream_reader = StreamReader(stream, BIG_ENDIAN)
        return cls(stream_reader)

    def iter_markers(self):
        """Generate a (marker_code, segment_offset) 2-tuple for each marker in the JPEG
        `stream`, in the order they occur in the stream."""
        marker_finder = _MarkerFinder.from_stream(self._stream)
        start = 0
        marker_code = None
        while marker_code != JPEG_MARKER_CODE.EOI:
            marker_code, segment_offset = marker_finder.next(start)
            marker = _MarkerFactory(marker_code, self._stream, segment_offset)
            yield marker
            start = segment_offset + marker.segment_length


class _MarkerFinder:
    """Service class that knows how to find the next JFIF marker in a stream."""

    def __init__(self, stream):
        super(_MarkerFinder, self).__init__()
        self._stream = stream

    @classmethod
    def from_stream(cls, stream):
        """Return a |_MarkerFinder| instance to find JFIF markers in `stream`."""
        return cls(stream)

    def next(self, start):
        """Return a (marker_code, segment_offset) 2-tuple identifying and locating the
        first marker in `stream` occuring after offset `start`.

        The returned `segment_offset` points to the position immediately following the
        2-byte marker code, the start of the marker segment, for those markers that have
        a segment.
        """
        position = start
        while True:
            # skip over any non-\xFF bytes
            position = self._offset_of_next_ff_byte(start=position)
            # skip over any \xFF padding bytes
            position, byte_ = self._next_non_ff_byte(start=position + 1)
            # 'FF 00' sequence is not a marker, start over if found
            if byte_ == b"\x00":
                continue
            # this is a marker, gather return values and break out of scan
            marker_code, segment_offset = byte_, position + 1
            break
        return marker_code, segment_offset

    def _next_non_ff_byte(self, start):
        """Return an offset, byte 2-tuple for the next byte in `stream` that is not
        '\xFF', starting with the byte at offset `start`.

        If the byte at offset `start` is not '\xFF', `start` and the returned `offset`
        will be the same.
        """
        self._stream.seek(start)
        byte_ = self._read_byte()
        while byte_ == b"\xFF":
            byte_ = self._read_byte()
        offset_of_non_ff_byte = self._stream.tell() - 1
        return offset_of_non_ff_byte, byte_

    def _offset_of_next_ff_byte(self, start):
        """Return the offset of the next '\xFF' byte in `stream` starting with the byte
        at offset `start`.

        Returns `start` if the byte at that offset is a hex 255; it does not necessarily
        advance in the stream.
        """
        self._stream.seek(start)
        byte_ = self._read_byte()
        while byte_ != b"\xFF":
            byte_ = self._read_byte()
        offset_of_ff_byte = self._stream.tell() - 1
        return offset_of_ff_byte

    def _read_byte(self):
        """Return the next byte read from stream.

        Raise Exception if stream is at end of file.
        """
        byte_ = self._stream.read(1)
        if not byte_:  # pragma: no cover
            raise Exception("unexpected end of file")
        return byte_


def _MarkerFactory(marker_code, stream, offset):
    """Return |_Marker| or subclass instance appropriate for marker at `offset` in
    `stream` having `marker_code`."""
    if marker_code == JPEG_MARKER_CODE.APP0:
        marker_cls = _App0Marker
    elif marker_code == JPEG_MARKER_CODE.APP1:
        marker_cls = _App1Marker
    elif marker_code in JPEG_MARKER_CODE.SOF_MARKER_CODES:
        marker_cls = _SofMarker
    else:
        marker_cls = _Marker
    return marker_cls.from_stream(stream, marker_code, offset)


class _Marker:
    """Base class for JFIF marker classes.

    Represents a marker and its segment occuring in a JPEG byte stream.
    """

    def __init__(self, marker_code, offset, segment_length):
        super(_Marker, self).__init__()
        self._marker_code = marker_code
        self._offset = offset
        self._segment_length = segment_length

    @classmethod
    def from_stream(cls, stream, marker_code, offset):
        """Return a generic |_Marker| instance for the marker at `offset` in `stream`
        having `marker_code`."""
        if JPEG_MARKER_CODE.is_standalone(marker_code):
            segment_length = 0
        else:
            segment_length = stream.read_short(offset)
        return cls(marker_code, offset, segment_length)

    @property
    def marker_code(self):
        """The single-byte code that identifies the type of this marker, e.g. ``'\xE0'``
        for start of image (SOI)."""
        return self._marker_code

    @property
    def name(self):  # pragma: no cover
        return JPEG_MARKER_CODE.marker_names[self._marker_code]

    @property
    def offset(self):  # pragma: no cover
        return self._offset

    @property
    def segment_length(self):
        """The length in bytes of this marker's segment."""
        return self._segment_length


class _App0Marker(_Marker):
    """Represents a JFIF APP0 marker segment."""

    def __init__(
        self, marker_code, offset, length, density_units, x_density, y_density
    ):
        super(_App0Marker, self).__init__(marker_code, offset, length)
        self._density_units = density_units
        self._x_density = x_density
        self._y_density = y_density

    @property
    def horz_dpi(self):
        """Horizontal dots per inch specified in this marker, defaults to 72 if not
        specified."""
        return self._dpi(self._x_density)

    @property
    def vert_dpi(self):
        """Vertical dots per inch specified in this marker, defaults to 72 if not
        specified."""
        return self._dpi(self._y_density)

    def _dpi(self, density):
        """Return dots per inch corresponding to `density` value."""
        if self._density_units == 1:
            dpi = density
        elif self._density_units == 2:
            dpi = int(round(density * 2.54))
        else:
            dpi = 72
        return dpi

    @classmethod
    def from_stream(cls, stream, marker_code, offset):
        """Return an |_App0Marker| instance for the APP0 marker at `offset` in
        `stream`."""
        # field               off  type   notes
        # ------------------  ---  -----  -------------------
        # segment length       0   short
        # JFIF identifier      2   5 chr  'JFIF\x00'
        # major JPEG version   7   byte   typically 1
        # minor JPEG version   8   byte   typically 1 or 2
        # density units        9   byte   1=inches, 2=cm
        # horz dots per unit  10   short
        # vert dots per unit  12   short
        # ------------------  ---  -----  -------------------
        segment_length = stream.read_short(offset)
        density_units = stream.read_byte(offset, 9)
        x_density = stream.read_short(offset, 10)
        y_density = stream.read_short(offset, 12)
        return cls(
            marker_code, offset, segment_length, density_units, x_density, y_density
        )


class _App1Marker(_Marker):
    """Represents a JFIF APP1 (Exif) marker segment."""

    def __init__(self, marker_code, offset, length, horz_dpi, vert_dpi):
        super(_App1Marker, self).__init__(marker_code, offset, length)
        self._horz_dpi = horz_dpi
        self._vert_dpi = vert_dpi

    @classmethod
    def from_stream(cls, stream, marker_code, offset):
        """Extract the horizontal and vertical dots-per-inch value from the APP1 header
        at `offset` in `stream`."""
        # field                 off  len  type   notes
        # --------------------  ---  ---  -----  ----------------------------
        # segment length         0    2   short
        # Exif identifier        2    6   6 chr  'Exif\x00\x00'
        # TIFF byte order        8    2   2 chr  'II'=little 'MM'=big endian
        # meaning of universe   10    2   2 chr  '*\x00' or '\x00*' depending
        # IFD0 off fr/II or MM  10   16   long   relative to ...?
        # --------------------  ---  ---  -----  ----------------------------
        segment_length = stream.read_short(offset)
        if cls._is_non_Exif_APP1_segment(stream, offset):
            return cls(marker_code, offset, segment_length, 72, 72)
        tiff = cls._tiff_from_exif_segment(stream, offset, segment_length)
        return cls(marker_code, offset, segment_length, tiff.horz_dpi, tiff.vert_dpi)

    @property
    def horz_dpi(self):
        """Horizontal dots per inch specified in this marker, defaults to 72 if not
        specified."""
        return self._horz_dpi

    @property
    def vert_dpi(self):
        """Vertical dots per inch specified in this marker, defaults to 72 if not
        specified."""
        return self._vert_dpi

    @classmethod
    def _is_non_Exif_APP1_segment(cls, stream, offset):
        """Return True if the APP1 segment at `offset` in `stream` is NOT an Exif
        segment, as determined by the ``'Exif\x00\x00'`` signature at offset 2 in the
        segment."""
        stream.seek(offset + 2)
        exif_signature = stream.read(6)
        return exif_signature != b"Exif\x00\x00"

    @classmethod
    def _tiff_from_exif_segment(cls, stream, offset, segment_length):
        """Return a |Tiff| instance parsed from the Exif APP1 segment of
        `segment_length` at `offset` in `stream`."""
        # wrap full segment in its own stream and feed to Tiff()
        stream.seek(offset + 8)
        segment_bytes = stream.read(segment_length - 8)
        substream = io.BytesIO(segment_bytes)
        return Tiff.from_stream(substream)


class _SofMarker(_Marker):
    """Represents a JFIF start of frame (SOFx) marker segment."""

    def __init__(self, marker_code, offset, segment_length, px_width, px_height):
        super(_SofMarker, self).__init__(marker_code, offset, segment_length)
        self._px_width = px_width
        self._px_height = px_height

    @classmethod
    def from_stream(cls, stream, marker_code, offset):
        """Return an |_SofMarker| instance for the SOFn marker at `offset` in stream."""
        # field                 off  type   notes
        # ------------------  ---  -----  ----------------------------
        # segment length       0   short
        # Data precision       2   byte
        # Vertical lines       3   short  px_height
        # Horizontal lines     5   short  px_width
        # ------------------  ---  -----  ----------------------------
        segment_length = stream.read_short(offset)
        px_height = stream.read_short(offset, 3)
        px_width = stream.read_short(offset, 5)
        return cls(marker_code, offset, segment_length, px_width, px_height)

    @property
    def px_height(self):
        """Image height in pixels."""
        return self._px_height

    @property
    def px_width(self):
        """Image width in pixels."""
        return self._px_width


# png.py
from .constants import MIME_TYPE, PNG_CHUNK_TYPE
from .exceptions import InvalidImageStreamError
from .helpers import BIG_ENDIAN, StreamReader
from .image import BaseImageHeader


class Png(BaseImageHeader):
    """Image header parser for PNG images."""

    @property
    def content_type(self):
        """MIME content type for this image, unconditionally `image/png` for PNG
        images."""
        return MIME_TYPE.PNG

    @property
    def default_ext(self):
        """Default filename extension, always 'png' for PNG images."""
        return "png"

    @classmethod
    def from_stream(cls, stream):
        """Return a |Png| instance having header properties parsed from image in
        `stream`."""
        parser = _PngParser.parse(stream)

        px_width = parser.px_width
        px_height = parser.px_height
        horz_dpi = parser.horz_dpi
        vert_dpi = parser.vert_dpi

        return cls(px_width, px_height, horz_dpi, vert_dpi)


class _PngParser:
    """Parses a PNG image stream to extract the image properties found in its chunks."""

    def __init__(self, chunks):
        super(_PngParser, self).__init__()
        self._chunks = chunks

    @classmethod
    def parse(cls, stream):
        """Return a |_PngParser| instance containing the header properties parsed from
        the PNG image in `stream`."""
        chunks = _Chunks.from_stream(stream)
        return cls(chunks)

    @property
    def px_width(self):
        """The number of pixels in each row of the image."""
        IHDR = self._chunks.IHDR
        return IHDR.px_width

    @property
    def px_height(self):
        """The number of stacked rows of pixels in the image."""
        IHDR = self._chunks.IHDR
        return IHDR.px_height

    @property
    def horz_dpi(self):
        """Integer dots per inch for the width of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        pHYs = self._chunks.pHYs
        if pHYs is None:
            return 72
        return self._dpi(pHYs.units_specifier, pHYs.horz_px_per_unit)

    @property
    def vert_dpi(self):
        """Integer dots per inch for the height of this image.

        Defaults to 72 when not present in the file, as is often the case.
        """
        pHYs = self._chunks.pHYs
        if pHYs is None:
            return 72
        return self._dpi(pHYs.units_specifier, pHYs.vert_px_per_unit)

    @staticmethod
    def _dpi(units_specifier, px_per_unit):
        """Return dots per inch value calculated from `units_specifier` and
        `px_per_unit`."""
        if units_specifier == 1 and px_per_unit:
            return int(round(px_per_unit * 0.0254))
        return 72


class _Chunks:
    """Collection of the chunks parsed from a PNG image stream."""

    def __init__(self, chunk_iterable):
        super(_Chunks, self).__init__()
        self._chunks = list(chunk_iterable)

    @classmethod
    def from_stream(cls, stream):
        """Return a |_Chunks| instance containing the PNG chunks in `stream`."""
        chunk_parser = _ChunkParser.from_stream(stream)
        chunks = list(chunk_parser.iter_chunks())
        return cls(chunks)

    @property
    def IHDR(self):
        """IHDR chunk in PNG image."""
        match = lambda chunk: chunk.type_name == PNG_CHUNK_TYPE.IHDR  # noqa
        IHDR = self._find_first(match)
        if IHDR is None:
            raise InvalidImageStreamError("no IHDR chunk in PNG image")
        return IHDR

    @property
    def pHYs(self):
        """PHYs chunk in PNG image, or |None| if not present."""
        match = lambda chunk: chunk.type_name == PNG_CHUNK_TYPE.pHYs  # noqa
        return self._find_first(match)

    def _find_first(self, match):
        """Return first chunk in stream order returning True for function `match`."""
        for chunk in self._chunks:
            if match(chunk):
                return chunk
        return None


class _ChunkParser:
    """Extracts chunks from a PNG image stream."""

    def __init__(self, stream_rdr):
        super(_ChunkParser, self).__init__()
        self._stream_rdr = stream_rdr

    @classmethod
    def from_stream(cls, stream):
        """Return a |_ChunkParser| instance that can extract the chunks from the PNG
        image in `stream`."""
        stream_rdr = StreamReader(stream, BIG_ENDIAN)
        return cls(stream_rdr)

    def iter_chunks(self):
        """Generate a |_Chunk| subclass instance for each chunk in this parser's PNG
        stream, in the order encountered in the stream."""
        for chunk_type, offset in self._iter_chunk_offsets():
            chunk = _ChunkFactory(chunk_type, self._stream_rdr, offset)
            yield chunk

    def _iter_chunk_offsets(self):
        """Generate a (chunk_type, chunk_offset) 2-tuple for each of the chunks in the
        PNG image stream.

        Iteration stops after the IEND chunk is returned.
        """
        chunk_offset = 8
        while True:
            chunk_data_len = self._stream_rdr.read_long(chunk_offset)
            chunk_type = self._stream_rdr.read_str(4, chunk_offset, 4)
            data_offset = chunk_offset + 8
            yield chunk_type, data_offset
            if chunk_type == "IEND":
                break
            # incr offset for chunk len long, chunk type, chunk data, and CRC
            chunk_offset += 4 + 4 + chunk_data_len + 4


def _ChunkFactory(chunk_type, stream_rdr, offset):
    """Return a |_Chunk| subclass instance appropriate to `chunk_type` parsed from
    `stream_rdr` at `offset`."""
    chunk_cls_map = {
        PNG_CHUNK_TYPE.IHDR: _IHDRChunk,
        PNG_CHUNK_TYPE.pHYs: _pHYsChunk,
    }
    chunk_cls = chunk_cls_map.get(chunk_type, _Chunk)
    return chunk_cls.from_offset(chunk_type, stream_rdr, offset)


class _Chunk:
    """Base class for specific chunk types.

    Also serves as the default chunk type.
    """

    def __init__(self, chunk_type):
        super(_Chunk, self).__init__()
        self._chunk_type = chunk_type

    @classmethod
    def from_offset(cls, chunk_type, stream_rdr, offset):
        """Return a default _Chunk instance that only knows its chunk type."""
        return cls(chunk_type)

    @property
    def type_name(self):
        """The chunk type name, e.g. 'IHDR', 'pHYs', etc."""
        return self._chunk_type


class _IHDRChunk(_Chunk):
    """IHDR chunk, contains the image dimensions."""

    def __init__(self, chunk_type, px_width, px_height):
        super(_IHDRChunk, self).__init__(chunk_type)
        self._px_width = px_width
        self._px_height = px_height

    @classmethod
    def from_offset(cls, chunk_type, stream_rdr, offset):
        """Return an _IHDRChunk instance containing the image dimensions extracted from
        the IHDR chunk in `stream` at `offset`."""
        px_width = stream_rdr.read_long(offset)
        px_height = stream_rdr.read_long(offset, 4)
        return cls(chunk_type, px_width, px_height)

    @property
    def px_width(self):
        return self._px_width

    @property
    def px_height(self):
        return self._px_height


class _pHYsChunk(_Chunk):
    """PYHs chunk, contains the image dpi information."""

    def __init__(self, chunk_type, horz_px_per_unit, vert_px_per_unit, units_specifier):
        super(_pHYsChunk, self).__init__(chunk_type)
        self._horz_px_per_unit = horz_px_per_unit
        self._vert_px_per_unit = vert_px_per_unit
        self._units_specifier = units_specifier

    @classmethod
    def from_offset(cls, chunk_type, stream_rdr, offset):
        """Return a _pHYsChunk instance containing the image resolution extracted from
        the pHYs chunk in `stream` at `offset`."""
        horz_px_per_unit = stream_rdr.read_long(offset)
        vert_px_per_unit = stream_rdr.read_long(offset, 4)
        units_specifier = stream_rdr.read_byte(offset, 8)
        return cls(chunk_type, horz_px_per_unit, vert_px_per_unit, units_specifier)

    @property
    def horz_px_per_unit(self):
        return self._horz_px_per_unit

    @property
    def vert_px_per_unit(self):
        return self._vert_px_per_unit

    @property
    def units_specifier(self):
        return self._units_specifier


# tiff.py
from .constants import MIME_TYPE, TIFF_FLD, TIFF_TAG
from .helpers import BIG_ENDIAN, LITTLE_ENDIAN, StreamReader
from .image import BaseImageHeader


class Tiff(BaseImageHeader):
    """Image header parser for TIFF images.

    Handles both big and little endian byte ordering.
    """

    @property
    def content_type(self):
        """Return the MIME type of this TIFF image, unconditionally the string
        ``image/tiff``."""
        return MIME_TYPE.TIFF

    @property
    def default_ext(self):
        """Default filename extension, always 'tiff' for TIFF images."""
        return "tiff"

    @classmethod
    def from_stream(cls, stream):
        """Return a |Tiff| instance containing the properties of the TIFF image in
        `stream`."""
        parser = _TiffParser.parse(stream)

        px_width = parser.px_width
        px_height = parser.px_height
        horz_dpi = parser.horz_dpi
        vert_dpi = parser.vert_dpi

        return cls(px_width, px_height, horz_dpi, vert_dpi)


class _TiffParser:
    """Parses a TIFF image stream to extract the image properties found in its main
    image file directory (IFD)"""

    def __init__(self, ifd_entries):
        super(_TiffParser, self).__init__()
        self._ifd_entries = ifd_entries

    @classmethod
    def parse(cls, stream):
        """Return an instance of |_TiffParser| containing the properties parsed from the
        TIFF image in `stream`."""
        stream_rdr = cls._make_stream_reader(stream)
        ifd0_offset = stream_rdr.read_long(4)
        ifd_entries = _IfdEntries.from_stream(stream_rdr, ifd0_offset)
        return cls(ifd_entries)

    @property
    def horz_dpi(self):
        """The horizontal dots per inch value calculated from the XResolution and
        ResolutionUnit tags of the IFD; defaults to 72 if those tags are not present."""
        return self._dpi(TIFF_TAG.X_RESOLUTION)

    @property
    def vert_dpi(self):
        """The vertical dots per inch value calculated from the XResolution and
        ResolutionUnit tags of the IFD; defaults to 72 if those tags are not present."""
        return self._dpi(TIFF_TAG.Y_RESOLUTION)

    @property
    def px_height(self):
        """The number of stacked rows of pixels in the image, |None| if the IFD contains
        no ``ImageLength`` tag, the expected case when the TIFF is embeded in an Exif
        image."""
        return self._ifd_entries.get(TIFF_TAG.IMAGE_LENGTH)

    @property
    def px_width(self):
        """The number of pixels in each row in the image, |None| if the IFD contains no
        ``ImageWidth`` tag, the expected case when the TIFF is embeded in an Exif
        image."""
        return self._ifd_entries.get(TIFF_TAG.IMAGE_WIDTH)

    @classmethod
    def _detect_endian(cls, stream):
        """Return either BIG_ENDIAN or LITTLE_ENDIAN depending on the endian indicator
        found in the TIFF `stream` header, either 'MM' or 'II'."""
        stream.seek(0)
        endian_str = stream.read(2)
        return BIG_ENDIAN if endian_str == b"MM" else LITTLE_ENDIAN

    def _dpi(self, resolution_tag):
        """Return the dpi value calculated for `resolution_tag`, which can be either
        TIFF_TAG.X_RESOLUTION or TIFF_TAG.Y_RESOLUTION.

        The calculation is based on the values of both that tag and the
        TIFF_TAG.RESOLUTION_UNIT tag in this parser's |_IfdEntries| instance.
        """
        ifd_entries = self._ifd_entries

        if resolution_tag not in ifd_entries:
            return 72

        # resolution unit defaults to inches (2)
        resolution_unit = ifd_entries.get(TIFF_TAG.RESOLUTION_UNIT, 2)

        if resolution_unit == 1:  # aspect ratio only
            return 72
        # resolution_unit == 2 for inches, 3 for centimeters
        units_per_inch = 1 if resolution_unit == 2 else 2.54
        dots_per_unit = ifd_entries[resolution_tag]
        return int(round(dots_per_unit * units_per_inch))

    @classmethod
    def _make_stream_reader(cls, stream):
        """Return a |StreamReader| instance with wrapping `stream` and having "endian-
        ness" determined by the 'MM' or 'II' indicator in the TIFF stream header."""
        endian = cls._detect_endian(stream)
        return StreamReader(stream, endian)


class _IfdEntries:
    """Image File Directory for a TIFF image, having mapping (dict) semantics allowing
    "tag" values to be retrieved by tag code."""

    def __init__(self, entries):
        super(_IfdEntries, self).__init__()
        self._entries = entries

    def __contains__(self, key):
        """Provides ``in`` operator, e.g. ``tag in ifd_entries``"""
        return self._entries.__contains__(key)

    def __getitem__(self, key):
        """Provides indexed access, e.g. ``tag_value = ifd_entries[tag_code]``"""
        return self._entries.__getitem__(key)

    @classmethod
    def from_stream(cls, stream, offset):
        """Return a new |_IfdEntries| instance parsed from `stream` starting at
        `offset`."""
        ifd_parser = _IfdParser(stream, offset)
        entries = {e.tag: e.value for e in ifd_parser.iter_entries()}
        return cls(entries)

    def get(self, tag_code, default=None):
        """Return value of IFD entry having tag matching `tag_code`, or `default` if no
        matching tag found."""
        return self._entries.get(tag_code, default)


class _IfdParser:
    """Service object that knows how to extract directory entries from an Image File
    Directory (IFD)"""

    def __init__(self, stream_rdr, offset):
        super(_IfdParser, self).__init__()
        self._stream_rdr = stream_rdr
        self._offset = offset

    def iter_entries(self):
        """Generate an |_IfdEntry| instance corresponding to each entry in the
        directory."""
        for idx in range(self._entry_count):
            dir_entry_offset = self._offset + 2 + (idx * 12)
            ifd_entry = _IfdEntryFactory(self._stream_rdr, dir_entry_offset)
            yield ifd_entry

    @property
    def _entry_count(self):
        """The count of directory entries, read from the top of the IFD header."""
        return self._stream_rdr.read_short(self._offset)


def _IfdEntryFactory(stream_rdr, offset):
    """Return an |_IfdEntry| subclass instance containing the value of the directory
    entry at `offset` in `stream_rdr`."""
    ifd_entry_classes = {
        TIFF_FLD.ASCII: _AsciiIfdEntry,
        TIFF_FLD.SHORT: _ShortIfdEntry,
        TIFF_FLD.LONG: _LongIfdEntry,
        TIFF_FLD.RATIONAL: _RationalIfdEntry,
    }
    field_type = stream_rdr.read_short(offset, 2)
    EntryCls = ifd_entry_classes.get(field_type, _IfdEntry)
    return EntryCls.from_stream(stream_rdr, offset)


class _IfdEntry:
    """Base class for IFD entry classes.

    Subclasses are differentiated by value type, e.g. ASCII, long int, etc.
    """

    def __init__(self, tag_code, value):
        super(_IfdEntry, self).__init__()
        self._tag_code = tag_code
        self._value = value

    @classmethod
    def from_stream(cls, stream_rdr, offset):
        """Return an |_IfdEntry| subclass instance containing the tag and value of the
        tag parsed from `stream_rdr` at `offset`.

        Note this method is common to all subclasses. Override the ``_parse_value()``
        method to provide distinctive behavior based on field type.
        """
        tag_code = stream_rdr.read_short(offset, 0)
        value_count = stream_rdr.read_long(offset, 4)
        value_offset = stream_rdr.read_long(offset, 8)
        value = cls._parse_value(stream_rdr, offset, value_count, value_offset)
        return cls(tag_code, value)

    @classmethod
    def _parse_value(cls, stream_rdr, offset, value_count, value_offset):
        """Return the value of this field parsed from `stream_rdr` at `offset`.

        Intended to be overridden by subclasses.
        """
        return "UNIMPLEMENTED FIELD TYPE"  # pragma: no cover

    @property
    def tag(self):
        """Short int code that identifies this IFD entry."""
        return self._tag_code

    @property
    def value(self):
        """Value of this tag, its type being dependent on the tag."""
        return self._value


class _AsciiIfdEntry(_IfdEntry):
    """IFD entry having the form of a NULL-terminated ASCII string."""

    @classmethod
    def _parse_value(cls, stream_rdr, offset, value_count, value_offset):
        """Return the ASCII string parsed from `stream_rdr` at `value_offset`.

        The length of the string, including a terminating '\x00' (NUL) character, is in
        `value_count`.
        """
        return stream_rdr.read_str(value_count - 1, value_offset)


class _ShortIfdEntry(_IfdEntry):
    """IFD entry expressed as a short (2-byte) integer."""

    @classmethod
    def _parse_value(cls, stream_rdr, offset, value_count, value_offset):
        """Return the short int value contained in the `value_offset` field of this
        entry.

        Only supports single values at present.
        """
        if value_count == 1:
            return stream_rdr.read_short(offset, 8)
        else:  # pragma: no cover
            return "Multi-value short integer NOT IMPLEMENTED"


class _LongIfdEntry(_IfdEntry):
    """IFD entry expressed as a long (4-byte) integer."""

    @classmethod
    def _parse_value(cls, stream_rdr, offset, value_count, value_offset):
        """Return the long int value contained in the `value_offset` field of this
        entry.

        Only supports single values at present.
        """
        if value_count == 1:
            return stream_rdr.read_long(offset, 8)
        else:  # pragma: no cover
            return "Multi-value long integer NOT IMPLEMENTED"


class _RationalIfdEntry(_IfdEntry):
    """IFD entry expressed as a numerator, denominator pair."""

    @classmethod
    def _parse_value(cls, stream_rdr, offset, value_count, value_offset):
        """Return the rational (numerator / denominator) value at `value_offset` in
        `stream_rdr` as a floating-point number.

        Only supports single values at present.
        """
        if value_count == 1:
            numerator = stream_rdr.read_long(value_offset)
            denominator = stream_rdr.read_long(value_offset, 4)
            return numerator / denominator
        else:  # pragma: no cover
            return "Multi-value Rational NOT IMPLEMENTED"


# __init__.py
"""Provides objects that can characterize image streams.

That characterization is as to content type and size, as a required step in including
them in a document.
"""

from docx.image.bmp import Bmp
from docx.image.gif import Gif
from docx.image.jpeg import Exif, Jfif
from docx.image.png import Png
from docx.image.tiff import Tiff

SIGNATURES = (
    # class, offset, signature_bytes
    (Png, 0, b"\x89PNG\x0D\x0A\x1A\x0A"),
    (Jfif, 6, b"JFIF"),
    (Exif, 6, b"Exif"),
    (Gif, 0, b"GIF87a"),
    (Gif, 0, b"GIF89a"),
    (Tiff, 0, b"MM\x00*"),  # big-endian (Motorola) TIFF
    (Tiff, 0, b"II*\x00"),  # little-endian (Intel) TIFF
    (Bmp, 0, b"BM"),
)


# constants.py
"""Constant values related to the Open Packaging Convention.

In particular it includes content types and relationship types.
"""


class CONTENT_TYPE:
    """Content type URIs (like MIME-types) that specify a part's format."""

    BMP = "image/bmp"
    DML_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
    DML_CHARTSHAPES = (
        "application/vnd.openxmlformats-officedocument.drawingml.chartshapes+xml"
    )
    DML_DIAGRAM_COLORS = (
        "application/vnd.openxmlformats-officedocument.drawingml.diagramColors+xml"
    )
    DML_DIAGRAM_DATA = (
        "application/vnd.openxmlformats-officedocument.drawingml.diagramData+xml"
    )
    DML_DIAGRAM_LAYOUT = (
        "application/vnd.openxmlformats-officedocument.drawingml.diagramLayout+xml"
    )
    DML_DIAGRAM_STYLE = (
        "application/vnd.openxmlformats-officedocument.drawingml.diagramStyle+xml"
    )
    GIF = "image/gif"
    JPEG = "image/jpeg"
    MS_PHOTO = "image/vnd.ms-photo"
    OFC_CUSTOM_PROPERTIES = (
        "application/vnd.openxmlformats-officedocument.custom-properties+xml"
    )
    OFC_CUSTOM_XML_PROPERTIES = (
        "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"
    )
    OFC_DRAWING = "application/vnd.openxmlformats-officedocument.drawing+xml"
    OFC_EXTENDED_PROPERTIES = (
        "application/vnd.openxmlformats-officedocument.extended-properties+xml"
    )
    OFC_OLE_OBJECT = "application/vnd.openxmlformats-officedocument.oleObject"
    OFC_PACKAGE = "application/vnd.openxmlformats-officedocument.package"
    OFC_THEME = "application/vnd.openxmlformats-officedocument.theme+xml"
    OFC_THEME_OVERRIDE = (
        "application/vnd.openxmlformats-officedocument.themeOverride+xml"
    )
    OFC_VML_DRAWING = "application/vnd.openxmlformats-officedocument.vmlDrawing"
    OPC_CORE_PROPERTIES = "application/vnd.openxmlformats-package.core-properties+xml"
    OPC_DIGITAL_SIGNATURE_CERTIFICATE = (
        "application/vnd.openxmlformats-package.digital-signature-certificate"
    )
    OPC_DIGITAL_SIGNATURE_ORIGIN = (
        "application/vnd.openxmlformats-package.digital-signature-origin"
    )
    OPC_DIGITAL_SIGNATURE_XMLSIGNATURE = (
        "application/vnd.openxmlformats-package.digital-signature-xmlsignature+xml"
    )
    OPC_RELATIONSHIPS = "application/vnd.openxmlformats-package.relationships+xml"
    PML_COMMENTS = (
        "application/vnd.openxmlformats-officedocument.presentationml.comments+xml"
    )
    PML_COMMENT_AUTHORS = (
        "application/vnd.openxmlformats-officedocument.presentationml.commen"
        "tAuthors+xml"
    )
    PML_HANDOUT_MASTER = (
        "application/vnd.openxmlformats-officedocument.presentationml.handou"
        "tMaster+xml"
    )
    PML_NOTES_MASTER = (
        "application/vnd.openxmlformats-officedocument.presentationml.notesM"
        "aster+xml"
    )
    PML_NOTES_SLIDE = (
        "application/vnd.openxmlformats-officedocument.presentationml.notesSlide+xml"
    )
    PML_PRESENTATION_MAIN = (
        "application/vnd.openxmlformats-officedocument.presentationml.presen"
        "tation.main+xml"
    )
    PML_PRES_PROPS = (
        "application/vnd.openxmlformats-officedocument.presentationml.presProps+xml"
    )
    PML_PRINTER_SETTINGS = (
        "application/vnd.openxmlformats-officedocument.presentationml.printe"
        "rSettings"
    )
    PML_SLIDE = "application/vnd.openxmlformats-officedocument.presentationml.slide+xml"
    PML_SLIDESHOW_MAIN = (
        "application/vnd.openxmlformats-officedocument.presentationml.slides"
        "how.main+xml"
    )
    PML_SLIDE_LAYOUT = (
        "application/vnd.openxmlformats-officedocument.presentationml.slideL"
        "ayout+xml"
    )
    PML_SLIDE_MASTER = (
        "application/vnd.openxmlformats-officedocument.presentationml.slideM"
        "aster+xml"
    )
    PML_SLIDE_UPDATE_INFO = (
        "application/vnd.openxmlformats-officedocument.presentationml.slideU"
        "pdateInfo+xml"
    )
    PML_TABLE_STYLES = (
        "application/vnd.openxmlformats-officedocument.presentationml.tableS"
        "tyles+xml"
    )
    PML_TAGS = "application/vnd.openxmlformats-officedocument.presentationml.tags+xml"
    PML_TEMPLATE_MAIN = (
        "application/vnd.openxmlformats-officedocument.presentationml.templa"
        "te.main+xml"
    )
    PML_VIEW_PROPS = (
        "application/vnd.openxmlformats-officedocument.presentationml.viewProps+xml"
    )
    PNG = "image/png"
    SML_CALC_CHAIN = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.calcChain+xml"
    )
    SML_CHARTSHEET = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.chartsheet+xml"
    )
    SML_COMMENTS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.comments+xml"
    )
    SML_CONNECTIONS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.connections+xml"
    )
    SML_CUSTOM_PROPERTY = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.customProperty"
    )
    SML_DIALOGSHEET = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.dialogsheet+xml"
    )
    SML_EXTERNAL_LINK = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.externa"
        "lLink+xml"
    )
    SML_PIVOT_CACHE_DEFINITION = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.pivotCa"
        "cheDefinition+xml"
    )
    SML_PIVOT_CACHE_RECORDS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.pivotCa"
        "cheRecords+xml"
    )
    SML_PIVOT_TABLE = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.pivotTable+xml"
    )
    SML_PRINTER_SETTINGS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.printerSettings"
    )
    SML_QUERY_TABLE = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.queryTable+xml"
    )
    SML_REVISION_HEADERS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.revisio"
        "nHeaders+xml"
    )
    SML_REVISION_LOG = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.revisionLog+xml"
    )
    SML_SHARED_STRINGS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sharedS"
        "trings+xml"
    )
    SML_SHEET = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    SML_SHEET_MAIN = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"
    )
    SML_SHEET_METADATA = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheetMe"
        "tadata+xml"
    )
    SML_STYLES = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"
    )
    SML_TABLE = "application/vnd.openxmlformats-officedocument.spreadsheetml.table+xml"
    SML_TABLE_SINGLE_CELLS = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.tableSi"
        "ngleCells+xml"
    )
    SML_TEMPLATE_MAIN = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.templat"
        "e.main+xml"
    )
    SML_USER_NAMES = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.userNames+xml"
    )
    SML_VOLATILE_DEPENDENCIES = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.volatil"
        "eDependencies+xml"
    )
    SML_WORKSHEET = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"
    )
    TIFF = "image/tiff"
    WML_COMMENTS = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    )
    WML_DOCUMENT = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    WML_DOCUMENT_GLOSSARY = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.docu"
        "ment.glossary+xml"
    )
    WML_DOCUMENT_MAIN = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.docu"
        "ment.main+xml"
    )
    WML_ENDNOTES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
    )
    WML_FONT_TABLE = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.font"
        "Table+xml"
    )
    WML_FOOTER = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
    )
    WML_FOOTNOTES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.foot"
        "notes+xml"
    )
    WML_HEADER = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
    )
    WML_NUMBERING = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.numb"
        "ering+xml"
    )
    WML_PRINTER_SETTINGS = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.prin"
        "terSettings"
    )
    WML_SETTINGS = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"
    )
    WML_STYLES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"
    )
    WML_WEB_SETTINGS = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.webS"
        "ettings+xml"
    )
    XML = "application/xml"
    X_EMF = "image/x-emf"
    X_FONTDATA = "application/x-fontdata"
    X_FONT_TTF = "application/x-font-ttf"
    X_WMF = "image/x-wmf"


class NAMESPACE:
    """Constant values for OPC XML namespaces."""

    DML_WORDPROCESSING_DRAWING = (
        "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    )
    OFC_RELATIONSHIPS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    )
    OPC_RELATIONSHIPS = "http://schemas.openxmlformats.org/package/2006/relationships"
    OPC_CONTENT_TYPES = "http://schemas.openxmlformats.org/package/2006/content-types"
    WML_MAIN = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class RELATIONSHIP_TARGET_MODE:
    """Open XML relationship target modes."""

    EXTERNAL = "External"
    INTERNAL = "Internal"


class RELATIONSHIP_TYPE:
    AUDIO = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/audio"
    A_F_CHUNK = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk"
    )
    CALC_CHAIN = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/calcChain"
    )
    CERTIFICATE = (
        "http://schemas.openxmlformats.org/package/2006/relationships/digita"
        "l-signature/certificate"
    )
    CHART = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
    CHARTSHEET = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/chartsheet"
    )
    CHART_USER_SHAPES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/chartUserShapes"
    )
    COMMENTS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/comments"
    )
    COMMENT_AUTHORS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/commentAuthors"
    )
    CONNECTIONS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/connections"
    )
    CONTROL = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/control"
    )
    CORE_PROPERTIES = (
        "http://schemas.openxmlformats.org/package/2006/relationships/metada"
        "ta/core-properties"
    )
    CUSTOM_PROPERTIES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/custom-properties"
    )
    CUSTOM_PROPERTY = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/customProperty"
    )
    CUSTOM_XML = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/customXml"
    )
    CUSTOM_XML_PROPS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/customXmlProps"
    )
    DIAGRAM_COLORS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/diagramColors"
    )
    DIAGRAM_DATA = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/diagramData"
    )
    DIAGRAM_LAYOUT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/diagramLayout"
    )
    DIAGRAM_QUICK_STYLE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/diagramQuickStyle"
    )
    DIALOGSHEET = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/dialogsheet"
    )
    DRAWING = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing"
    )
    ENDNOTES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/endnotes"
    )
    EXTENDED_PROPERTIES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/extended-properties"
    )
    EXTERNAL_LINK = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/externalLink"
    )
    FONT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
    FONT_TABLE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/fontTable"
    )
    FOOTER = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
    )
    FOOTNOTES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/footnotes"
    )
    GLOSSARY_DOCUMENT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/glossaryDocument"
    )
    HANDOUT_MASTER = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/handoutMaster"
    )
    HEADER = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
    )
    HYPERLINK = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/hyperlink"
    )
    IMAGE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    NOTES_MASTER = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/notesMaster"
    )
    NOTES_SLIDE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/notesSlide"
    )
    NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/numbering"
    )
    OFFICE_DOCUMENT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/officeDocument"
    )
    OLE_OBJECT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/oleObject"
    )
    ORIGIN = (
        "http://schemas.openxmlformats.org/package/2006/relationships/digita"
        "l-signature/origin"
    )
    PACKAGE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"
    )
    PIVOT_CACHE_DEFINITION = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/pivotCacheDefinition"
    )
    PIVOT_CACHE_RECORDS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/spreadsheetml/pivotCacheRecords"
    )
    PIVOT_TABLE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/pivotTable"
    )
    PRES_PROPS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/presProps"
    )
    PRINTER_SETTINGS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/printerSettings"
    )
    QUERY_TABLE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/queryTable"
    )
    REVISION_HEADERS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/revisionHeaders"
    )
    REVISION_LOG = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/revisionLog"
    )
    SETTINGS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/settings"
    )
    SHARED_STRINGS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/sharedStrings"
    )
    SHEET_METADATA = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/sheetMetadata"
    )
    SIGNATURE = (
        "http://schemas.openxmlformats.org/package/2006/relationships/digita"
        "l-signature/signature"
    )
    SLIDE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"
    SLIDE_LAYOUT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/slideLayout"
    )
    SLIDE_MASTER = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/slideMaster"
    )
    SLIDE_UPDATE_INFO = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/slideUpdateInfo"
    )
    STYLES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"
    )
    TABLE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/table"
    TABLE_SINGLE_CELLS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/tableSingleCells"
    )
    TABLE_STYLES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/tableStyles"
    )
    TAGS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/tags"
    THEME = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
    THEME_OVERRIDE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/themeOverride"
    )
    THUMBNAIL = (
        "http://schemas.openxmlformats.org/package/2006/relationships/metada"
        "ta/thumbnail"
    )
    USERNAMES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/usernames"
    )
    VIDEO = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/video"
    VIEW_PROPS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/viewProps"
    )
    VML_DRAWING = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/vmlDrawing"
    )
    VOLATILE_DEPENDENCIES = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/volatileDependencies"
    )
    WEB_SETTINGS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/webSettings"
    )
    WORKSHEET_SOURCE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        "/worksheetSource"
    )
    XML_MAPS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/xmlMaps"
    )


# coreprops.py
"""Provides CoreProperties, Dublin-Core attributes of the document.

These are broadly-standardized attributes like author, last-modified, etc.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.oxml.coreprops import CT_CoreProperties

if TYPE_CHECKING:
    from docx.oxml.coreprops import CT_CoreProperties


class CoreProperties:
    """Corresponds to part named ``/docProps/core.xml``, containing the core document
    properties for this document package."""

    def __init__(self, element: CT_CoreProperties):
        self._element = element

    @property
    def author(self):
        return self._element.author_text

    @author.setter
    def author(self, value: str):
        self._element.author_text = value

    @property
    def category(self):
        return self._element.category_text

    @category.setter
    def category(self, value: str):
        self._element.category_text = value

    @property
    def comments(self):
        return self._element.comments_text

    @comments.setter
    def comments(self, value: str):
        self._element.comments_text = value

    @property
    def content_status(self):
        return self._element.contentStatus_text

    @content_status.setter
    def content_status(self, value: str):
        self._element.contentStatus_text = value

    @property
    def created(self):
        return self._element.created_datetime

    @created.setter
    def created(self, value):
        self._element.created_datetime = value

    @property
    def identifier(self):
        return self._element.identifier_text

    @identifier.setter
    def identifier(self, value: str):
        self._element.identifier_text = value

    @property
    def keywords(self):
        return self._element.keywords_text

    @keywords.setter
    def keywords(self, value: str):
        self._element.keywords_text = value

    @property
    def language(self):
        return self._element.language_text

    @language.setter
    def language(self, value: str):
        self._element.language_text = value

    @property
    def last_modified_by(self):
        return self._element.lastModifiedBy_text

    @last_modified_by.setter
    def last_modified_by(self, value: str):
        self._element.lastModifiedBy_text = value

    @property
    def last_printed(self):
        return self._element.lastPrinted_datetime

    @last_printed.setter
    def last_printed(self, value):
        self._element.lastPrinted_datetime = value

    @property
    def modified(self):
        return self._element.modified_datetime

    @modified.setter
    def modified(self, value):
        self._element.modified_datetime = value

    @property
    def revision(self):
        return self._element.revision_number

    @revision.setter
    def revision(self, value):
        self._element.revision_number = value

    @property
    def subject(self):
        return self._element.subject_text

    @subject.setter
    def subject(self, value: str):
        self._element.subject_text = value

    @property
    def title(self):
        return self._element.title_text

    @title.setter
    def title(self, value: str):
        self._element.title_text = value

    @property
    def version(self):
        return self._element.version_text

    @version.setter
    def version(self, value: str):
        self._element.version_text = value


# exceptions.py
"""Exceptions specific to python-opc.

The base exception class is OpcError.
"""


class OpcError(Exception):
    """Base error class for python-opc."""


class PackageNotFoundError(OpcError):
    """Raised when a package cannot be found at the specified path."""


# oxml.py
# pyright: reportPrivateUsage=false

"""Temporary stand-in for main oxml module.

This module came across with the PackageReader transplant. Probably much will get
replaced with objects from the pptx.oxml.core and then this module will either get
deleted or only hold the package related custom element classes.
"""

from __future__ import annotations

from typing import cast

from lxml import etree

from docx.opc.constants import NAMESPACE as NS
from docx.opc.constants import RELATIONSHIP_TARGET_MODE as RTM

# configure XML parser
element_class_lookup = etree.ElementNamespaceClassLookup()
oxml_parser = etree.XMLParser(remove_blank_text=True, resolve_entities=False)
oxml_parser.set_element_class_lookup(element_class_lookup)

nsmap = {
    "ct": NS.OPC_CONTENT_TYPES,
    "pr": NS.OPC_RELATIONSHIPS,
    "r": NS.OFC_RELATIONSHIPS,
}


# ===========================================================================
# functions
# ===========================================================================


def parse_xml(text: str) -> etree._Element:
    """`etree.fromstring()` replacement that uses oxml parser."""
    return etree.fromstring(text, oxml_parser)


def qn(tag):
    """Stands for "qualified name", a utility function to turn a namespace prefixed tag
    name into a Clark-notation qualified tag name for lxml.

    For
    example, ``qn('p:cSld')`` returns ``'{http://schemas.../main}cSld'``.
    """
    prefix, tagroot = tag.split(":")
    uri = nsmap[prefix]
    return "{%s}%s" % (uri, tagroot)


def serialize_part_xml(part_elm: etree._Element):
    """Serialize `part_elm` etree element to XML suitable for storage as an XML part.

    That is to say, no insignificant whitespace added for readability, and an
    appropriate XML declaration added with UTF-8 encoding specified.
    """
    return etree.tostring(part_elm, encoding="UTF-8", standalone=True)


def serialize_for_reading(element):
    """Serialize `element` to human-readable XML suitable for tests.

    No XML declaration.
    """
    return etree.tostring(element, encoding="unicode", pretty_print=True)


# ===========================================================================
# Custom element classes
# ===========================================================================


class BaseOxmlElement(etree.ElementBase):
    """Base class for all custom element classes, to add standardized behavior to all
    classes in one place."""

    @property
    def xml(self):
        """Return XML string for this element, suitable for testing purposes.

        Pretty printed for readability and without an XML declaration at the top.
        """
        return serialize_for_reading(self)


class CT_Default(BaseOxmlElement):
    """``<Default>`` element, specifying the default content type to be applied to a
    part with the specified extension."""

    @property
    def content_type(self):
        """String held in the ``ContentType`` attribute of this ``<Default>``
        element."""
        return self.get("ContentType")

    @property
    def extension(self):
        """String held in the ``Extension`` attribute of this ``<Default>`` element."""
        return self.get("Extension")

    @staticmethod
    def new(ext, content_type):
        """Return a new ``<Default>`` element with attributes set to parameter
        values."""
        xml = '<Default xmlns="%s"/>' % nsmap["ct"]
        default = parse_xml(xml)
        default.set("Extension", ext)
        default.set("ContentType", content_type)
        return default


class CT_Override(BaseOxmlElement):
    """``<Override>`` element, specifying the content type to be applied for a part with
    the specified partname."""

    @property
    def content_type(self):
        """String held in the ``ContentType`` attribute of this ``<Override>``
        element."""
        return self.get("ContentType")

    @staticmethod
    def new(partname, content_type):
        """Return a new ``<Override>`` element with attributes set to parameter
        values."""
        xml = '<Override xmlns="%s"/>' % nsmap["ct"]
        override = parse_xml(xml)
        override.set("PartName", partname)
        override.set("ContentType", content_type)
        return override

    @property
    def partname(self):
        """String held in the ``PartName`` attribute of this ``<Override>`` element."""
        return self.get("PartName")


class CT_Relationship(BaseOxmlElement):
    """``<Relationship>`` element, representing a single relationship from a source to a
    target part."""

    @staticmethod
    def new(rId: str, reltype: str, target: str, target_mode: str = RTM.INTERNAL):
        """Return a new ``<Relationship>`` element."""
        xml = '<Relationship xmlns="%s"/>' % nsmap["pr"]
        relationship = parse_xml(xml)
        relationship.set("Id", rId)
        relationship.set("Type", reltype)
        relationship.set("Target", target)
        if target_mode == RTM.EXTERNAL:
            relationship.set("TargetMode", RTM.EXTERNAL)
        return relationship

    @property
    def rId(self):
        """String held in the ``Id`` attribute of this ``<Relationship>`` element."""
        return self.get("Id")

    @property
    def reltype(self):
        """String held in the ``Type`` attribute of this ``<Relationship>`` element."""
        return self.get("Type")

    @property
    def target_ref(self):
        """String held in the ``Target`` attribute of this ``<Relationship>``
        element."""
        return self.get("Target")

    @property
    def target_mode(self):
        """String held in the ``TargetMode`` attribute of this ``<Relationship>``
        element, either ``Internal`` or ``External``.

        Defaults to ``Internal``.
        """
        return self.get("TargetMode", RTM.INTERNAL)


class CT_Relationships(BaseOxmlElement):
    """``<Relationships>`` element, the root element in a .rels file."""

    def add_rel(self, rId: str, reltype: str, target: str, is_external: bool = False):
        """Add a child ``<Relationship>`` element with attributes set according to
        parameter values."""
        target_mode = RTM.EXTERNAL if is_external else RTM.INTERNAL
        relationship = CT_Relationship.new(rId, reltype, target, target_mode)
        self.append(relationship)

    @staticmethod
    def new() -> CT_Relationships:
        """Return a new ``<Relationships>`` element."""
        xml = '<Relationships xmlns="%s"/>' % nsmap["pr"]
        return cast(CT_Relationships, parse_xml(xml))

    @property
    def Relationship_lst(self):
        """Return a list containing all the ``<Relationship>`` child elements."""
        return self.findall(qn("pr:Relationship"))

    @property
    def xml(self):
        """Return XML string for this element, suitable for saving in a .rels stream,
        not pretty printed and with an XML declaration at the top."""
        return serialize_part_xml(self)


class CT_Types(BaseOxmlElement):
    """``<Types>`` element, the container element for Default and Override elements in
    [Content_Types].xml."""

    def add_default(self, ext, content_type):
        """Add a child ``<Default>`` element with attributes set to parameter values."""
        default = CT_Default.new(ext, content_type)
        self.append(default)

    def add_override(self, partname, content_type):
        """Add a child ``<Override>`` element with attributes set to parameter
        values."""
        override = CT_Override.new(partname, content_type)
        self.append(override)

    @property
    def defaults(self):
        return self.findall(qn("ct:Default"))

    @staticmethod
    def new():
        """Return a new ``<Types>`` element."""
        xml = '<Types xmlns="%s"/>' % nsmap["ct"]
        types = parse_xml(xml)
        return types

    @property
    def overrides(self):
        return self.findall(qn("ct:Override"))


ct_namespace = element_class_lookup.get_namespace(nsmap["ct"])
ct_namespace["Default"] = CT_Default
ct_namespace["Override"] = CT_Override
ct_namespace["Types"] = CT_Types

pr_namespace = element_class_lookup.get_namespace(nsmap["pr"])
pr_namespace["Relationship"] = CT_Relationship
pr_namespace["Relationships"] = CT_Relationships


# package.py
"""Objects that implement reading and writing OPC packages."""

from __future__ import annotations

from typing import IO, TYPE_CHECKING, Iterator, cast

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PACKAGE_URI, PackURI
from docx.opc.part import PartFactory
from docx.opc.parts.coreprops import CorePropertiesPart
from docx.opc.pkgreader import PackageReader
from docx.opc.pkgwriter import PackageWriter
from docx.opc.rel import Relationships
from docx.shared import lazyproperty

if TYPE_CHECKING:
    from docx.opc.coreprops import CoreProperties
    from docx.opc.part import Part
    from docx.opc.rel import _Relationship  # pyright: ignore[reportPrivateUsage]


class OpcPackage:
    """Main API class for |python-opc|.

    A new instance is constructed by calling the :meth:`open` class method with a path
    to a package file or file-like object containing one.
    """

    def __init__(self):
        super(OpcPackage, self).__init__()

    def after_unmarshal(self):
        """Entry point for any post-unmarshaling processing.

        May be overridden by subclasses without forwarding call to super.
        """
        # don't place any code here, just catch call if not overridden by
        # subclass
        pass

    @property
    def core_properties(self) -> CoreProperties:
        """|CoreProperties| object providing read/write access to the Dublin Core
        properties for this document."""
        return self._core_properties_part.core_properties

    def iter_rels(self) -> Iterator[_Relationship]:
        """Generate exactly one reference to each relationship in the package by
        performing a depth-first traversal of the rels graph."""

        def walk_rels(
            source: OpcPackage | Part, visited: list[Part] | None = None
        ) -> Iterator[_Relationship]:
            visited = [] if visited is None else visited
            for rel in source.rels.values():
                yield rel
                if rel.is_external:
                    continue
                part = rel.target_part
                if part in visited:
                    continue
                visited.append(part)
                new_source = part
                for rel in walk_rels(new_source, visited):
                    yield rel

        for rel in walk_rels(self):
            yield rel

    def iter_parts(self) -> Iterator[Part]:
        """Generate exactly one reference to each of the parts in the package by
        performing a depth-first traversal of the rels graph."""

        def walk_parts(source, visited=[]):
            for rel in source.rels.values():
                if rel.is_external:
                    continue
                part = rel.target_part
                if part in visited:
                    continue
                visited.append(part)
                yield part
                new_source = part
                for part in walk_parts(new_source, visited):
                    yield part

        for part in walk_parts(self):
            yield part

    def load_rel(self, reltype: str, target: Part | str, rId: str, is_external: bool = False):
        """Return newly added |_Relationship| instance of `reltype` between this part
        and `target` with key `rId`.

        Target mode is set to ``RTM.EXTERNAL`` if `is_external` is |True|. Intended for
        use during load from a serialized package, where the rId is well known. Other
        methods exist for adding a new relationship to the package during processing.
        """
        return self.rels.add_relationship(reltype, target, rId, is_external)

    @property
    def main_document_part(self):
        """Return a reference to the main document part for this package.

        Examples include a document part for a WordprocessingML package, a presentation
        part for a PresentationML package, or a workbook part for a SpreadsheetML
        package.
        """
        return self.part_related_by(RT.OFFICE_DOCUMENT)

    def next_partname(self, template: str) -> PackURI:
        """Return a |PackURI| instance representing partname matching `template`.

        The returned part-name has the next available numeric suffix to distinguish it
        from other parts of its type. `template` is a printf (%)-style template string
        containing a single replacement item, a '%d' to be used to insert the integer
        portion of the partname. Example: "/word/header%d.xml"
        """
        partnames = {part.partname for part in self.iter_parts()}
        for n in range(1, len(partnames) + 2):
            candidate_partname = template % n
            if candidate_partname not in partnames:
                return PackURI(candidate_partname)

    @classmethod
    def open(cls, pkg_file: str | IO[bytes]) -> OpcPackage:
        """Return an |OpcPackage| instance loaded with the contents of `pkg_file`."""
        pkg_reader = PackageReader.from_file(pkg_file)
        package = cls()
        Unmarshaller.unmarshal(pkg_reader, package, PartFactory)
        return package

    def part_related_by(self, reltype: str) -> Part:
        """Return part to which this package has a relationship of `reltype`.

        Raises |KeyError| if no such relationship is found and |ValueError| if more than
        one such relationship is found.
        """
        return self.rels.part_with_reltype(reltype)

    @property
    def parts(self) -> list[Part]:
        """Return a list containing a reference to each of the parts in this package."""
        return list(self.iter_parts())

    def relate_to(self, part: Part, reltype: str):
        """Return rId key of new or existing relationship to `part`.

        If a relationship of `reltype` to `part` already exists, its rId is returned. Otherwise a
        new relationship is created and that rId is returned.
        """
        rel = self.rels.get_or_add(reltype, part)
        return rel.rId

    @lazyproperty
    def rels(self):
        """Return a reference to the |Relationships| instance holding the collection of
        relationships for this package."""
        return Relationships(PACKAGE_URI.baseURI)

    def save(self, pkg_file: str | IO[bytes]):
        """Save this package to `pkg_file`.

        `pkg_file` can be either a file-path or a file-like object.
        """
        for part in self.parts:
            part.before_marshal()
        PackageWriter.write(pkg_file, self.rels, self.parts)

    @property
    def _core_properties_part(self) -> CorePropertiesPart:
        """|CorePropertiesPart| object related to this package.

        Creates a default core properties part if one is not present (not common).
        """
        try:
            return cast(CorePropertiesPart, self.part_related_by(RT.CORE_PROPERTIES))
        except KeyError:
            core_properties_part = CorePropertiesPart.default(self)
            self.relate_to(core_properties_part, RT.CORE_PROPERTIES)
            return core_properties_part


class Unmarshaller:
    """Hosts static methods for unmarshalling a package from a |PackageReader|."""

    @staticmethod
    def unmarshal(pkg_reader, package, part_factory):
        """Construct graph of parts and realized relationships based on the contents of
        `pkg_reader`, delegating construction of each part to `part_factory`.

        Package relationships are added to `pkg`.
        """
        parts = Unmarshaller._unmarshal_parts(pkg_reader, package, part_factory)
        Unmarshaller._unmarshal_relationships(pkg_reader, package, parts)
        for part in parts.values():
            part.after_unmarshal()
        package.after_unmarshal()

    @staticmethod
    def _unmarshal_parts(pkg_reader, package, part_factory):
        """Return a dictionary of |Part| instances unmarshalled from `pkg_reader`, keyed
        by partname.

        Side-effect is that each part in `pkg_reader` is constructed using
        `part_factory`.
        """
        parts = {}
        for partname, content_type, reltype, blob in pkg_reader.iter_sparts():
            parts[partname] = part_factory(partname, content_type, reltype, blob, package)
        return parts

    @staticmethod
    def _unmarshal_relationships(pkg_reader, package, parts):
        """Add a relationship to the source object corresponding to each of the
        relationships in `pkg_reader` with its target_part set to the actual target part
        in `parts`."""
        for source_uri, srel in pkg_reader.iter_srels():
            source = package if source_uri == "/" else parts[source_uri]
            target = srel.target_ref if srel.is_external else parts[srel.target_partname]
            source.load_rel(srel.reltype, target, srel.rId, srel.is_external)


# packuri.py
"""Provides the PackURI value type.

Also some useful known pack URI strings such as PACKAGE_URI.
"""

from __future__ import annotations

import posixpath
import re


class PackURI(str):
    """Provides access to pack URI components such as the baseURI and the filename
    slice.

    Behaves as |str| otherwise.
    """

    _filename_re = re.compile("([a-zA-Z]+)([1-9][0-9]*)?")

    def __new__(cls, pack_uri_str: str):
        if pack_uri_str[0] != "/":
            tmpl = "PackURI must begin with slash, got '%s'"
            raise ValueError(tmpl % pack_uri_str)
        return str.__new__(cls, pack_uri_str)

    @staticmethod
    def from_rel_ref(baseURI: str, relative_ref: str) -> PackURI:
        """The absolute PackURI formed by translating `relative_ref` onto `baseURI`."""
        joined_uri = posixpath.join(baseURI, relative_ref)
        abs_uri = posixpath.abspath(joined_uri)
        return PackURI(abs_uri)

    @property
    def baseURI(self) -> str:
        """The base URI of this pack URI, the directory portion, roughly speaking.

        E.g. ``'/ppt/slides'`` for ``'/ppt/slides/slide1.xml'``. For the package pseudo-
        partname '/', baseURI is '/'.
        """
        return posixpath.split(self)[0]

    @property
    def ext(self) -> str:
        """The extension portion of this pack URI, e.g. ``'xml'`` for ``'/word/document.xml'``.

        Note the period is not included.
        """
        # raw_ext is either empty string or starts with period, e.g. '.xml'
        raw_ext = posixpath.splitext(self)[1]
        return raw_ext[1:] if raw_ext.startswith(".") else raw_ext

    @property
    def filename(self):
        """The "filename" portion of this pack URI, e.g. ``'slide1.xml'`` for
        ``'/ppt/slides/slide1.xml'``.

        For the package pseudo-partname '/', filename is ''.
        """
        return posixpath.split(self)[1]

    @property
    def idx(self):
        """Return partname index as integer for tuple partname or None for singleton
        partname, e.g. ``21`` for ``'/ppt/slides/slide21.xml'`` and |None| for
        ``'/ppt/presentation.xml'``."""
        filename = self.filename
        if not filename:
            return None
        name_part = posixpath.splitext(filename)[0]  # filename w/ext removed
        match = self._filename_re.match(name_part)
        if match is None:
            return None
        if match.group(2):
            return int(match.group(2))
        return None

    @property
    def membername(self):
        """The pack URI with the leading slash stripped off, the form used as the Zip
        file membername for the package item.

        Returns '' for the package pseudo-partname '/'.
        """
        return self[1:]

    def relative_ref(self, baseURI: str):
        """Return string containing relative reference to package item from `baseURI`.

        E.g. PackURI('/ppt/slideLayouts/slideLayout1.xml') would return
        '../slideLayouts/slideLayout1.xml' for baseURI '/ppt/slides'.
        """
        # workaround for posixpath bug in 2.6, doesn't generate correct
        # relative path when `start` (second) parameter is root ('/')
        return self[1:] if baseURI == "/" else posixpath.relpath(self, baseURI)

    @property
    def rels_uri(self):
        """The pack URI of the .rels part corresponding to the current pack URI.

        Only produces sensible output if the pack URI is a partname or the package
        pseudo-partname '/'.
        """
        rels_filename = "%s.rels" % self.filename
        rels_uri_str = posixpath.join(self.baseURI, "_rels", rels_filename)
        return PackURI(rels_uri_str)


PACKAGE_URI = PackURI("/")
CONTENT_TYPES_URI = PackURI("/[Content_Types].xml")


# part.py
# pyright: reportImportCycles=false

"""Open Packaging Convention (OPC) objects related to package parts."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable, Type, cast

from docx.opc.oxml import serialize_part_xml
from docx.opc.packuri import PackURI
from docx.opc.rel import Relationships
from docx.opc.shared import cls_method_fn
from docx.oxml.parser import parse_xml
from docx.shared import lazyproperty

if TYPE_CHECKING:
    from docx.oxml.xmlchemy import BaseOxmlElement
    from docx.package import Package


class Part:
    """Base class for package parts.

    Provides common properties and methods, but intended to be subclassed in client code
    to implement specific part behaviors.
    """

    def __init__(
        self,
        partname: PackURI,
        content_type: str,
        blob: bytes | None = None,
        package: Package | None = None,
    ):
        super(Part, self).__init__()
        self._partname = partname
        self._content_type = content_type
        self._blob = blob
        self._package = package

    def after_unmarshal(self):
        """Entry point for post-unmarshaling processing, for example to parse the part
        XML.

        May be overridden by subclasses without forwarding call to super.
        """
        # don't place any code here, just catch call if not overridden by
        # subclass
        pass

    def before_marshal(self):
        """Entry point for pre-serialization processing, for example to finalize part
        naming if necessary.

        May be overridden by subclasses without forwarding call to super.
        """
        # don't place any code here, just catch call if not overridden by
        # subclass
        pass

    @property
    def blob(self) -> bytes:
        """Contents of this package part as a sequence of bytes.

        May be text or binary. Intended to be overridden by subclasses. Default behavior
        is to return load blob.
        """
        return self._blob or b""

    @property
    def content_type(self):
        """Content type of this part."""
        return self._content_type

    def drop_rel(self, rId: str):
        """Remove the relationship identified by `rId` if its reference count is less
        than 2.

        Relationships with a reference count of 0 are implicit relationships.
        """
        if self._rel_ref_count(rId) < 2:
            del self.rels[rId]

    @classmethod
    def load(cls, partname: PackURI, content_type: str, blob: bytes, package: Package):
        return cls(partname, content_type, blob, package)

    def load_rel(self, reltype: str, target: Part | str, rId: str, is_external: bool = False):
        """Return newly added |_Relationship| instance of `reltype`.

        The new relationship relates the `target` part to this part with key `rId`.

        Target mode is set to ``RTM.EXTERNAL`` if `is_external` is |True|. Intended for
        use during load from a serialized package, where the rId is well-known. Other
        methods exist for adding a new relationship to a part when manipulating a part.
        """
        return self.rels.add_relationship(reltype, target, rId, is_external)

    @property
    def package(self):
        """|OpcPackage| instance this part belongs to."""
        return self._package

    @property
    def partname(self):
        """|PackURI| instance holding partname of this part, e.g.
        '/ppt/slides/slide1.xml'."""
        return self._partname

    @partname.setter
    def partname(self, partname: str):
        if not isinstance(partname, PackURI):
            tmpl = "partname must be instance of PackURI, got '%s'"
            raise TypeError(tmpl % type(partname).__name__)
        self._partname = partname

    def part_related_by(self, reltype: str) -> Part:
        """Return part to which this part has a relationship of `reltype`.

        Raises |KeyError| if no such relationship is found and |ValueError| if more than
        one such relationship is found. Provides ability to resolve implicitly related
        part, such as Slide -> SlideLayout.
        """
        return self.rels.part_with_reltype(reltype)

    def relate_to(self, target: Part | str, reltype: str, is_external: bool = False) -> str:
        """Return rId key of relationship of `reltype` to `target`.

        The returned `rId` is from an existing relationship if there is one, otherwise a
        new relationship is created.
        """
        if is_external:
            return self.rels.get_or_add_ext_rel(reltype, cast(str, target))
        else:
            rel = self.rels.get_or_add(reltype, cast(Part, target))
            return rel.rId

    @property
    def related_parts(self):
        """Dictionary mapping related parts by rId, so child objects can resolve
        explicit relationships present in the part XML, e.g. sldIdLst to a specific
        |Slide| instance."""
        return self.rels.related_parts

    @lazyproperty
    def rels(self):
        """|Relationships| instance holding the relationships for this part."""
        # -- prevent breakage in `python-docx-template` by retaining legacy `._rels` attribute --
        self._rels = Relationships(self._partname.baseURI)
        return self._rels

    def target_ref(self, rId: str) -> str:
        """Return URL contained in target ref of relationship identified by `rId`."""
        rel = self.rels[rId]
        return rel.target_ref

    def _rel_ref_count(self, rId: str) -> int:
        """Return the count of references in this part to the relationship identified by `rId`.

        Only an XML part can contain references, so this is 0 for `Part`.
        """
        return 0


class PartFactory:
    """Provides a way for client code to specify a subclass of |Part| to be constructed
    by |Unmarshaller| based on its content type and/or a custom callable.

    Setting ``PartFactory.part_class_selector`` to a callable object will cause that
    object to be called with the parameters ``content_type, reltype``, once for each
    part in the package. If the callable returns an object, it is used as the class for
    that part. If it returns |None|, part class selection falls back to the content type
    map defined in ``PartFactory.part_type_for``. If no class is returned from either of
    these, the class contained in ``PartFactory.default_part_type`` is used to construct
    the part, which is by default ``opc.package.Part``.
    """

    part_class_selector: Callable[[str, str], Type[Part] | None] | None
    part_type_for: dict[str, Type[Part]] = {}
    default_part_type = Part

    def __new__(
        cls,
        partname: PackURI,
        content_type: str,
        reltype: str,
        blob: bytes,
        package: Package,
    ):
        PartClass: Type[Part] | None = None
        if cls.part_class_selector is not None:
            part_class_selector = cls_method_fn(cls, "part_class_selector")
            PartClass = part_class_selector(content_type, reltype)
        if PartClass is None:
            PartClass = cls._part_cls_for(content_type)
        return PartClass.load(partname, content_type, blob, package)

    @classmethod
    def _part_cls_for(cls, content_type: str):
        """Return the custom part class registered for `content_type`, or the default
        part class if no custom class is registered for `content_type`."""
        if content_type in cls.part_type_for:
            return cls.part_type_for[content_type]
        return cls.default_part_type


class XmlPart(Part):
    """Base class for package parts containing an XML payload, which is most of them.

    Provides additional methods to the |Part| base class that take care of parsing and
    reserializing the XML payload and managing relationships to other parts.
    """

    def __init__(
        self, partname: PackURI, content_type: str, element: BaseOxmlElement, package: Package
    ):
        super(XmlPart, self).__init__(partname, content_type, package=package)
        self._element = element

    @property
    def blob(self):
        return serialize_part_xml(self._element)

    @property
    def element(self):
        """The root XML element of this XML part."""
        return self._element

    @classmethod
    def load(cls, partname: PackURI, content_type: str, blob: bytes, package: Package):
        element = parse_xml(blob)
        return cls(partname, content_type, element, package)

    @property
    def part(self):
        """Part of the parent protocol, "children" of the document will not know the
        part that contains them so must ask their parent object.

        That chain of delegation ends here for child objects.
        """
        return self

    def _rel_ref_count(self, rId: str) -> int:
        """Return the count of references in this part's XML to the relationship
        identified by `rId`."""
        rIds = cast("list[str]", self._element.xpath("//@r:id"))
        return len([_rId for _rId in rIds if _rId == rId])


# coreprops.py
"""Core properties part, corresponds to ``/docProps/core.xml`` part in package."""

from __future__ import annotations

import datetime as dt
from typing import TYPE_CHECKING

from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.coreprops import CoreProperties
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml.coreprops import CT_CoreProperties

if TYPE_CHECKING:
    from docx.opc.package import OpcPackage


class CorePropertiesPart(XmlPart):
    """Corresponds to part named ``/docProps/core.xml``.

    The "core" is short for "Dublin Core" and contains document metadata relatively common across
    documents of all types, not just DOCX.
    """

    @classmethod
    def default(cls, package: OpcPackage):
        """Return a new |CorePropertiesPart| object initialized with default values for
        its base properties."""
        core_properties_part = cls._new(package)
        core_properties = core_properties_part.core_properties
        core_properties.title = "Word Document"
        core_properties.last_modified_by = "python-docx"
        core_properties.revision = 1
        core_properties.modified = dt.datetime.now(dt.timezone.utc)
        return core_properties_part

    @property
    def core_properties(self):
        """A |CoreProperties| object providing read/write access to the core properties
        contained in this core properties part."""
        return CoreProperties(self.element)

    @classmethod
    def _new(cls, package: OpcPackage) -> CorePropertiesPart:
        partname = PackURI("/docProps/core.xml")
        content_type = CT.OPC_CORE_PROPERTIES
        coreProperties = CT_CoreProperties.new()
        return CorePropertiesPart(partname, content_type, coreProperties, package)


# __init__.py


# phys_pkg.py
"""Provides a general interface to a `physical` OPC package, such as a zip file."""

import os
from zipfile import ZIP_DEFLATED, ZipFile, is_zipfile

from docx.opc.exceptions import PackageNotFoundError
from docx.opc.packuri import CONTENT_TYPES_URI


class PhysPkgReader:
    """Factory for physical package reader objects."""

    def __new__(cls, pkg_file):
        # if `pkg_file` is a string, treat it as a path
        if isinstance(pkg_file, str):
            if os.path.isdir(pkg_file):
                reader_cls = _DirPkgReader
            elif is_zipfile(pkg_file):
                reader_cls = _ZipPkgReader
            else:
                raise PackageNotFoundError("Package not found at '%s'" % pkg_file)
        else:  # assume it's a stream and pass it to Zip reader to sort out
            reader_cls = _ZipPkgReader

        return super(PhysPkgReader, cls).__new__(reader_cls)


class PhysPkgWriter:
    """Factory for physical package writer objects."""

    def __new__(cls, pkg_file):
        return super(PhysPkgWriter, cls).__new__(_ZipPkgWriter)


class _DirPkgReader(PhysPkgReader):
    """Implements |PhysPkgReader| interface for an OPC package extracted into a
    directory."""

    def __init__(self, path):
        """`path` is the path to a directory containing an expanded package."""
        super(_DirPkgReader, self).__init__()
        self._path = os.path.abspath(path)

    def blob_for(self, pack_uri):
        """Return contents of file corresponding to `pack_uri` in package directory."""
        path = os.path.join(self._path, pack_uri.membername)
        with open(path, "rb") as f:
            blob = f.read()
        return blob

    def close(self):
        """Provides interface consistency with |ZipFileSystem|, but does nothing, a
        directory file system doesn't need closing."""
        pass

    @property
    def content_types_xml(self):
        """Return the `[Content_Types].xml` blob from the package."""
        return self.blob_for(CONTENT_TYPES_URI)

    def rels_xml_for(self, source_uri):
        """Return rels item XML for source with `source_uri`, or None if the item has no
        rels item."""
        try:
            rels_xml = self.blob_for(source_uri.rels_uri)
        except IOError:
            rels_xml = None
        return rels_xml


class _ZipPkgReader(PhysPkgReader):
    """Implements |PhysPkgReader| interface for a zip file OPC package."""

    def __init__(self, pkg_file):
        super(_ZipPkgReader, self).__init__()
        self._zipf = ZipFile(pkg_file, "r")

    def blob_for(self, pack_uri):
        """Return blob corresponding to `pack_uri`.

        Raises |ValueError| if no matching member is present in zip archive.
        """
        return self._zipf.read(pack_uri.membername)

    def close(self):
        """Close the zip archive, releasing any resources it is using."""
        self._zipf.close()

    @property
    def content_types_xml(self):
        """Return the `[Content_Types].xml` blob from the zip package."""
        return self.blob_for(CONTENT_TYPES_URI)

    def rels_xml_for(self, source_uri):
        """Return rels item XML for source with `source_uri` or None if no rels item is
        present."""
        try:
            rels_xml = self.blob_for(source_uri.rels_uri)
        except KeyError:
            rels_xml = None
        return rels_xml


class _ZipPkgWriter(PhysPkgWriter):
    """Implements |PhysPkgWriter| interface for a zip file OPC package."""

    def __init__(self, pkg_file):
        super(_ZipPkgWriter, self).__init__()
        self._zipf = ZipFile(pkg_file, "w", compression=ZIP_DEFLATED)

    def close(self):
        """Close the zip archive, flushing any pending physical writes and releasing any
        resources it's using."""
        self._zipf.close()

    def write(self, pack_uri, blob):
        """Write `blob` to this zip package with the membername corresponding to
        `pack_uri`."""
        self._zipf.writestr(pack_uri.membername, blob)


# pkgreader.py
"""Low-level, read-only API to a serialized Open Packaging Convention (OPC) package."""

from docx.opc.constants import RELATIONSHIP_TARGET_MODE as RTM
from docx.opc.oxml import parse_xml
from docx.opc.packuri import PACKAGE_URI, PackURI
from docx.opc.phys_pkg import PhysPkgReader
from docx.opc.shared import CaseInsensitiveDict


class PackageReader:
    """Provides access to the contents of a zip-format OPC package via its
    :attr:`serialized_parts` and :attr:`pkg_srels` attributes."""

    def __init__(self, content_types, pkg_srels, sparts):
        super(PackageReader, self).__init__()
        self._pkg_srels = pkg_srels
        self._sparts = sparts

    @staticmethod
    def from_file(pkg_file):
        """Return a |PackageReader| instance loaded with contents of `pkg_file`."""
        phys_reader = PhysPkgReader(pkg_file)
        content_types = _ContentTypeMap.from_xml(phys_reader.content_types_xml)
        pkg_srels = PackageReader._srels_for(phys_reader, PACKAGE_URI)
        sparts = PackageReader._load_serialized_parts(
            phys_reader, pkg_srels, content_types
        )
        phys_reader.close()
        return PackageReader(content_types, pkg_srels, sparts)

    def iter_sparts(self):
        """Generate a 4-tuple `(partname, content_type, reltype, blob)` for each of the
        serialized parts in the package."""
        for s in self._sparts:
            yield (s.partname, s.content_type, s.reltype, s.blob)

    def iter_srels(self):
        """Generate a 2-tuple `(source_uri, srel)` for each of the relationships in the
        package."""
        for srel in self._pkg_srels:
            yield (PACKAGE_URI, srel)
        for spart in self._sparts:
            for srel in spart.srels:
                yield (spart.partname, srel)

    @staticmethod
    def _load_serialized_parts(phys_reader, pkg_srels, content_types):
        """Return a list of |_SerializedPart| instances corresponding to the parts in
        `phys_reader` accessible by walking the relationship graph starting with
        `pkg_srels`."""
        sparts = []
        part_walker = PackageReader._walk_phys_parts(phys_reader, pkg_srels)
        for partname, blob, reltype, srels in part_walker:
            content_type = content_types[partname]
            spart = _SerializedPart(partname, content_type, reltype, blob, srels)
            sparts.append(spart)
        return tuple(sparts)

    @staticmethod
    def _srels_for(phys_reader, source_uri):
        """Return |_SerializedRelationships| instance populated with relationships for
        source identified by `source_uri`."""
        rels_xml = phys_reader.rels_xml_for(source_uri)
        return _SerializedRelationships.load_from_xml(source_uri.baseURI, rels_xml)

    @staticmethod
    def _walk_phys_parts(phys_reader, srels, visited_partnames=None):
        """Generate a 4-tuple `(partname, blob, reltype, srels)` for each of the parts
        in `phys_reader` by walking the relationship graph rooted at srels."""
        if visited_partnames is None:
            visited_partnames = []
        for srel in srels:
            if srel.is_external:
                continue
            partname = srel.target_partname
            if partname in visited_partnames:
                continue
            visited_partnames.append(partname)
            reltype = srel.reltype
            part_srels = PackageReader._srels_for(phys_reader, partname)
            blob = phys_reader.blob_for(partname)
            yield (partname, blob, reltype, part_srels)
            next_walker = PackageReader._walk_phys_parts(
                phys_reader, part_srels, visited_partnames
            )
            for partname, blob, reltype, srels in next_walker:
                yield (partname, blob, reltype, srels)


class _ContentTypeMap:
    """Value type providing dictionary semantics for looking up content type by part
    name, e.g. ``content_type = cti['/ppt/presentation.xml']``."""

    def __init__(self):
        super(_ContentTypeMap, self).__init__()
        self._overrides = CaseInsensitiveDict()
        self._defaults = CaseInsensitiveDict()

    def __getitem__(self, partname):
        """Return content type for part identified by `partname`."""
        if not isinstance(partname, PackURI):
            tmpl = "_ContentTypeMap key must be <type 'PackURI'>, got %s"
            raise KeyError(tmpl % type(partname))
        if partname in self._overrides:
            return self._overrides[partname]
        if partname.ext in self._defaults:
            return self._defaults[partname.ext]
        tmpl = "no content type for partname '%s' in [Content_Types].xml"
        raise KeyError(tmpl % partname)

    @staticmethod
    def from_xml(content_types_xml):
        """Return a new |_ContentTypeMap| instance populated with the contents of
        `content_types_xml`."""
        types_elm = parse_xml(content_types_xml)
        ct_map = _ContentTypeMap()
        for o in types_elm.overrides:
            ct_map._add_override(o.partname, o.content_type)
        for d in types_elm.defaults:
            ct_map._add_default(d.extension, d.content_type)
        return ct_map

    def _add_default(self, extension, content_type):
        """Add the default mapping of `extension` to `content_type` to this content type
        mapping."""
        self._defaults[extension] = content_type

    def _add_override(self, partname, content_type):
        """Add the default mapping of `partname` to `content_type` to this content type
        mapping."""
        self._overrides[partname] = content_type


class _SerializedPart:
    """Value object for an OPC package part.

    Provides access to the partname, content type, blob, and serialized relationships
    for the part.
    """

    def __init__(self, partname, content_type, reltype, blob, srels):
        super(_SerializedPart, self).__init__()
        self._partname = partname
        self._content_type = content_type
        self._reltype = reltype
        self._blob = blob
        self._srels = srels

    @property
    def partname(self):
        return self._partname

    @property
    def content_type(self):
        return self._content_type

    @property
    def blob(self):
        return self._blob

    @property
    def reltype(self):
        """The referring relationship type of this part."""
        return self._reltype

    @property
    def srels(self):
        return self._srels


class _SerializedRelationship:
    """Value object representing a serialized relationship in an OPC package.

    Serialized, in this case, means any target part is referred to via its partname
    rather than a direct link to an in-memory |Part| object.
    """

    def __init__(self, baseURI, rel_elm):
        super(_SerializedRelationship, self).__init__()
        self._baseURI = baseURI
        self._rId = rel_elm.rId
        self._reltype = rel_elm.reltype
        self._target_mode = rel_elm.target_mode
        self._target_ref = rel_elm.target_ref

    @property
    def is_external(self):
        """True if target_mode is ``RTM.EXTERNAL``"""
        return self._target_mode == RTM.EXTERNAL

    @property
    def reltype(self):
        """Relationship type, like ``RT.OFFICE_DOCUMENT``"""
        return self._reltype

    @property
    def rId(self):
        """Relationship id, like 'rId9', corresponds to the ``Id`` attribute on the
        ``CT_Relationship`` element."""
        return self._rId

    @property
    def target_mode(self):
        """String in ``TargetMode`` attribute of ``CT_Relationship`` element, one of
        ``RTM.INTERNAL`` or ``RTM.EXTERNAL``."""
        return self._target_mode

    @property
    def target_ref(self):
        """String in ``Target`` attribute of ``CT_Relationship`` element, a relative
        part reference for internal target mode or an arbitrary URI, e.g. an HTTP URL,
        for external target mode."""
        return self._target_ref

    @property
    def target_partname(self):
        """|PackURI| instance containing partname targeted by this relationship.

        Raises ``ValueError`` on reference if target_mode is ``'External'``. Use
        :attr:`target_mode` to check before referencing.
        """
        if self.is_external:
            msg = (
                "target_partname attribute on Relationship is undefined w"
                'here TargetMode == "External"'
            )
            raise ValueError(msg)
        # lazy-load _target_partname attribute
        if not hasattr(self, "_target_partname"):
            self._target_partname = PackURI.from_rel_ref(self._baseURI, self.target_ref)
        return self._target_partname


class _SerializedRelationships:
    """Read-only sequence of |_SerializedRelationship| instances corresponding to the
    relationships item XML passed to constructor."""

    def __init__(self):
        super(_SerializedRelationships, self).__init__()
        self._srels = []

    def __iter__(self):
        """Support iteration, e.g. 'for x in srels:'."""
        return self._srels.__iter__()

    @staticmethod
    def load_from_xml(baseURI, rels_item_xml):
        """Return |_SerializedRelationships| instance loaded with the relationships
        contained in `rels_item_xml`.

        Returns an empty collection if `rels_item_xml` is |None|.
        """
        srels = _SerializedRelationships()
        if rels_item_xml is not None:
            rels_elm = parse_xml(rels_item_xml)
            for rel_elm in rels_elm.Relationship_lst:
                srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
        return srels


# pkgwriter.py
"""Provides low-level, write-only API to serialized (OPC) package.

OPC stands for Open Packaging Convention. This is e, essentially an implementation of
OpcPackage.save().
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Iterable

from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.oxml import CT_Types, serialize_part_xml
from docx.opc.packuri import CONTENT_TYPES_URI, PACKAGE_URI
from docx.opc.phys_pkg import PhysPkgWriter
from docx.opc.shared import CaseInsensitiveDict
from docx.opc.spec import default_content_types

if TYPE_CHECKING:
    from docx.opc.part import Part


class PackageWriter:
    """Writes a zip-format OPC package to `pkg_file`, where `pkg_file` can be either a
    path to a zip file (a string) or a file-like object.

    Its single API method, :meth:`write`, is static, so this class is not intended to be
    instantiated.
    """

    @staticmethod
    def write(pkg_file, pkg_rels, parts):
        """Write a physical package (.pptx file) to `pkg_file` containing `pkg_rels` and
        `parts` and a content types stream based on the content types of the parts."""
        phys_writer = PhysPkgWriter(pkg_file)
        PackageWriter._write_content_types_stream(phys_writer, parts)
        PackageWriter._write_pkg_rels(phys_writer, pkg_rels)
        PackageWriter._write_parts(phys_writer, parts)
        phys_writer.close()

    @staticmethod
    def _write_content_types_stream(phys_writer, parts):
        """Write ``[Content_Types].xml`` part to the physical package with an
        appropriate content type lookup target for each part in `parts`."""
        cti = _ContentTypesItem.from_parts(parts)
        phys_writer.write(CONTENT_TYPES_URI, cti.blob)

    @staticmethod
    def _write_parts(phys_writer: PhysPkgWriter, parts: Iterable[Part]):
        """Write the blob of each part in `parts` to the package, along with a rels item
        for its relationships if and only if it has any."""
        for part in parts:
            phys_writer.write(part.partname, part.blob)
            if len(part.rels):
                phys_writer.write(part.partname.rels_uri, part.rels.xml)

    @staticmethod
    def _write_pkg_rels(phys_writer, pkg_rels):
        """Write the XML rels item for `pkg_rels` ('/_rels/.rels') to the package."""
        phys_writer.write(PACKAGE_URI.rels_uri, pkg_rels.xml)


class _ContentTypesItem:
    """Service class that composes a content types item ([Content_Types].xml) based on a
    list of parts.

    Not meant to be instantiated directly, its single interface method is xml_for(),
    e.g. ``_ContentTypesItem.xml_for(parts)``.
    """

    def __init__(self):
        self._defaults = CaseInsensitiveDict()
        self._overrides = {}

    @property
    def blob(self):
        """Return XML form of this content types item, suitable for storage as
        ``[Content_Types].xml`` in an OPC package."""
        return serialize_part_xml(self._element)

    @classmethod
    def from_parts(cls, parts):
        """Return content types XML mapping each part in `parts` to the appropriate
        content type and suitable for storage as ``[Content_Types].xml`` in an OPC
        package."""
        cti = cls()
        cti._defaults["rels"] = CT.OPC_RELATIONSHIPS
        cti._defaults["xml"] = CT.XML
        for part in parts:
            cti._add_content_type(part.partname, part.content_type)
        return cti

    def _add_content_type(self, partname, content_type):
        """Add a content type for the part with `partname` and `content_type`, using a
        default or override as appropriate."""
        ext = partname.ext
        if (ext.lower(), content_type) in default_content_types:
            self._defaults[ext] = content_type
        else:
            self._overrides[partname] = content_type

    @property
    def _element(self):
        """Return XML form of this content types item, suitable for storage as
        ``[Content_Types].xml`` in an OPC package.

        Although the sequence of elements is not strictly significant, as an aid to
        testing and readability Default elements are sorted by extension and Override
        elements are sorted by partname.
        """
        _types_elm = CT_Types.new()
        for ext in sorted(self._defaults.keys()):
            _types_elm.add_default(ext, self._defaults[ext])
        for partname in sorted(self._overrides.keys()):
            _types_elm.add_override(partname, self._overrides[partname])
        return _types_elm


# rel.py
"""Relationship-related objects."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, Dict, cast

from docx.opc.oxml import CT_Relationships

if TYPE_CHECKING:
    from docx.opc.part import Part


class Relationships(Dict[str, "_Relationship"]):
    """Collection object for |_Relationship| instances, having list semantics."""

    def __init__(self, baseURI: str):
        super(Relationships, self).__init__()
        self._baseURI = baseURI
        self._target_parts_by_rId: dict[str, Any] = {}

    def add_relationship(
        self, reltype: str, target: Part | str, rId: str, is_external: bool = False
    ) -> "_Relationship":
        """Return a newly added |_Relationship| instance."""
        rel = _Relationship(rId, reltype, target, self._baseURI, is_external)
        self[rId] = rel
        if not is_external:
            self._target_parts_by_rId[rId] = target
        return rel

    def get_or_add(self, reltype: str, target_part: Part) -> _Relationship:
        """Return relationship of `reltype` to `target_part`, newly added if not already
        present in collection."""
        rel = self._get_matching(reltype, target_part)
        if rel is None:
            rId = self._next_rId
            rel = self.add_relationship(reltype, target_part, rId)
        return rel

    def get_or_add_ext_rel(self, reltype: str, target_ref: str) -> str:
        """Return rId of external relationship of `reltype` to `target_ref`, newly added
        if not already present in collection."""
        rel = self._get_matching(reltype, target_ref, is_external=True)
        if rel is None:
            rId = self._next_rId
            rel = self.add_relationship(reltype, target_ref, rId, is_external=True)
        return rel.rId

    def part_with_reltype(self, reltype: str) -> Part:
        """Return target part of rel with matching `reltype`, raising |KeyError| if not
        found and |ValueError| if more than one matching relationship is found."""
        rel = self._get_rel_of_type(reltype)
        return rel.target_part

    @property
    def related_parts(self):
        """Dict mapping rIds to target parts for all the internal relationships in the
        collection."""
        return self._target_parts_by_rId

    @property
    def xml(self) -> str:
        """Serialize this relationship collection into XML suitable for storage as a
        .rels file in an OPC package."""
        rels_elm = CT_Relationships.new()
        for rel in self.values():
            rels_elm.add_rel(rel.rId, rel.reltype, rel.target_ref, rel.is_external)
        return rels_elm.xml

    def _get_matching(
        self, reltype: str, target: Part | str, is_external: bool = False
    ) -> _Relationship | None:
        """Return relationship of matching `reltype`, `target`, and `is_external` from
        collection, or None if not found."""

        def matches(rel: _Relationship, reltype: str, target: Part | str, is_external: bool):
            if rel.reltype != reltype:
                return False
            if rel.is_external != is_external:
                return False
            rel_target = rel.target_ref if rel.is_external else rel.target_part
            if rel_target != target:
                return False
            return True

        for rel in self.values():
            if matches(rel, reltype, target, is_external):
                return rel
        return None

    def _get_rel_of_type(self, reltype: str):
        """Return single relationship of type `reltype` from the collection.

        Raises |KeyError| if no matching relationship is found. Raises |ValueError| if
        more than one matching relationship is found.
        """
        matching = [rel for rel in self.values() if rel.reltype == reltype]
        if len(matching) == 0:
            tmpl = "no relationship of type '%s' in collection"
            raise KeyError(tmpl % reltype)
        if len(matching) > 1:
            tmpl = "multiple relationships of type '%s' in collection"
            raise ValueError(tmpl % reltype)
        return matching[0]

    @property
    def _next_rId(self) -> str:  # pyright: ignore[reportReturnType]
        """Next available rId in collection, starting from 'rId1' and making use of any
        gaps in numbering, e.g. 'rId2' for rIds ['rId1', 'rId3']."""
        for n in range(1, len(self) + 2):
            rId_candidate = "rId%d" % n  # like 'rId19'
            if rId_candidate not in self:
                return rId_candidate


class _Relationship:
    """Value object for relationship to part."""

    def __init__(
        self, rId: str, reltype: str, target: Part | str, baseURI: str, external: bool = False
    ):
        super(_Relationship, self).__init__()
        self._rId = rId
        self._reltype = reltype
        self._target = target
        self._baseURI = baseURI
        self._is_external = bool(external)

    @property
    def is_external(self) -> bool:
        return self._is_external

    @property
    def reltype(self) -> str:
        return self._reltype

    @property
    def rId(self) -> str:
        return self._rId

    @property
    def target_part(self) -> Part:
        if self._is_external:
            raise ValueError(
                "target_part property on _Relationship is undef" "ined when target mode is External"
            )
        return cast("Part", self._target)

    @property
    def target_ref(self) -> str:
        if self._is_external:
            return cast(str, self._target)
        else:
            target = cast("Part", self._target)
            return target.partname.relative_ref(self._baseURI)


# shared.py
"""Objects shared by opc modules."""

from __future__ import annotations

from typing import Any, Dict, TypeVar

_T = TypeVar("_T")


class CaseInsensitiveDict(Dict[str, Any]):
    """Mapping type that behaves like dict except that it matches without respect to the
    case of the key.

    E.g. cid['A'] == cid['a']. Note this is not general-purpose, just complete enough to
    satisfy opc package needs. It assumes str keys, and that it is created empty; keys
    passed in constructor are not accounted for
    """

    def __contains__(self, key):
        return super(CaseInsensitiveDict, self).__contains__(key.lower())

    def __getitem__(self, key):
        return super(CaseInsensitiveDict, self).__getitem__(key.lower())

    def __setitem__(self, key, value):
        return super(CaseInsensitiveDict, self).__setitem__(key.lower(), value)


def cls_method_fn(cls: type, method_name: str):
    """Return method of `cls` having `method_name`."""
    return getattr(cls, method_name)


# spec.py
"""Provides mappings that embody aspects of the Open XML spec ISO/IEC 29500."""

from docx.opc.constants import CONTENT_TYPE as CT

default_content_types = (
    ("bin", CT.PML_PRINTER_SETTINGS),
    ("bin", CT.SML_PRINTER_SETTINGS),
    ("bin", CT.WML_PRINTER_SETTINGS),
    ("bmp", CT.BMP),
    ("emf", CT.X_EMF),
    ("fntdata", CT.X_FONTDATA),
    ("gif", CT.GIF),
    ("jpe", CT.JPEG),
    ("jpeg", CT.JPEG),
    ("jpg", CT.JPEG),
    ("png", CT.PNG),
    ("rels", CT.OPC_RELATIONSHIPS),
    ("tif", CT.TIFF),
    ("tiff", CT.TIFF),
    ("wdp", CT.MS_PHOTO),
    ("wmf", CT.X_WMF),
    ("xlsx", CT.SML_SHEET),
    ("xml", CT.XML),
)


# __init__.py


# coreprops.py
"""Custom element classes for core properties-related XML elements."""

from __future__ import annotations

import datetime as dt
import re
from typing import TYPE_CHECKING, Any, Callable

from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne

if TYPE_CHECKING:
    from lxml.etree import _Element as etree_Element  # pyright: ignore[reportPrivateUsage]


class CT_CoreProperties(BaseOxmlElement):
    """`<cp:coreProperties>` element, the root element of the Core Properties part.

    Stored as `/docProps/core.xml`. Implements many of the Dublin Core document metadata
    elements. String elements resolve to an empty string ("") if the element is not
    present in the XML. String elements are limited in length to 255 unicode characters.
    """

    get_or_add_revision: Callable[[], etree_Element]

    category = ZeroOrOne("cp:category", successors=())
    contentStatus = ZeroOrOne("cp:contentStatus", successors=())
    created = ZeroOrOne("dcterms:created", successors=())
    creator = ZeroOrOne("dc:creator", successors=())
    description = ZeroOrOne("dc:description", successors=())
    identifier = ZeroOrOne("dc:identifier", successors=())
    keywords = ZeroOrOne("cp:keywords", successors=())
    language = ZeroOrOne("dc:language", successors=())
    lastModifiedBy = ZeroOrOne("cp:lastModifiedBy", successors=())
    lastPrinted = ZeroOrOne("cp:lastPrinted", successors=())
    modified = ZeroOrOne("dcterms:modified", successors=())
    revision: etree_Element | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "cp:revision", successors=()
    )
    subject = ZeroOrOne("dc:subject", successors=())
    title = ZeroOrOne("dc:title", successors=())
    version = ZeroOrOne("cp:version", successors=())

    _coreProperties_tmpl = "<cp:coreProperties %s/>\n" % nsdecls("cp", "dc", "dcterms")

    @classmethod
    def new(cls):
        """Return a new `<cp:coreProperties>` element."""
        xml = cls._coreProperties_tmpl
        coreProperties = parse_xml(xml)
        return coreProperties

    @property
    def author_text(self):
        """The text in the `dc:creator` child element."""
        return self._text_of_element("creator")

    @author_text.setter
    def author_text(self, value: str):
        self._set_element_text("creator", value)

    @property
    def category_text(self) -> str:
        return self._text_of_element("category")

    @category_text.setter
    def category_text(self, value: str):
        self._set_element_text("category", value)

    @property
    def comments_text(self) -> str:
        return self._text_of_element("description")

    @comments_text.setter
    def comments_text(self, value: str):
        self._set_element_text("description", value)

    @property
    def contentStatus_text(self):
        return self._text_of_element("contentStatus")

    @contentStatus_text.setter
    def contentStatus_text(self, value: str):
        self._set_element_text("contentStatus", value)

    @property
    def created_datetime(self):
        return self._datetime_of_element("created")

    @created_datetime.setter
    def created_datetime(self, value: dt.datetime):
        self._set_element_datetime("created", value)

    @property
    def identifier_text(self):
        return self._text_of_element("identifier")

    @identifier_text.setter
    def identifier_text(self, value: str):
        self._set_element_text("identifier", value)

    @property
    def keywords_text(self):
        return self._text_of_element("keywords")

    @keywords_text.setter
    def keywords_text(self, value: str):
        self._set_element_text("keywords", value)

    @property
    def language_text(self):
        return self._text_of_element("language")

    @language_text.setter
    def language_text(self, value: str):
        self._set_element_text("language", value)

    @property
    def lastModifiedBy_text(self):
        return self._text_of_element("lastModifiedBy")

    @lastModifiedBy_text.setter
    def lastModifiedBy_text(self, value: str):
        self._set_element_text("lastModifiedBy", value)

    @property
    def lastPrinted_datetime(self):
        return self._datetime_of_element("lastPrinted")

    @lastPrinted_datetime.setter
    def lastPrinted_datetime(self, value: dt.datetime):
        self._set_element_datetime("lastPrinted", value)

    @property
    def modified_datetime(self) -> dt.datetime | None:
        return self._datetime_of_element("modified")

    @modified_datetime.setter
    def modified_datetime(self, value: dt.datetime):
        self._set_element_datetime("modified", value)

    @property
    def revision_number(self):
        """Integer value of revision property."""
        revision = self.revision
        if revision is None:
            return 0
        revision_str = str(revision.text)
        try:
            revision = int(revision_str)
        except ValueError:
            # non-integer revision strings also resolve to 0
            revision = 0
        # as do negative integers
        if revision < 0:
            revision = 0
        return revision

    @revision_number.setter
    def revision_number(self, value: int):
        """Set revision property to string value of integer `value`."""
        if not isinstance(value, int) or value < 1:  # pyright: ignore[reportUnnecessaryIsInstance]
            tmpl = "revision property requires positive int, got '%s'"
            raise ValueError(tmpl % value)
        revision = self.get_or_add_revision()
        revision.text = str(value)

    @property
    def subject_text(self):
        return self._text_of_element("subject")

    @subject_text.setter
    def subject_text(self, value: str):
        self._set_element_text("subject", value)

    @property
    def title_text(self):
        return self._text_of_element("title")

    @title_text.setter
    def title_text(self, value: str):
        self._set_element_text("title", value)

    @property
    def version_text(self):
        return self._text_of_element("version")

    @version_text.setter
    def version_text(self, value: str):
        self._set_element_text("version", value)

    def _datetime_of_element(self, property_name: str) -> dt.datetime | None:
        element = getattr(self, property_name)
        if element is None:
            return None
        datetime_str = element.text
        try:
            return self._parse_W3CDTF_to_datetime(datetime_str)
        except ValueError:
            # invalid datetime strings are ignored
            return None

    def _get_or_add(self, prop_name: str) -> BaseOxmlElement:
        """Return element returned by "get_or_add_" method for `prop_name`."""
        get_or_add_method_name = "get_or_add_%s" % prop_name
        get_or_add_method = getattr(self, get_or_add_method_name)
        element = get_or_add_method()
        return element

    @classmethod
    def _offset_dt(cls, dt_: dt.datetime, offset_str: str) -> dt.datetime:
        """A |datetime| instance offset from `dt_` by timezone offset in `offset_str`.

        `offset_str` is like `"-07:00"`.
        """
        match = cls._offset_pattern.match(offset_str)
        if match is None:
            raise ValueError("'%s' is not a valid offset string" % offset_str)
        sign, hours_str, minutes_str = match.groups()
        sign_factor = -1 if sign == "+" else 1
        hours = int(hours_str) * sign_factor
        minutes = int(minutes_str) * sign_factor
        td = dt.timedelta(hours=hours, minutes=minutes)
        return dt_ + td

    _offset_pattern = re.compile(r"([+-])(\d\d):(\d\d)")

    @classmethod
    def _parse_W3CDTF_to_datetime(cls, w3cdtf_str: str) -> dt.datetime:
        # valid W3CDTF date cases:
        # yyyy e.g. "2003"
        # yyyy-mm e.g. "2003-12"
        # yyyy-mm-dd e.g. "2003-12-31"
        # UTC timezone e.g. "2003-12-31T10:14:55Z"
        # numeric timezone e.g. "2003-12-31T10:14:55-08:00"
        templates = (
            "%Y-%m-%dT%H:%M:%S",
            "%Y-%m-%d",
            "%Y-%m",
            "%Y",
        )
        # strptime isn't smart enough to parse literal timezone offsets like
        # "-07:30", so we have to do it ourselves
        parseable_part = w3cdtf_str[:19]
        offset_str = w3cdtf_str[19:]
        dt_ = None
        for tmpl in templates:
            try:
                dt_ = dt.datetime.strptime(parseable_part, tmpl)
            except ValueError:
                continue
        if dt_ is None:
            tmpl = "could not parse W3CDTF datetime string '%s'"
            raise ValueError(tmpl % w3cdtf_str)
        if len(offset_str) == 6:
            dt_ = cls._offset_dt(dt_, offset_str)
        return dt_.replace(tzinfo=dt.timezone.utc)

    def _set_element_datetime(self, prop_name: str, value: dt.datetime):
        """Set date/time value of child element having `prop_name` to `value`."""
        if not isinstance(value, dt.datetime):  # pyright: ignore[reportUnnecessaryIsInstance]
            tmpl = "property requires <type 'datetime.datetime'> object, got %s"
            raise ValueError(tmpl % type(value))
        element = self._get_or_add(prop_name)
        dt_str = value.strftime("%Y-%m-%dT%H:%M:%SZ")
        element.text = dt_str
        if prop_name in ("created", "modified"):
            # These two require an explicit "xsi:type="dcterms:W3CDTF""
            # attribute. The first and last line are a hack required to add
            # the xsi namespace to the root element rather than each child
            # element in which it is referenced
            self.set(qn("xsi:foo"), "bar")
            element.set(qn("xsi:type"), "dcterms:W3CDTF")
            del self.attrib[qn("xsi:foo")]

    def _set_element_text(self, prop_name: str, value: Any) -> None:
        """Set string value of `name` property to `value`."""
        if not isinstance(value, str):
            value = str(value)

        if len(value) > 255:
            tmpl = "exceeded 255 char limit for property, got:\n\n'%s'"
            raise ValueError(tmpl % value)
        element = self._get_or_add(prop_name)
        element.text = value

    def _text_of_element(self, property_name: str) -> str:
        """The text in the element matching `property_name`.

        The empty string if the element is not present or contains no text.
        """
        element = getattr(self, property_name)
        if element is None:
            return ""
        if element.text is None:
            return ""
        return element.text


# document.py
"""Custom element classes that correspond to the document part, e.g. <w:document>."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable, List

from docx.oxml.section import CT_SectPr
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrMore, ZeroOrOne

if TYPE_CHECKING:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P


class CT_Document(BaseOxmlElement):
    """``<w:document>`` element, the root element of a document.xml file."""

    body: CT_Body = ZeroOrOne("w:body")  # pyright: ignore[reportAssignmentType]

    @property
    def sectPr_lst(self) -> List[CT_SectPr]:
        """All `w:sectPr` elements directly accessible from document element.

        Note this does not include a `sectPr` child in a paragraphs wrapped in
        revision marks or other intervening layer, perhaps `w:sdt` or customXml
        elements.

        `w:sectPr` elements appear in document order. The last one is always
        `w:body/w:sectPr`, all preceding are `w:p/w:pPr/w:sectPr`.
        """
        xpath = "./w:body/w:p/w:pPr/w:sectPr | ./w:body/w:sectPr"
        return self.xpath(xpath)


class CT_Body(BaseOxmlElement):
    """`w:body`, the container element for the main document story in `document.xml`."""

    add_p: Callable[[], CT_P]
    get_or_add_sectPr: Callable[[], CT_SectPr]
    p_lst: List[CT_P]
    tbl_lst: List[CT_Tbl]

    _insert_tbl: Callable[[CT_Tbl], CT_Tbl]

    p = ZeroOrMore("w:p", successors=("w:sectPr",))
    tbl = ZeroOrMore("w:tbl", successors=("w:sectPr",))
    sectPr: CT_SectPr | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:sectPr", successors=()
    )

    def add_section_break(self) -> CT_SectPr:
        """Return `w:sectPr` element for new section added at end of document.

        The last `w:sectPr` becomes the second-to-last, with the new `w:sectPr` being an
        exact clone of the previous one, except that all header and footer references
        are removed (and are therefore now "inherited" from the prior section).

        A copy of the previously-last `w:sectPr` will now appear in a new `w:p` at the
        end of the document. The returned `w:sectPr` is the sentinel `w:sectPr` for the
        document (and as implemented, `is` the prior sentinel `w:sectPr` with headers
        and footers removed).
        """
        # ---get the sectPr at file-end, which controls last section (sections[-1])---
        sentinel_sectPr = self.get_or_add_sectPr()
        # ---add exact copy to new `w:p` element; that is now second-to last section---
        self.add_p().set_sectPr(sentinel_sectPr.clone())
        # ---remove any header or footer references from "new" last section---
        for hdrftr_ref in sentinel_sectPr.xpath("w:headerReference|w:footerReference"):
            sentinel_sectPr.remove(hdrftr_ref)
        # ---the sentinel `w:sectPr` now controls the new last section---
        return sentinel_sectPr

    def clear_content(self):
        """Remove all content child elements from this <w:body> element.

        Leave the <w:sectPr> element if it is present.
        """
        for content_elm in self.xpath("./*[not(self::w:sectPr)]"):
            self.remove(content_elm)

    @property
    def inner_content_elements(self) -> List[CT_P | CT_Tbl]:
        """Generate all `w:p` and `w:tbl` elements in this document-body.

        Elements appear in document order. Elements shaded by nesting in a `w:ins` or
        other "wrapper" element will not be included.
        """
        return self.xpath("./w:p | ./w:tbl")


# drawing.py
"""Custom element-classes for DrawingML-related elements like `<w:drawing>`.

For legacy reasons, many DrawingML-related elements are in `docx.oxml.shape`. Expect
those to move over here as we have reason to touch them.
"""

from docx.oxml.xmlchemy import BaseOxmlElement


class CT_Drawing(BaseOxmlElement):
    """`<w:drawing>` element, containing a DrawingML object like a picture or chart."""


# exceptions.py
"""Exceptions for oxml sub-package."""


class XmlchemyError(Exception):
    """Generic error class."""


class InvalidXmlError(XmlchemyError):
    """Raised when invalid XML is encountered, such as on attempt to access a missing
    required child element."""


# ns.py
"""Namespace-related objects."""

from __future__ import annotations

from typing import Any, Dict

nsmap = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcmitype": "http://purl.org/dc/dcmitype/",
    "dcterms": "http://purl.org/dc/terms/",
    "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "sl": "http://schemas.openxmlformats.org/schemaLibrary/2006/main",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "xml": "http://www.w3.org/XML/1998/namespace",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
}

pfxmap = {value: key for key, value in nsmap.items()}


class NamespacePrefixedTag(str):
    """Value object that knows the semantics of an XML tag having a namespace prefix."""

    def __new__(cls, nstag: str, *args: Any):
        return super(NamespacePrefixedTag, cls).__new__(cls, nstag)

    def __init__(self, nstag: str):
        self._pfx, self._local_part = nstag.split(":")
        self._ns_uri = nsmap[self._pfx]

    @property
    def clark_name(self) -> str:
        return "{%s}%s" % (self._ns_uri, self._local_part)

    @classmethod
    def from_clark_name(cls, clark_name: str) -> NamespacePrefixedTag:
        nsuri, local_name = clark_name[1:].split("}")
        nstag = "%s:%s" % (pfxmap[nsuri], local_name)
        return cls(nstag)

    @property
    def local_part(self) -> str:
        """The local part of this tag.

        E.g. "foobar" is returned for tag "f:foobar".
        """
        return self._local_part

    @property
    def nsmap(self) -> Dict[str, str]:
        """Single-member dict mapping prefix of this tag to it's namespace name.

        Example: `{"f": "http://foo/bar"}`. This is handy for passing to xpath calls
        and other uses.
        """
        return {self._pfx: self._ns_uri}

    @property
    def nspfx(self) -> str:
        """The namespace-prefix for this tag.

        For example, "f" is returned for tag "f:foobar".
        """
        return self._pfx

    @property
    def nsuri(self) -> str:
        """The namespace URI for this tag.

        For example, "http://foo/bar" would be returned for tag "f:foobar" if the "f"
        prefix maps to "http://foo/bar" in nsmap.
        """
        return self._ns_uri


def nsdecls(*prefixes: str) -> str:
    """Namespace declaration including each namespace-prefix in `prefixes`.

    Handy for adding required namespace declarations to a tree root element.
    """
    return " ".join(['xmlns:%s="%s"' % (pfx, nsmap[pfx]) for pfx in prefixes])


def nspfxmap(*nspfxs: str) -> Dict[str, str]:
    """Subset namespace-prefix mappings specified by *nspfxs*.

    Any number of namespace prefixes can be supplied, e.g. namespaces("a", "r", "p").
    """
    return {pfx: nsmap[pfx] for pfx in nspfxs}


def qn(tag: str) -> str:
    """Stands for "qualified name".

    This utility function converts a familiar namespace-prefixed tag name like "w:p"
    into a Clark-notation qualified tag name for lxml. For example, `qn("w:p")` returns
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p".
    """
    prefix, tagroot = tag.split(":")
    uri = nsmap[prefix]
    return "{%s}%s" % (uri, tagroot)


# numbering.py
"""Custom element classes related to the numbering part."""

from docx.oxml.parser import OxmlElement
from docx.oxml.shared import CT_DecimalNumber
from docx.oxml.simpletypes import ST_DecimalNumber
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OneAndOnlyOne,
    RequiredAttribute,
    ZeroOrMore,
    ZeroOrOne,
)


class CT_Num(BaseOxmlElement):
    """``<w:num>`` element, which represents a concrete list definition instance, having
    a required child <w:abstractNumId> that references an abstract numbering definition
    that defines most of the formatting details."""

    abstractNumId = OneAndOnlyOne("w:abstractNumId")
    lvlOverride = ZeroOrMore("w:lvlOverride")
    numId = RequiredAttribute("w:numId", ST_DecimalNumber)

    def add_lvlOverride(self, ilvl):
        """Return a newly added CT_NumLvl (<w:lvlOverride>) element having its ``ilvl``
        attribute set to `ilvl`."""
        return self._add_lvlOverride(ilvl=ilvl)

    @classmethod
    def new(cls, num_id, abstractNum_id):
        """Return a new ``<w:num>`` element having numId of `num_id` and having a
        ``<w:abstractNumId>`` child with val attribute set to `abstractNum_id`."""
        num = OxmlElement("w:num")
        num.numId = num_id
        abstractNumId = CT_DecimalNumber.new("w:abstractNumId", abstractNum_id)
        num.append(abstractNumId)
        return num


class CT_NumLvl(BaseOxmlElement):
    """``<w:lvlOverride>`` element, which identifies a level in a list definition to
    override with settings it contains."""

    startOverride = ZeroOrOne("w:startOverride", successors=("w:lvl",))
    ilvl = RequiredAttribute("w:ilvl", ST_DecimalNumber)

    def add_startOverride(self, val):
        """Return a newly added CT_DecimalNumber element having tagname
        ``w:startOverride`` and ``val`` attribute set to `val`."""
        return self._add_startOverride(val=val)


class CT_NumPr(BaseOxmlElement):
    """A ``<w:numPr>`` element, a container for numbering properties applied to a
    paragraph."""

    ilvl = ZeroOrOne("w:ilvl", successors=("w:numId", "w:numberingChange", "w:ins"))
    numId = ZeroOrOne("w:numId", successors=("w:numberingChange", "w:ins"))

    # @ilvl.setter
    # def _set_ilvl(self, val):
    #     """
    #     Get or add a <w:ilvl> child and set its ``w:val`` attribute to `val`.
    #     """
    #     ilvl = self.get_or_add_ilvl()
    #     ilvl.val = val

    # @numId.setter
    # def numId(self, val):
    #     """
    #     Get or add a <w:numId> child and set its ``w:val`` attribute to
    #     `val`.
    #     """
    #     numId = self.get_or_add_numId()
    #     numId.val = val


class CT_Numbering(BaseOxmlElement):
    """``<w:numbering>`` element, the root element of a numbering part, i.e.
    numbering.xml."""

    num = ZeroOrMore("w:num", successors=("w:numIdMacAtCleanup",))

    def add_num(self, abstractNum_id):
        """Return a newly added CT_Num (<w:num>) element referencing the abstract
        numbering definition identified by `abstractNum_id`."""
        next_num_id = self._next_numId
        num = CT_Num.new(next_num_id, abstractNum_id)
        return self._insert_num(num)

    def num_having_numId(self, numId):
        """Return the ``<w:num>`` child element having ``numId`` attribute matching
        `numId`."""
        xpath = './w:num[@w:numId="%d"]' % numId
        try:
            return self.xpath(xpath)[0]
        except IndexError:
            raise KeyError("no <w:num> element with numId %d" % numId)

    @property
    def _next_numId(self):
        """The first ``numId`` unused by a ``<w:num>`` element, starting at 1 and
        filling any gaps in numbering between existing ``<w:num>`` elements."""
        numId_strs = self.xpath("./w:num/@w:numId")
        num_ids = [int(numId_str) for numId_str in numId_strs]
        for num in range(1, len(num_ids) + 2):
            if num not in num_ids:
                break
        return num


# parser.py
# pyright: reportImportCycles=false

"""XML parser for python-docx."""

from __future__ import annotations

from typing import TYPE_CHECKING, Dict, Type, cast

from lxml import etree

from docx.oxml.ns import NamespacePrefixedTag, nsmap

if TYPE_CHECKING:
    from docx.oxml.xmlchemy import BaseOxmlElement


# -- configure XML parser --
element_class_lookup = etree.ElementNamespaceClassLookup()
oxml_parser = etree.XMLParser(remove_blank_text=True, resolve_entities=False)
oxml_parser.set_element_class_lookup(element_class_lookup)


def parse_xml(xml: str | bytes) -> "BaseOxmlElement":
    """Root lxml element obtained by parsing XML character string `xml`.

    The custom parser is used, so custom element classes are produced for elements in
    `xml` that have them.
    """
    return cast("BaseOxmlElement", etree.fromstring(xml, oxml_parser))


def register_element_cls(tag: str, cls: Type["BaseOxmlElement"]):
    """Register an lxml custom element-class to use for `tag`.

    A instance of `cls` to be constructed when the oxml parser encounters an element
    with matching `tag`. `tag` is a string of the form `nspfx:tagroot`, e.g.
    `'w:document'`.
    """
    nspfx, tagroot = tag.split(":")
    namespace = element_class_lookup.get_namespace(nsmap[nspfx])
    namespace[tagroot] = cls


def OxmlElement(
    nsptag_str: str,
    attrs: Dict[str, str] | None = None,
    nsdecls: Dict[str, str] | None = None,
) -> BaseOxmlElement | etree._Element:  # pyright: ignore[reportPrivateUsage]
    """Return a 'loose' lxml element having the tag specified by `nsptag_str`.

    The tag in `nsptag_str` must contain the standard namespace prefix, e.g. `a:tbl`.
    The resulting element is an instance of the custom element class for this tag name
    if one is defined. A dictionary of attribute values may be provided as `attrs`; they
    are set if present. All namespaces defined in the dict `nsdecls` are declared in the
    element using the key as the prefix and the value as the namespace name. If
    `nsdecls` is not provided, a single namespace declaration is added based on the
    prefix on `nsptag_str`.
    """
    nsptag = NamespacePrefixedTag(nsptag_str)
    if nsdecls is None:
        nsdecls = nsptag.nsmap
    return oxml_parser.makeelement(nsptag.clark_name, attrib=attrs, nsmap=nsdecls)


# section.py
"""Section-related custom element classes."""

from __future__ import annotations

from copy import deepcopy
from typing import Callable, Iterator, List, Sequence, cast

from lxml import etree
from typing_extensions import TypeAlias

from docx.enum.section import WD_HEADER_FOOTER, WD_ORIENTATION, WD_SECTION_START
from docx.oxml.ns import nsmap
from docx.oxml.shared import CT_OnOff
from docx.oxml.simpletypes import ST_SignedTwipsMeasure, ST_TwipsMeasure, XsdString
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrMore,
    ZeroOrOne,
)
from docx.shared import Length, lazyproperty

BlockElement: TypeAlias = "CT_P | CT_Tbl"


class CT_HdrFtr(BaseOxmlElement):
    """`w:hdr` and `w:ftr`, the root element for header and footer part respectively."""

    add_p: Callable[[], CT_P]
    p_lst: List[CT_P]
    tbl_lst: List[CT_Tbl]

    _insert_tbl: Callable[[CT_Tbl], CT_Tbl]

    p = ZeroOrMore("w:p", successors=())
    tbl = ZeroOrMore("w:tbl", successors=())

    @property
    def inner_content_elements(self) -> List[CT_P | CT_Tbl]:
        """Generate all `w:p` and `w:tbl` elements in this header or footer.

        Elements appear in document order. Elements shaded by nesting in a `w:ins` or
        other "wrapper" element will not be included.
        """
        return self.xpath("./w:p | ./w:tbl")


class CT_HdrFtrRef(BaseOxmlElement):
    """`w:headerReference` and `w:footerReference` elements."""

    type_: WD_HEADER_FOOTER = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "w:type", WD_HEADER_FOOTER
    )
    rId: str = RequiredAttribute("r:id", XsdString)  # pyright: ignore[reportAssignmentType]


class CT_PageMar(BaseOxmlElement):
    """``<w:pgMar>`` element, defining page margins."""

    top: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:top", ST_SignedTwipsMeasure
    )
    right: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:right", ST_TwipsMeasure
    )
    bottom: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:bottom", ST_SignedTwipsMeasure
    )
    left: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:left", ST_TwipsMeasure
    )
    header: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:header", ST_TwipsMeasure
    )
    footer: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:footer", ST_TwipsMeasure
    )
    gutter: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:gutter", ST_TwipsMeasure
    )


class CT_PageSz(BaseOxmlElement):
    """``<w:pgSz>`` element, defining page dimensions and orientation."""

    w: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:w", ST_TwipsMeasure
    )
    h: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:h", ST_TwipsMeasure
    )
    orient: WD_ORIENTATION = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:orient", WD_ORIENTATION, default=WD_ORIENTATION.PORTRAIT
    )


class CT_SectPr(BaseOxmlElement):
    """`w:sectPr` element, the container element for section properties."""

    get_or_add_pgMar: Callable[[], CT_PageMar]
    get_or_add_pgSz: Callable[[], CT_PageSz]
    get_or_add_titlePg: Callable[[], CT_OnOff]
    get_or_add_type: Callable[[], CT_SectType]
    _add_footerReference: Callable[[], CT_HdrFtrRef]
    _add_headerReference: Callable[[], CT_HdrFtrRef]
    _remove_titlePg: Callable[[], None]
    _remove_type: Callable[[], None]

    _tag_seq = (
        "w:footnotePr",
        "w:endnotePr",
        "w:type",
        "w:pgSz",
        "w:pgMar",
        "w:paperSrc",
        "w:pgBorders",
        "w:lnNumType",
        "w:pgNumType",
        "w:cols",
        "w:formProt",
        "w:vAlign",
        "w:noEndnote",
        "w:titlePg",
        "w:textDirection",
        "w:bidi",
        "w:rtlGutter",
        "w:docGrid",
        "w:printerSettings",
        "w:sectPrChange",
    )
    headerReference = ZeroOrMore("w:headerReference", successors=_tag_seq)
    footerReference = ZeroOrMore("w:footerReference", successors=_tag_seq)
    type: CT_SectType | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:type", successors=_tag_seq[3:]
    )
    pgSz: CT_PageSz | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:pgSz", successors=_tag_seq[4:]
    )
    pgMar: CT_PageMar | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:pgMar", successors=_tag_seq[5:]
    )
    titlePg: CT_OnOff | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:titlePg", successors=_tag_seq[14:]
    )
    del _tag_seq

    def add_footerReference(self, type_: WD_HEADER_FOOTER, rId: str) -> CT_HdrFtrRef:
        """Return newly added CT_HdrFtrRef element of `type_` with `rId`.

        The element tag is `w:footerReference`.
        """
        footerReference = self._add_footerReference()
        footerReference.type_ = type_
        footerReference.rId = rId
        return footerReference

    def add_headerReference(self, type_: WD_HEADER_FOOTER, rId: str) -> CT_HdrFtrRef:
        """Return newly added CT_HdrFtrRef element of `type_` with `rId`.

        The element tag is `w:headerReference`.
        """
        headerReference = self._add_headerReference()
        headerReference.type_ = type_
        headerReference.rId = rId
        return headerReference

    @property
    def bottom_margin(self) -> Length | None:
        """Value of the `w:bottom` attr of `<w:pgMar>` child element, as |Length|.

        |None| when either the element or the attribute is not present.
        """
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.bottom

    @bottom_margin.setter
    def bottom_margin(self, value: int | Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.bottom = value if value is None or isinstance(value, Length) else Length(value)

    def clone(self) -> CT_SectPr:
        """Return an exact duplicate of this ``<w:sectPr>`` element tree suitable for
        use in adding a section break.

        All rsid* attributes are removed from the root ``<w:sectPr>`` element.
        """
        cloned_sectPr = deepcopy(self)
        cloned_sectPr.attrib.clear()
        return cloned_sectPr

    @property
    def footer(self) -> Length | None:
        """Distance from bottom edge of page to bottom edge of the footer.

        This is the value of the `w:footer` attribute in the `w:pgMar` child element,
        as a |Length| object, or |None| if either the element or the attribute is not
        present.
        """
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.footer

    @footer.setter
    def footer(self, value: int | Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.footer = value if value is None or isinstance(value, Length) else Length(value)

    def get_footerReference(self, type_: WD_HEADER_FOOTER) -> CT_HdrFtrRef | None:
        """Return footerReference element of `type_` or None if not present."""
        path = "./w:footerReference[@w:type='%s']" % WD_HEADER_FOOTER.to_xml(type_)
        footerReferences = self.xpath(path)
        if not footerReferences:
            return None
        return footerReferences[0]

    def get_headerReference(self, type_: WD_HEADER_FOOTER) -> CT_HdrFtrRef | None:
        """Return headerReference element of `type_` or None if not present."""
        matching_headerReferences = self.xpath(
            "./w:headerReference[@w:type='%s']" % WD_HEADER_FOOTER.to_xml(type_)
        )
        if len(matching_headerReferences) == 0:
            return None
        return matching_headerReferences[0]

    @property
    def gutter(self) -> Length | None:
        """The value of the ``w:gutter`` attribute in the ``<w:pgMar>`` child element,
        as a |Length| object, or |None| if either the element or the attribute is not
        present."""
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.gutter

    @gutter.setter
    def gutter(self, value: int | Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.gutter = value if value is None or isinstance(value, Length) else Length(value)

    @property
    def header(self) -> Length | None:
        """Distance from top edge of page to top edge of header.

        This value comes from the `w:header` attribute on the `w:pgMar` child element.
        |None| if either the element or the attribute is not present.
        """
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.header

    @header.setter
    def header(self, value: int | Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.header = value if value is None or isinstance(value, Length) else Length(value)

    def iter_inner_content(self) -> Iterator[CT_P | CT_Tbl]:
        """Generate all `w:p` and `w:tbl` elements in this section.

        Elements appear in document order. Elements shaded by nesting in a `w:ins` or
        other "wrapper" element will not be included.
        """
        return _SectBlockElementIterator.iter_sect_block_elements(self)

    @property
    def left_margin(self) -> Length | None:
        """The value of the ``w:left`` attribute in the ``<w:pgMar>`` child element, as
        a |Length| object, or |None| if either the element or the attribute is not
        present."""
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.left

    @left_margin.setter
    def left_margin(self, value: int | Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.left = value if value is None or isinstance(value, Length) else Length(value)

    @property
    def orientation(self) -> WD_ORIENTATION:
        """`WD_ORIENTATION` member indicating page-orientation for this section.

        This is the value of the `orient` attribute on the `w:pgSz` child, or
        `WD_ORIENTATION.PORTRAIT` if not present.
        """
        pgSz = self.pgSz
        if pgSz is None:
            return WD_ORIENTATION.PORTRAIT
        return pgSz.orient

    @orientation.setter
    def orientation(self, value: WD_ORIENTATION | None):
        pgSz = self.get_or_add_pgSz()
        pgSz.orient = value if value else WD_ORIENTATION.PORTRAIT

    @property
    def page_height(self) -> Length | None:
        """Value in EMU of the `h` attribute of the `w:pgSz` child element.

        |None| if not present.
        """
        pgSz = self.pgSz
        if pgSz is None:
            return None
        return pgSz.h

    @page_height.setter
    def page_height(self, value: Length | None):
        pgSz = self.get_or_add_pgSz()
        pgSz.h = value

    @property
    def page_width(self) -> Length | None:
        """Value in EMU of the ``w`` attribute of the ``<w:pgSz>`` child element.

        |None| if not present.
        """
        pgSz = self.pgSz
        if pgSz is None:
            return None
        return pgSz.w

    @page_width.setter
    def page_width(self, value: Length | None):
        pgSz = self.get_or_add_pgSz()
        pgSz.w = value

    @property
    def preceding_sectPr(self) -> CT_SectPr | None:
        """SectPr immediately preceding this one or None if this is the first."""
        # -- [1] predicate returns list of zero or one value --
        preceding_sectPrs = self.xpath("./preceding::w:sectPr[1]")
        return preceding_sectPrs[0] if len(preceding_sectPrs) > 0 else None

    def remove_footerReference(self, type_: WD_HEADER_FOOTER) -> str:
        """Return rId of w:footerReference child of `type_` after removing it."""
        footerReference = self.get_footerReference(type_)
        if footerReference is None:
            # -- should never happen, but to satisfy type-check and just in case --
            raise ValueError("CT_SectPr has no footer reference")
        rId = footerReference.rId
        self.remove(footerReference)
        return rId

    def remove_headerReference(self, type_: WD_HEADER_FOOTER):
        """Return rId of w:headerReference child of `type_` after removing it."""
        headerReference = self.get_headerReference(type_)
        if headerReference is None:
            # -- should never happen, but to satisfy type-check and just in case --
            raise ValueError("CT_SectPr has no header reference")
        rId = headerReference.rId
        self.remove(headerReference)
        return rId

    @property
    def right_margin(self) -> Length | None:
        """The value of the ``w:right`` attribute in the ``<w:pgMar>`` child element, as
        a |Length| object, or |None| if either the element or the attribute is not
        present."""
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.right

    @right_margin.setter
    def right_margin(self, value: Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.right = value

    @property
    def start_type(self) -> WD_SECTION_START:
        """The member of the ``WD_SECTION_START`` enumeration corresponding to the value
        of the ``val`` attribute of the ``<w:type>`` child element, or
        ``WD_SECTION_START.NEW_PAGE`` if not present."""
        type = self.type
        if type is None or type.val is None:
            return WD_SECTION_START.NEW_PAGE
        return type.val

    @start_type.setter
    def start_type(self, value: WD_SECTION_START | None):
        if value is None or value is WD_SECTION_START.NEW_PAGE:
            self._remove_type()
            return
        type = self.get_or_add_type()
        type.val = value

    @property
    def titlePg_val(self) -> bool:
        """Value of `w:titlePg/@val` or |False| if `./w:titlePg` is not present."""
        titlePg = self.titlePg
        if titlePg is None:
            return False
        return titlePg.val

    @titlePg_val.setter
    def titlePg_val(self, value: bool | None):
        if value in [None, False]:
            self._remove_titlePg()
        else:
            self.get_or_add_titlePg().val = True

    @property
    def top_margin(self) -> Length | None:
        """The value of the ``w:top`` attribute in the ``<w:pgMar>`` child element, as a
        |Length| object, or |None| if either the element or the attribute is not
        present."""
        pgMar = self.pgMar
        if pgMar is None:
            return None
        return pgMar.top

    @top_margin.setter
    def top_margin(self, value: Length | None):
        pgMar = self.get_or_add_pgMar()
        pgMar.top = value


class CT_SectType(BaseOxmlElement):
    """``<w:sectType>`` element, defining the section start type."""

    val: WD_SECTION_START | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", WD_SECTION_START
    )


# == HELPERS =========================================================================


class _SectBlockElementIterator:
    """Generates the block-item XML elements in a section.

    A block-item element is a `CT_P` (paragraph) or a `CT_Tbl` (table).
    """

    _compiled_blocks_xpath: etree.XPath | None = None
    _compiled_count_xpath: etree.XPath | None = None

    def __init__(self, sectPr: CT_SectPr):
        self._sectPr = sectPr

    @classmethod
    def iter_sect_block_elements(cls, sectPr: CT_SectPr) -> Iterator[BlockElement]:
        """Generate each CT_P or CT_Tbl element within extents governed by `sectPr`."""
        return cls(sectPr)._iter_sect_block_elements()

    def _iter_sect_block_elements(self) -> Iterator[BlockElement]:
        """Generate each CT_P or CT_Tbl element in section."""
        # -- General strategy is to get all block (<w;p> and <w:tbl>) elements from
        # -- start of doc to and including this section, then compute the count of those
        # -- elements that came from prior sections and skip that many to leave only the
        # -- ones in this section. It's possible to express this "between here and
        # -- there" (end of prior section and end of this one) concept in XPath, but it
        # -- would be harder to follow because there are special cases (e.g. no prior
        # -- section) and the boundary expressions are fairly hairy. I also believe it
        # -- would be computationally more expensive than doing it this straighforward
        # -- albeit (theoretically) slightly wasteful way.

        sectPr, sectPrs = self._sectPr, self._sectPrs
        sectPr_idx = sectPrs.index(sectPr)

        # -- count block items belonging to prior sections --
        n_blks_to_skip = (
            0
            if sectPr_idx == 0
            else self._count_of_blocks_in_and_above_section(sectPrs[sectPr_idx - 1])
        )

        # -- and skip those in set of all blks from doc start to end of this section --
        for element in self._blocks_in_and_above_section(sectPr)[n_blks_to_skip:]:
            yield element

    def _blocks_in_and_above_section(self, sectPr: CT_SectPr) -> Sequence[BlockElement]:
        """All ps and tbls in section defined by `sectPr` and all prior sections."""
        if self._compiled_blocks_xpath is None:
            self._compiled_blocks_xpath = etree.XPath(
                self._blocks_in_and_above_section_xpath,
                namespaces=nsmap,
                regexp=False,
            )
        xpath = self._compiled_blocks_xpath
        # -- XPath callable results are Any (basically), so need a cast. --
        return cast(Sequence[BlockElement], xpath(sectPr))

    @lazyproperty
    def _blocks_in_and_above_section_xpath(self) -> str:
        """XPath expr for ps and tbls in context of a sectPr and all prior sectPrs."""
        # -- "p_sect" is a section with sectPr located at w:p/w:pPr/w:sectPr.
        # -- "body_sect" is a section with sectPr located at w:body/w:sectPr. The last
        # -- section in the document is a "body_sect". All others are of the "p_sect"
        # -- variety. "term" means "terminal", like the last p or tbl in the section.
        # -- "pred" means "predecessor", like a preceding p or tbl in the section.

        # -- the terminal block in a p-based sect is the p the sectPr appears in --
        p_sect_term_block = "./parent::w:pPr/parent::w:p"
        # -- the terminus of a body-based sect is the sectPr itself (not a block) --
        body_sect_term = "self::w:sectPr[parent::w:body]"
        # -- all the ps and tbls preceding (but not including) the context node --
        pred_ps_and_tbls = "preceding-sibling::*[self::w:p | self::w:tbl]"

        # -- p_sect_term_block and body_sect_term(inus) are mutually exclusive. So the
        # -- result is either the union of nodes found by the first two selectors or the
        # -- nodes found by the last selector, never both.
        return (
            # -- include the p containing a sectPr --
            f"{p_sect_term_block}"
            # -- along with all the blocks that precede it --
            f" | {p_sect_term_block}/{pred_ps_and_tbls}"
            # -- or all the preceding blocks if sectPr is body-based (last sectPr) --
            f" | {body_sect_term}/{pred_ps_and_tbls}"
        )

    def _count_of_blocks_in_and_above_section(self, sectPr: CT_SectPr) -> int:
        """All ps and tbls in section defined by `sectPr` and all prior sections."""
        if self._compiled_count_xpath is None:
            self._compiled_count_xpath = etree.XPath(
                f"count({self._blocks_in_and_above_section_xpath})",
                namespaces=nsmap,
                regexp=False,
            )
        xpath = self._compiled_count_xpath
        # -- numeric XPath results are always float, so need an int() conversion --
        return int(cast(float, xpath(sectPr)))

    @lazyproperty
    def _sectPrs(self) -> Sequence[CT_SectPr]:
        """All w:sectPr elements in document, in document-order."""
        return self._sectPr.xpath(
            "/w:document/w:body/w:p/w:pPr/w:sectPr | /w:document/w:body/w:sectPr",
        )


# settings.py
"""Custom element classes related to document settings."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable

from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne

if TYPE_CHECKING:
    from docx.oxml.shared import CT_OnOff


class CT_Settings(BaseOxmlElement):
    """`w:settings` element, root element for the settings part."""

    get_or_add_evenAndOddHeaders: Callable[[], CT_OnOff]
    _remove_evenAndOddHeaders: Callable[[], None]

    _tag_seq = (
        "w:writeProtection",
        "w:view",
        "w:zoom",
        "w:removePersonalInformation",
        "w:removeDateAndTime",
        "w:doNotDisplayPageBoundaries",
        "w:displayBackgroundShape",
        "w:printPostScriptOverText",
        "w:printFractionalCharacterWidth",
        "w:printFormsData",
        "w:embedTrueTypeFonts",
        "w:embedSystemFonts",
        "w:saveSubsetFonts",
        "w:saveFormsData",
        "w:mirrorMargins",
        "w:alignBordersAndEdges",
        "w:bordersDoNotSurroundHeader",
        "w:bordersDoNotSurroundFooter",
        "w:gutterAtTop",
        "w:hideSpellingErrors",
        "w:hideGrammaticalErrors",
        "w:activeWritingStyle",
        "w:proofState",
        "w:formsDesign",
        "w:attachedTemplate",
        "w:linkStyles",
        "w:stylePaneFormatFilter",
        "w:stylePaneSortMethod",
        "w:documentType",
        "w:mailMerge",
        "w:revisionView",
        "w:trackRevisions",
        "w:doNotTrackMoves",
        "w:doNotTrackFormatting",
        "w:documentProtection",
        "w:autoFormatOverride",
        "w:styleLockTheme",
        "w:styleLockQFSet",
        "w:defaultTabStop",
        "w:autoHyphenation",
        "w:consecutiveHyphenLimit",
        "w:hyphenationZone",
        "w:doNotHyphenateCaps",
        "w:showEnvelope",
        "w:summaryLength",
        "w:clickAndTypeStyle",
        "w:defaultTableStyle",
        "w:evenAndOddHeaders",
        "w:bookFoldRevPrinting",
        "w:bookFoldPrinting",
        "w:bookFoldPrintingSheets",
        "w:drawingGridHorizontalSpacing",
        "w:drawingGridVerticalSpacing",
        "w:displayHorizontalDrawingGridEvery",
        "w:displayVerticalDrawingGridEvery",
        "w:doNotUseMarginsForDrawingGridOrigin",
        "w:drawingGridHorizontalOrigin",
        "w:drawingGridVerticalOrigin",
        "w:doNotShadeFormData",
        "w:noPunctuationKerning",
        "w:characterSpacingControl",
        "w:printTwoOnOne",
        "w:strictFirstAndLastChars",
        "w:noLineBreaksAfter",
        "w:noLineBreaksBefore",
        "w:savePreviewPicture",
        "w:doNotValidateAgainstSchema",
        "w:saveInvalidXml",
        "w:ignoreMixedContent",
        "w:alwaysShowPlaceholderText",
        "w:doNotDemarcateInvalidXml",
        "w:saveXmlDataOnly",
        "w:useXSLTWhenSaving",
        "w:saveThroughXslt",
        "w:showXMLTags",
        "w:alwaysMergeEmptyNamespace",
        "w:updateFields",
        "w:hdrShapeDefaults",
        "w:footnotePr",
        "w:endnotePr",
        "w:compat",
        "w:docVars",
        "w:rsids",
        "m:mathPr",
        "w:attachedSchema",
        "w:themeFontLang",
        "w:clrSchemeMapping",
        "w:doNotIncludeSubdocsInStats",
        "w:doNotAutoCompressPictures",
        "w:forceUpgrade",
        "w:captions",
        "w:readModeInkLockDown",
        "w:smartTagType",
        "sl:schemaLibrary",
        "w:shapeDefaults",
        "w:doNotEmbedSmartTags",
        "w:decimalSymbol",
        "w:listSeparator",
    )
    evenAndOddHeaders: CT_OnOff | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:evenAndOddHeaders", successors=_tag_seq[48:]
    )
    del _tag_seq

    @property
    def evenAndOddHeaders_val(self) -> bool:
        """Value of `w:evenAndOddHeaders/@w:val` or |None| if not present."""
        evenAndOddHeaders = self.evenAndOddHeaders
        if evenAndOddHeaders is None:
            return False
        return evenAndOddHeaders.val

    @evenAndOddHeaders_val.setter
    def evenAndOddHeaders_val(self, value: bool | None):
        if value is None or value is False:
            self._remove_evenAndOddHeaders()
            return

        self.get_or_add_evenAndOddHeaders().val = value


# shape.py
"""Custom element classes for shape-related elements like `<w:inline>`."""

from __future__ import annotations

from typing import TYPE_CHECKING, cast

from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from docx.oxml.simpletypes import (
    ST_Coordinate,
    ST_DrawingElementId,
    ST_PositiveCoordinate,
    ST_RelationshipId,
    XsdString,
    XsdToken,
)
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OneAndOnlyOne,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrOne,
)

if TYPE_CHECKING:
    from docx.shared import Length


class CT_Anchor(BaseOxmlElement):
    """`<wp:anchor>` element, container for a "floating" shape."""


class CT_Blip(BaseOxmlElement):
    """``<a:blip>`` element, specifies image source and adjustments such as alpha and
    tint."""

    embed: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "r:embed", ST_RelationshipId
    )
    link: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "r:link", ST_RelationshipId
    )


class CT_BlipFillProperties(BaseOxmlElement):
    """``<pic:blipFill>`` element, specifies picture properties."""

    blip: CT_Blip = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "a:blip", successors=("a:srcRect", "a:tile", "a:stretch")
    )


class CT_GraphicalObject(BaseOxmlElement):
    """``<a:graphic>`` element, container for a DrawingML object."""

    graphicData: CT_GraphicalObjectData = OneAndOnlyOne(  # pyright: ignore[reportAssignmentType]
        "a:graphicData"
    )


class CT_GraphicalObjectData(BaseOxmlElement):
    """``<a:graphicData>`` element, container for the XML of a DrawingML object."""

    pic: CT_Picture = ZeroOrOne("pic:pic")  # pyright: ignore[reportAssignmentType]
    uri: str = RequiredAttribute("uri", XsdToken)  # pyright: ignore[reportAssignmentType]


class CT_Inline(BaseOxmlElement):
    """`<wp:inline>` element, container for an inline shape."""

    extent: CT_PositiveSize2D = OneAndOnlyOne("wp:extent")  # pyright: ignore[reportAssignmentType]
    docPr: CT_NonVisualDrawingProps = OneAndOnlyOne(  # pyright: ignore[reportAssignmentType]
        "wp:docPr"
    )
    graphic: CT_GraphicalObject = OneAndOnlyOne(  # pyright: ignore[reportAssignmentType]
        "a:graphic"
    )

    @classmethod
    def new(cls, cx: Length, cy: Length, shape_id: int, pic: CT_Picture) -> CT_Inline:
        """Return a new ``<wp:inline>`` element populated with the values passed as
        parameters."""
        inline = cast(CT_Inline, parse_xml(cls._inline_xml()))
        inline.extent.cx = cx
        inline.extent.cy = cy
        inline.docPr.id = shape_id
        inline.docPr.name = "Picture %d" % shape_id
        inline.graphic.graphicData.uri = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        inline.graphic.graphicData._insert_pic(pic)
        return inline

    @classmethod
    def new_pic_inline(
        cls, shape_id: int, rId: str, filename: str, cx: Length, cy: Length
    ) -> CT_Inline:
        """Create `wp:inline` element containing a `pic:pic` element.

        The contents of the `pic:pic` element is taken from the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        inline = cls.new(cx, cy, shape_id, pic)
        inline.graphic.graphicData._insert_pic(pic)
        return inline

    @classmethod
    def _inline_xml(cls):
        return (
            "<wp:inline %s>\n"
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            "  <wp:cNvGraphicFramePr>\n"
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            "  </wp:cNvGraphicFramePr>\n"
            "  <a:graphic>\n"
            '    <a:graphicData uri="URI not set"/>\n'
            "  </a:graphic>\n"
            "</wp:inline>" % nsdecls("wp", "a", "pic", "r")
        )


class CT_NonVisualDrawingProps(BaseOxmlElement):
    """Used for ``<wp:docPr>`` element, and perhaps others.

    Specifies the id and name of a DrawingML drawing.
    """

    id = RequiredAttribute("id", ST_DrawingElementId)
    name = RequiredAttribute("name", XsdString)


class CT_NonVisualPictureProperties(BaseOxmlElement):
    """``<pic:cNvPicPr>`` element, specifies picture locking and resize behaviors."""


class CT_Picture(BaseOxmlElement):
    """``<pic:pic>`` element, a DrawingML picture."""

    nvPicPr: CT_PictureNonVisual = OneAndOnlyOne(  # pyright: ignore[reportAssignmentType]
        "pic:nvPicPr"
    )
    blipFill: CT_BlipFillProperties = OneAndOnlyOne(  # pyright: ignore[reportAssignmentType]
        "pic:blipFill"
    )
    spPr: CT_ShapeProperties = OneAndOnlyOne("pic:spPr")  # pyright: ignore[reportAssignmentType]

    @classmethod
    def new(cls, pic_id, filename, rId, cx, cy):
        """Return a new ``<pic:pic>`` element populated with the minimal contents
        required to define a viable picture element, based on the values passed as
        parameters."""
        pic = parse_xml(cls._pic_xml())
        pic.nvPicPr.cNvPr.id = pic_id
        pic.nvPicPr.cNvPr.name = filename
        pic.blipFill.blip.embed = rId
        pic.spPr.cx = cx
        pic.spPr.cy = cy
        return pic

    @classmethod
    def _pic_xml(cls):
        return (
            "<pic:pic %s>\n"
            "  <pic:nvPicPr>\n"
            '    <pic:cNvPr id="666" name="unnamed"/>\n'
            "    <pic:cNvPicPr/>\n"
            "  </pic:nvPicPr>\n"
            "  <pic:blipFill>\n"
            "    <a:blip/>\n"
            "    <a:stretch>\n"
            "      <a:fillRect/>\n"
            "    </a:stretch>\n"
            "  </pic:blipFill>\n"
            "  <pic:spPr>\n"
            "    <a:xfrm>\n"
            '      <a:off x="0" y="0"/>\n'
            '      <a:ext cx="914400" cy="914400"/>\n'
            "    </a:xfrm>\n"
            '    <a:prstGeom prst="rect"/>\n'
            "  </pic:spPr>\n"
            "</pic:pic>" % nsdecls("pic", "a", "r")
        )


class CT_PictureNonVisual(BaseOxmlElement):
    """``<pic:nvPicPr>`` element, non-visual picture properties."""

    cNvPr = OneAndOnlyOne("pic:cNvPr")


class CT_Point2D(BaseOxmlElement):
    """Used for ``<a:off>`` element, and perhaps others.

    Specifies an x, y coordinate (point).
    """

    x = RequiredAttribute("x", ST_Coordinate)
    y = RequiredAttribute("y", ST_Coordinate)


class CT_PositiveSize2D(BaseOxmlElement):
    """Used for ``<wp:extent>`` element, and perhaps others later.

    Specifies the size of a DrawingML drawing.
    """

    cx: Length = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "cx", ST_PositiveCoordinate
    )
    cy: Length = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "cy", ST_PositiveCoordinate
    )


class CT_PresetGeometry2D(BaseOxmlElement):
    """``<a:prstGeom>`` element, specifies an preset autoshape geometry, such as
    ``rect``."""


class CT_RelativeRect(BaseOxmlElement):
    """``<a:fillRect>`` element, specifying picture should fill containing rectangle
    shape."""


class CT_ShapeProperties(BaseOxmlElement):
    """``<pic:spPr>`` element, specifies size and shape of picture container."""

    xfrm = ZeroOrOne(
        "a:xfrm",
        successors=(
            "a:custGeom",
            "a:prstGeom",
            "a:ln",
            "a:effectLst",
            "a:effectDag",
            "a:scene3d",
            "a:sp3d",
            "a:extLst",
        ),
    )

    @property
    def cx(self):
        """Shape width as an instance of Emu, or None if not present."""
        xfrm = self.xfrm
        if xfrm is None:
            return None
        return xfrm.cx

    @cx.setter
    def cx(self, value):
        xfrm = self.get_or_add_xfrm()
        xfrm.cx = value

    @property
    def cy(self):
        """Shape height as an instance of Emu, or None if not present."""
        xfrm = self.xfrm
        if xfrm is None:
            return None
        return xfrm.cy

    @cy.setter
    def cy(self, value):
        xfrm = self.get_or_add_xfrm()
        xfrm.cy = value


class CT_StretchInfoProperties(BaseOxmlElement):
    """``<a:stretch>`` element, specifies how picture should fill its containing
    shape."""


class CT_Transform2D(BaseOxmlElement):
    """``<a:xfrm>`` element, specifies size and shape of picture container."""

    off = ZeroOrOne("a:off", successors=("a:ext",))
    ext = ZeroOrOne("a:ext", successors=())

    @property
    def cx(self):
        ext = self.ext
        if ext is None:
            return None
        return ext.cx

    @cx.setter
    def cx(self, value):
        ext = self.get_or_add_ext()
        ext.cx = value

    @property
    def cy(self):
        ext = self.ext
        if ext is None:
            return None
        return ext.cy

    @cy.setter
    def cy(self, value):
        ext = self.get_or_add_ext()
        ext.cy = value


# shared.py
"""Objects shared by modules in the docx.oxml subpackage."""

from __future__ import annotations

from typing import cast

from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.simpletypes import ST_DecimalNumber, ST_OnOff, ST_String
from docx.oxml.xmlchemy import BaseOxmlElement, OptionalAttribute, RequiredAttribute


class CT_DecimalNumber(BaseOxmlElement):
    """Used for ``<w:numId>``, ``<w:ilvl>``, ``<w:abstractNumId>`` and several others,
    containing a text representation of a decimal number (e.g. 42) in its ``val``
    attribute."""

    val: int = RequiredAttribute("w:val", ST_DecimalNumber)  # pyright: ignore[reportAssignmentType]

    @classmethod
    def new(cls, nsptagname: str, val: int):
        """Return a new ``CT_DecimalNumber`` element having tagname `nsptagname` and
        ``val`` attribute set to `val`."""
        return OxmlElement(nsptagname, attrs={qn("w:val"): str(val)})


class CT_OnOff(BaseOxmlElement):
    """Used for `w:b`, `w:i` elements and others.

    Contains a bool-ish string in its `val` attribute, xsd:boolean plus "on" and
    "off". Defaults to `True`, so `<w:b>` for example means "bold is turned on".
    """

    val: bool = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", ST_OnOff, default=True
    )


class CT_String(BaseOxmlElement):
    """Used for `w:pStyle` and `w:tblStyle` elements and others.

    In those cases, it containing a style name in its `val` attribute.
    """

    val: str = RequiredAttribute("w:val", ST_String)  # pyright: ignore[reportAssignmentType]

    @classmethod
    def new(cls, nsptagname: str, val: str):
        """Return a new ``CT_String`` element with tagname `nsptagname` and ``val``
        attribute set to `val`."""
        elm = cast(CT_String, OxmlElement(nsptagname))
        elm.val = val
        return elm


# simpletypes.py
# pyright: reportImportCycles=false

"""Simple-type classes, corresponding to ST_* schema items.

These provide validation and format translation for values stored in XML element
attributes. Naming generally corresponds to the simple type in the associated XML
schema.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, Tuple

from docx.exceptions import InvalidXmlError
from docx.shared import Emu, Pt, RGBColor, Twips

if TYPE_CHECKING:
    from docx.shared import Length


class BaseSimpleType:
    """Base class for simple-types."""

    @classmethod
    def from_xml(cls, xml_value: str) -> Any:
        return cls.convert_from_xml(xml_value)

    @classmethod
    def to_xml(cls, value: Any) -> str:
        cls.validate(value)
        str_value = cls.convert_to_xml(value)
        return str_value

    @classmethod
    def convert_from_xml(cls, str_value: str) -> Any:
        return int(str_value)

    @classmethod
    def convert_to_xml(cls, value: Any) -> str: ...

    @classmethod
    def validate(cls, value: Any) -> None: ...

    @classmethod
    def validate_int(cls, value: object):
        if not isinstance(value, int):
            raise TypeError("value must be <type 'int'>, got %s" % type(value))

    @classmethod
    def validate_int_in_range(cls, value: int, min_inclusive: int, max_inclusive: int) -> None:
        cls.validate_int(value)
        if value < min_inclusive or value > max_inclusive:
            raise ValueError(
                "value must be in range %d to %d inclusive, got %d"
                % (min_inclusive, max_inclusive, value)
            )

    @classmethod
    def validate_string(cls, value: Any) -> str:
        if not isinstance(value, str):
            raise TypeError("value must be a string, got %s" % type(value))
        return value


class BaseIntType(BaseSimpleType):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> int:
        return int(str_value)

    @classmethod
    def convert_to_xml(cls, value: int) -> str:
        return str(value)

    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int(value)


class BaseStringType(BaseSimpleType):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> str:
        return str_value

    @classmethod
    def convert_to_xml(cls, value: str) -> str:
        return value

    @classmethod
    def validate(cls, value: str):
        cls.validate_string(value)


class BaseStringEnumerationType(BaseStringType):
    _members: Tuple[str, ...]

    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_string(value)
        if value not in cls._members:
            raise ValueError("must be one of %s, got '%s'" % (cls._members, value))


class XsdAnyUri(BaseStringType):
    """There's a regex in the spec this is supposed to meet...

    but current assessment is that spending cycles on validating wouldn't be worth it
    for the number of programming errors it would catch.
    """


class XsdBoolean(BaseSimpleType):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> bool:
        if str_value not in ("1", "0", "true", "false"):
            raise InvalidXmlError(
                "value must be one of '1', '0', 'true' or 'false', got '%s'" % str_value
            )
        return str_value in ("1", "true")

    @classmethod
    def convert_to_xml(cls, value: bool) -> str:
        return {True: "1", False: "0"}[value]

    @classmethod
    def validate(cls, value: Any) -> None:
        if value not in (True, False):
            raise TypeError(
                "only True or False (and possibly None) may be assigned, got" " '%s'" % value
            )


class XsdId(BaseStringType):
    """String that must begin with a letter or underscore and cannot contain any colons.

    Not fully validated because not used in external API.
    """

    pass


class XsdInt(BaseIntType):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, -2147483648, 2147483647)


class XsdLong(BaseIntType):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, -9223372036854775808, 9223372036854775807)


class XsdString(BaseStringType):
    pass


class XsdStringEnumeration(BaseStringEnumerationType):
    """Set of enumerated xsd:string values."""


class XsdToken(BaseStringType):
    """Xsd:string with whitespace collapsing, e.g. multiple spaces reduced to one,
    leading and trailing space stripped."""

    pass


class XsdUnsignedInt(BaseIntType):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, 0, 4294967295)


class XsdUnsignedLong(BaseIntType):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, 0, 18446744073709551615)


class ST_BrClear(XsdString):
    @classmethod
    def validate(cls, value: str) -> None:
        cls.validate_string(value)
        valid_values = ("none", "left", "right", "all")
        if value not in valid_values:
            raise ValueError("must be one of %s, got '%s'" % (valid_values, value))


class ST_BrType(XsdString):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_string(value)
        valid_values = ("page", "column", "textWrapping")
        if value not in valid_values:
            raise ValueError("must be one of %s, got '%s'" % (valid_values, value))


class ST_Coordinate(BaseIntType):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> Length:
        if "i" in str_value or "m" in str_value or "p" in str_value:
            return ST_UniversalMeasure.convert_from_xml(str_value)
        return Emu(int(str_value))

    @classmethod
    def validate(cls, value: Any) -> None:
        ST_CoordinateUnqualified.validate(value)


class ST_CoordinateUnqualified(XsdLong):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, -27273042329600, 27273042316900)


class ST_DecimalNumber(XsdInt):
    pass


class ST_DrawingElementId(XsdUnsignedInt):
    pass


class ST_HexColor(BaseStringType):
    @classmethod
    def convert_from_xml(  # pyright: ignore[reportIncompatibleMethodOverride]
        cls, str_value: str
    ) -> RGBColor | str:
        if str_value == "auto":
            return ST_HexColorAuto.AUTO
        return RGBColor.from_string(str_value)

    @classmethod
    def convert_to_xml(  # pyright: ignore[reportIncompatibleMethodOverride]
        cls, value: RGBColor
    ) -> str:
        """Keep alpha hex numerals all uppercase just for consistency."""
        # expecting 3-tuple of ints in range 0-255
        return "%02X%02X%02X" % value

    @classmethod
    def validate(cls, value: Any) -> None:
        # must be an RGBColor object ---
        if not isinstance(value, RGBColor):
            raise ValueError(
                "rgb color value must be RGBColor object, got %s %s" % (type(value), value)
            )


class ST_HexColorAuto(XsdStringEnumeration):
    """Value for `w:color/[@val="auto"] attribute setting."""

    AUTO = "auto"

    _members = (AUTO,)


class ST_HpsMeasure(XsdUnsignedLong):
    """Half-point measure, e.g. 24.0 represents 12.0 points."""

    @classmethod
    def convert_from_xml(cls, str_value: str) -> Length:
        if "m" in str_value or "n" in str_value or "p" in str_value:
            return ST_UniversalMeasure.convert_from_xml(str_value)
        return Pt(int(str_value) / 2.0)

    @classmethod
    def convert_to_xml(cls, value: int | Length) -> str:
        emu = Emu(value)
        half_points = int(emu.pt * 2)
        return str(half_points)


class ST_Merge(XsdStringEnumeration):
    """Valid values for <w:xMerge val=""> attribute."""

    CONTINUE = "continue"
    RESTART = "restart"

    _members = (CONTINUE, RESTART)


class ST_OnOff(XsdBoolean):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> bool:
        if str_value not in ("1", "0", "true", "false", "on", "off"):
            raise InvalidXmlError(
                "value must be one of '1', '0', 'true', 'false', 'on', or 'o"
                "ff', got '%s'" % str_value
            )
        return str_value in ("1", "true", "on")


class ST_PositiveCoordinate(XsdLong):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> Length:
        return Emu(int(str_value))

    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_int_in_range(value, 0, 27273042316900)


class ST_RelationshipId(XsdString):
    pass


class ST_SignedTwipsMeasure(XsdInt):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> Length:
        if "i" in str_value or "m" in str_value or "p" in str_value:
            return ST_UniversalMeasure.convert_from_xml(str_value)
        return Twips(int(round(float(str_value))))

    @classmethod
    def convert_to_xml(cls, value: int | Length) -> str:
        emu = Emu(value)
        twips = emu.twips
        return str(twips)


class ST_String(XsdString):
    pass


class ST_TblLayoutType(XsdString):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_string(value)
        valid_values = ("fixed", "autofit")
        if value not in valid_values:
            raise ValueError("must be one of %s, got '%s'" % (valid_values, value))


class ST_TblWidth(XsdString):
    @classmethod
    def validate(cls, value: Any) -> None:
        cls.validate_string(value)
        valid_values = ("auto", "dxa", "nil", "pct")
        if value not in valid_values:
            raise ValueError("must be one of %s, got '%s'" % (valid_values, value))


class ST_TwipsMeasure(XsdUnsignedLong):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> Length:
        if "i" in str_value or "m" in str_value or "p" in str_value:
            return ST_UniversalMeasure.convert_from_xml(str_value)
        return Twips(int(str_value))

    @classmethod
    def convert_to_xml(cls, value: int | Length) -> str:
        emu = Emu(value)
        twips = emu.twips
        return str(twips)


class ST_UniversalMeasure(BaseSimpleType):
    @classmethod
    def convert_from_xml(cls, str_value: str) -> Emu:
        float_part, units_part = str_value[:-2], str_value[-2:]
        quantity = float(float_part)
        multiplier = {
            "mm": 36000,
            "cm": 360000,
            "in": 914400,
            "pt": 12700,
            "pc": 152400,
            "pi": 152400,
        }[units_part]
        return Emu(int(round(quantity * multiplier)))


class ST_VerticalAlignRun(XsdStringEnumeration):
    """Valid values for `w:vertAlign/@val`."""

    BASELINE = "baseline"
    SUPERSCRIPT = "superscript"
    SUBSCRIPT = "subscript"

    _members = (BASELINE, SUPERSCRIPT, SUBSCRIPT)


# styles.py
"""Custom element classes related to the styles part."""

from __future__ import annotations

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.simpletypes import ST_DecimalNumber, ST_OnOff, ST_String
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrMore,
    ZeroOrOne,
)


def styleId_from_name(name):
    """Return the style id corresponding to `name`, taking into account special-case
    names such as 'Heading 1'."""
    return {
        "caption": "Caption",
        "heading 1": "Heading1",
        "heading 2": "Heading2",
        "heading 3": "Heading3",
        "heading 4": "Heading4",
        "heading 5": "Heading5",
        "heading 6": "Heading6",
        "heading 7": "Heading7",
        "heading 8": "Heading8",
        "heading 9": "Heading9",
    }.get(name, name.replace(" ", ""))


class CT_LatentStyles(BaseOxmlElement):
    """`w:latentStyles` element, defining behavior defaults for latent styles and
    containing `w:lsdException` child elements that each override those defaults for a
    named latent style."""

    lsdException = ZeroOrMore("w:lsdException", successors=())

    count = OptionalAttribute("w:count", ST_DecimalNumber)
    defLockedState = OptionalAttribute("w:defLockedState", ST_OnOff)
    defQFormat = OptionalAttribute("w:defQFormat", ST_OnOff)
    defSemiHidden = OptionalAttribute("w:defSemiHidden", ST_OnOff)
    defUIPriority = OptionalAttribute("w:defUIPriority", ST_DecimalNumber)
    defUnhideWhenUsed = OptionalAttribute("w:defUnhideWhenUsed", ST_OnOff)

    def bool_prop(self, attr_name):
        """Return the boolean value of the attribute having `attr_name`, or |False| if
        not present."""
        value = getattr(self, attr_name)
        if value is None:
            return False
        return value

    def get_by_name(self, name):
        """Return the `w:lsdException` child having `name`, or |None| if not found."""
        found = self.xpath('w:lsdException[@w:name="%s"]' % name)
        if not found:
            return None
        return found[0]

    def set_bool_prop(self, attr_name, value):
        """Set the on/off attribute having `attr_name` to `value`."""
        setattr(self, attr_name, bool(value))


class CT_LsdException(BaseOxmlElement):
    """``<w:lsdException>`` element, defining override visibility behaviors for a named
    latent style."""

    locked = OptionalAttribute("w:locked", ST_OnOff)
    name = RequiredAttribute("w:name", ST_String)
    qFormat = OptionalAttribute("w:qFormat", ST_OnOff)
    semiHidden = OptionalAttribute("w:semiHidden", ST_OnOff)
    uiPriority = OptionalAttribute("w:uiPriority", ST_DecimalNumber)
    unhideWhenUsed = OptionalAttribute("w:unhideWhenUsed", ST_OnOff)

    def delete(self):
        """Remove this `w:lsdException` element from the XML document."""
        self.getparent().remove(self)

    def on_off_prop(self, attr_name):
        """Return the boolean value of the attribute having `attr_name`, or |None| if
        not present."""
        return getattr(self, attr_name)

    def set_on_off_prop(self, attr_name, value):
        """Set the on/off attribute having `attr_name` to `value`."""
        setattr(self, attr_name, value)


class CT_Style(BaseOxmlElement):
    """A ``<w:style>`` element, representing a style definition."""

    _tag_seq = (
        "w:name",
        "w:aliases",
        "w:basedOn",
        "w:next",
        "w:link",
        "w:autoRedefine",
        "w:hidden",
        "w:uiPriority",
        "w:semiHidden",
        "w:unhideWhenUsed",
        "w:qFormat",
        "w:locked",
        "w:personal",
        "w:personalCompose",
        "w:personalReply",
        "w:rsid",
        "w:pPr",
        "w:rPr",
        "w:tblPr",
        "w:trPr",
        "w:tcPr",
        "w:tblStylePr",
    )
    name = ZeroOrOne("w:name", successors=_tag_seq[1:])
    basedOn = ZeroOrOne("w:basedOn", successors=_tag_seq[3:])
    next = ZeroOrOne("w:next", successors=_tag_seq[4:])
    uiPriority = ZeroOrOne("w:uiPriority", successors=_tag_seq[8:])
    semiHidden = ZeroOrOne("w:semiHidden", successors=_tag_seq[9:])
    unhideWhenUsed = ZeroOrOne("w:unhideWhenUsed", successors=_tag_seq[10:])
    qFormat = ZeroOrOne("w:qFormat", successors=_tag_seq[11:])
    locked = ZeroOrOne("w:locked", successors=_tag_seq[12:])
    pPr = ZeroOrOne("w:pPr", successors=_tag_seq[17:])
    rPr = ZeroOrOne("w:rPr", successors=_tag_seq[18:])
    del _tag_seq

    type: WD_STYLE_TYPE | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:type", WD_STYLE_TYPE
    )
    styleId: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:styleId", ST_String
    )
    default = OptionalAttribute("w:default", ST_OnOff)
    customStyle = OptionalAttribute("w:customStyle", ST_OnOff)

    @property
    def basedOn_val(self):
        """Value of `w:basedOn/@w:val` or |None| if not present."""
        basedOn = self.basedOn
        if basedOn is None:
            return None
        return basedOn.val

    @basedOn_val.setter
    def basedOn_val(self, value):
        if value is None:
            self._remove_basedOn()
        else:
            self.get_or_add_basedOn().val = value

    @property
    def base_style(self):
        """Sibling CT_Style element this style is based on or |None| if no base style or
        base style not found."""
        basedOn = self.basedOn
        if basedOn is None:
            return None
        styles = self.getparent()
        base_style = styles.get_by_id(basedOn.val)
        if base_style is None:
            return None
        return base_style

    def delete(self):
        """Remove this `w:style` element from its parent `w:styles` element."""
        self.getparent().remove(self)

    @property
    def locked_val(self):
        """Value of `w:locked/@w:val` or |False| if not present."""
        locked = self.locked
        if locked is None:
            return False
        return locked.val

    @locked_val.setter
    def locked_val(self, value):
        self._remove_locked()
        if bool(value) is True:
            locked = self._add_locked()
            locked.val = value

    @property
    def name_val(self):
        """Value of ``<w:name>`` child or |None| if not present."""
        name = self.name
        if name is None:
            return None
        return name.val

    @name_val.setter
    def name_val(self, value):
        self._remove_name()
        if value is not None:
            name = self._add_name()
            name.val = value

    @property
    def next_style(self):
        """Sibling CT_Style element identified by the value of `w:name/@w:val` or |None|
        if no value is present or no style with that style id is found."""
        next = self.next
        if next is None:
            return None
        styles = self.getparent()
        return styles.get_by_id(next.val)  # None if not found

    @property
    def qFormat_val(self):
        """Value of `w:qFormat/@w:val` or |False| if not present."""
        qFormat = self.qFormat
        if qFormat is None:
            return False
        return qFormat.val

    @qFormat_val.setter
    def qFormat_val(self, value):
        self._remove_qFormat()
        if bool(value):
            self._add_qFormat()

    @property
    def semiHidden_val(self):
        """Value of ``<w:semiHidden>`` child or |False| if not present."""
        semiHidden = self.semiHidden
        if semiHidden is None:
            return False
        return semiHidden.val

    @semiHidden_val.setter
    def semiHidden_val(self, value):
        self._remove_semiHidden()
        if bool(value) is True:
            semiHidden = self._add_semiHidden()
            semiHidden.val = value

    @property
    def uiPriority_val(self):
        """Value of ``<w:uiPriority>`` child or |None| if not present."""
        uiPriority = self.uiPriority
        if uiPriority is None:
            return None
        return uiPriority.val

    @uiPriority_val.setter
    def uiPriority_val(self, value):
        self._remove_uiPriority()
        if value is not None:
            uiPriority = self._add_uiPriority()
            uiPriority.val = value

    @property
    def unhideWhenUsed_val(self):
        """Value of `w:unhideWhenUsed/@w:val` or |False| if not present."""
        unhideWhenUsed = self.unhideWhenUsed
        if unhideWhenUsed is None:
            return False
        return unhideWhenUsed.val

    @unhideWhenUsed_val.setter
    def unhideWhenUsed_val(self, value):
        self._remove_unhideWhenUsed()
        if bool(value) is True:
            unhideWhenUsed = self._add_unhideWhenUsed()
            unhideWhenUsed.val = value


class CT_Styles(BaseOxmlElement):
    """``<w:styles>`` element, the root element of a styles part, i.e. styles.xml."""

    _tag_seq = ("w:docDefaults", "w:latentStyles", "w:style")
    latentStyles = ZeroOrOne("w:latentStyles", successors=_tag_seq[2:])
    style = ZeroOrMore("w:style", successors=())
    del _tag_seq

    def add_style_of_type(self, name, style_type, builtin):
        """Return a newly added `w:style` element having `name` and `style_type`.

        `w:style/@customStyle` is set based on the value of `builtin`.
        """
        style = self.add_style()
        style.type = style_type
        style.customStyle = None if builtin else True
        style.styleId = styleId_from_name(name)
        style.name_val = name
        return style

    def default_for(self, style_type):
        """Return `w:style[@w:type="*{style_type}*][-1]` or |None| if not found."""
        default_styles_for_type = [
            s for s in self._iter_styles() if s.type == style_type and s.default
        ]
        if not default_styles_for_type:
            return None
        # spec calls for last default in document order
        return default_styles_for_type[-1]

    def get_by_id(self, styleId: str) -> CT_Style | None:
        """`w:style` child where @styleId = `styleId`.

        |None| if not found.
        """
        xpath = f'w:style[@w:styleId="{styleId}"]'
        return next(iter(self.xpath(xpath)), None)

    def get_by_name(self, name: str) -> CT_Style | None:
        """`w:style` child with `w:name` grandchild having value `name`.

        |None| if not found.
        """
        xpath = 'w:style[w:name/@w:val="%s"]' % name
        return next(iter(self.xpath(xpath)), None)

    def _iter_styles(self):
        """Generate each of the `w:style` child elements in document order."""
        return (style for style in self.xpath("w:style"))


# table.py
"""Custom element classes for tables."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable, cast

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_DIRECTION
from docx.exceptions import InvalidSpanError
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.oxml.shared import CT_DecimalNumber
from docx.oxml.simpletypes import (
    ST_Merge,
    ST_TblLayoutType,
    ST_TblWidth,
    ST_TwipsMeasure,
    XsdInt,
)
from docx.oxml.text.paragraph import CT_P
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OneAndOnlyOne,
    OneOrMore,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrMore,
    ZeroOrOne,
)
from docx.shared import Emu, Length, Twips

if TYPE_CHECKING:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import CT_OnOff, CT_String
    from docx.oxml.text.parfmt import CT_Jc


class CT_Height(BaseOxmlElement):
    """Used for `w:trHeight` to specify a row height and row height rule."""

    val: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", ST_TwipsMeasure
    )
    hRule: WD_ROW_HEIGHT_RULE | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:hRule", WD_ROW_HEIGHT_RULE
    )


class CT_Row(BaseOxmlElement):
    """``<w:tr>`` element."""

    add_tc: Callable[[], CT_Tc]
    get_or_add_trPr: Callable[[], CT_TrPr]
    _add_trPr: Callable[[], CT_TrPr]

    tc_lst: list[CT_Tc]
    # -- custom inserter below --
    tblPrEx: CT_TblPrEx | None = ZeroOrOne("w:tblPrEx")  # pyright: ignore[reportAssignmentType]
    # -- custom inserter below --
    trPr: CT_TrPr | None = ZeroOrOne("w:trPr")  # pyright: ignore[reportAssignmentType]
    tc = ZeroOrMore("w:tc")

    @property
    def grid_after(self) -> int:
        """The number of unpopulated layout-grid cells at the end of this row."""
        trPr = self.trPr
        if trPr is None:
            return 0
        return trPr.grid_after

    @property
    def grid_before(self) -> int:
        """The number of unpopulated layout-grid cells at the start of this row."""
        trPr = self.trPr
        if trPr is None:
            return 0
        return trPr.grid_before

    def tc_at_grid_offset(self, grid_offset: int) -> CT_Tc:
        """The `tc` element in this tr at exact `grid offset`.

        Raises ValueError when this `w:tr` contains no `w:tc` with exact starting `grid_offset`.
        """
        # -- account for omitted cells at the start of the row --
        remaining_offset = grid_offset - self.grid_before

        for tc in self.tc_lst:
            # -- We've gone past grid_offset without finding a tc, no sense searching further. --
            if remaining_offset < 0:
                break
            # -- We've arrived at grid_offset, this is the `w:tc` we're looking for. --
            if remaining_offset == 0:
                return tc
            # -- We're not there yet, skip forward the number of layout-grid cells this cell
            # -- occupies.
            remaining_offset -= tc.grid_span

        raise ValueError(f"no `tc` element at grid_offset={grid_offset}")

    @property
    def tr_idx(self) -> int:
        """Index of this `w:tr` element within its parent `w:tbl` element."""
        tbl = cast(CT_Tbl, self.getparent())
        return tbl.tr_lst.index(self)

    @property
    def trHeight_hRule(self) -> WD_ROW_HEIGHT_RULE | None:
        """The value of `./w:trPr/w:trHeight/@w:hRule`, or |None| if not present."""
        trPr = self.trPr
        if trPr is None:
            return None
        return trPr.trHeight_hRule

    @trHeight_hRule.setter
    def trHeight_hRule(self, value: WD_ROW_HEIGHT_RULE | None):
        trPr = self.get_or_add_trPr()
        trPr.trHeight_hRule = value

    @property
    def trHeight_val(self):
        """Return the value of `w:trPr/w:trHeight@w:val`, or |None| if not present."""
        trPr = self.trPr
        if trPr is None:
            return None
        return trPr.trHeight_val

    @trHeight_val.setter
    def trHeight_val(self, value: Length | None):
        trPr = self.get_or_add_trPr()
        trPr.trHeight_val = value

    def _insert_tblPrEx(self, tblPrEx: CT_TblPrEx):
        self.insert(0, tblPrEx)

    def _insert_trPr(self, trPr: CT_TrPr):
        tblPrEx = self.tblPrEx
        if tblPrEx is not None:
            tblPrEx.addnext(trPr)
        else:
            self.insert(0, trPr)

    def _new_tc(self):
        return CT_Tc.new()


class CT_Tbl(BaseOxmlElement):
    """``<w:tbl>`` element."""

    add_tr: Callable[[], CT_Row]
    tr_lst: list[CT_Row]

    tblPr: CT_TblPr = OneAndOnlyOne("w:tblPr")  # pyright: ignore[reportAssignmentType]
    tblGrid: CT_TblGrid = OneAndOnlyOne("w:tblGrid")  # pyright: ignore[reportAssignmentType]
    tr = ZeroOrMore("w:tr")

    @property
    def bidiVisual_val(self) -> bool | None:
        """Value of `./w:tblPr/w:bidiVisual/@w:val` or |None| if not present.

        Controls whether table cells are displayed right-to-left or left-to-right.
        """
        bidiVisual = self.tblPr.bidiVisual
        if bidiVisual is None:
            return None
        return bidiVisual.val

    @bidiVisual_val.setter
    def bidiVisual_val(self, value: WD_TABLE_DIRECTION | None):
        tblPr = self.tblPr
        if value is None:
            tblPr._remove_bidiVisual()  # pyright: ignore[reportPrivateUsage]
        else:
            tblPr.get_or_add_bidiVisual().val = bool(value)

    @property
    def col_count(self):
        """The number of grid columns in this table."""
        return len(self.tblGrid.gridCol_lst)

    def iter_tcs(self):
        """Generate each of the `w:tc` elements in this table, left to right and top to
        bottom.

        Each cell in the first row is generated, followed by each cell in the second
        row, etc.
        """
        for tr in self.tr_lst:
            for tc in tr.tc_lst:
                yield tc

    @classmethod
    def new_tbl(cls, rows: int, cols: int, width: Length) -> CT_Tbl:
        """Return a new `w:tbl` element having `rows` rows and `cols` columns.

        `width` is distributed evenly between the columns.
        """
        return cast(CT_Tbl, parse_xml(cls._tbl_xml(rows, cols, width)))

    @property
    def tblStyle_val(self) -> str | None:
        """`w:tblPr/w:tblStyle/@w:val` (a table style id) or |None| if not present."""
        tblStyle = self.tblPr.tblStyle
        if tblStyle is None:
            return None
        return tblStyle.val

    @tblStyle_val.setter
    def tblStyle_val(self, styleId: str | None) -> None:
        """Set the value of `w:tblPr/w:tblStyle/@w:val` (a table style id) to `styleId`.

        If `styleId` is None, remove the `w:tblStyle` element.
        """
        tblPr = self.tblPr
        tblPr._remove_tblStyle()  # pyright: ignore[reportPrivateUsage]
        if styleId is None:
            return
        tblPr._add_tblStyle().val = styleId  # pyright: ignore[reportPrivateUsage]

    @classmethod
    def _tbl_xml(cls, rows: int, cols: int, width: Length) -> str:
        col_width = Emu(width // cols) if cols > 0 else Emu(0)
        return (
            f"<w:tbl {nsdecls('w')}>\n"
            f"  <w:tblPr>\n"
            f'    <w:tblW w:type="auto" w:w="0"/>\n'
            f'    <w:tblLook w:firstColumn="1" w:firstRow="1"\n'
            f'               w:lastColumn="0" w:lastRow="0" w:noHBand="0"\n'
            f'               w:noVBand="1" w:val="04A0"/>\n'
            f"  </w:tblPr>\n"
            f"{cls._tblGrid_xml(cols, col_width)}"
            f"{cls._trs_xml(rows, cols, col_width)}"
            f"</w:tbl>\n"
        )

    @classmethod
    def _tblGrid_xml(cls, col_count: int, col_width: Length) -> str:
        xml = "  <w:tblGrid>\n"
        for _ in range(col_count):
            xml += '    <w:gridCol w:w="%d"/>\n' % col_width.twips
        xml += "  </w:tblGrid>\n"
        return xml

    @classmethod
    def _trs_xml(cls, row_count: int, col_count: int, col_width: Length) -> str:
        return f"  <w:tr>\n{cls._tcs_xml(col_count, col_width)}  </w:tr>\n" * row_count

    @classmethod
    def _tcs_xml(cls, col_count: int, col_width: Length) -> str:
        return (
            f"    <w:tc>\n"
            f"      <w:tcPr>\n"
            f'        <w:tcW w:type="dxa" w:w="{col_width.twips}"/>\n'
            f"      </w:tcPr>\n"
            f"      <w:p/>\n"
            f"    </w:tc>\n"
        ) * col_count


class CT_TblGrid(BaseOxmlElement):
    """`w:tblGrid` element.

    Child of `w:tbl`, holds `w:gridCol> elements that define column count, width, etc.
    """

    add_gridCol: Callable[[], CT_TblGridCol]
    gridCol_lst: list[CT_TblGridCol]

    gridCol = ZeroOrMore("w:gridCol", successors=("w:tblGridChange",))


class CT_TblGridCol(BaseOxmlElement):
    """`w:gridCol` element, child of `w:tblGrid`, defines a table column."""

    w: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:w", ST_TwipsMeasure
    )

    @property
    def gridCol_idx(self) -> int:
        """Index of this `w:gridCol` element within its parent `w:tblGrid` element."""
        tblGrid = cast(CT_TblGrid, self.getparent())
        return tblGrid.gridCol_lst.index(self)


class CT_TblLayoutType(BaseOxmlElement):
    """`w:tblLayout` element.

    Specifies whether column widths are fixed or can be automatically adjusted based on
    content.
    """

    type: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:type", ST_TblLayoutType
    )


class CT_TblPr(BaseOxmlElement):
    """``<w:tblPr>`` element, child of ``<w:tbl>``, holds child elements that define
    table properties such as style and borders."""

    get_or_add_bidiVisual: Callable[[], CT_OnOff]
    get_or_add_jc: Callable[[], CT_Jc]
    get_or_add_tblLayout: Callable[[], CT_TblLayoutType]
    _add_tblStyle: Callable[[], CT_String]
    _remove_bidiVisual: Callable[[], None]
    _remove_jc: Callable[[], None]
    _remove_tblStyle: Callable[[], None]

    _tag_seq = (
        "w:tblStyle",
        "w:tblpPr",
        "w:tblOverlap",
        "w:bidiVisual",
        "w:tblStyleRowBandSize",
        "w:tblStyleColBandSize",
        "w:tblW",
        "w:jc",
        "w:tblCellSpacing",
        "w:tblInd",
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    tblStyle: CT_String | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:tblStyle", successors=_tag_seq[1:]
    )
    bidiVisual: CT_OnOff | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:bidiVisual", successors=_tag_seq[4:]
    )
    jc: CT_Jc | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:jc", successors=_tag_seq[8:]
    )
    tblLayout: CT_TblLayoutType | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:tblLayout", successors=_tag_seq[13:]
    )
    del _tag_seq

    @property
    def alignment(self) -> WD_TABLE_ALIGNMENT | None:
        """Horizontal alignment of table, |None| if `./w:jc` is not present."""
        jc = self.jc
        if jc is None:
            return None
        return cast("WD_TABLE_ALIGNMENT | None", jc.val)

    @alignment.setter
    def alignment(self, value: WD_TABLE_ALIGNMENT | None):
        self._remove_jc()
        if value is None:
            return
        jc = self.get_or_add_jc()
        jc.val = cast("WD_ALIGN_PARAGRAPH", value)

    @property
    def autofit(self) -> bool:
        """|False| when there is a `w:tblLayout` child with `@w:type="fixed"`.

        Otherwise |True|.
        """
        tblLayout = self.tblLayout
        return True if tblLayout is None else tblLayout.type != "fixed"

    @autofit.setter
    def autofit(self, value: bool):
        tblLayout = self.get_or_add_tblLayout()
        tblLayout.type = "autofit" if value else "fixed"

    @property
    def style(self):
        """Return the value of the ``val`` attribute of the ``<w:tblStyle>`` child or
        |None| if not present."""
        tblStyle = self.tblStyle
        if tblStyle is None:
            return None
        return tblStyle.val

    @style.setter
    def style(self, value: str | None):
        self._remove_tblStyle()
        if value is None:
            return
        self._add_tblStyle().val = value


class CT_TblPrEx(BaseOxmlElement):
    """`w:tblPrEx` element, exceptions to table-properties.

    Applied at a lower level, like a `w:tr` to modify the appearance. Possibly used when
    two tables are merged. For more see:
    http://officeopenxml.com/WPtablePropertyExceptions.php
    """


class CT_TblWidth(BaseOxmlElement):
    """Used for `w:tblW` and `w:tcW` and others, specifies a table-related width."""

    # the type for `w` attr is actually ST_MeasurementOrPercent, but using
    # XsdInt for now because only dxa (twips) values are being used. It's not
    # entirely clear what the semantics are for other values like -01.4mm
    w: int = RequiredAttribute("w:w", XsdInt)  # pyright: ignore[reportAssignmentType]
    type = RequiredAttribute("w:type", ST_TblWidth)

    @property
    def width(self) -> Length | None:
        """EMU length indicated by the combined `w:w` and `w:type` attrs."""
        if self.type != "dxa":
            return None
        return Twips(self.w)

    @width.setter
    def width(self, value: Length):
        self.type = "dxa"
        self.w = Emu(value).twips


class CT_Tc(BaseOxmlElement):
    """`w:tc` table cell element."""

    add_p: Callable[[], CT_P]
    get_or_add_tcPr: Callable[[], CT_TcPr]
    p_lst: list[CT_P]
    tbl_lst: list[CT_Tbl]
    _insert_tbl: Callable[[CT_Tbl], CT_Tbl]
    _new_p: Callable[[], CT_P]

    # -- tcPr has many successors, `._insert_tcPr()` is overridden below --
    tcPr: CT_TcPr | None = ZeroOrOne("w:tcPr")  # pyright: ignore[reportAssignmentType]
    p = OneOrMore("w:p")
    tbl = OneOrMore("w:tbl")

    @property
    def bottom(self) -> int:
        """The row index that marks the bottom extent of the vertical span of this cell.

        This is one greater than the index of the bottom-most row of the span, similar
        to how a slice of the cell's rows would be specified.
        """
        if self.vMerge is not None:
            tc_below = self._tc_below
            if tc_below is not None and tc_below.vMerge == ST_Merge.CONTINUE:
                return tc_below.bottom
        return self._tr_idx + 1

    def clear_content(self):
        """Remove all content elements, preserving `w:tcPr` element if present.

        Note that this leaves the `w:tc` element in an invalid state because it doesn't
        contain at least one block-level element. It's up to the caller to add a
        `w:p`child element as the last content element.
        """
        # -- remove all cell inner-content except a `w:tcPr` when present. --
        for e in self.xpath("./*[not(self::w:tcPr)]"):
            self.remove(e)

    @property
    def grid_offset(self) -> int:
        """Starting offset of `tc` in the layout-grid columns of its table.

        A cell in the leftmost grid-column has offset 0.
        """
        grid_before = self._tr.grid_before
        preceding_tc_grid_spans = sum(
            tc.grid_span for tc in self.xpath("./preceding-sibling::w:tc")
        )
        return grid_before + preceding_tc_grid_spans

    @property
    def grid_span(self) -> int:
        """The integer number of columns this cell spans.

        Determined by ./w:tcPr/w:gridSpan/@val, it defaults to 1.
        """
        tcPr = self.tcPr
        return 1 if tcPr is None else tcPr.grid_span

    @grid_span.setter
    def grid_span(self, value: int):
        tcPr = self.get_or_add_tcPr()
        tcPr.grid_span = value

    @property
    def inner_content_elements(self) -> list[CT_P | CT_Tbl]:
        """Generate all `w:p` and `w:tbl` elements in this document-body.

        Elements appear in document order. Elements shaded by nesting in a `w:ins` or
        other "wrapper" element will not be included.
        """
        return self.xpath("./w:p | ./w:tbl")

    def iter_block_items(self):
        """Generate a reference to each of the block-level content elements in this
        cell, in the order they appear."""
        block_item_tags = (qn("w:p"), qn("w:tbl"), qn("w:sdt"))
        for child in self:
            if child.tag in block_item_tags:
                yield child

    @property
    def left(self) -> int:
        """The grid column index at which this ``<w:tc>`` element appears."""
        return self.grid_offset

    def merge(self, other_tc: CT_Tc) -> CT_Tc:
        """Return top-left `w:tc` element of a new span.

        Span is formed by merging the rectangular region defined by using this tc
        element and `other_tc` as diagonal corners.
        """
        top, left, height, width = self._span_dimensions(other_tc)
        top_tc = self._tbl.tr_lst[top].tc_at_grid_offset(left)
        top_tc._grow_to(width, height)
        return top_tc

    @classmethod
    def new(cls) -> CT_Tc:
        """A new `w:tc` element, containing an empty paragraph as the required EG_BlockLevelElt."""
        return cast(CT_Tc, parse_xml("<w:tc %s>\n" "  <w:p/>\n" "</w:tc>" % nsdecls("w")))

    @property
    def right(self) -> int:
        """The grid column index that marks the right-side extent of the horizontal span
        of this cell.

        This is one greater than the index of the right-most column of the span, similar
        to how a slice of the cell's columns would be specified.
        """
        return self.grid_offset + self.grid_span

    @property
    def top(self) -> int:
        """The top-most row index in the vertical span of this cell."""
        if self.vMerge is None or self.vMerge == ST_Merge.RESTART:
            return self._tr_idx
        return self._tc_above.top

    @property
    def vMerge(self) -> str | None:
        """Value of ./w:tcPr/w:vMerge/@val, |None| if w:vMerge is not present."""
        tcPr = self.tcPr
        if tcPr is None:
            return None
        return tcPr.vMerge_val

    @vMerge.setter
    def vMerge(self, value: str | None):
        tcPr = self.get_or_add_tcPr()
        tcPr.vMerge_val = value

    @property
    def width(self) -> Length | None:
        """EMU length represented in `./w:tcPr/w:tcW` or |None| if not present."""
        tcPr = self.tcPr
        if tcPr is None:
            return None
        return tcPr.width

    @width.setter
    def width(self, value: Length):
        tcPr = self.get_or_add_tcPr()
        tcPr.width = value

    def _add_width_of(self, other_tc: CT_Tc):
        """Add the width of `other_tc` to this cell.

        Does nothing if either this tc or `other_tc` does not have a specified width.
        """
        if self.width and other_tc.width:
            self.width = Length(self.width + other_tc.width)

    def _grow_to(self, width: int, height: int, top_tc: CT_Tc | None = None):
        """Grow this cell to `width` grid columns and `height` rows.

        This is accomplished by expanding horizontal spans and creating continuation
        cells to form vertical spans.
        """

        def vMerge_val(top_tc: CT_Tc):
            return (
                ST_Merge.CONTINUE
                if top_tc is not self
                else None if height == 1 else ST_Merge.RESTART
            )

        top_tc = self if top_tc is None else top_tc
        self._span_to_width(width, top_tc, vMerge_val(top_tc))
        if height > 1:
            tc_below = self._tc_below
            assert tc_below is not None
            tc_below._grow_to(width, height - 1, top_tc)

    def _insert_tcPr(self, tcPr: CT_TcPr) -> CT_TcPr:
        """Override default `._insert_tcPr()`."""
        # -- `tcPr`` has a large number of successors, but always comes first if it appears,
        # -- so just using insert(0, ...) rather than spelling out successors.
        self.insert(0, tcPr)
        return tcPr

    @property
    def _is_empty(self) -> bool:
        """True if this cell contains only a single empty `w:p` element."""
        block_items = list(self.iter_block_items())
        if len(block_items) > 1:
            return False
        # -- cell must include at least one block item but can be a `w:tbl`, `w:sdt`,
        # -- `w:customXml` or a `w:p`
        only_item = block_items[0]
        if isinstance(only_item, CT_P) and len(only_item.r_lst) == 0:
            return True
        return False

    def _move_content_to(self, other_tc: CT_Tc):
        """Append the content of this cell to `other_tc`.

        Leaves this cell with a single empty ``<w:p>`` element.
        """
        if other_tc is self:
            return
        if self._is_empty:
            return
        other_tc._remove_trailing_empty_p()
        # -- appending moves each element from self to other_tc --
        for block_element in self.iter_block_items():
            other_tc.append(block_element)
        # -- add back the required minimum single empty <w:p> element --
        self.append(self._new_p())

    def _new_tbl(self) -> None:
        raise NotImplementedError(
            "use CT_Tbl.new_tbl() to add a new table, specifying rows and columns"
        )

    @property
    def _next_tc(self) -> CT_Tc | None:
        """The `w:tc` element immediately following this one in this row, or |None| if
        this is the last `w:tc` element in the row."""
        following_tcs = self.xpath("./following-sibling::w:tc")
        return following_tcs[0] if following_tcs else None

    def _remove(self):
        """Remove this `w:tc` element from the XML tree."""
        parent_element = self.getparent()
        assert parent_element is not None
        parent_element.remove(self)

    def _remove_trailing_empty_p(self):
        """Remove last content element from this cell if it's an empty `w:p` element."""
        block_items = list(self.iter_block_items())
        last_content_elm = block_items[-1]
        if not isinstance(last_content_elm, CT_P):
            return
        p = last_content_elm
        if len(p.r_lst) > 0:
            return
        self.remove(p)

    def _span_dimensions(self, other_tc: CT_Tc) -> tuple[int, int, int, int]:
        """Return a (top, left, height, width) 4-tuple specifying the extents of the
        merged cell formed by using this tc and `other_tc` as opposite corner
        extents."""

        def raise_on_inverted_L(a: CT_Tc, b: CT_Tc):
            if a.top == b.top and a.bottom != b.bottom:
                raise InvalidSpanError("requested span not rectangular")
            if a.left == b.left and a.right != b.right:
                raise InvalidSpanError("requested span not rectangular")

        def raise_on_tee_shaped(a: CT_Tc, b: CT_Tc):
            top_most, other = (a, b) if a.top < b.top else (b, a)
            if top_most.top < other.top and top_most.bottom > other.bottom:
                raise InvalidSpanError("requested span not rectangular")

            left_most, other = (a, b) if a.left < b.left else (b, a)
            if left_most.left < other.left and left_most.right > other.right:
                raise InvalidSpanError("requested span not rectangular")

        raise_on_inverted_L(self, other_tc)
        raise_on_tee_shaped(self, other_tc)

        top = min(self.top, other_tc.top)
        left = min(self.left, other_tc.left)
        bottom = max(self.bottom, other_tc.bottom)
        right = max(self.right, other_tc.right)

        return top, left, bottom - top, right - left

    def _span_to_width(self, grid_width: int, top_tc: CT_Tc, vMerge: str | None):
        """Incorporate `w:tc` elements to the right until this cell spans `grid_width`.

        Incorporated `w:tc` elements are removed (replaced by gridSpan value).

        Raises |ValueError| if `grid_width` cannot be exactly achieved, such as when a
        merged cell would drive the span width greater than `grid_width` or if not
        enough grid columns are available to make this cell that wide. All content from
        incorporated cells is appended to `top_tc`. The val attribute of the vMerge
        element on the single remaining cell is set to `vMerge`. If `vMerge` is |None|,
        the vMerge element is removed if present.
        """
        self._move_content_to(top_tc)
        while self.grid_span < grid_width:
            self._swallow_next_tc(grid_width, top_tc)
        self.vMerge = vMerge

    def _swallow_next_tc(self, grid_width: int, top_tc: CT_Tc):
        """Extend the horizontal span of this `w:tc` element to incorporate the
        following `w:tc` element in the row and then delete that following `w:tc`
        element.

        Any content in the following `w:tc` element is appended to the content of
        `top_tc`. The width of the following `w:tc` element is added to this one, if
        present. Raises |InvalidSpanError| if the width of the resulting cell is greater
        than `grid_width` or if there is no next `<w:tc>` element in the row.
        """

        def raise_on_invalid_swallow(next_tc: CT_Tc | None):
            if next_tc is None:
                raise InvalidSpanError("not enough grid columns")
            if self.grid_span + next_tc.grid_span > grid_width:
                raise InvalidSpanError("span is not rectangular")

        next_tc = self._next_tc
        raise_on_invalid_swallow(next_tc)
        assert next_tc is not None
        next_tc._move_content_to(top_tc)
        self._add_width_of(next_tc)
        self.grid_span += next_tc.grid_span
        next_tc._remove()

    @property
    def _tbl(self) -> CT_Tbl:
        """The tbl element this tc element appears in."""
        return cast(CT_Tbl, self.xpath("./ancestor::w:tbl[position()=1]")[0])

    @property
    def _tc_above(self) -> CT_Tc:
        """The `w:tc` element immediately above this one in its grid column."""
        return self._tr_above.tc_at_grid_offset(self.grid_offset)

    @property
    def _tc_below(self) -> CT_Tc | None:
        """The tc element immediately below this one in its grid column."""
        tr_below = self._tr_below
        if tr_below is None:
            return None
        return tr_below.tc_at_grid_offset(self.grid_offset)

    @property
    def _tr(self) -> CT_Row:
        """The tr element this tc element appears in."""
        return cast(CT_Row, self.xpath("./ancestor::w:tr[position()=1]")[0])

    @property
    def _tr_above(self) -> CT_Row:
        """The tr element prior in sequence to the tr this cell appears in.

        Raises |ValueError| if called on a cell in the top-most row.
        """
        tr_aboves = self.xpath("./ancestor::w:tr[position()=1]/preceding-sibling::w:tr[1]")
        if not tr_aboves:
            raise ValueError("no tr above topmost tr in w:tbl")
        return tr_aboves[0]

    @property
    def _tr_below(self) -> CT_Row | None:
        """The tr element next in sequence after the tr this cell appears in, or |None|
        if this cell appears in the last row."""
        tr_lst = self._tbl.tr_lst
        tr_idx = tr_lst.index(self._tr)
        try:
            return tr_lst[tr_idx + 1]
        except IndexError:
            return None

    @property
    def _tr_idx(self) -> int:
        """The row index of the tr element this tc element appears in."""
        return self._tbl.tr_lst.index(self._tr)


class CT_TcPr(BaseOxmlElement):
    """``<w:tcPr>`` element, defining table cell properties."""

    get_or_add_gridSpan: Callable[[], CT_DecimalNumber]
    get_or_add_tcW: Callable[[], CT_TblWidth]
    get_or_add_vAlign: Callable[[], CT_VerticalJc]
    _add_vMerge: Callable[[], CT_VMerge]
    _remove_gridSpan: Callable[[], None]
    _remove_vAlign: Callable[[], None]
    _remove_vMerge: Callable[[], None]

    _tag_seq = (
        "w:cnfStyle",
        "w:tcW",
        "w:gridSpan",
        "w:hMerge",
        "w:vMerge",
        "w:tcBorders",
        "w:shd",
        "w:noWrap",
        "w:tcMar",
        "w:textDirection",
        "w:tcFitText",
        "w:vAlign",
        "w:hideMark",
        "w:headers",
        "w:cellIns",
        "w:cellDel",
        "w:cellMerge",
        "w:tcPrChange",
    )
    tcW: CT_TblWidth | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:tcW", successors=_tag_seq[2:]
    )
    gridSpan: CT_DecimalNumber | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:gridSpan", successors=_tag_seq[3:]
    )
    vMerge: CT_VMerge | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:vMerge", successors=_tag_seq[5:]
    )
    vAlign: CT_VerticalJc | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:vAlign", successors=_tag_seq[12:]
    )
    del _tag_seq

    @property
    def grid_span(self) -> int:
        """The integer number of columns this cell spans.

        Determined by ./w:gridSpan/@val, it defaults to 1.
        """
        gridSpan = self.gridSpan
        return 1 if gridSpan is None else gridSpan.val

    @grid_span.setter
    def grid_span(self, value: int):
        self._remove_gridSpan()
        if value > 1:
            self.get_or_add_gridSpan().val = value

    @property
    def vAlign_val(self):
        """Value of `w:val` attribute on  `w:vAlign` child.

        Value is |None| if `w:vAlign` child is not present. The `w:val` attribute on
        `w:vAlign` is required.
        """
        vAlign = self.vAlign
        if vAlign is None:
            return None
        return vAlign.val

    @vAlign_val.setter
    def vAlign_val(self, value: WD_CELL_VERTICAL_ALIGNMENT | None):
        if value is None:
            self._remove_vAlign()
            return
        self.get_or_add_vAlign().val = value

    @property
    def vMerge_val(self):
        """The value of the ./w:vMerge/@val attribute, or |None| if the w:vMerge element
        is not present."""
        vMerge = self.vMerge
        if vMerge is None:
            return None
        return vMerge.val

    @vMerge_val.setter
    def vMerge_val(self, value: str | None):
        self._remove_vMerge()
        if value is not None:
            self._add_vMerge().val = value

    @property
    def width(self) -> Length | None:
        """EMU length in `./w:tcW` or |None| if not present or its type is not 'dxa'."""
        tcW = self.tcW
        if tcW is None:
            return None
        return tcW.width

    @width.setter
    def width(self, value: Length):
        tcW = self.get_or_add_tcW()
        tcW.width = value


class CT_TrPr(BaseOxmlElement):
    """``<w:trPr>`` element, defining table row properties."""

    get_or_add_trHeight: Callable[[], CT_Height]

    _tag_seq = (
        "w:cnfStyle",
        "w:divId",
        "w:gridBefore",
        "w:gridAfter",
        "w:wBefore",
        "w:wAfter",
        "w:cantSplit",
        "w:trHeight",
        "w:tblHeader",
        "w:tblCellSpacing",
        "w:jc",
        "w:hidden",
        "w:ins",
        "w:del",
        "w:trPrChange",
    )
    gridAfter: CT_DecimalNumber | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:gridAfter", successors=_tag_seq[4:]
    )
    gridBefore: CT_DecimalNumber | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:gridBefore", successors=_tag_seq[3:]
    )
    trHeight: CT_Height | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:trHeight", successors=_tag_seq[8:]
    )
    del _tag_seq

    @property
    def grid_after(self) -> int:
        """The number of unpopulated layout-grid cells at the end of this row."""
        gridAfter = self.gridAfter
        return 0 if gridAfter is None else gridAfter.val

    @property
    def grid_before(self) -> int:
        """The number of unpopulated layout-grid cells at the start of this row."""
        gridBefore = self.gridBefore
        return 0 if gridBefore is None else gridBefore.val

    @property
    def trHeight_hRule(self) -> WD_ROW_HEIGHT_RULE | None:
        """Return the value of `w:trHeight@w:hRule`, or |None| if not present."""
        trHeight = self.trHeight
        return None if trHeight is None else trHeight.hRule

    @trHeight_hRule.setter
    def trHeight_hRule(self, value: WD_ROW_HEIGHT_RULE | None):
        if value is None and self.trHeight is None:
            return
        trHeight = self.get_or_add_trHeight()
        trHeight.hRule = value

    @property
    def trHeight_val(self):
        """Return the value of `w:trHeight@w:val`, or |None| if not present."""
        trHeight = self.trHeight
        return None if trHeight is None else trHeight.val

    @trHeight_val.setter
    def trHeight_val(self, value: Length | None):
        if value is None and self.trHeight is None:
            return
        trHeight = self.get_or_add_trHeight()
        trHeight.val = value


class CT_VerticalJc(BaseOxmlElement):
    """`w:vAlign` element, specifying vertical alignment of cell."""

    val: WD_CELL_VERTICAL_ALIGNMENT = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", WD_CELL_VERTICAL_ALIGNMENT
    )


class CT_VMerge(BaseOxmlElement):
    """``<w:vMerge>`` element, specifying vertical merging behavior of a cell."""

    val: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", ST_Merge, default=ST_Merge.CONTINUE
    )


# font.py
"""Custom element classes related to run properties (font)."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable

from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from docx.oxml.simpletypes import (
    ST_HexColor,
    ST_HpsMeasure,
    ST_String,
    ST_VerticalAlignRun,
)
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrOne,
)

if TYPE_CHECKING:
    from docx.oxml.shared import CT_OnOff, CT_String
    from docx.shared import Length


class CT_Color(BaseOxmlElement):
    """`w:color` element, specifying the color of a font and perhaps other objects."""

    val = RequiredAttribute("w:val", ST_HexColor)
    themeColor = OptionalAttribute("w:themeColor", MSO_THEME_COLOR)


class CT_Fonts(BaseOxmlElement):
    """`<w:rFonts>` element.

    Specifies typeface name for the various language types.
    """

    ascii: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:ascii", ST_String
    )
    hAnsi: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:hAnsi", ST_String
    )


class CT_Highlight(BaseOxmlElement):
    """`w:highlight` element, specifying font highlighting/background color."""

    val: WD_COLOR_INDEX = RequiredAttribute(  # pyright: ignore[reportGeneralTypeIssues]
        "w:val", WD_COLOR_INDEX
    )


class CT_HpsMeasure(BaseOxmlElement):
    """Used for `<w:sz>` element and others, specifying font size in half-points."""

    val: Length = RequiredAttribute(  # pyright: ignore[reportGeneralTypeIssues]
        "w:val", ST_HpsMeasure
    )


class CT_RPr(BaseOxmlElement):
    """`<w:rPr>` element, containing the properties for a run."""

    get_or_add_highlight: Callable[[], CT_Highlight]
    get_or_add_rFonts: Callable[[], CT_Fonts]
    get_or_add_sz: Callable[[], CT_HpsMeasure]
    get_or_add_vertAlign: Callable[[], CT_VerticalAlignRun]
    _add_rStyle: Callable[..., CT_String]
    _add_u: Callable[[], CT_Underline]
    _remove_highlight: Callable[[], None]
    _remove_rFonts: Callable[[], None]
    _remove_rStyle: Callable[[], None]
    _remove_sz: Callable[[], None]
    _remove_u: Callable[[], None]
    _remove_vertAlign: Callable[[], None]

    _tag_seq = (
        "w:rStyle",
        "w:rFonts",
        "w:b",
        "w:bCs",
        "w:i",
        "w:iCs",
        "w:caps",
        "w:smallCaps",
        "w:strike",
        "w:dstrike",
        "w:outline",
        "w:shadow",
        "w:emboss",
        "w:imprint",
        "w:noProof",
        "w:snapToGrid",
        "w:vanish",
        "w:webHidden",
        "w:color",
        "w:spacing",
        "w:w",
        "w:kern",
        "w:position",
        "w:sz",
        "w:szCs",
        "w:highlight",
        "w:u",
        "w:effect",
        "w:bdr",
        "w:shd",
        "w:fitText",
        "w:vertAlign",
        "w:rtl",
        "w:cs",
        "w:em",
        "w:lang",
        "w:eastAsianLayout",
        "w:specVanish",
        "w:oMath",
    )
    rStyle: CT_String | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:rStyle", successors=_tag_seq[1:]
    )
    rFonts: CT_Fonts | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:rFonts", successors=_tag_seq[2:]
    )
    b: CT_OnOff | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:b", successors=_tag_seq[3:]
    )
    bCs = ZeroOrOne("w:bCs", successors=_tag_seq[4:])
    i = ZeroOrOne("w:i", successors=_tag_seq[5:])
    iCs = ZeroOrOne("w:iCs", successors=_tag_seq[6:])
    caps = ZeroOrOne("w:caps", successors=_tag_seq[7:])
    smallCaps = ZeroOrOne("w:smallCaps", successors=_tag_seq[8:])
    strike = ZeroOrOne("w:strike", successors=_tag_seq[9:])
    dstrike = ZeroOrOne("w:dstrike", successors=_tag_seq[10:])
    outline = ZeroOrOne("w:outline", successors=_tag_seq[11:])
    shadow = ZeroOrOne("w:shadow", successors=_tag_seq[12:])
    emboss = ZeroOrOne("w:emboss", successors=_tag_seq[13:])
    imprint = ZeroOrOne("w:imprint", successors=_tag_seq[14:])
    noProof = ZeroOrOne("w:noProof", successors=_tag_seq[15:])
    snapToGrid = ZeroOrOne("w:snapToGrid", successors=_tag_seq[16:])
    vanish = ZeroOrOne("w:vanish", successors=_tag_seq[17:])
    webHidden = ZeroOrOne("w:webHidden", successors=_tag_seq[18:])
    color = ZeroOrOne("w:color", successors=_tag_seq[19:])
    sz: CT_HpsMeasure | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:sz", successors=_tag_seq[24:]
    )
    highlight: CT_Highlight | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:highlight", successors=_tag_seq[26:]
    )
    u: CT_Underline | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:u", successors=_tag_seq[27:]
    )
    vertAlign: CT_VerticalAlignRun | None = ZeroOrOne(  # pyright: ignore[reportGeneralTypeIssues]
        "w:vertAlign", successors=_tag_seq[32:]
    )
    rtl = ZeroOrOne("w:rtl", successors=_tag_seq[33:])
    cs = ZeroOrOne("w:cs", successors=_tag_seq[34:])
    specVanish = ZeroOrOne("w:specVanish", successors=_tag_seq[38:])
    oMath = ZeroOrOne("w:oMath", successors=_tag_seq[39:])
    del _tag_seq

    def _new_color(self):
        """Override metaclass method to set `w:color/@val` to RGB black on create."""
        return parse_xml('<w:color %s w:val="000000"/>' % nsdecls("w"))

    @property
    def highlight_val(self) -> WD_COLOR_INDEX | None:
        """Value of `./w:highlight/@val`.

        Specifies font's highlight color, or `None` if the text is not highlighted.
        """
        highlight = self.highlight
        if highlight is None:
            return None
        return highlight.val

    @highlight_val.setter
    def highlight_val(self, value: WD_COLOR_INDEX | None) -> None:
        if value is None:
            self._remove_highlight()
            return
        highlight = self.get_or_add_highlight()
        highlight.val = value

    @property
    def rFonts_ascii(self) -> str | None:
        """The value of `w:rFonts/@w:ascii` or |None| if not present.

        Represents the assigned typeface name. The rFonts element also specifies other
        special-case typeface names; this method handles the case where just the common
        name is required.
        """
        rFonts = self.rFonts
        if rFonts is None:
            return None
        return rFonts.ascii

    @rFonts_ascii.setter
    def rFonts_ascii(self, value: str | None) -> None:
        if value is None:
            self._remove_rFonts()
            return
        rFonts = self.get_or_add_rFonts()
        rFonts.ascii = value

    @property
    def rFonts_hAnsi(self) -> str | None:
        """The value of `w:rFonts/@w:hAnsi` or |None| if not present."""
        rFonts = self.rFonts
        if rFonts is None:
            return None
        return rFonts.hAnsi

    @rFonts_hAnsi.setter
    def rFonts_hAnsi(self, value: str | None):
        if value is None and self.rFonts is None:
            return
        rFonts = self.get_or_add_rFonts()
        rFonts.hAnsi = value

    @property
    def style(self) -> str | None:
        """String in `./w:rStyle/@val`, or None if `w:rStyle` is not present."""
        rStyle = self.rStyle
        if rStyle is None:
            return None
        return rStyle.val

    @style.setter
    def style(self, style: str | None) -> None:
        """Set `./w:rStyle/@val` to `style`, adding the `w:rStyle` element if necessary.

        If `style` is |None|, remove `w:rStyle` element if present.
        """
        if style is None:
            self._remove_rStyle()
        elif self.rStyle is None:
            self._add_rStyle(val=style)
        else:
            self.rStyle.val = style

    @property
    def subscript(self) -> bool | None:
        """|True| if `./w:vertAlign/@w:val` is "subscript".

        |False| if `w:vertAlign/@w:val` contains any other value. |None| if
        `w:vertAlign` is not present.
        """
        vertAlign = self.vertAlign
        if vertAlign is None:
            return None
        if vertAlign.val == ST_VerticalAlignRun.SUBSCRIPT:
            return True
        return False

    @subscript.setter
    def subscript(self, value: bool | None) -> None:
        if value is None:
            self._remove_vertAlign()
        elif bool(value) is True:
            self.get_or_add_vertAlign().val = ST_VerticalAlignRun.SUBSCRIPT
        # -- assert bool(value) is False --
        elif self.vertAlign is not None and self.vertAlign.val == ST_VerticalAlignRun.SUBSCRIPT:
            self._remove_vertAlign()

    @property
    def superscript(self) -> bool | None:
        """|True| if `w:vertAlign/@w:val` is 'superscript'.

        |False| if `w:vertAlign/@w:val` contains any other value. |None| if
        `w:vertAlign` is not present.
        """
        vertAlign = self.vertAlign
        if vertAlign is None:
            return None
        if vertAlign.val == ST_VerticalAlignRun.SUPERSCRIPT:
            return True
        return False

    @superscript.setter
    def superscript(self, value: bool | None):
        if value is None:
            self._remove_vertAlign()
        elif bool(value) is True:
            self.get_or_add_vertAlign().val = ST_VerticalAlignRun.SUPERSCRIPT
        # -- assert bool(value) is False --
        elif self.vertAlign is not None and self.vertAlign.val == ST_VerticalAlignRun.SUPERSCRIPT:
            self._remove_vertAlign()

    @property
    def sz_val(self) -> Length | None:
        """The value of `w:sz/@w:val` or |None| if not present."""
        sz = self.sz
        if sz is None:
            return None
        return sz.val

    @sz_val.setter
    def sz_val(self, value: Length | None):
        if value is None:
            self._remove_sz()
            return
        sz = self.get_or_add_sz()
        sz.val = value

    @property
    def u_val(self) -> WD_UNDERLINE | None:
        """Value of `w:u/@val`, or None if not present.

        Values `WD_UNDERLINE.SINGLE` and `WD_UNDERLINE.NONE` are mapped to `True` and
        `False` respectively.
        """
        u = self.u
        if u is None:
            return None
        return u.val

    @u_val.setter
    def u_val(self, value: WD_UNDERLINE | None):
        self._remove_u()
        if value is not None:
            self._add_u().val = value

    def _get_bool_val(self, name: str) -> bool | None:
        """Value of boolean child with `name`, e.g. "w:b", "w:i", and "w:smallCaps"."""
        element = getattr(self, name)
        if element is None:
            return None
        return element.val

    def _set_bool_val(self, name: str, value: bool | None):
        if value is None:
            getattr(self, "_remove_%s" % name)()
            return
        element = getattr(self, "get_or_add_%s" % name)()
        element.val = value


class CT_Underline(BaseOxmlElement):
    """`<w:u>` element, specifying the underlining style for a run."""

    val: WD_UNDERLINE | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", WD_UNDERLINE
    )


class CT_VerticalAlignRun(BaseOxmlElement):
    """`<w:vertAlign>` element, specifying subscript or superscript."""

    val: str = RequiredAttribute(  # pyright: ignore[reportGeneralTypeIssues]
        "w:val", ST_VerticalAlignRun
    )


# hyperlink.py
"""Custom element classes related to hyperlinks (CT_Hyperlink)."""

from __future__ import annotations

from typing import TYPE_CHECKING, List

from docx.oxml.simpletypes import ST_OnOff, ST_String, XsdString
from docx.oxml.text.run import CT_R
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OptionalAttribute,
    ZeroOrMore,
)

if TYPE_CHECKING:
    from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak


class CT_Hyperlink(BaseOxmlElement):
    """`<w:hyperlink>` element, containing the text and address for a hyperlink."""

    r_lst: List[CT_R]

    rId: str | None = OptionalAttribute("r:id", XsdString)  # pyright: ignore[reportAssignmentType]
    anchor: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:anchor", ST_String
    )
    history: bool = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:history", ST_OnOff, default=True
    )

    r = ZeroOrMore("w:r")

    @property
    def lastRenderedPageBreaks(self) -> List[CT_LastRenderedPageBreak]:
        """All `w:lastRenderedPageBreak` descendants of this hyperlink."""
        return self.xpath("./w:r/w:lastRenderedPageBreak")

    @property
    def text(self) -> str:  # pyright: ignore[reportIncompatibleMethodOverride]
        """The textual content of this hyperlink.

        `CT_Hyperlink` stores the hyperlink-text as one or more `w:r` children.
        """
        return "".join(r.text for r in self.xpath("w:r"))


# pagebreak.py
"""Custom element class for rendered page-break (CT_LastRenderedPageBreak)."""

from __future__ import annotations

import copy
from typing import TYPE_CHECKING

from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import lazyproperty

if TYPE_CHECKING:
    from docx.oxml.text.hyperlink import CT_Hyperlink
    from docx.oxml.text.paragraph import CT_P


class CT_LastRenderedPageBreak(BaseOxmlElement):
    """`<w:lastRenderedPageBreak>` element, indicating page break inserted by renderer.

    A rendered page-break is one inserted by the renderer when it runs out of room on a
    page. It is an empty element (no attrs or children) and is a child of CT_R, peer to
    CT_Text.

    NOTE: this complex-type name does not exist in the schema, where
    `w:lastRenderedPageBreak` maps to `CT_Empty`. This name was added to give it
    distinguished behavior. CT_Empty is used for many elements.
    """

    @property
    def following_fragment_p(self) -> CT_P:
        """A "loose" `CT_P` containing only the paragraph content before this break.

        Raises `ValueError` if this `w:lastRenderedPageBreak` is not the first rendered
        page-break in its paragraph.

        The returned `CT_P` is a "clone" (deepcopy) of the `w:p` ancestor of this
        page-break with this `w:lastRenderedPageBreak` element and all content preceding
        it removed.

        NOTE: this `w:p` can itself contain one or more `w:renderedPageBreak` elements
        (when the paragraph contained more than one). While this is rare, the caller
        should treat this paragraph the same as other paragraphs and split it if
        necessary in a folloing step or recursion.
        """
        if not self == self._first_lrpb_in_p(self._enclosing_p):
            raise ValueError("only defined on first rendered page-break in paragraph")

        # -- splitting approach is different when break is inside a hyperlink --
        return (
            self._following_frag_in_hlink
            if self._is_in_hyperlink
            else self._following_frag_in_run
        )

    @property
    def follows_all_content(self) -> bool:
        """True when this page-break element is the last "content" in the paragraph.

        This is very uncommon case and may only occur in contrived or cases where the
        XML is edited by hand, but it is not precluded by the spec.
        """
        # -- a page-break inside a hyperlink never meets these criteria (for our
        # -- purposes at least) because it is considered "atomic" and always associated
        # -- with the page it starts on.
        if self._is_in_hyperlink:
            return False

        return bool(
            # -- XPath will match zero-or-one w:lastRenderedPageBreak element --
            self._enclosing_p.xpath(
                # -- in first run of paragraph --
                f"(./w:r)[last()]"
                # -- all page-breaks --
                f"/w:lastRenderedPageBreak"
                # -- that are not preceded by any content-bearing elements --
                f"[not(following-sibling::*[{self._run_inner_content_xpath}])]"
            )
        )

    @property
    def precedes_all_content(self) -> bool:
        """True when a `w:lastRenderedPageBreak` precedes all paragraph content.

        This is a common case; it occurs whenever the page breaks on an even paragraph
        boundary.
        """
        # -- a page-break inside a hyperlink never meets these criteria because there
        # -- is always part of the hyperlink text before the page-break.
        if self._is_in_hyperlink:
            return False

        return bool(
            # -- XPath will match zero-or-one w:lastRenderedPageBreak element --
            self._enclosing_p.xpath(
                # -- in first run of paragraph --
                f"./w:r[1]"
                # -- all page-breaks --
                f"/w:lastRenderedPageBreak"
                # -- that are not preceded by any content-bearing elements --
                f"[not(preceding-sibling::*[{self._run_inner_content_xpath}])]"
            )
        )

    @property
    def preceding_fragment_p(self) -> CT_P:
        """A "loose" `CT_P` containing only the paragraph content before this break.

        Raises `ValueError` if this `w:lastRenderedPageBreak` is not the first rendered
        paragraph in its paragraph.

        The returned `CT_P` is a "clone" (deepcopy) of the `w:p` ancestor of this
        page-break with this `w:lastRenderedPageBreak` element and all its following
        siblings removed.
        """
        if not self == self._first_lrpb_in_p(self._enclosing_p):
            raise ValueError("only defined on first rendered page-break in paragraph")

        # -- splitting approach is different when break is inside a hyperlink --
        return (
            self._preceding_frag_in_hlink
            if self._is_in_hyperlink
            else self._preceding_frag_in_run
        )

    def _enclosing_hyperlink(self, lrpb: CT_LastRenderedPageBreak) -> CT_Hyperlink:
        """The `w:hyperlink` grandparent of this `w:lastRenderedPageBreak`.

        Raises `IndexError` when this page-break has a `w:p` grandparent, so only call
        when `._is_in_hyperlink` is True.
        """
        return lrpb.xpath("./parent::w:r/parent::w:hyperlink")[0]

    @property
    def _enclosing_p(self) -> CT_P:
        """The `w:p` element parent or grandparent of this `w:lastRenderedPageBreak`."""
        return self.xpath("./ancestor::w:p[1]")[0]

    def _first_lrpb_in_p(self, p: CT_P) -> CT_LastRenderedPageBreak:
        """The first `w:lastRenderedPageBreak` element in `p`.

        Raises `ValueError` if there are no rendered page-breaks in `p`.
        """
        lrpbs = p.xpath(
            "./w:r/w:lastRenderedPageBreak | ./w:hyperlink/w:r/w:lastRenderedPageBreak"
        )
        if not lrpbs:
            raise ValueError("no rendered page-breaks in paragraph element")
        return lrpbs[0]

    @lazyproperty
    def _following_frag_in_hlink(self) -> CT_P:
        """Following CT_P fragment when break occurs within a hyperlink.

        Note this is a *partial-function* and raises when `lrpb` is not inside a
        hyperlink.
        """
        if not self._is_in_hyperlink:
            raise ValueError("only defined on a rendered page-break in a hyperlink")

        # -- work on a clone `w:p` so our mutations don't persist --
        p = copy.deepcopy(self._enclosing_p)

        # -- get this `w:lastRenderedPageBreak` in the cloned `w:p` (not self) --
        lrpb = self._first_lrpb_in_p(p)

        # -- locate `w:hyperlink` in which this `w:lastRenderedPageBreak` is found --
        hyperlink = lrpb._enclosing_hyperlink(lrpb)

        # -- delete all w:p inner-content preceding the hyperlink --
        for e in hyperlink.xpath("./preceding-sibling::*[not(self::w:pPr)]"):
            p.remove(e)

        # -- remove the whole hyperlink, it belongs to the preceding-fragment-p --
        hyperlink.getparent().remove(hyperlink)

        # -- that's it, return the remaining fragment of `w:p` clone --
        return p

    @lazyproperty
    def _following_frag_in_run(self) -> CT_P:
        """following CT_P fragment when break does not occur in a hyperlink.

        Note this is a *partial-function* and raises when `lrpb` is inside a hyperlink.
        """
        if self._is_in_hyperlink:
            raise ValueError("only defined on a rendered page-break not in a hyperlink")

        # -- work on a clone `w:p` so our mutations don't persist --
        p = copy.deepcopy(self._enclosing_p)

        # -- get this `w:lastRenderedPageBreak` in the cloned `w:p` (not self) --
        lrpb = self._first_lrpb_in_p(p)

        # -- locate `w:r` in which this `w:lastRenderedPageBreak` is found --
        enclosing_r = lrpb.xpath("./parent::w:r")[0]

        # -- delete all w:p inner-content preceding that run (but not w:pPr) --
        for e in enclosing_r.xpath("./preceding-sibling::*[not(self::w:pPr)]"):
            p.remove(e)

        # -- then remove all run inner-content preceding this lrpb in its run (but not
        # -- the `w:rPr`) and also remove the page-break itself
        for e in lrpb.xpath("./preceding-sibling::*[not(self::w:rPr)]"):
            enclosing_r.remove(e)
        enclosing_r.remove(lrpb)

        return p

    @lazyproperty
    def _is_in_hyperlink(self) -> bool:
        """True when this page-break is embedded in a hyperlink run."""
        return bool(self.xpath("./parent::w:r/parent::w:hyperlink"))

    @lazyproperty
    def _preceding_frag_in_hlink(self) -> CT_P:
        """Preceding CT_P fragment when break occurs within a hyperlink.

        Note this is a *partial-function* and raises when `lrpb` is not inside a
        hyperlink.
        """
        if not self._is_in_hyperlink:
            raise ValueError("only defined on a rendered page-break in a hyperlink")

        # -- work on a clone `w:p` so our mutations don't persist --
        p = copy.deepcopy(self._enclosing_p)

        # -- get this `w:lastRenderedPageBreak` in the cloned `w:p` (not self) --
        lrpb = self._first_lrpb_in_p(p)

        # -- locate `w:hyperlink` in which this `w:lastRenderedPageBreak` is found --
        hyperlink = lrpb._enclosing_hyperlink(lrpb)

        # -- delete all w:p inner-content following the hyperlink --
        for e in hyperlink.xpath("./following-sibling::*"):
            p.remove(e)

        # -- remove this page-break from inside the hyperlink --
        lrpb.getparent().remove(lrpb)

        # -- that's it, the entire hyperlink goes into the preceding fragment so
        # -- the hyperlink is not "split".
        return p

    @lazyproperty
    def _preceding_frag_in_run(self) -> CT_P:
        """Preceding CT_P fragment when break does not occur in a hyperlink.

        Note this is a *partial-function* and raises when `lrpb` is inside a hyperlink.
        """
        if self._is_in_hyperlink:
            raise ValueError("only defined on a rendered page-break not in a hyperlink")

        # -- work on a clone `w:p` so our mutations don't persist --
        p = copy.deepcopy(self._enclosing_p)

        # -- get this `w:lastRenderedPageBreak` in the cloned `w:p` (not self) --
        lrpb = self._first_lrpb_in_p(p)

        # -- locate `w:r` in which this `w:lastRenderedPageBreak` is found --
        enclosing_r = lrpb.xpath("./parent::w:r")[0]

        # -- delete all `w:p` inner-content following that run --
        for e in enclosing_r.xpath("./following-sibling::*"):
            p.remove(e)

        # -- then delete all `w:r` inner-content following this lrpb in its run and
        # -- also remove the page-break itself
        for e in lrpb.xpath("./following-sibling::*"):
            enclosing_r.remove(e)
        enclosing_r.remove(lrpb)

        return p

    @lazyproperty
    def _run_inner_content_xpath(self) -> str:
        """XPath fragment matching any run inner-content elements."""
        return (
            "self::w:br"
            " | self::w:cr"
            " | self::w:drawing"
            " | self::w:noBreakHyphen"
            " | self::w:ptab"
            " | self::w:t"
            " | self::w:tab"
        )


# paragraph.py
# pyright: reportPrivateUsage=false

"""Custom element classes related to paragraphs (CT_P)."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable, List, cast

from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrMore, ZeroOrOne

if TYPE_CHECKING:
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.section import CT_SectPr
    from docx.oxml.text.hyperlink import CT_Hyperlink
    from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak
    from docx.oxml.text.parfmt import CT_PPr
    from docx.oxml.text.run import CT_R


class CT_P(BaseOxmlElement):
    """`<w:p>` element, containing the properties and text for a paragraph."""

    add_r: Callable[[], CT_R]
    get_or_add_pPr: Callable[[], CT_PPr]
    hyperlink_lst: List[CT_Hyperlink]
    r_lst: List[CT_R]

    pPr: CT_PPr | None = ZeroOrOne("w:pPr")  # pyright: ignore[reportAssignmentType]
    hyperlink = ZeroOrMore("w:hyperlink")
    r = ZeroOrMore("w:r")

    def add_p_before(self) -> CT_P:
        """Return a new `<w:p>` element inserted directly prior to this one."""
        new_p = cast(CT_P, OxmlElement("w:p"))
        self.addprevious(new_p)
        return new_p

    @property
    def alignment(self) -> WD_PARAGRAPH_ALIGNMENT | None:
        """The value of the `<w:jc>` grandchild element or |None| if not present."""
        pPr = self.pPr
        if pPr is None:
            return None
        return pPr.jc_val

    @alignment.setter
    def alignment(self, value: WD_PARAGRAPH_ALIGNMENT):
        pPr = self.get_or_add_pPr()
        pPr.jc_val = value

    def clear_content(self):
        """Remove all child elements, except the `<w:pPr>` element if present."""
        for child in self.xpath("./*[not(self::w:pPr)]"):
            self.remove(child)

    @property
    def inner_content_elements(self) -> List[CT_R | CT_Hyperlink]:
        """Run and hyperlink children of the `w:p` element, in document order."""
        return self.xpath("./w:r | ./w:hyperlink")

    @property
    def lastRenderedPageBreaks(self) -> List[CT_LastRenderedPageBreak]:
        """All `w:lastRenderedPageBreak` descendants of this paragraph.

        Rendered page-breaks commonly occur in a run but can also occur in a run inside
        a hyperlink. This returns both.
        """
        return self.xpath(
            "./w:r/w:lastRenderedPageBreak | ./w:hyperlink/w:r/w:lastRenderedPageBreak"
        )

    def set_sectPr(self, sectPr: CT_SectPr):
        """Unconditionally replace or add `sectPr` as grandchild in correct sequence."""
        pPr = self.get_or_add_pPr()
        pPr._remove_sectPr()
        pPr._insert_sectPr(sectPr)

    @property
    def style(self) -> str | None:
        """String contained in `w:val` attribute of `./w:pPr/w:pStyle` grandchild.

        |None| if not present.
        """
        pPr = self.pPr
        if pPr is None:
            return None
        return pPr.style

    @style.setter
    def style(self, style: str | None):
        pPr = self.get_or_add_pPr()
        pPr.style = style

    @property
    def text(self):  # pyright: ignore[reportIncompatibleMethodOverride]
        """The textual content of this paragraph.

        Inner-content child elements like `w:r` and `w:hyperlink` are translated to
        their text equivalent.
        """
        return "".join(e.text for e in self.xpath("w:r | w:hyperlink"))

    def _insert_pPr(self, pPr: CT_PPr) -> CT_PPr:
        self.insert(0, pPr)
        return pPr


# parfmt.py
"""Custom element classes related to paragraph properties (CT_PPr)."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml.simpletypes import ST_SignedTwipsMeasure, ST_TwipsMeasure
from docx.oxml.xmlchemy import (
    BaseOxmlElement,
    OneOrMore,
    OptionalAttribute,
    RequiredAttribute,
    ZeroOrOne,
)
from docx.shared import Length

if TYPE_CHECKING:
    from docx.oxml.section import CT_SectPr
    from docx.oxml.shared import CT_String


class CT_Ind(BaseOxmlElement):
    """``<w:ind>`` element, specifying paragraph indentation."""

    left: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:left", ST_SignedTwipsMeasure
    )
    right: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:right", ST_SignedTwipsMeasure
    )
    firstLine: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:firstLine", ST_TwipsMeasure
    )
    hanging: Length | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:hanging", ST_TwipsMeasure
    )


class CT_Jc(BaseOxmlElement):
    """``<w:jc>`` element, specifying paragraph justification."""

    val: WD_ALIGN_PARAGRAPH = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", WD_ALIGN_PARAGRAPH
    )


class CT_PPr(BaseOxmlElement):
    """``<w:pPr>`` element, containing the properties for a paragraph."""

    get_or_add_ind: Callable[[], CT_Ind]
    get_or_add_pStyle: Callable[[], CT_String]
    _insert_sectPr: Callable[[CT_SectPr], None]
    _remove_pStyle: Callable[[], None]
    _remove_sectPr: Callable[[], None]

    _tag_seq = (
        "w:pStyle",
        "w:keepNext",
        "w:keepLines",
        "w:pageBreakBefore",
        "w:framePr",
        "w:widowControl",
        "w:numPr",
        "w:suppressLineNumbers",
        "w:pBdr",
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    pStyle: CT_String | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:pStyle", successors=_tag_seq[1:]
    )
    keepNext = ZeroOrOne("w:keepNext", successors=_tag_seq[2:])
    keepLines = ZeroOrOne("w:keepLines", successors=_tag_seq[3:])
    pageBreakBefore = ZeroOrOne("w:pageBreakBefore", successors=_tag_seq[4:])
    widowControl = ZeroOrOne("w:widowControl", successors=_tag_seq[6:])
    numPr = ZeroOrOne("w:numPr", successors=_tag_seq[7:])
    tabs = ZeroOrOne("w:tabs", successors=_tag_seq[11:])
    spacing = ZeroOrOne("w:spacing", successors=_tag_seq[22:])
    ind: CT_Ind | None = ZeroOrOne(  # pyright: ignore[reportAssignmentType]
        "w:ind", successors=_tag_seq[23:]
    )
    jc = ZeroOrOne("w:jc", successors=_tag_seq[27:])
    sectPr = ZeroOrOne("w:sectPr", successors=_tag_seq[35:])
    del _tag_seq

    @property
    def first_line_indent(self) -> Length | None:
        """A |Length| value calculated from the values of `w:ind/@w:firstLine` and
        `w:ind/@w:hanging`.

        Returns |None| if the `w:ind` child is not present.
        """
        ind = self.ind
        if ind is None:
            return None
        hanging = ind.hanging
        if hanging is not None:
            return Length(-hanging)
        firstLine = ind.firstLine
        if firstLine is None:
            return None
        return firstLine

    @first_line_indent.setter
    def first_line_indent(self, value: Length | None):
        if self.ind is None and value is None:
            return
        ind = self.get_or_add_ind()
        ind.firstLine = ind.hanging = None
        if value is None:
            return
        elif value < 0:
            ind.hanging = -value
        else:
            ind.firstLine = value

    @property
    def ind_left(self) -> Length | None:
        """The value of `w:ind/@w:left` or |None| if not present."""
        ind = self.ind
        if ind is None:
            return None
        return ind.left

    @ind_left.setter
    def ind_left(self, value: Length | None):
        if value is None and self.ind is None:
            return
        ind = self.get_or_add_ind()
        ind.left = value

    @property
    def ind_right(self) -> Length | None:
        """The value of `w:ind/@w:right` or |None| if not present."""
        ind = self.ind
        if ind is None:
            return None
        return ind.right

    @ind_right.setter
    def ind_right(self, value: Length | None):
        if value is None and self.ind is None:
            return
        ind = self.get_or_add_ind()
        ind.right = value

    @property
    def jc_val(self) -> WD_ALIGN_PARAGRAPH | None:
        """Value of the `<w:jc>` child element or |None| if not present."""
        return self.jc.val if self.jc is not None else None

    @jc_val.setter
    def jc_val(self, value):
        if value is None:
            self._remove_jc()
            return
        self.get_or_add_jc().val = value

    @property
    def keepLines_val(self):
        """The value of `keepLines/@val` or |None| if not present."""
        keepLines = self.keepLines
        if keepLines is None:
            return None
        return keepLines.val

    @keepLines_val.setter
    def keepLines_val(self, value):
        if value is None:
            self._remove_keepLines()
        else:
            self.get_or_add_keepLines().val = value

    @property
    def keepNext_val(self):
        """The value of `keepNext/@val` or |None| if not present."""
        keepNext = self.keepNext
        if keepNext is None:
            return None
        return keepNext.val

    @keepNext_val.setter
    def keepNext_val(self, value):
        if value is None:
            self._remove_keepNext()
        else:
            self.get_or_add_keepNext().val = value

    @property
    def pageBreakBefore_val(self):
        """The value of `pageBreakBefore/@val` or |None| if not present."""
        pageBreakBefore = self.pageBreakBefore
        if pageBreakBefore is None:
            return None
        return pageBreakBefore.val

    @pageBreakBefore_val.setter
    def pageBreakBefore_val(self, value):
        if value is None:
            self._remove_pageBreakBefore()
        else:
            self.get_or_add_pageBreakBefore().val = value

    @property
    def spacing_after(self):
        """The value of `w:spacing/@w:after` or |None| if not present."""
        spacing = self.spacing
        if spacing is None:
            return None
        return spacing.after

    @spacing_after.setter
    def spacing_after(self, value):
        if value is None and self.spacing is None:
            return
        self.get_or_add_spacing().after = value

    @property
    def spacing_before(self):
        """The value of `w:spacing/@w:before` or |None| if not present."""
        spacing = self.spacing
        if spacing is None:
            return None
        return spacing.before

    @spacing_before.setter
    def spacing_before(self, value):
        if value is None and self.spacing is None:
            return
        self.get_or_add_spacing().before = value

    @property
    def spacing_line(self):
        """The value of `w:spacing/@w:line` or |None| if not present."""
        spacing = self.spacing
        if spacing is None:
            return None
        return spacing.line

    @spacing_line.setter
    def spacing_line(self, value):
        if value is None and self.spacing is None:
            return
        self.get_or_add_spacing().line = value

    @property
    def spacing_lineRule(self):
        """The value of `w:spacing/@w:lineRule` as a member of the :ref:`WdLineSpacing`
        enumeration.

        Only the `MULTIPLE`, `EXACTLY`, and `AT_LEAST` members are used. It is the
        responsibility of the client to calculate the use of `SINGLE`, `DOUBLE`, and
        `MULTIPLE` based on the value of `w:spacing/@w:line` if that behavior is
        desired.
        """
        spacing = self.spacing
        if spacing is None:
            return None
        lineRule = spacing.lineRule
        if lineRule is None and spacing.line is not None:
            return WD_LINE_SPACING.MULTIPLE
        return lineRule

    @spacing_lineRule.setter
    def spacing_lineRule(self, value):
        if value is None and self.spacing is None:
            return
        self.get_or_add_spacing().lineRule = value

    @property
    def style(self) -> str | None:
        """String contained in `./w:pStyle/@val`, or None if child is not present."""
        pStyle = self.pStyle
        if pStyle is None:
            return None
        return pStyle.val

    @style.setter
    def style(self, style: str | None):
        """Set `./w:pStyle/@val` `style`, adding a new element if necessary.

        If `style` is |None|, remove `./w:pStyle` when present.
        """
        if style is None:
            self._remove_pStyle()
            return
        pStyle = self.get_or_add_pStyle()
        pStyle.val = style

    @property
    def widowControl_val(self):
        """The value of `widowControl/@val` or |None| if not present."""
        widowControl = self.widowControl
        if widowControl is None:
            return None
        return widowControl.val

    @widowControl_val.setter
    def widowControl_val(self, value):
        if value is None:
            self._remove_widowControl()
        else:
            self.get_or_add_widowControl().val = value


class CT_Spacing(BaseOxmlElement):
    """``<w:spacing>`` element, specifying paragraph spacing attributes such as space
    before and line spacing."""

    after = OptionalAttribute("w:after", ST_TwipsMeasure)
    before = OptionalAttribute("w:before", ST_TwipsMeasure)
    line = OptionalAttribute("w:line", ST_SignedTwipsMeasure)
    lineRule = OptionalAttribute("w:lineRule", WD_LINE_SPACING)


class CT_TabStop(BaseOxmlElement):
    """`<w:tab>` element, representing an individual tab stop.

    Overloaded to use for a tab-character in a run, which also uses the w:tab tag but
    only needs a __str__ method.
    """

    val: WD_TAB_ALIGNMENT = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "w:val", WD_TAB_ALIGNMENT
    )
    leader: WD_TAB_LEADER | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:leader", WD_TAB_LEADER, default=WD_TAB_LEADER.SPACES
    )
    pos: Length = RequiredAttribute(  # pyright: ignore[reportAssignmentType]
        "w:pos", ST_SignedTwipsMeasure
    )

    def __str__(self) -> str:
        """Text equivalent of a `w:tab` element appearing in a run.

        Allows text of run inner-content to be accessed consistently across all text
        inner-content.
        """
        return "\t"


class CT_TabStops(BaseOxmlElement):
    """``<w:tabs>`` element, container for a sorted sequence of tab stops."""

    tab = OneOrMore("w:tab", successors=())

    def insert_tab_in_order(self, pos, align, leader):
        """Insert a newly created `w:tab` child element in `pos` order."""
        new_tab = self._new_tab()
        new_tab.pos, new_tab.val, new_tab.leader = pos, align, leader
        for tab in self.tab_lst:
            if new_tab.pos < tab.pos:
                tab.addprevious(new_tab)
                return new_tab
        self.append(new_tab)
        return new_tab


# run.py
"""Custom element classes related to text runs (CT_R)."""

from __future__ import annotations

from typing import TYPE_CHECKING, Callable, Iterator, List

from docx.oxml.drawing import CT_Drawing
from docx.oxml.ns import qn
from docx.oxml.simpletypes import ST_BrClear, ST_BrType
from docx.oxml.text.font import CT_RPr
from docx.oxml.xmlchemy import BaseOxmlElement, OptionalAttribute, ZeroOrMore, ZeroOrOne
from docx.shared import TextAccumulator

if TYPE_CHECKING:
    from docx.oxml.shape import CT_Anchor, CT_Inline
    from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak
    from docx.oxml.text.parfmt import CT_TabStop

# ------------------------------------------------------------------------------------
# Run-level elements


class CT_R(BaseOxmlElement):
    """`<w:r>` element, containing the properties and text for a run."""

    add_br: Callable[[], CT_Br]
    add_tab: Callable[[], CT_TabStop]
    get_or_add_rPr: Callable[[], CT_RPr]
    _add_drawing: Callable[[], CT_Drawing]
    _add_t: Callable[..., CT_Text]

    rPr: CT_RPr | None = ZeroOrOne("w:rPr")  # pyright: ignore[reportAssignmentType]
    br = ZeroOrMore("w:br")
    cr = ZeroOrMore("w:cr")
    drawing = ZeroOrMore("w:drawing")
    t = ZeroOrMore("w:t")
    tab = ZeroOrMore("w:tab")

    def add_t(self, text: str) -> CT_Text:
        """Return a newly added `<w:t>` element containing `text`."""
        t = self._add_t(text=text)
        if len(text.strip()) < len(text):
            t.set(qn("xml:space"), "preserve")
        return t

    def add_drawing(self, inline_or_anchor: CT_Inline | CT_Anchor) -> CT_Drawing:
        """Return newly appended `CT_Drawing` (`w:drawing`) child element.

        The `w:drawing` element has `inline_or_anchor` as its child.
        """
        drawing = self._add_drawing()
        drawing.append(inline_or_anchor)
        return drawing

    def clear_content(self) -> None:
        """Remove all child elements except a `w:rPr` element if present."""
        # -- remove all run inner-content except a `w:rPr` when present. --
        for e in self.xpath("./*[not(self::w:rPr)]"):
            self.remove(e)

    @property
    def inner_content_items(self) -> List[str | CT_Drawing | CT_LastRenderedPageBreak]:
        """Text of run, possibly punctuated by `w:lastRenderedPageBreak` elements."""
        from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak

        accum = TextAccumulator()

        def iter_items() -> Iterator[str | CT_Drawing | CT_LastRenderedPageBreak]:
            for e in self.xpath(
                "w:br"
                " | w:cr"
                " | w:drawing"
                " | w:lastRenderedPageBreak"
                " | w:noBreakHyphen"
                " | w:ptab"
                " | w:t"
                " | w:tab"
            ):
                if isinstance(e, (CT_Drawing, CT_LastRenderedPageBreak)):
                    yield from accum.pop()
                    yield e
                else:
                    accum.push(str(e))

            # -- don't forget the "tail" string --
            yield from accum.pop()

        return list(iter_items())

    @property
    def lastRenderedPageBreaks(self) -> List[CT_LastRenderedPageBreak]:
        """All `w:lastRenderedPageBreaks` descendants of this run."""
        return self.xpath("./w:lastRenderedPageBreak")

    @property
    def style(self) -> str | None:
        """String contained in `w:val` attribute of `w:rStyle` grandchild.

        |None| if that element is not present.
        """
        rPr = self.rPr
        if rPr is None:
            return None
        return rPr.style

    @style.setter
    def style(self, style: str | None):
        """Set character style of this `w:r` element to `style`.

        If `style` is None, remove the style element.
        """
        rPr = self.get_or_add_rPr()
        rPr.style = style

    @property
    def text(self) -> str:
        """The textual content of this run.

        Inner-content child elements like `w:tab` are translated to their text
        equivalent.
        """
        return "".join(
            str(e) for e in self.xpath("w:br | w:cr | w:noBreakHyphen | w:ptab | w:t | w:tab")
        )

    @text.setter
    def text(self, text: str):  # pyright: ignore[reportIncompatibleMethodOverride]
        self.clear_content()
        _RunContentAppender.append_to_run_from_text(self, text)

    def _insert_rPr(self, rPr: CT_RPr) -> CT_RPr:
        self.insert(0, rPr)
        return rPr


# ------------------------------------------------------------------------------------
# Run inner-content elements


class CT_Br(BaseOxmlElement):
    """`<w:br>` element, indicating a line, page, or column break in a run."""

    type: str | None = OptionalAttribute(  # pyright: ignore[reportAssignmentType]
        "w:type", ST_BrType, default="textWrapping"
    )
    clear: str | None = OptionalAttribute("w:clear", ST_BrClear)  # pyright: ignore

    def __str__(self) -> str:
        """Text equivalent of this element. Actual value depends on break type.

        A line break is translated as "\n". Column and page breaks produce the empty
        string ("").

        This allows the text of run inner-content to be accessed in a consistent way
        for all run inner-context text elements.
        """
        return "\n" if self.type == "textWrapping" else ""


class CT_Cr(BaseOxmlElement):
    """`<w:cr>` element, representing a carriage-return (0x0D) character within a run.

    In Word, this represents a "soft carriage-return" in the sense that it does not end
    the paragraph the way pressing Enter (aka. Return) on the keyboard does. Here the
    text equivalent is considered to be newline ("\n") since in plain-text that's the
    closest Python equivalent.

    NOTE: this complex-type name does not exist in the schema, where `w:tab` maps to
    `CT_Empty`. This name was added to give it distinguished behavior. CT_Empty is used
    for many elements.
    """

    def __str__(self) -> str:
        """Text equivalent of this element, a single newline ("\n")."""
        return "\n"


class CT_NoBreakHyphen(BaseOxmlElement):
    """`<w:noBreakHyphen>` element, a hyphen ineligible for a line-wrap position.

    This maps to a plain-text dash ("-").

    NOTE: this complex-type name does not exist in the schema, where `w:noBreakHyphen`
    maps to `CT_Empty`. This name was added to give it behavior distinguished from the
    many other elements represented in the schema by CT_Empty.
    """

    def __str__(self) -> str:
        """Text equivalent of this element, a single dash character ("-")."""
        return "-"


class CT_PTab(BaseOxmlElement):
    """`<w:ptab>` element, representing an absolute-position tab character within a run.

    This character advances the rendering position to the specified position regardless
    of any tab-stops, perhaps for layout of a table-of-contents (TOC) or similar.
    """

    def __str__(self) -> str:
        """Text equivalent of this element, a single tab ("\t") character.

        This allows the text of run inner-content to be accessed in a consistent way
        for all run inner-context text elements.
        """
        return "\t"


# -- CT_Tab functionality is provided by CT_TabStop which also uses `w:tab` tag. That
# -- element class provides the __str__() method for this empty element, unconditionally
# -- returning "\t".


class CT_Text(BaseOxmlElement):
    """`<w:t>` element, containing a sequence of characters within a run."""

    def __str__(self) -> str:
        """Text contained in this element, the empty string if it has no content.

        This property allows this run inner-content element to be queried for its text
        the same way as other run-content elements are. In particular, this never
        returns None, as etree._Element does when there is no content.
        """
        return self.text or ""


# ------------------------------------------------------------------------------------
# Utility


class _RunContentAppender:
    """Translates a Python string into run content elements appended in a `w:r` element.

    Contiguous sequences of regular characters are appended in a single `<w:t>` element.
    Each tab character ('\t') causes a `<w:tab/>` element to be appended. Likewise a
    newline or carriage return character ('\n', '\r') causes a `<w:cr>` element to be
    appended.
    """

    def __init__(self, r: CT_R):
        self._r = r
        self._bfr: List[str] = []

    @classmethod
    def append_to_run_from_text(cls, r: CT_R, text: str):
        """Append inner-content elements for `text` to `r` element."""
        appender = cls(r)
        appender.add_text(text)

    def add_text(self, text: str):
        """Append inner-content elements for `text` to the `w:r` element."""
        for char in text:
            self.add_char(char)
        self.flush()

    def add_char(self, char: str):
        """Process next character of input through finite state maching (FSM).

        There are two possible states, buffer pending and not pending, but those are
        hidden behind the `.flush()` method which must be called at the end of text to
        ensure any pending `<w:t>` element is written.
        """
        if char == "\t":
            self.flush()
            self._r.add_tab()
        elif char in "\r\n":
            self.flush()
            self._r.add_br()
        else:
            self._bfr.append(char)

    def flush(self):
        text = "".join(self._bfr)
        if text:
            self._r.add_t(text)
        self._bfr.clear()


# __init__.py


# xmlchemy.py
# pyright: reportImportCycles=false

"""Enabling declarative definition of lxml custom element classes."""

from __future__ import annotations

import re
from typing import (
    TYPE_CHECKING,
    Any,
    Callable,
    Dict,
    List,
    Sequence,
    Tuple,
    Type,
    TypeVar,
)

from lxml import etree
from lxml.etree import ElementBase, _Element  # pyright: ignore[reportPrivateUsage]

from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import NamespacePrefixedTag, nsmap, qn
from docx.shared import lazyproperty

if TYPE_CHECKING:
    from docx.enum.base import BaseXmlEnum
    from docx.oxml.simpletypes import BaseSimpleType


def serialize_for_reading(element: ElementBase):
    """Serialize `element` to human-readable XML suitable for tests.

    No XML declaration.
    """
    xml = etree.tostring(element, encoding="unicode", pretty_print=True)
    return XmlString(xml)


class XmlString(str):
    """Provides string comparison override suitable for serialized XML that is useful
    for tests."""

    # '    <w:xyz xmlns:a="http://ns/decl/a" attr_name="val">text</w:xyz>'
    # |          |                                          ||           |
    # +----------+------------------------------------------++-----------+
    #  front      attrs                                     | text
    #                                                     close

    _xml_elm_line_patt = re.compile(r"( *</?[\w:]+)(.*?)(/?>)([^<]*</[\w:]+>)?$")

    def __eq__(self, other: object) -> bool:
        if not isinstance(other, str):
            return False
        lines = self.splitlines()
        lines_other = other.splitlines()
        if len(lines) != len(lines_other):
            return False
        for line, line_other in zip(lines, lines_other):
            if not self._eq_elm_strs(line, line_other):
                return False
        return True

    def __ne__(self, other: object) -> bool:
        return not self.__eq__(other)

    def _attr_seq(self, attrs: str) -> List[str]:
        """Return a sequence of attribute strings parsed from `attrs`.

        Each attribute string is stripped of whitespace on both ends.
        """
        attrs = attrs.strip()
        attr_lst = attrs.split()
        return sorted(attr_lst)

    def _eq_elm_strs(self, line: str, line_2: str):
        """Return True if the element in `line_2` is XML equivalent to the element in
        `line`."""
        front, attrs, close, text = self._parse_line(line)
        front_2, attrs_2, close_2, text_2 = self._parse_line(line_2)
        if front != front_2:
            return False
        if self._attr_seq(attrs) != self._attr_seq(attrs_2):
            return False
        if close != close_2:
            return False
        if text != text_2:
            return False
        return True

    @classmethod
    def _parse_line(cls, line: str) -> Tuple[str, str, str, str]:
        """(front, attrs, close, text) 4-tuple result of parsing XML element `line`."""
        match = cls._xml_elm_line_patt.match(line)
        if match is None:
            return "", "", "", ""
        front, attrs, close, text = [match.group(n) for n in range(1, 5)]
        return front, attrs, close, text


_T = TypeVar("_T")


class MetaOxmlElement(type):
    """Metaclass for BaseOxmlElement."""

    def __init__(cls, clsname: str, bases: Tuple[type, ...], namespace: Dict[str, Any]):
        dispatchable = (
            OneAndOnlyOne,
            OneOrMore,
            OptionalAttribute,
            RequiredAttribute,
            ZeroOrMore,
            ZeroOrOne,
            ZeroOrOneChoice,
        )
        for key, value in namespace.items():
            if isinstance(value, dispatchable):
                value.populate_class_members(cls, key)


class BaseAttribute:
    """Base class for OptionalAttribute and RequiredAttribute.

    Provides common methods.
    """

    def __init__(self, attr_name: str, simple_type: Type[BaseXmlEnum] | Type[BaseSimpleType]):
        super(BaseAttribute, self).__init__()
        self._attr_name = attr_name
        self._simple_type = simple_type

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        self._element_cls = element_cls
        self._prop_name = prop_name

        self._add_attr_property()

    def _add_attr_property(self):
        """Add a read/write `.{prop_name}` property to the element class.

        The property returns the interpreted value of this attribute on access and
        changes the attribute value to its ST_* counterpart on assignment.
        """
        property_ = property(self._getter, self._setter, None)
        # -- assign unconditionally to overwrite element name definition --
        setattr(self._element_cls, self._prop_name, property_)

    @property
    def _clark_name(self):
        if ":" in self._attr_name:
            return qn(self._attr_name)
        return self._attr_name

    @property
    def _getter(self) -> Callable[[BaseOxmlElement], Any | None]: ...

    @property
    def _setter(
        self,
    ) -> Callable[[BaseOxmlElement, Any | None], None]: ...


class OptionalAttribute(BaseAttribute):
    """Defines an optional attribute on a custom element class.

    An optional attribute returns a default value when not present for reading. When
    assigned |None|, the attribute is removed, but still returns the default value when
    one is specified.
    """

    def __init__(
        self,
        attr_name: str,
        simple_type: Type[BaseXmlEnum] | Type[BaseSimpleType],
        default: BaseXmlEnum | BaseSimpleType | str | bool | None = None,
    ):
        super(OptionalAttribute, self).__init__(attr_name, simple_type)
        self._default = default

    @property
    def _docstring(self):
        """String to use as `__doc__` attribute of attribute property."""
        return (
            f"{self._simple_type.__name__} type-converted value of"
            f" ``{self._attr_name}`` attribute, or |None| (or specified default"
            f" value) if not present. Assigning the default value causes the"
            f" attribute to be removed from the element."
        )

    @property
    def _getter(
        self,
    ) -> Callable[[BaseOxmlElement], Any | None]:
        """Function suitable for `__get__()` method on attribute property descriptor."""

        def get_attr_value(
            obj: BaseOxmlElement,
        ) -> Any | None:
            attr_str_value = obj.get(self._clark_name)
            if attr_str_value is None:
                return self._default
            return self._simple_type.from_xml(attr_str_value)

        get_attr_value.__doc__ = self._docstring
        return get_attr_value

    @property
    def _setter(self) -> Callable[[BaseOxmlElement, Any], None]:
        """Function suitable for `__set__()` method on attribute property descriptor."""

        def set_attr_value(obj: BaseOxmlElement, value: Any | None):
            if value is None or value == self._default:
                if self._clark_name in obj.attrib:
                    del obj.attrib[self._clark_name]
                return
            str_value = self._simple_type.to_xml(value)
            if str_value is None:
                if self._clark_name in obj.attrib:
                    del obj.attrib[self._clark_name]
                return
            obj.set(self._clark_name, str_value)

        return set_attr_value


class RequiredAttribute(BaseAttribute):
    """Defines a required attribute on a custom element class.

    A required attribute is assumed to be present for reading, so does not have a
    default value; its actual value is always used. If missing on read, an
    |InvalidXmlError| is raised. It also does not remove the attribute if |None| is
    assigned. Assigning |None| raises |TypeError| or |ValueError|, depending on the
    simple type of the attribute.
    """

    @property
    def _docstring(self):
        """Return the string to use as the ``__doc__`` attribute of the property for
        this attribute."""
        return "%s type-converted value of ``%s`` attribute." % (
            self._simple_type.__name__,
            self._attr_name,
        )

    @property
    def _getter(self) -> Callable[[BaseOxmlElement], Any]:
        """function object suitable for "get" side of attr property descriptor."""

        def get_attr_value(obj: BaseOxmlElement) -> Any | None:
            attr_str_value = obj.get(self._clark_name)
            if attr_str_value is None:
                raise InvalidXmlError(
                    "required '%s' attribute not present on element %s" % (self._attr_name, obj.tag)
                )
            return self._simple_type.from_xml(attr_str_value)

        get_attr_value.__doc__ = self._docstring
        return get_attr_value

    @property
    def _setter(self) -> Callable[[BaseOxmlElement, Any], None]:
        """function object suitable for "set" side of attribute property descriptor."""

        def set_attr_value(obj: BaseOxmlElement, value: Any):
            str_value = self._simple_type.to_xml(value)
            if str_value is None:
                raise ValueError(f"cannot assign {value} to this required attribute")
            obj.set(self._clark_name, str_value)

        return set_attr_value


class _BaseChildElement:
    """Base class for the child-element classes.

    The child-element sub-classes correspond to varying cardinalities, such as ZeroOrOne
    and ZeroOrMore.
    """

    def __init__(self, nsptagname: str, successors: Tuple[str, ...] = ()):
        super(_BaseChildElement, self).__init__()
        self._nsptagname = nsptagname
        self._successors = successors

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Baseline behavior for adding the appropriate methods to `element_cls`."""
        self._element_cls = element_cls
        self._prop_name = prop_name

    def _add_adder(self):
        """Add an ``_add_x()`` method to the element class for this child element."""

        def _add_child(obj: BaseOxmlElement, **attrs: Any):
            new_method = getattr(obj, self._new_method_name)
            child = new_method()
            for key, value in attrs.items():
                setattr(child, key, value)
            insert_method = getattr(obj, self._insert_method_name)
            insert_method(child)
            return child

        _add_child.__doc__ = (
            "Add a new ``<%s>`` child element unconditionally, inserted in t"
            "he correct sequence." % self._nsptagname
        )
        self._add_to_class(self._add_method_name, _add_child)

    def _add_creator(self):
        """Add a ``_new_{prop_name}()`` method to the element class that creates a new,
        empty element of the correct type, having no attributes."""
        creator = self._creator
        creator.__doc__ = (
            'Return a "loose", newly created ``<%s>`` element having no attri'
            "butes, text, or children." % self._nsptagname
        )
        self._add_to_class(self._new_method_name, creator)

    def _add_getter(self):
        """Add a read-only ``{prop_name}`` property to the element class for this child
        element."""
        property_ = property(self._getter, None, None)
        # -- assign unconditionally to overwrite element name definition --
        setattr(self._element_cls, self._prop_name, property_)

    def _add_inserter(self):
        """Add an ``_insert_x()`` method to the element class for this child element."""

        def _insert_child(obj: BaseOxmlElement, child: BaseOxmlElement):
            obj.insert_element_before(child, *self._successors)
            return child

        _insert_child.__doc__ = (
            "Return the passed ``<%s>`` element after inserting it as a chil"
            "d in the correct sequence." % self._nsptagname
        )
        self._add_to_class(self._insert_method_name, _insert_child)

    def _add_list_getter(self):
        """Add a read-only ``{prop_name}_lst`` property to the element class to retrieve
        a list of child elements matching this type."""
        prop_name = "%s_lst" % self._prop_name
        property_ = property(self._list_getter, None, None)
        setattr(self._element_cls, prop_name, property_)

    @lazyproperty
    def _add_method_name(self):
        return "_add_%s" % self._prop_name

    def _add_public_adder(self):
        """Add a public ``add_x()`` method to the parent element class."""

        def add_child(obj: BaseOxmlElement):
            private_add_method = getattr(obj, self._add_method_name)
            child = private_add_method()
            return child

        add_child.__doc__ = (
            "Add a new ``<%s>`` child element unconditionally, inserted in t"
            "he correct sequence." % self._nsptagname
        )
        self._add_to_class(self._public_add_method_name, add_child)

    def _add_to_class(self, name: str, method: Callable[..., Any]):
        """Add `method` to the target class as `name`, unless `name` is already defined
        on the class."""
        if hasattr(self._element_cls, name):
            return
        setattr(self._element_cls, name, method)

    @property
    def _creator(self) -> Callable[[BaseOxmlElement], BaseOxmlElement]:
        """Callable that creates an empty element of the right type, with no attrs."""
        from docx.oxml.parser import OxmlElement

        def new_child_element(obj: BaseOxmlElement):
            return OxmlElement(self._nsptagname)

        return new_child_element

    @property
    def _getter(self):
        """Return a function object suitable for the "get" side of the property
        descriptor.

        This default getter returns the child element with matching tag name or |None|
        if not present.
        """

        def get_child_element(obj: BaseOxmlElement):
            return obj.find(qn(self._nsptagname))

        get_child_element.__doc__ = (
            "``<%s>`` child element or |None| if not present." % self._nsptagname
        )
        return get_child_element

    @lazyproperty
    def _insert_method_name(self):
        return "_insert_%s" % self._prop_name

    @property
    def _list_getter(self):
        """Return a function object suitable for the "get" side of a list property
        descriptor."""

        def get_child_element_list(obj: BaseOxmlElement):
            return obj.findall(qn(self._nsptagname))

        get_child_element_list.__doc__ = (
            "A list containing each of the ``<%s>`` child elements, in the o"
            "rder they appear." % self._nsptagname
        )
        return get_child_element_list

    @lazyproperty
    def _public_add_method_name(self):
        """add_childElement() is public API for a repeating element, allowing new
        elements to be added to the sequence.

        May be overridden to provide a friendlier API to clients having domain
        appropriate parameter names for required attributes.
        """
        return "add_%s" % self._prop_name

    @lazyproperty
    def _remove_method_name(self):
        return "_remove_%s" % self._prop_name

    @lazyproperty
    def _new_method_name(self):
        return "_new_%s" % self._prop_name


class Choice(_BaseChildElement):
    """Defines a child element belonging to a group, only one of which may appear as a
    child."""

    @property
    def nsptagname(self):
        return self._nsptagname

    def populate_class_members(  # pyright: ignore[reportIncompatibleMethodOverride]
        self,
        element_cls: MetaOxmlElement,
        group_prop_name: str,
        successors: Tuple[str, ...],
    ) -> None:
        """Add the appropriate methods to `element_cls`."""
        self._element_cls = element_cls
        self._group_prop_name = group_prop_name
        self._successors = successors

        self._add_getter()
        self._add_creator()
        self._add_inserter()
        self._add_adder()
        self._add_get_or_change_to_method()

    def _add_get_or_change_to_method(self):
        """Add a ``get_or_change_to_x()`` method to the element class for this child
        element."""

        def get_or_change_to_child(obj: BaseOxmlElement):
            child = getattr(obj, self._prop_name)
            if child is not None:
                return child
            remove_group_method = getattr(obj, self._remove_group_method_name)
            remove_group_method()
            add_method = getattr(obj, self._add_method_name)
            child = add_method()
            return child

        get_or_change_to_child.__doc__ = (
            "Return the ``<%s>`` child, replacing any other group element if" " found."
        ) % self._nsptagname
        self._add_to_class(self._get_or_change_to_method_name, get_or_change_to_child)

    @property
    def _prop_name(self):
        """Property name computed from tag name, e.g. a:schemeClr -> schemeClr."""
        start = self._nsptagname.index(":") + 1 if ":" in self._nsptagname else 0
        return self._nsptagname[start:]

    @lazyproperty
    def _get_or_change_to_method_name(self):
        return "get_or_change_to_%s" % self._prop_name

    @lazyproperty
    def _remove_group_method_name(self):
        return "_remove_%s" % self._group_prop_name


class OneAndOnlyOne(_BaseChildElement):
    """Defines a required child element for MetaOxmlElement."""

    def __init__(self, nsptagname: str):
        super(OneAndOnlyOne, self).__init__(nsptagname, ())

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        super(OneAndOnlyOne, self).populate_class_members(element_cls, prop_name)
        self._add_getter()

    @property
    def _getter(self):
        """Return a function object suitable for the "get" side of the property
        descriptor."""

        def get_child_element(obj: BaseOxmlElement):
            child = obj.find(qn(self._nsptagname))
            if child is None:
                raise InvalidXmlError(
                    "required ``<%s>`` child element not present" % self._nsptagname
                )
            return child

        get_child_element.__doc__ = "Required ``<%s>`` child element." % self._nsptagname
        return get_child_element


class OneOrMore(_BaseChildElement):
    """Defines a repeating child element for MetaOxmlElement that must appear at least
    once."""

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        super(OneOrMore, self).populate_class_members(element_cls, prop_name)
        self._add_list_getter()
        self._add_creator()
        self._add_inserter()
        self._add_adder()
        self._add_public_adder()
        delattr(element_cls, prop_name)


class ZeroOrMore(_BaseChildElement):
    """Defines an optional repeating child element for MetaOxmlElement."""

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        super(ZeroOrMore, self).populate_class_members(element_cls, prop_name)
        self._add_list_getter()
        self._add_creator()
        self._add_inserter()
        self._add_adder()
        self._add_public_adder()
        delattr(element_cls, prop_name)


class ZeroOrOne(_BaseChildElement):
    """Defines an optional child element for MetaOxmlElement."""

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        super(ZeroOrOne, self).populate_class_members(element_cls, prop_name)
        self._add_getter()
        self._add_creator()
        self._add_inserter()
        self._add_adder()
        self._add_get_or_adder()
        self._add_remover()

    def _add_get_or_adder(self):
        """Add a ``get_or_add_x()`` method to the element class for this child
        element."""

        def get_or_add_child(obj: BaseOxmlElement):
            child = getattr(obj, self._prop_name)
            if child is None:
                add_method = getattr(obj, self._add_method_name)
                child = add_method()
            return child

        get_or_add_child.__doc__ = (
            "Return the ``<%s>`` child element, newly added if not present."
        ) % self._nsptagname
        self._add_to_class(self._get_or_add_method_name, get_or_add_child)

    def _add_remover(self):
        """Add a ``_remove_x()`` method to the element class for this child element."""

        def _remove_child(obj: BaseOxmlElement):
            obj.remove_all(self._nsptagname)

        _remove_child.__doc__ = ("Remove all ``<%s>`` child elements.") % self._nsptagname
        self._add_to_class(self._remove_method_name, _remove_child)

    @lazyproperty
    def _get_or_add_method_name(self):
        return "get_or_add_%s" % self._prop_name


class ZeroOrOneChoice(_BaseChildElement):
    """Correspondes to an ``EG_*`` element group where at most one of its members may
    appear as a child."""

    def __init__(self, choices: Sequence[Choice], successors: Tuple[str, ...] = ()):
        self._choices = choices
        self._successors = successors

    def populate_class_members(self, element_cls: MetaOxmlElement, prop_name: str) -> None:
        """Add the appropriate methods to `element_cls`."""
        super(ZeroOrOneChoice, self).populate_class_members(element_cls, prop_name)
        self._add_choice_getter()
        for choice in self._choices:
            choice.populate_class_members(element_cls, self._prop_name, self._successors)
        self._add_group_remover()

    def _add_choice_getter(self):
        """Add a read-only ``{prop_name}`` property to the element class that returns
        the present member of this group, or |None| if none are present."""
        property_ = property(self._choice_getter, None, None)
        # assign unconditionally to overwrite element name definition
        setattr(self._element_cls, self._prop_name, property_)

    def _add_group_remover(self):
        """Add a ``_remove_eg_x()`` method to the element class for this choice
        group."""

        def _remove_choice_group(obj: BaseOxmlElement):
            for tagname in self._member_nsptagnames:
                obj.remove_all(tagname)

        _remove_choice_group.__doc__ = "Remove the current choice group child element if present."
        self._add_to_class(self._remove_choice_group_method_name, _remove_choice_group)

    @property
    def _choice_getter(self):
        """Return a function object suitable for the "get" side of the property
        descriptor."""

        def get_group_member_element(obj: BaseOxmlElement):
            return obj.first_child_found_in(*self._member_nsptagnames)

        get_group_member_element.__doc__ = (
            "Return the child element belonging to this element group, or "
            "|None| if no member child is present."
        )
        return get_group_member_element

    @lazyproperty
    def _member_nsptagnames(self):
        """Sequence of namespace-prefixed tagnames, one for each of the member elements
        of this choice group."""
        return [choice.nsptagname for choice in self._choices]

    @lazyproperty
    def _remove_choice_group_method_name(self):
        return "_remove_%s" % self._prop_name


# -- lxml typing isn't quite right here, just ignore this error on _Element --
class BaseOxmlElement(etree.ElementBase, metaclass=MetaOxmlElement):
    """Effective base class for all custom element classes.

    Adds standardized behavior to all classes in one place.
    """

    def __repr__(self):
        return "<%s '<%s>' at 0x%0x>" % (
            self.__class__.__name__,
            self._nsptag,
            id(self),
        )

    def first_child_found_in(self, *tagnames: str) -> _Element | None:
        """First child with tag in `tagnames`, or None if not found."""
        for tagname in tagnames:
            child = self.find(qn(tagname))
            if child is not None:
                return child
        return None

    def insert_element_before(self, elm: ElementBase, *tagnames: str):
        successor = self.first_child_found_in(*tagnames)
        if successor is not None:
            successor.addprevious(elm)
        else:
            self.append(elm)
        return elm

    def remove_all(self, *tagnames: str) -> None:
        """Remove child elements with tagname (e.g. "a:p") in `tagnames`."""
        for tagname in tagnames:
            matching = self.findall(qn(tagname))
            for child in matching:
                self.remove(child)

    @property
    def xml(self) -> str:
        """XML string for this element, suitable for testing purposes.

        Pretty printed for readability and without an XML declaration at the top.
        """
        return serialize_for_reading(self)

    def xpath(self, xpath_str: str) -> Any:  # pyright: ignore[reportIncompatibleMethodOverride]
        """Override of `lxml` _Element.xpath() method.

        Provides standard Open XML namespace mapping (`nsmap`) in centralized location.
        """
        return super().xpath(xpath_str, namespaces=nsmap)

    @property
    def _nsptag(self) -> str:
        return NamespacePrefixedTag.from_clark_name(self.tag)


# __init__.py
"""Initializes oxml sub-package.

This including registering custom element classes corresponding to Open XML elements.
"""

from __future__ import annotations

from docx.oxml.drawing import CT_Drawing
from docx.oxml.parser import OxmlElement, parse_xml, register_element_cls
from docx.oxml.shape import (
    CT_Anchor,
    CT_Blip,
    CT_BlipFillProperties,
    CT_GraphicalObject,
    CT_GraphicalObjectData,
    CT_Inline,
    CT_NonVisualDrawingProps,
    CT_Picture,
    CT_PictureNonVisual,
    CT_Point2D,
    CT_PositiveSize2D,
    CT_ShapeProperties,
    CT_Transform2D,
)
from docx.oxml.shared import CT_DecimalNumber, CT_OnOff, CT_String
from docx.oxml.text.hyperlink import CT_Hyperlink
from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak
from docx.oxml.text.run import (
    CT_R,
    CT_Br,
    CT_Cr,
    CT_NoBreakHyphen,
    CT_PTab,
    CT_Text,
)

# -- `OxmlElement` and `parse_xml()` are not used in this module but several downstream
# -- "extension" packages expect to find them here and there's no compelling reason
# -- not to republish them here so those keep working.
__all__ = ["OxmlElement", "parse_xml"]

# ---------------------------------------------------------------------------
# DrawingML-related elements

register_element_cls("a:blip", CT_Blip)
register_element_cls("a:ext", CT_PositiveSize2D)
register_element_cls("a:graphic", CT_GraphicalObject)
register_element_cls("a:graphicData", CT_GraphicalObjectData)
register_element_cls("a:off", CT_Point2D)
register_element_cls("a:xfrm", CT_Transform2D)
register_element_cls("pic:blipFill", CT_BlipFillProperties)
register_element_cls("pic:cNvPr", CT_NonVisualDrawingProps)
register_element_cls("pic:nvPicPr", CT_PictureNonVisual)
register_element_cls("pic:pic", CT_Picture)
register_element_cls("pic:spPr", CT_ShapeProperties)
register_element_cls("w:drawing", CT_Drawing)
register_element_cls("wp:anchor", CT_Anchor)
register_element_cls("wp:docPr", CT_NonVisualDrawingProps)
register_element_cls("wp:extent", CT_PositiveSize2D)
register_element_cls("wp:inline", CT_Inline)

# ---------------------------------------------------------------------------
# hyperlink-related elements

register_element_cls("w:hyperlink", CT_Hyperlink)

# ---------------------------------------------------------------------------
# text-related elements

register_element_cls("w:br", CT_Br)
register_element_cls("w:cr", CT_Cr)
register_element_cls("w:lastRenderedPageBreak", CT_LastRenderedPageBreak)
register_element_cls("w:noBreakHyphen", CT_NoBreakHyphen)
register_element_cls("w:ptab", CT_PTab)
register_element_cls("w:r", CT_R)
register_element_cls("w:t", CT_Text)

# ---------------------------------------------------------------------------
# header/footer-related mappings

register_element_cls("w:evenAndOddHeaders", CT_OnOff)
register_element_cls("w:titlePg", CT_OnOff)

# ---------------------------------------------------------------------------
# other custom element class mappings

from .coreprops import CT_CoreProperties  # noqa

register_element_cls("cp:coreProperties", CT_CoreProperties)

from .document import CT_Body, CT_Document  # noqa

register_element_cls("w:body", CT_Body)
register_element_cls("w:document", CT_Document)

from .numbering import CT_Num, CT_Numbering, CT_NumLvl, CT_NumPr  # noqa

register_element_cls("w:abstractNumId", CT_DecimalNumber)
register_element_cls("w:ilvl", CT_DecimalNumber)
register_element_cls("w:lvlOverride", CT_NumLvl)
register_element_cls("w:num", CT_Num)
register_element_cls("w:numId", CT_DecimalNumber)
register_element_cls("w:numPr", CT_NumPr)
register_element_cls("w:numbering", CT_Numbering)
register_element_cls("w:startOverride", CT_DecimalNumber)

from .section import (  # noqa
    CT_HdrFtr,
    CT_HdrFtrRef,
    CT_PageMar,
    CT_PageSz,
    CT_SectPr,
    CT_SectType,
)

register_element_cls("w:footerReference", CT_HdrFtrRef)
register_element_cls("w:ftr", CT_HdrFtr)
register_element_cls("w:hdr", CT_HdrFtr)
register_element_cls("w:headerReference", CT_HdrFtrRef)
register_element_cls("w:pgMar", CT_PageMar)
register_element_cls("w:pgSz", CT_PageSz)
register_element_cls("w:sectPr", CT_SectPr)
register_element_cls("w:type", CT_SectType)

from .settings import CT_Settings  # noqa

register_element_cls("w:settings", CT_Settings)

from .styles import CT_LatentStyles, CT_LsdException, CT_Style, CT_Styles  # noqa

register_element_cls("w:basedOn", CT_String)
register_element_cls("w:latentStyles", CT_LatentStyles)
register_element_cls("w:locked", CT_OnOff)
register_element_cls("w:lsdException", CT_LsdException)
register_element_cls("w:name", CT_String)
register_element_cls("w:next", CT_String)
register_element_cls("w:qFormat", CT_OnOff)
register_element_cls("w:semiHidden", CT_OnOff)
register_element_cls("w:style", CT_Style)
register_element_cls("w:styles", CT_Styles)
register_element_cls("w:uiPriority", CT_DecimalNumber)
register_element_cls("w:unhideWhenUsed", CT_OnOff)

from .table import (  # noqa
    CT_Height,
    CT_Row,
    CT_Tbl,
    CT_TblGrid,
    CT_TblGridCol,
    CT_TblLayoutType,
    CT_TblPr,
    CT_TblPrEx,
    CT_TblWidth,
    CT_Tc,
    CT_TcPr,
    CT_TrPr,
    CT_VMerge,
    CT_VerticalJc,
)

register_element_cls("w:bidiVisual", CT_OnOff)
register_element_cls("w:gridAfter", CT_DecimalNumber)
register_element_cls("w:gridBefore", CT_DecimalNumber)
register_element_cls("w:gridCol", CT_TblGridCol)
register_element_cls("w:gridSpan", CT_DecimalNumber)
register_element_cls("w:tbl", CT_Tbl)
register_element_cls("w:tblGrid", CT_TblGrid)
register_element_cls("w:tblLayout", CT_TblLayoutType)
register_element_cls("w:tblPr", CT_TblPr)
register_element_cls("w:tblPrEx", CT_TblPrEx)
register_element_cls("w:tblStyle", CT_String)
register_element_cls("w:tc", CT_Tc)
register_element_cls("w:tcPr", CT_TcPr)
register_element_cls("w:tcW", CT_TblWidth)
register_element_cls("w:tr", CT_Row)
register_element_cls("w:trHeight", CT_Height)
register_element_cls("w:trPr", CT_TrPr)
register_element_cls("w:vAlign", CT_VerticalJc)
register_element_cls("w:vMerge", CT_VMerge)

from .text.font import (  # noqa
    CT_Color,
    CT_Fonts,
    CT_Highlight,
    CT_HpsMeasure,
    CT_RPr,
    CT_Underline,
    CT_VerticalAlignRun,
)

register_element_cls("w:b", CT_OnOff)
register_element_cls("w:bCs", CT_OnOff)
register_element_cls("w:caps", CT_OnOff)
register_element_cls("w:color", CT_Color)
register_element_cls("w:cs", CT_OnOff)
register_element_cls("w:dstrike", CT_OnOff)
register_element_cls("w:emboss", CT_OnOff)
register_element_cls("w:highlight", CT_Highlight)
register_element_cls("w:i", CT_OnOff)
register_element_cls("w:iCs", CT_OnOff)
register_element_cls("w:imprint", CT_OnOff)
register_element_cls("w:noProof", CT_OnOff)
register_element_cls("w:oMath", CT_OnOff)
register_element_cls("w:outline", CT_OnOff)
register_element_cls("w:rFonts", CT_Fonts)
register_element_cls("w:rPr", CT_RPr)
register_element_cls("w:rStyle", CT_String)
register_element_cls("w:rtl", CT_OnOff)
register_element_cls("w:shadow", CT_OnOff)
register_element_cls("w:smallCaps", CT_OnOff)
register_element_cls("w:snapToGrid", CT_OnOff)
register_element_cls("w:specVanish", CT_OnOff)
register_element_cls("w:strike", CT_OnOff)
register_element_cls("w:sz", CT_HpsMeasure)
register_element_cls("w:u", CT_Underline)
register_element_cls("w:vanish", CT_OnOff)
register_element_cls("w:vertAlign", CT_VerticalAlignRun)
register_element_cls("w:webHidden", CT_OnOff)

from .text.paragraph import CT_P  # noqa

register_element_cls("w:p", CT_P)

from .text.parfmt import (  # noqa
    CT_Ind,
    CT_Jc,
    CT_PPr,
    CT_Spacing,
    CT_TabStop,
    CT_TabStops,
)

register_element_cls("w:ind", CT_Ind)
register_element_cls("w:jc", CT_Jc)
register_element_cls("w:keepLines", CT_OnOff)
register_element_cls("w:keepNext", CT_OnOff)
register_element_cls("w:pageBreakBefore", CT_OnOff)
register_element_cls("w:pPr", CT_PPr)
register_element_cls("w:pStyle", CT_String)
register_element_cls("w:spacing", CT_Spacing)
register_element_cls("w:tab", CT_TabStop)
register_element_cls("w:tabs", CT_TabStops)
register_element_cls("w:widowControl", CT_OnOff)


# package.py
"""WordprocessingML Package class and related objects."""

from __future__ import annotations

from typing import IO, cast

from docx.image.image import Image
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.package import OpcPackage
from docx.opc.packuri import PackURI
from docx.parts.image import ImagePart
from docx.shared import lazyproperty


class Package(OpcPackage):
    """Customizations specific to a WordprocessingML package."""

    def after_unmarshal(self):
        """Called by loading code after all parts and relationships have been loaded.

        This method affords the opportunity for any required post-processing.
        """
        self._gather_image_parts()

    def get_or_add_image_part(self, image_descriptor: str | IO[bytes]) -> ImagePart:
        """Return |ImagePart| containing image specified by `image_descriptor`.

        The image-part is newly created if a matching one is not already present in the
        collection.
        """
        return self.image_parts.get_or_add_image_part(image_descriptor)

    @lazyproperty
    def image_parts(self) -> ImageParts:
        """|ImageParts| collection object for this package."""
        return ImageParts()

    def _gather_image_parts(self):
        """Load the image part collection with all the image parts in package."""
        for rel in self.iter_rels():
            if rel.is_external:
                continue
            if rel.reltype != RT.IMAGE:
                continue
            if rel.target_part in self.image_parts:
                continue
            self.image_parts.append(cast("ImagePart", rel.target_part))


class ImageParts:
    """Collection of |ImagePart| objects corresponding to images in the package."""

    def __init__(self):
        self._image_parts: list[ImagePart] = []

    def __contains__(self, item: object):
        return self._image_parts.__contains__(item)

    def __iter__(self):
        return self._image_parts.__iter__()

    def __len__(self):
        return self._image_parts.__len__()

    def append(self, item: ImagePart):
        self._image_parts.append(item)

    def get_or_add_image_part(self, image_descriptor: str | IO[bytes]) -> ImagePart:
        """Return |ImagePart| object containing image identified by `image_descriptor`.

        The image-part is newly created if a matching one is not present in the
        collection.
        """
        image = Image.from_file(image_descriptor)
        matching_image_part = self._get_by_sha1(image.sha1)
        if matching_image_part is not None:
            return matching_image_part
        return self._add_image_part(image)

    def _add_image_part(self, image: Image):
        """Return |ImagePart| instance newly created from `image` and appended to the collection."""
        partname = self._next_image_partname(image.ext)
        image_part = ImagePart.from_image(image, partname)
        self.append(image_part)
        return image_part

    def _get_by_sha1(self, sha1: str) -> ImagePart | None:
        """Return the image part in this collection having a SHA1 hash matching `sha1`,
        or |None| if not found."""
        for image_part in self._image_parts:
            if image_part.sha1 == sha1:
                return image_part
        return None

    def _next_image_partname(self, ext: str) -> PackURI:
        """The next available image partname, starting from ``/word/media/image1.{ext}``
        where unused numbers are reused.

        The partname is unique by number, without regard to the extension. `ext` does
        not include the leading period.
        """

        def image_partname(n: int) -> PackURI:
            return PackURI("/word/media/image%d.%s" % (n, ext))

        used_numbers = [image_part.partname.idx for image_part in self]
        for n in range(1, len(self) + 1):
            if n not in used_numbers:
                return image_partname(n)
        return image_partname(len(self) + 1)


# document.py
"""|DocumentPart| and closely related objects."""

from __future__ import annotations

from typing import IO, TYPE_CHECKING, cast

from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.parts.numbering import NumberingPart
from docx.parts.settings import SettingsPart
from docx.parts.story import StoryPart
from docx.parts.styles import StylesPart
from docx.shape import InlineShapes
from docx.shared import lazyproperty

if TYPE_CHECKING:
    from docx.opc.coreprops import CoreProperties
    from docx.settings import Settings
    from docx.styles.style import BaseStyle


class DocumentPart(StoryPart):
    """Main document part of a WordprocessingML (WML) package, aka a .docx file.

    Acts as broker to other parts such as image, core properties, and style parts. It
    also acts as a convenient delegate when a mid-document object needs a service
    involving a remote ancestor. The `Parented.part` property inherited by many content
    objects provides access to this part object for that purpose.
    """

    def add_footer_part(self):
        """Return (footer_part, rId) pair for newly-created footer part."""
        footer_part = FooterPart.new(self.package)
        rId = self.relate_to(footer_part, RT.FOOTER)
        return footer_part, rId

    def add_header_part(self):
        """Return (header_part, rId) pair for newly-created header part."""
        header_part = HeaderPart.new(self.package)
        rId = self.relate_to(header_part, RT.HEADER)
        return header_part, rId

    @property
    def core_properties(self) -> CoreProperties:
        """A |CoreProperties| object providing read/write access to the core properties
        of this document."""
        return self.package.core_properties

    @property
    def document(self):
        """A |Document| object providing access to the content of this document."""
        return Document(self._element, self)

    def drop_header_part(self, rId: str) -> None:
        """Remove related header part identified by `rId`."""
        self.drop_rel(rId)

    def footer_part(self, rId: str):
        """Return |FooterPart| related by `rId`."""
        return self.related_parts[rId]

    def get_style(self, style_id: str | None, style_type: WD_STYLE_TYPE) -> BaseStyle:
        """Return the style in this document matching `style_id`.

        Returns the default style for `style_type` if `style_id` is |None| or does not
        match a defined style of `style_type`.
        """
        return self.styles.get_by_id(style_id, style_type)

    def get_style_id(self, style_or_name, style_type):
        """Return the style_id (|str|) of the style of `style_type` matching
        `style_or_name`.

        Returns |None| if the style resolves to the default style for `style_type` or if
        `style_or_name` is itself |None|. Raises if `style_or_name` is a style of the
        wrong type or names a style not present in the document.
        """
        return self.styles.get_style_id(style_or_name, style_type)

    def header_part(self, rId: str):
        """Return |HeaderPart| related by `rId`."""
        return self.related_parts[rId]

    @lazyproperty
    def inline_shapes(self):
        """The |InlineShapes| instance containing the inline shapes in the document."""
        return InlineShapes(self._element.body, self)

    @lazyproperty
    def numbering_part(self):
        """A |NumberingPart| object providing access to the numbering definitions for
        this document.

        Creates an empty numbering part if one is not present.
        """
        try:
            return self.part_related_by(RT.NUMBERING)
        except KeyError:
            numbering_part = NumberingPart.new()
            self.relate_to(numbering_part, RT.NUMBERING)
            return numbering_part

    def save(self, path_or_stream: str | IO[bytes]):
        """Save this document to `path_or_stream`, which can be either a path to a
        filesystem location (a string) or a file-like object."""
        self.package.save(path_or_stream)

    @property
    def settings(self) -> Settings:
        """A |Settings| object providing access to the settings in the settings part of
        this document."""
        return self._settings_part.settings

    @property
    def styles(self):
        """A |Styles| object providing access to the styles in the styles part of this
        document."""
        return self._styles_part.styles

    @property
    def _settings_part(self) -> SettingsPart:
        """A |SettingsPart| object providing access to the document-level settings for
        this document.

        Creates a default settings part if one is not present.
        """
        try:
            return cast(SettingsPart, self.part_related_by(RT.SETTINGS))
        except KeyError:
            settings_part = SettingsPart.default(self.package)
            self.relate_to(settings_part, RT.SETTINGS)
            return settings_part

    @property
    def _styles_part(self) -> StylesPart:
        """Instance of |StylesPart| for this document.

        Creates an empty styles part if one is not present.
        """
        try:
            return cast(StylesPart, self.part_related_by(RT.STYLES))
        except KeyError:
            package = self.package
            assert package is not None
            styles_part = StylesPart.default(package)
            self.relate_to(styles_part, RT.STYLES)
            return styles_part


# hdrftr.py
"""Header and footer part objects."""

from __future__ import annotations

import os
from typing import TYPE_CHECKING

from docx.opc.constants import CONTENT_TYPE as CT
from docx.oxml.parser import parse_xml
from docx.parts.story import StoryPart

if TYPE_CHECKING:
    from docx.package import Package


class FooterPart(StoryPart):
    """Definition of a section footer."""

    @classmethod
    def new(cls, package: Package):
        """Return newly created footer part."""
        partname = package.next_partname("/word/footer%d.xml")
        content_type = CT.WML_FOOTER
        element = parse_xml(cls._default_footer_xml())
        return cls(partname, content_type, element, package)

    @classmethod
    def _default_footer_xml(cls):
        """Return bytes containing XML for a default footer part."""
        path = os.path.join(os.path.split(__file__)[0], "..", "templates", "default-footer.xml")
        with open(path, "rb") as f:
            xml_bytes = f.read()
        return xml_bytes


class HeaderPart(StoryPart):
    """Definition of a section header."""

    @classmethod
    def new(cls, package: Package):
        """Return newly created header part."""
        partname = package.next_partname("/word/header%d.xml")
        content_type = CT.WML_HEADER
        element = parse_xml(cls._default_header_xml())
        return cls(partname, content_type, element, package)

    @classmethod
    def _default_header_xml(cls):
        """Return bytes containing XML for a default header part."""
        path = os.path.join(os.path.split(__file__)[0], "..", "templates", "default-header.xml")
        with open(path, "rb") as f:
            xml_bytes = f.read()
        return xml_bytes


# image.py
"""The proxy class for an image part, and related objects."""

from __future__ import annotations

import hashlib
from typing import TYPE_CHECKING

from docx.image.image import Image
from docx.opc.part import Part
from docx.shared import Emu, Inches

if TYPE_CHECKING:
    from docx.opc.package import OpcPackage
    from docx.opc.packuri import PackURI


class ImagePart(Part):
    """An image part.

    Corresponds to the target part of a relationship with type RELATIONSHIP_TYPE.IMAGE.
    """

    def __init__(
        self, partname: PackURI, content_type: str, blob: bytes, image: Image | None = None
    ):
        super(ImagePart, self).__init__(partname, content_type, blob)
        self._image = image

    @property
    def default_cx(self):
        """Native width of this image, calculated from its width in pixels and
        horizontal dots per inch (dpi)."""
        px_width = self.image.px_width
        horz_dpi = self.image.horz_dpi
        width_in_inches = px_width / horz_dpi
        return Inches(width_in_inches)

    @property
    def default_cy(self):
        """Native height of this image, calculated from its height in pixels and
        vertical dots per inch (dpi)."""
        px_height = self.image.px_height
        horz_dpi = self.image.horz_dpi
        height_in_emu = int(round(914400 * px_height / horz_dpi))
        return Emu(height_in_emu)

    @property
    def filename(self):
        """Filename from which this image part was originally created.

        A generic name, e.g. 'image.png', is substituted if no name is available, for
        example when the image was loaded from an unnamed stream. In that case a default
        extension is applied based on the detected MIME type of the image.
        """
        if self._image is not None:
            return self._image.filename
        return "image.%s" % self.partname.ext

    @classmethod
    def from_image(cls, image: Image, partname: PackURI):
        """Return an |ImagePart| instance newly created from `image` and assigned
        `partname`."""
        return ImagePart(partname, image.content_type, image.blob, image)

    @property
    def image(self) -> Image:
        if self._image is None:
            self._image = Image.from_blob(self.blob)
        return self._image

    @classmethod
    def load(cls, partname: PackURI, content_type: str, blob: bytes, package: OpcPackage):
        """Called by ``docx.opc.package.PartFactory`` to load an image part from a
        package being opened by ``Document(...)`` call."""
        return cls(partname, content_type, blob)

    @property
    def sha1(self):
        """SHA1 hash digest of the blob of this image part."""
        return hashlib.sha1(self.blob).hexdigest()


# numbering.py
"""|NumberingPart| and closely related objects."""

from ..opc.part import XmlPart
from ..shared import lazyproperty


class NumberingPart(XmlPart):
    """Proxy for the numbering.xml part containing numbering definitions for a document
    or glossary."""

    @classmethod
    def new(cls):
        """Return newly created empty numbering part, containing only the root
        ``<w:numbering>`` element."""
        raise NotImplementedError

    @lazyproperty
    def numbering_definitions(self):
        """The |_NumberingDefinitions| instance containing the numbering definitions
        (<w:num> element proxies) for this numbering part."""
        return _NumberingDefinitions(self._element)


class _NumberingDefinitions:
    """Collection of |_NumberingDefinition| instances corresponding to the ``<w:num>``
    elements in a numbering part."""

    def __init__(self, numbering_elm):
        super(_NumberingDefinitions, self).__init__()
        self._numbering = numbering_elm

    def __len__(self):
        return len(self._numbering.num_lst)


# settings.py
"""|SettingsPart| and closely related objects."""

from __future__ import annotations

import os
from typing import TYPE_CHECKING, cast

from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml.parser import parse_xml
from docx.settings import Settings

if TYPE_CHECKING:
    from docx.oxml.settings import CT_Settings
    from docx.package import Package


class SettingsPart(XmlPart):
    """Document-level settings part of a WordprocessingML (WML) package."""

    def __init__(
        self, partname: PackURI, content_type: str, element: CT_Settings, package: Package
    ):
        super().__init__(partname, content_type, element, package)
        self._settings = element

    @classmethod
    def default(cls, package: Package):
        """Return a newly created settings part, containing a default `w:settings`
        element tree."""
        partname = PackURI("/word/settings.xml")
        content_type = CT.WML_SETTINGS
        element = cast("CT_Settings", parse_xml(cls._default_settings_xml()))
        return cls(partname, content_type, element, package)

    @property
    def settings(self) -> Settings:
        """A |Settings| proxy object for the `w:settings` element in this part.

        Contains the document-level settings for this document.
        """
        return Settings(self._settings)

    @classmethod
    def _default_settings_xml(cls):
        """Return a bytestream containing XML for a default settings part."""
        path = os.path.join(os.path.split(__file__)[0], "..", "templates", "default-settings.xml")
        with open(path, "rb") as f:
            xml_bytes = f.read()
        return xml_bytes


# story.py
"""|StoryPart| and related objects."""

from __future__ import annotations

from typing import IO, TYPE_CHECKING, Tuple, cast

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import XmlPart
from docx.oxml.shape import CT_Inline
from docx.shared import Length, lazyproperty

if TYPE_CHECKING:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.image.image import Image
    from docx.parts.document import DocumentPart
    from docx.styles.style import BaseStyle


class StoryPart(XmlPart):
    """Base class for story parts.

    A story part is one that can contain textual content, such as the document-part and
    header or footer parts. These all share content behaviors like `.paragraphs`,
    `.add_paragraph()`, `.add_table()` etc.
    """

    def get_or_add_image(self, image_descriptor: str | IO[bytes]) -> Tuple[str, Image]:
        """Return (rId, image) pair for image identified by `image_descriptor`.

        `rId` is the str key (often like "rId7") for the relationship between this story
        part and the image part, reused if already present, newly created if not.
        `image` is an |Image| instance providing access to the properties of the image,
        such as dimensions and image type.
        """
        package = self._package
        assert package is not None
        image_part = package.get_or_add_image_part(image_descriptor)
        rId = self.relate_to(image_part, RT.IMAGE)
        return rId, image_part.image

    def get_style(self, style_id: str | None, style_type: WD_STYLE_TYPE) -> BaseStyle:
        """Return the style in this document matching `style_id`.

        Returns the default style for `style_type` if `style_id` is |None| or does not
        match a defined style of `style_type`.
        """
        return self._document_part.get_style(style_id, style_type)

    def get_style_id(
        self, style_or_name: BaseStyle | str | None, style_type: WD_STYLE_TYPE
    ) -> str | None:
        """Return str style_id for `style_or_name` of `style_type`.

        Returns |None| if the style resolves to the default style for `style_type` or if
        `style_or_name` is itself |None|. Raises if `style_or_name` is a style of the
        wrong type or names a style not present in the document.
        """
        return self._document_part.get_style_id(style_or_name, style_type)

    def new_pic_inline(
        self,
        image_descriptor: str | IO[bytes],
        width: int | Length | None = None,
        height: int | Length | None = None,
    ) -> CT_Inline:
        """Return a newly-created `w:inline` element.

        The element contains the image specified by `image_descriptor` and is scaled
        based on the values of `width` and `height`.
        """
        rId, image = self.get_or_add_image(image_descriptor)
        cx, cy = image.scaled_dimensions(width, height)
        shape_id, filename = self.next_id, image.filename
        return CT_Inline.new_pic_inline(shape_id, rId, filename, cx, cy)

    @property
    def next_id(self) -> int:
        """Next available positive integer id value in this story XML document.

        The value is determined by incrementing the maximum existing id value. Gaps in
        the existing id sequence are not filled. The id attribute value is unique in the
        document, without regard to the element type it appears on.
        """
        id_str_lst = self._element.xpath("//@id")
        used_ids = [int(id_str) for id_str in id_str_lst if id_str.isdigit()]
        if not used_ids:
            return 1
        return max(used_ids) + 1

    @lazyproperty
    def _document_part(self) -> DocumentPart:
        """|DocumentPart| object for this package."""
        package = self.package
        assert package is not None
        return cast("DocumentPart", package.main_document_part)


# styles.py
"""Provides StylesPart and related objects."""

from __future__ import annotations

import os
from typing import TYPE_CHECKING

from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml.parser import parse_xml
from docx.styles.styles import Styles

if TYPE_CHECKING:
    from docx.opc.package import OpcPackage


class StylesPart(XmlPart):
    """Proxy for the styles.xml part containing style definitions for a document or
    glossary."""

    @classmethod
    def default(cls, package: OpcPackage) -> StylesPart:
        """Return a newly created styles part, containing a default set of elements."""
        partname = PackURI("/word/styles.xml")
        content_type = CT.WML_STYLES
        element = parse_xml(cls._default_styles_xml())
        return cls(partname, content_type, element, package)

    @property
    def styles(self):
        """The |_Styles| instance containing the styles (<w:style> element proxies) for
        this styles part."""
        return Styles(self.element)

    @classmethod
    def _default_styles_xml(cls):
        """Return a bytestream containing XML for a default styles part."""
        path = os.path.join(
            os.path.split(__file__)[0], "..", "templates", "default-styles.xml"
        )
        with open(path, "rb") as f:
            xml_bytes = f.read()
        return xml_bytes


# __init__.py


# section.py
"""The |Section| object and related proxy classes."""

from __future__ import annotations

from typing import TYPE_CHECKING, Iterator, List, Sequence, overload

from docx.blkcntnr import BlockItemContainer
from docx.enum.section import WD_HEADER_FOOTER
from docx.oxml.text.paragraph import CT_P
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.shared import lazyproperty
from docx.table import Table
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
    from docx.oxml.document import CT_Document
    from docx.oxml.section import CT_SectPr
    from docx.parts.document import DocumentPart
    from docx.parts.story import StoryPart
    from docx.shared import Length


class Section:
    """Document section, providing access to section and page setup settings.

    Also provides access to headers and footers.
    """

    def __init__(self, sectPr: CT_SectPr, document_part: DocumentPart):
        super(Section, self).__init__()
        self._sectPr = sectPr
        self._document_part = document_part

    @property
    def bottom_margin(self) -> Length | None:
        """Read/write. Bottom margin for pages in this section, in EMU.

        `None` when no bottom margin has been specified. Assigning |None| removes any
        bottom-margin setting.
        """
        return self._sectPr.bottom_margin

    @bottom_margin.setter
    def bottom_margin(self, value: int | Length | None):
        self._sectPr.bottom_margin = value

    @property
    def different_first_page_header_footer(self) -> bool:
        """True if this section displays a distinct first-page header and footer.

        Read/write. The definition of the first-page header and footer are accessed
        using :attr:`.first_page_header` and :attr:`.first_page_footer` respectively.
        """
        return self._sectPr.titlePg_val

    @different_first_page_header_footer.setter
    def different_first_page_header_footer(self, value: bool):
        self._sectPr.titlePg_val = value

    @property
    def even_page_footer(self) -> _Footer:
        """|_Footer| object defining footer content for even pages.

        The content of this footer definition is ignored unless the document setting
        :attr:`~.Settings.odd_and_even_pages_header_footer` is set True.
        """
        return _Footer(self._sectPr, self._document_part, WD_HEADER_FOOTER.EVEN_PAGE)

    @property
    def even_page_header(self) -> _Header:
        """|_Header| object defining header content for even pages.

        The content of this header definition is ignored unless the document setting
        :attr:`~.Settings.odd_and_even_pages_header_footer` is set True.
        """
        return _Header(self._sectPr, self._document_part, WD_HEADER_FOOTER.EVEN_PAGE)

    @property
    def first_page_footer(self) -> _Footer:
        """|_Footer| object defining footer content for the first page of this section.

        The content of this footer definition is ignored unless the property
        :attr:`.different_first_page_header_footer` is set True.
        """
        return _Footer(self._sectPr, self._document_part, WD_HEADER_FOOTER.FIRST_PAGE)

    @property
    def first_page_header(self) -> _Header:
        """|_Header| object defining header content for the first page of this section.

        The content of this header definition is ignored unless the property
        :attr:`.different_first_page_header_footer` is set True.
        """
        return _Header(self._sectPr, self._document_part, WD_HEADER_FOOTER.FIRST_PAGE)

    @lazyproperty
    def footer(self) -> _Footer:
        """|_Footer| object representing default page footer for this section.

        The default footer is used for odd-numbered pages when separate odd/even footers
        are enabled. It is used for both odd and even-numbered pages otherwise.
        """
        return _Footer(self._sectPr, self._document_part, WD_HEADER_FOOTER.PRIMARY)

    @property
    def footer_distance(self) -> Length | None:
        """Distance from bottom edge of page to bottom edge of the footer.

        Read/write. |None| if no setting is present in the XML.
        """
        return self._sectPr.footer

    @footer_distance.setter
    def footer_distance(self, value: int | Length | None):
        self._sectPr.footer = value

    @property
    def gutter(self) -> Length | None:
        """|Length| object representing page gutter size in English Metric Units.

        Read/write. The page gutter is extra spacing added to the `inner` margin to
        ensure even margins after page binding. Generally only used in book-bound
        documents with double-sided and facing pages.

        This setting applies to all pages in this section.

        """
        return self._sectPr.gutter

    @gutter.setter
    def gutter(self, value: int | Length | None):
        self._sectPr.gutter = value

    @lazyproperty
    def header(self) -> _Header:
        """|_Header| object representing default page header for this section.

        The default header is used for odd-numbered pages when separate odd/even headers
        are enabled. It is used for both odd and even-numbered pages otherwise.
        """
        return _Header(self._sectPr, self._document_part, WD_HEADER_FOOTER.PRIMARY)

    @property
    def header_distance(self) -> Length | None:
        """Distance from top edge of page to top edge of header.

        Read/write. |None| if no setting is present in the XML. Assigning |None| causes
        default value to be used.
        """
        return self._sectPr.header

    @header_distance.setter
    def header_distance(self, value: int | Length | None):
        self._sectPr.header = value

    def iter_inner_content(self) -> Iterator[Paragraph | Table]:
        """Generate each Paragraph or Table object in this `section`.

        Items appear in document order.
        """
        for element in self._sectPr.iter_inner_content():
            yield (Paragraph(element, self) if isinstance(element, CT_P) else Table(element, self))

    @property
    def left_margin(self) -> Length | None:
        """|Length| object representing the left margin for all pages in this section in
        English Metric Units."""
        return self._sectPr.left_margin

    @left_margin.setter
    def left_margin(self, value: int | Length | None):
        self._sectPr.left_margin = value

    @property
    def orientation(self) -> WD_ORIENTATION:
        """:ref:`WdOrientation` member specifying page orientation for this section.

        One of ``WD_ORIENT.PORTRAIT`` or ``WD_ORIENT.LANDSCAPE``.
        """
        return self._sectPr.orientation

    @orientation.setter
    def orientation(self, value: WD_ORIENTATION | None):
        self._sectPr.orientation = value

    @property
    def page_height(self) -> Length | None:
        """Total page height used for this section.

        This value is inclusive of all edge spacing values such as margins.

        Page orientation is taken into account, so for example, its expected value
        would be ``Inches(8.5)`` for letter-sized paper when orientation is landscape.
        """
        return self._sectPr.page_height

    @page_height.setter
    def page_height(self, value: Length | None):
        self._sectPr.page_height = value

    @property
    def page_width(self) -> Length | None:
        """Total page width used for this section.

        This value is like "paper size" and includes all edge spacing values such as
        margins.

        Page orientation is taken into account, so for example, its expected value
        would be ``Inches(11)`` for letter-sized paper when orientation is landscape.
        """
        return self._sectPr.page_width

    @page_width.setter
    def page_width(self, value: Length | None):
        self._sectPr.page_width = value

    @property
    def part(self) -> StoryPart:
        return self._document_part

    @property
    def right_margin(self) -> Length | None:
        """|Length| object representing the right margin for all pages in this section
        in English Metric Units."""
        return self._sectPr.right_margin

    @right_margin.setter
    def right_margin(self, value: Length | None):
        self._sectPr.right_margin = value

    @property
    def start_type(self) -> WD_SECTION_START:
        """Type of page-break (if any) inserted at the start of this section.

        For exmple, ``WD_SECTION_START.ODD_PAGE`` if the section should begin on the
        next odd page, possibly inserting two page-breaks instead of one.
        """
        return self._sectPr.start_type

    @start_type.setter
    def start_type(self, value: WD_SECTION_START | None):
        self._sectPr.start_type = value

    @property
    def top_margin(self) -> Length | None:
        """|Length| object representing the top margin for all pages in this section in
        English Metric Units."""
        return self._sectPr.top_margin

    @top_margin.setter
    def top_margin(self, value: Length | None):
        self._sectPr.top_margin = value


class Sections(Sequence[Section]):
    """Sequence of |Section| objects corresponding to the sections in the document.

    Supports ``len()``, iteration, and indexed access.
    """

    def __init__(self, document_elm: CT_Document, document_part: DocumentPart):
        super(Sections, self).__init__()
        self._document_elm = document_elm
        self._document_part = document_part

    @overload
    def __getitem__(self, key: int) -> Section: ...

    @overload
    def __getitem__(self, key: slice) -> List[Section]: ...

    def __getitem__(self, key: int | slice) -> Section | List[Section]:
        if isinstance(key, slice):
            return [
                Section(sectPr, self._document_part)
                for sectPr in self._document_elm.sectPr_lst[key]
            ]
        return Section(self._document_elm.sectPr_lst[key], self._document_part)

    def __iter__(self) -> Iterator[Section]:
        for sectPr in self._document_elm.sectPr_lst:
            yield Section(sectPr, self._document_part)

    def __len__(self) -> int:
        return len(self._document_elm.sectPr_lst)


class _BaseHeaderFooter(BlockItemContainer):
    """Base class for header and footer classes."""

    def __init__(
        self,
        sectPr: CT_SectPr,
        document_part: DocumentPart,
        header_footer_index: WD_HEADER_FOOTER,
    ):
        self._sectPr = sectPr
        self._document_part = document_part
        self._hdrftr_index = header_footer_index

    @property
    def is_linked_to_previous(self) -> bool:
        """``True`` if this header/footer uses the definition from the prior section.

        ``False`` if this header/footer has an explicit definition.

        Assigning ``True`` to this property removes the header/footer definition for
        this section, causing it to "inherit" the corresponding definition of the prior
        section. Assigning ``False`` causes a new, empty definition to be added for this
        section, but only if no definition is already present.
        """
        # ---absence of a header/footer part indicates "linked" behavior---
        return not self._has_definition

    @is_linked_to_previous.setter
    def is_linked_to_previous(self, value: bool) -> None:
        new_state = bool(value)
        # ---do nothing when value is not being changed---
        if new_state == self.is_linked_to_previous:
            return
        if new_state is True:
            self._drop_definition()
        else:
            self._add_definition()

    @property
    def part(self) -> HeaderPart | FooterPart:
        """The |HeaderPart| or |FooterPart| for this header/footer.

        This overrides `BlockItemContainer.part` and is required to support image
        insertion and perhaps other content like hyperlinks.
        """
        # ---should not appear in documentation;
        # ---not an interface property, even though public
        return self._get_or_add_definition()

    def _add_definition(self) -> HeaderPart | FooterPart:
        """Return newly-added header/footer part."""
        raise NotImplementedError("must be implemented by each subclass")

    @property
    def _definition(self) -> HeaderPart | FooterPart:
        """|HeaderPart| or |FooterPart| object containing header/footer content."""
        raise NotImplementedError("must be implemented by each subclass")

    def _drop_definition(self) -> None:
        """Remove header/footer part containing the definition of this header/footer."""
        raise NotImplementedError("must be implemented by each subclass")

    @property
    def _element(self):
        """`w:hdr` or `w:ftr` element, root of header/footer part."""
        return self._get_or_add_definition().element

    def _get_or_add_definition(self) -> HeaderPart | FooterPart:
        """Return HeaderPart or FooterPart object for this section.

        If this header/footer inherits its content, the part for the prior header/footer
        is returned; this process continue recursively until a definition is found. If
        the definition cannot be inherited (because the header/footer belongs to the
        first section), a new definition is added for that first section and then
        returned.
        """
        # ---note this method is called recursively to access inherited definitions---
        # ---case-1: definition is not inherited---
        if self._has_definition:
            return self._definition
        # ---case-2: definition is inherited and belongs to second-or-later section---
        prior_headerfooter = self._prior_headerfooter
        if prior_headerfooter:
            return prior_headerfooter._get_or_add_definition()
        # ---case-3: definition is inherited, but belongs to first section---
        return self._add_definition()

    @property
    def _has_definition(self) -> bool:
        """True if this header/footer has a related part containing its definition."""
        raise NotImplementedError("must be implemented by each subclass")

    @property
    def _prior_headerfooter(self) -> _Header | _Footer | None:
        """|_Header| or |_Footer| proxy on prior sectPr element.

        Returns None if this is first section.
        """
        raise NotImplementedError("must be implemented by each subclass")


class _Footer(_BaseHeaderFooter):
    """Page footer, used for all three types (default, even-page, and first-page).

    Note that, like a document or table cell, a footer must contain a minimum of one
    paragraph and a new or otherwise "empty" footer contains a single empty paragraph.
    This first paragraph can be accessed as `footer.paragraphs[0]` for purposes of
    adding content to it. Using :meth:`add_paragraph()` by itself to add content will
    leave an empty paragraph above the newly added one.
    """

    def _add_definition(self) -> FooterPart:
        """Return newly-added footer part."""
        footer_part, rId = self._document_part.add_footer_part()
        self._sectPr.add_footerReference(self._hdrftr_index, rId)
        return footer_part

    @property
    def _definition(self):
        """|FooterPart| object containing content of this footer."""
        footerReference = self._sectPr.get_footerReference(self._hdrftr_index)
        # -- currently this is never called when `._has_definition` evaluates False --
        assert footerReference is not None
        return self._document_part.footer_part(footerReference.rId)

    def _drop_definition(self):
        """Remove footer definition (footer part) associated with this section."""
        rId = self._sectPr.remove_footerReference(self._hdrftr_index)
        self._document_part.drop_rel(rId)

    @property
    def _has_definition(self) -> bool:
        """True if a footer is defined for this section."""
        footerReference = self._sectPr.get_footerReference(self._hdrftr_index)
        return footerReference is not None

    @property
    def _prior_headerfooter(self):
        """|_Footer| proxy on prior sectPr element or None if this is first section."""
        preceding_sectPr = self._sectPr.preceding_sectPr
        return (
            None
            if preceding_sectPr is None
            else _Footer(preceding_sectPr, self._document_part, self._hdrftr_index)
        )


class _Header(_BaseHeaderFooter):
    """Page header, used for all three types (default, even-page, and first-page).

    Note that, like a document or table cell, a header must contain a minimum of one
    paragraph and a new or otherwise "empty" header contains a single empty paragraph.
    This first paragraph can be accessed as `header.paragraphs[0]` for purposes of
    adding content to it. Using :meth:`add_paragraph()` by itself to add content will
    leave an empty paragraph above the newly added one.
    """

    def _add_definition(self):
        """Return newly-added header part."""
        header_part, rId = self._document_part.add_header_part()
        self._sectPr.add_headerReference(self._hdrftr_index, rId)
        return header_part

    @property
    def _definition(self):
        """|HeaderPart| object containing content of this header."""
        headerReference = self._sectPr.get_headerReference(self._hdrftr_index)
        # -- currently this is never called when `._has_definition` evaluates False --
        assert headerReference is not None
        return self._document_part.header_part(headerReference.rId)

    def _drop_definition(self):
        """Remove header definition associated with this section."""
        rId = self._sectPr.remove_headerReference(self._hdrftr_index)
        self._document_part.drop_header_part(rId)

    @property
    def _has_definition(self) -> bool:
        """True if a header is explicitly defined for this section."""
        headerReference = self._sectPr.get_headerReference(self._hdrftr_index)
        return headerReference is not None

    @property
    def _prior_headerfooter(self):
        """|_Header| proxy on prior sectPr element or None if this is first section."""
        preceding_sectPr = self._sectPr.preceding_sectPr
        return (
            None
            if preceding_sectPr is None
            else _Header(preceding_sectPr, self._document_part, self._hdrftr_index)
        )


# settings.py
"""Settings object, providing access to document-level settings."""

from __future__ import annotations

from typing import TYPE_CHECKING, cast

from docx.shared import ElementProxy

if TYPE_CHECKING:
    import docx.types as t
    from docx.oxml.settings import CT_Settings
    from docx.oxml.xmlchemy import BaseOxmlElement


class Settings(ElementProxy):
    """Provides access to document-level settings for a document.

    Accessed using the :attr:`.Document.settings` property.
    """

    def __init__(self, element: BaseOxmlElement, parent: t.ProvidesXmlPart | None = None):
        super().__init__(element, parent)
        self._settings = cast("CT_Settings", element)

    @property
    def odd_and_even_pages_header_footer(self) -> bool:
        """True if this document has distinct odd and even page headers and footers.

        Read/write.
        """
        return self._settings.evenAndOddHeaders_val

    @odd_and_even_pages_header_footer.setter
    def odd_and_even_pages_header_footer(self, value: bool):
        self._settings.evenAndOddHeaders_val = value


# shape.py
"""Objects related to shapes.

A shape is a visual object that appears on the drawing layer of a document.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml.ns import nsmap
from docx.shared import Parented

if TYPE_CHECKING:
    from docx.oxml.document import CT_Body
    from docx.oxml.shape import CT_Inline
    from docx.parts.story import StoryPart
    from docx.shared import Length


class InlineShapes(Parented):
    """Sequence of |InlineShape| instances, supporting len(), iteration, and indexed access."""

    def __init__(self, body_elm: CT_Body, parent: StoryPart):
        super(InlineShapes, self).__init__(parent)
        self._body = body_elm

    def __getitem__(self, idx: int):
        """Provide indexed access, e.g. 'inline_shapes[idx]'."""
        try:
            inline = self._inline_lst[idx]
        except IndexError:
            msg = "inline shape index [%d] out of range" % idx
            raise IndexError(msg)

        return InlineShape(inline)

    def __iter__(self):
        return (InlineShape(inline) for inline in self._inline_lst)

    def __len__(self):
        return len(self._inline_lst)

    @property
    def _inline_lst(self):
        body = self._body
        xpath = "//w:p/w:r/w:drawing/wp:inline"
        return body.xpath(xpath)


class InlineShape:
    """Proxy for an ``<wp:inline>`` element, representing the container for an inline
    graphical object."""

    def __init__(self, inline: CT_Inline):
        super(InlineShape, self).__init__()
        self._inline = inline

    @property
    def height(self) -> Length:
        """Read/write.

        The display height of this inline shape as an |Emu| instance.
        """
        return self._inline.extent.cy

    @height.setter
    def height(self, cy: Length):
        self._inline.extent.cy = cy
        self._inline.graphic.graphicData.pic.spPr.cy = cy

    @property
    def type(self):
        """The type of this inline shape as a member of
        ``docx.enum.shape.WD_INLINE_SHAPE``, e.g. ``LINKED_PICTURE``.

        Read-only.
        """
        graphicData = self._inline.graphic.graphicData
        uri = graphicData.uri
        if uri == nsmap["pic"]:
            blip = graphicData.pic.blipFill.blip
            if blip.link is not None:
                return WD_INLINE_SHAPE.LINKED_PICTURE
            return WD_INLINE_SHAPE.PICTURE
        if uri == nsmap["c"]:
            return WD_INLINE_SHAPE.CHART
        if uri == nsmap["dgm"]:
            return WD_INLINE_SHAPE.SMART_ART
        return WD_INLINE_SHAPE.NOT_IMPLEMENTED

    @property
    def width(self):
        """Read/write.

        The display width of this inline shape as an |Emu| instance.
        """
        return self._inline.extent.cx

    @width.setter
    def width(self, cx: Length):
        self._inline.extent.cx = cx
        self._inline.graphic.graphicData.pic.spPr.cx = cx


# shared.py
"""Objects shared by docx modules."""

from __future__ import annotations

import functools
from typing import (
    TYPE_CHECKING,
    Any,
    Callable,
    Generic,
    Iterator,
    List,
    Tuple,
    TypeVar,
    cast,
)

if TYPE_CHECKING:
    import docx.types as t
    from docx.opc.part import XmlPart
    from docx.oxml.xmlchemy import BaseOxmlElement
    from docx.parts.story import StoryPart


class Length(int):
    """Base class for length constructor classes Inches, Cm, Mm, Px, and Emu.

    Behaves as an int count of English Metric Units, 914,400 to the inch, 36,000 to the
    mm. Provides convenience unit conversion methods in the form of read-only
    properties. Immutable.
    """

    _EMUS_PER_INCH = 914400
    _EMUS_PER_CM = 360000
    _EMUS_PER_MM = 36000
    _EMUS_PER_PT = 12700
    _EMUS_PER_TWIP = 635

    def __new__(cls, emu: int):
        return int.__new__(cls, emu)

    @property
    def cm(self):
        """The equivalent length expressed in centimeters (float)."""
        return self / float(self._EMUS_PER_CM)

    @property
    def emu(self):
        """The equivalent length expressed in English Metric Units (int)."""
        return self

    @property
    def inches(self):
        """The equivalent length expressed in inches (float)."""
        return self / float(self._EMUS_PER_INCH)

    @property
    def mm(self):
        """The equivalent length expressed in millimeters (float)."""
        return self / float(self._EMUS_PER_MM)

    @property
    def pt(self):
        """Floating point length in points."""
        return self / float(self._EMUS_PER_PT)

    @property
    def twips(self):
        """The equivalent length expressed in twips (int)."""
        return int(round(self / float(self._EMUS_PER_TWIP)))


class Inches(Length):
    """Convenience constructor for length in inches, e.g. ``width = Inches(0.5)``."""

    def __new__(cls, inches: float):
        emu = int(inches * Length._EMUS_PER_INCH)
        return Length.__new__(cls, emu)


class Cm(Length):
    """Convenience constructor for length in centimeters, e.g. ``height = Cm(12)``."""

    def __new__(cls, cm: float):
        emu = int(cm * Length._EMUS_PER_CM)
        return Length.__new__(cls, emu)


class Emu(Length):
    """Convenience constructor for length in English Metric Units, e.g. ``width =
    Emu(457200)``."""

    def __new__(cls, emu: int):
        return Length.__new__(cls, int(emu))


class Mm(Length):
    """Convenience constructor for length in millimeters, e.g. ``width = Mm(240.5)``."""

    def __new__(cls, mm: float):
        emu = int(mm * Length._EMUS_PER_MM)
        return Length.__new__(cls, emu)


class Pt(Length):
    """Convenience value class for specifying a length in points."""

    def __new__(cls, points: float):
        emu = int(points * Length._EMUS_PER_PT)
        return Length.__new__(cls, emu)


class Twips(Length):
    """Convenience constructor for length in twips, e.g. ``width = Twips(42)``.

    A twip is a twentieth of a point, 635 EMU.
    """

    def __new__(cls, twips: float):
        emu = int(twips * Length._EMUS_PER_TWIP)
        return Length.__new__(cls, emu)


class RGBColor(Tuple[int, int, int]):
    """Immutable value object defining a particular RGB color."""

    def __new__(cls, r: int, g: int, b: int):
        msg = "RGBColor() takes three integer values 0-255"
        for val in (r, g, b):
            if (
                not isinstance(val, int)  # pyright: ignore[reportUnnecessaryIsInstance]
                or val < 0
                or val > 255
            ):
                raise ValueError(msg)
        return super(RGBColor, cls).__new__(cls, (r, g, b))

    def __repr__(self):
        return "RGBColor(0x%02x, 0x%02x, 0x%02x)" % self

    def __str__(self):
        """Return a hex string rgb value, like '3C2F80'."""
        return "%02X%02X%02X" % self

    @classmethod
    def from_string(cls, rgb_hex_str: str) -> RGBColor:
        """Return a new instance from an RGB color hex string like ``'3C2F80'``."""
        r = int(rgb_hex_str[:2], 16)
        g = int(rgb_hex_str[2:4], 16)
        b = int(rgb_hex_str[4:], 16)
        return cls(r, g, b)


T = TypeVar("T")


class lazyproperty(Generic[T]):
    """Decorator like @property, but evaluated only on first access.

    Like @property, this can only be used to decorate methods having only a `self`
    parameter, and is accessed like an attribute on an instance, i.e. trailing
    parentheses are not used. Unlike @property, the decorated method is only evaluated
    on first access; the resulting value is cached and that same value returned on
    second and later access without re-evaluation of the method.

    Like @property, this class produces a *data descriptor* object, which is stored in
    the __dict__ of the *class* under the name of the decorated method ('fget'
    nominally). The cached value is stored in the __dict__ of the *instance* under that
    same name.

    Because it is a data descriptor (as opposed to a *non-data descriptor*), its
    `__get__()` method is executed on each access of the decorated attribute; the
    __dict__ item of the same name is "shadowed" by the descriptor.

    While this may represent a performance improvement over a property, its greater
    benefit may be its other characteristics. One common use is to construct
    collaborator objects, removing that "real work" from the constructor, while still
    only executing once. It also de-couples client code from any sequencing
    considerations; if it's accessed from more than one location, it's assured it will
    be ready whenever needed.

    Loosely based on: https://stackoverflow.com/a/6849299/1902513.

    A lazyproperty is read-only. There is no counterpart to the optional "setter" (or
    deleter) behavior of an @property. This is critically important to maintaining its
    immutability and idempotence guarantees. Attempting to assign to a lazyproperty
    raises AttributeError unconditionally.

    The parameter names in the methods below correspond to this usage example::

        class Obj(object)

            @lazyproperty
            def fget(self):
                return 'some result'

        obj = Obj()

    Not suitable for wrapping a function (as opposed to a method) because it is not
    callable."""

    def __init__(self, fget: Callable[..., T]) -> None:
        """*fget* is the decorated method (a "getter" function).

        A lazyproperty is read-only, so there is only an *fget* function (a regular
        @property can also have an fset and fdel function). This name was chosen for
        consistency with Python's `property` class which uses this name for the
        corresponding parameter.
        """
        # --- maintain a reference to the wrapped getter method
        self._fget = fget
        # --- and store the name of that decorated method
        self._name = fget.__name__
        # --- adopt fget's __name__, __doc__, and other attributes
        functools.update_wrapper(self, fget)  # pyright: ignore

    def __get__(self, obj: Any, type: Any = None) -> T:
        """Called on each access of 'fget' attribute on class or instance.

        *self* is this instance of a lazyproperty descriptor "wrapping" the property
        method it decorates (`fget`, nominally).

        *obj* is the "host" object instance when the attribute is accessed from an
        object instance, e.g. `obj = Obj(); obj.fget`. *obj* is None when accessed on
        the class, e.g. `Obj.fget`.

        *type* is the class hosting the decorated getter method (`fget`) on both class
        and instance attribute access.
        """
        # --- when accessed on class, e.g. Obj.fget, just return this descriptor
        # --- instance (patched above to look like fget).
        if obj is None:
            return self  # type: ignore

        # --- when accessed on instance, start by checking instance __dict__ for
        # --- item with key matching the wrapped function's name
        value = obj.__dict__.get(self._name)
        if value is None:
            # --- on first access, the __dict__ item will be absent. Evaluate fget()
            # --- and store that value in the (otherwise unused) host-object
            # --- __dict__ value of same name ('fget' nominally)
            value = self._fget(obj)
            obj.__dict__[self._name] = value
        return cast(T, value)

    def __set__(self, obj: Any, value: Any) -> None:
        """Raises unconditionally, to preserve read-only behavior.

        This decorator is intended to implement immutable (and idempotent) object
        attributes. For that reason, assignment to this property must be explicitly
        prevented.

        If this __set__ method was not present, this descriptor would become a
        *non-data descriptor*. That would be nice because the cached value would be
        accessed directly once set (__dict__ attrs have precedence over non-data
        descriptors on instance attribute lookup). The problem is, there would be
        nothing to stop assignment to the cached value, which would overwrite the result
        of `fget()` and break both the immutability and idempotence guarantees of this
        decorator.

        The performance with this __set__() method in place was roughly 0.4 usec per
        access when measured on a 2.8GHz development machine; so quite snappy and
        probably not a rich target for optimization efforts.
        """
        raise AttributeError("can't set attribute")


def write_only_property(f: Callable[[Any, Any], None]):
    """@write_only_property decorator.

    Creates a property (descriptor attribute) that accepts assignment, but not getattr
    (use in an expression).
    """
    docstring = f.__doc__

    return property(fset=f, doc=docstring)


class ElementProxy:
    """Base class for lxml element proxy classes.

    An element proxy class is one whose primary responsibilities are fulfilled by
    manipulating the attributes and child elements of an XML element. They are the most
    common type of class in python-docx other than custom element (oxml) classes.
    """

    def __init__(self, element: BaseOxmlElement, parent: t.ProvidesXmlPart | None = None):
        self._element = element
        self._parent = parent

    def __eq__(self, other: object):
        """Return |True| if this proxy object refers to the same oxml element as does
        `other`.

        ElementProxy objects are value objects and should maintain no mutable local
        state. Equality for proxy objects is defined as referring to the same XML
        element, whether or not they are the same proxy object instance.
        """
        if not isinstance(other, ElementProxy):
            return False
        return self._element is other._element

    def __ne__(self, other: object):
        if not isinstance(other, ElementProxy):
            return True
        return self._element is not other._element

    @property
    def element(self):
        """The lxml element proxied by this object."""
        return self._element

    @property
    def part(self) -> XmlPart:
        """The package part containing this object."""
        if self._parent is None:
            raise ValueError("part is not accessible from this element")
        return self._parent.part


class Parented:
    """Provides common services for document elements that occur below a part but may
    occasionally require an ancestor object to provide a service, such as add or drop a
    relationship.

    Provides ``self._parent`` attribute to subclasses.
    """

    def __init__(self, parent: t.ProvidesXmlPart):
        self._parent = parent

    @property
    def part(self):
        """The package part containing this object."""
        return self._parent.part


class StoryChild:
    """A document element within a story part.

    Story parts include DocumentPart and Header/FooterPart and can contain block items
    (paragraphs and tables). Items from the block-item subtree occasionally require an
    ancestor object to provide access to part-level or package-level items like styles
    or images or to add or drop a relationship.

    Provides `self._parent` attribute to subclasses.
    """

    def __init__(self, parent: t.ProvidesStoryPart):
        self._parent = parent

    @property
    def part(self) -> StoryPart:
        """The package part containing this object."""
        return self._parent.part


class TextAccumulator:
    """Accepts `str` fragments and joins them together, in order, on `.pop().

    Handy when text in a stream is broken up arbitrarily and you want to join it back
    together within certain bounds. The optional `separator` argument determines how
    the text fragments are punctuated, defaulting to the empty string.
    """

    def __init__(self, separator: str = ""):
        self._separator = separator
        self._texts: List[str] = []

    def push(self, text: str) -> None:
        """Add a text fragment to the accumulator."""
        self._texts.append(text)

    def pop(self) -> Iterator[str]:
        """Generate sero-or-one str from those accumulated.

        Using `yield from accum.pop()` in a generator setting avoids producing an empty
        string when no text is in the accumulator.
        """
        if not self._texts:
            return
        text = self._separator.join(self._texts)
        self._texts.clear()
        yield text


# latent.py
"""Latent style-related objects."""

from docx.shared import ElementProxy
from docx.styles import BabelFish


class LatentStyles(ElementProxy):
    """Provides access to the default behaviors for latent styles in this document and
    to the collection of |_LatentStyle| objects that define overrides of those defaults
    for a particular named latent style."""

    def __getitem__(self, key):
        """Enables dictionary-style access to a latent style by name."""
        style_name = BabelFish.ui2internal(key)
        lsdException = self._element.get_by_name(style_name)
        if lsdException is None:
            raise KeyError("no latent style with name '%s'" % key)
        return _LatentStyle(lsdException)

    def __iter__(self):
        return (_LatentStyle(ls) for ls in self._element.lsdException_lst)

    def __len__(self):
        return len(self._element.lsdException_lst)

    def add_latent_style(self, name):
        """Return a newly added |_LatentStyle| object to override the inherited defaults
        defined in this latent styles object for the built-in style having `name`."""
        lsdException = self._element.add_lsdException()
        lsdException.name = BabelFish.ui2internal(name)
        return _LatentStyle(lsdException)

    @property
    def default_priority(self):
        """Integer between 0 and 99 inclusive specifying the default sort order for
        latent styles in style lists and the style gallery.

        |None| if no value is assigned, which causes Word to use the default value 99.
        """
        return self._element.defUIPriority

    @default_priority.setter
    def default_priority(self, value):
        self._element.defUIPriority = value

    @property
    def default_to_hidden(self):
        """Boolean specifying whether the default behavior for latent styles is to be
        hidden.

        A hidden style does not appear in the recommended list or in the style gallery.
        """
        return self._element.bool_prop("defSemiHidden")

    @default_to_hidden.setter
    def default_to_hidden(self, value):
        self._element.set_bool_prop("defSemiHidden", value)

    @property
    def default_to_locked(self):
        """Boolean specifying whether the default behavior for latent styles is to be
        locked.

        A locked style does not appear in the styles panel or the style gallery and
        cannot be applied to document content. This behavior is only active when
        formatting protection is turned on for the document (via the Developer menu).
        """
        return self._element.bool_prop("defLockedState")

    @default_to_locked.setter
    def default_to_locked(self, value):
        self._element.set_bool_prop("defLockedState", value)

    @property
    def default_to_quick_style(self):
        """Boolean specifying whether the default behavior for latent styles is to
        appear in the style gallery when not hidden."""
        return self._element.bool_prop("defQFormat")

    @default_to_quick_style.setter
    def default_to_quick_style(self, value):
        self._element.set_bool_prop("defQFormat", value)

    @property
    def default_to_unhide_when_used(self):
        """Boolean specifying whether the default behavior for latent styles is to be
        unhidden when first applied to content."""
        return self._element.bool_prop("defUnhideWhenUsed")

    @default_to_unhide_when_used.setter
    def default_to_unhide_when_used(self, value):
        self._element.set_bool_prop("defUnhideWhenUsed", value)

    @property
    def load_count(self):
        """Integer specifying the number of built-in styles to initialize to the
        defaults specified in this |LatentStyles| object.

        |None| if there is no setting in the XML (very uncommon). The default Word 2011
        template sets this value to 276, accounting for the built-in styles in Word
        2010.
        """
        return self._element.count

    @load_count.setter
    def load_count(self, value):
        self._element.count = value


class _LatentStyle(ElementProxy):
    """Proxy for an `w:lsdException` element, which specifies display behaviors for a
    built-in style when no definition for that style is stored yet in the `styles.xml`
    part.

    The values in this element override the defaults specified in the parent
    `w:latentStyles` element.
    """

    def delete(self):
        """Remove this latent style definition such that the defaults defined in the
        containing |LatentStyles| object provide the effective value for each of its
        attributes.

        Attempting to access any attributes on this object after calling this method
        will raise |AttributeError|.
        """
        self._element.delete()
        self._element = None

    @property
    def hidden(self):
        """Tri-state value specifying whether this latent style should appear in the
        recommended list.

        |None| indicates the effective value is inherited from the parent
        ``<w:latentStyles>`` element.
        """
        return self._element.on_off_prop("semiHidden")

    @hidden.setter
    def hidden(self, value):
        self._element.set_on_off_prop("semiHidden", value)

    @property
    def locked(self):
        """Tri-state value specifying whether this latent styles is locked.

        A locked style does not appear in the styles panel or the style gallery and
        cannot be applied to document content. This behavior is only active when
        formatting protection is turned on for the document (via the Developer menu).
        """
        return self._element.on_off_prop("locked")

    @locked.setter
    def locked(self, value):
        self._element.set_on_off_prop("locked", value)

    @property
    def name(self):
        """The name of the built-in style this exception applies to."""
        return BabelFish.internal2ui(self._element.name)

    @property
    def priority(self):
        """The integer sort key for this latent style in the Word UI."""
        return self._element.uiPriority

    @priority.setter
    def priority(self, value):
        self._element.uiPriority = value

    @property
    def quick_style(self):
        """Tri-state value specifying whether this latent style should appear in the
        Word styles gallery when not hidden.

        |None| indicates the effective value should be inherited from the default values
        in its parent |LatentStyles| object.
        """
        return self._element.on_off_prop("qFormat")

    @quick_style.setter
    def quick_style(self, value):
        self._element.set_on_off_prop("qFormat", value)

    @property
    def unhide_when_used(self):
        """Tri-state value specifying whether this style should have its :attr:`hidden`
        attribute set |False| the next time the style is applied to content.

        |None| indicates the effective value should be inherited from the default
        specified by its parent |LatentStyles| object.
        """
        return self._element.on_off_prop("unhideWhenUsed")

    @unhide_when_used.setter
    def unhide_when_used(self, value):
        self._element.set_on_off_prop("unhideWhenUsed", value)


# style.py
"""Style object hierarchy."""

from __future__ import annotations

from typing import Type

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.styles import CT_Style
from docx.shared import ElementProxy
from docx.styles import BabelFish
from docx.text.font import Font
from docx.text.parfmt import ParagraphFormat


def StyleFactory(style_elm: CT_Style) -> BaseStyle:
    """Return `Style` object of appropriate |BaseStyle| subclass for `style_elm`."""
    style_cls: Type[BaseStyle] = {
        WD_STYLE_TYPE.PARAGRAPH: ParagraphStyle,
        WD_STYLE_TYPE.CHARACTER: CharacterStyle,
        WD_STYLE_TYPE.TABLE: _TableStyle,
        WD_STYLE_TYPE.LIST: _NumberingStyle,
    }[style_elm.type]

    return style_cls(style_elm)


class BaseStyle(ElementProxy):
    """Base class for the various types of style object, paragraph, character, table,
    and numbering.

    These properties and methods are inherited by all style objects.
    """

    def __init__(self, style_elm: CT_Style):
        super().__init__(style_elm)
        self._style_elm = style_elm

    @property
    def builtin(self):
        """Read-only.

        |True| if this style is a built-in style. |False| indicates it is a custom
        (user-defined) style. Note this value is based on the presence of a
        `customStyle` attribute in the XML, not on specific knowledge of which styles
        are built into Word.
        """
        return not self._element.customStyle

    def delete(self):
        """Remove this style definition from the document.

        Note that calling this method does not remove or change the style applied to any
        document content. Content items having the deleted style will be rendered using
        the default style, as is any content with a style not defined in the document.
        """
        self._element.delete()
        self._element = None

    @property
    def hidden(self):
        """|True| if display of this style in the style gallery and list of recommended
        styles is suppressed.

        |False| otherwise. In order to be shown in the style gallery, this value must be
        |False| and :attr:`.quick_style` must be |True|.
        """
        return self._element.semiHidden_val

    @hidden.setter
    def hidden(self, value):
        self._element.semiHidden_val = value

    @property
    def locked(self):
        """Read/write Boolean.

        |True| if this style is locked. A locked style does not appear in the styles
        panel or the style gallery and cannot be applied to document content. This
        behavior is only active when formatting protection is turned on for the document
        (via the Developer menu).
        """
        return self._element.locked_val

    @locked.setter
    def locked(self, value):
        self._element.locked_val = value

    @property
    def name(self):
        """The UI name of this style."""
        name = self._element.name_val
        if name is None:
            return None
        return BabelFish.internal2ui(name)

    @name.setter
    def name(self, value):
        self._element.name_val = value

    @property
    def priority(self):
        """The integer sort key governing display sequence of this style in the Word UI.

        |None| indicates no setting is defined, causing Word to use the default value of
        0. Style name is used as a secondary sort key to resolve ordering of styles
        having the same priority value.
        """
        return self._element.uiPriority_val

    @priority.setter
    def priority(self, value):
        self._element.uiPriority_val = value

    @property
    def quick_style(self):
        """|True| if this style should be displayed in the style gallery when
        :attr:`.hidden` is |False|.

        Read/write Boolean.
        """
        return self._element.qFormat_val

    @quick_style.setter
    def quick_style(self, value):
        self._element.qFormat_val = value

    @property
    def style_id(self) -> str:
        """The unique key name (string) for this style.

        This value is subject to rewriting by Word and should generally not be changed
        unless you are familiar with the internals involved.
        """
        return self._style_elm.styleId

    @style_id.setter
    def style_id(self, value):
        self._element.styleId = value

    @property
    def type(self):
        """Member of :ref:`WdStyleType` corresponding to the type of this style, e.g.
        ``WD_STYLE_TYPE.PARAGRAPH``."""
        type = self._style_elm.type
        if type is None:
            return WD_STYLE_TYPE.PARAGRAPH
        return type

    @property
    def unhide_when_used(self):
        """|True| if an application should make this style visible the next time it is
        applied to content.

        False otherwise. Note that |docx| does not automatically unhide a style having
        |True| for this attribute when it is applied to content.
        """
        return self._element.unhideWhenUsed_val

    @unhide_when_used.setter
    def unhide_when_used(self, value):
        self._element.unhideWhenUsed_val = value


class CharacterStyle(BaseStyle):
    """A character style.

    A character style is applied to a |Run| object and primarily provides character-
    level formatting via the |Font| object in its :attr:`.font` property.
    """

    @property
    def base_style(self):
        """Style object this style inherits from or |None| if this style is not based on
        another style."""
        base_style = self._element.base_style
        if base_style is None:
            return None
        return StyleFactory(base_style)

    @base_style.setter
    def base_style(self, style):
        style_id = style.style_id if style is not None else None
        self._element.basedOn_val = style_id

    @property
    def font(self):
        """The |Font| object providing access to the character formatting properties for
        this style, such as font name and size."""
        return Font(self._element)


# -- just in case someone uses the old name in an extension function --
_CharacterStyle = CharacterStyle


class ParagraphStyle(CharacterStyle):
    """A paragraph style.

    A paragraph style provides both character formatting and paragraph formatting such
    as indentation and line-spacing.
    """

    def __repr__(self):
        return "_ParagraphStyle('%s') id: %s" % (self.name, id(self))

    @property
    def next_paragraph_style(self):
        """|_ParagraphStyle| object representing the style to be applied automatically
        to a new paragraph inserted after a paragraph of this style.

        Returns self if no next paragraph style is defined. Assigning |None| or `self`
        removes the setting such that new paragraphs are created using this same style.
        """
        next_style_elm = self._element.next_style
        if next_style_elm is None:
            return self
        if next_style_elm.type != WD_STYLE_TYPE.PARAGRAPH:
            return self
        return StyleFactory(next_style_elm)

    @next_paragraph_style.setter
    def next_paragraph_style(self, style):
        if style is None or style.style_id == self.style_id:
            self._element._remove_next()
        else:
            self._element.get_or_add_next().val = style.style_id

    @property
    def paragraph_format(self):
        """The |ParagraphFormat| object providing access to the paragraph formatting
        properties for this style such as indentation."""
        return ParagraphFormat(self._element)


# -- just in case someone uses the old name in an extension function --
_ParagraphStyle = ParagraphStyle


class _TableStyle(ParagraphStyle):
    """A table style.

    A table style provides character and paragraph formatting for its contents as well
    as special table formatting properties.
    """

    def __repr__(self):
        return "_TableStyle('%s') id: %s" % (self.name, id(self))


class _NumberingStyle(BaseStyle):
    """A numbering style.

    Not yet implemented.
    """


# styles.py
"""Styles object, container for all objects in the styles part."""

from __future__ import annotations

from warnings import warn

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.styles import CT_Styles
from docx.shared import ElementProxy
from docx.styles import BabelFish
from docx.styles.latent import LatentStyles
from docx.styles.style import BaseStyle, StyleFactory


class Styles(ElementProxy):
    """Provides access to the styles defined in a document.

    Accessed using the :attr:`.Document.styles` property. Supports ``len()``, iteration,
    and dictionary-style access by style name.
    """

    def __init__(self, styles: CT_Styles):
        super().__init__(styles)
        self._element = styles

    def __contains__(self, name):
        """Enables `in` operator on style name."""
        internal_name = BabelFish.ui2internal(name)
        return any(style.name_val == internal_name for style in self._element.style_lst)

    def __getitem__(self, key: str):
        """Enables dictionary-style access by UI name.

        Lookup by style id is deprecated, triggers a warning, and will be removed in a
        near-future release.
        """
        style_elm = self._element.get_by_name(BabelFish.ui2internal(key))
        if style_elm is not None:
            return StyleFactory(style_elm)

        style_elm = self._element.get_by_id(key)
        if style_elm is not None:
            msg = (
                "style lookup by style_id is deprecated. Use style name as "
                "key instead."
            )
            warn(msg, UserWarning, stacklevel=2)
            return StyleFactory(style_elm)

        raise KeyError("no style with name '%s'" % key)

    def __iter__(self):
        return (StyleFactory(style) for style in self._element.style_lst)

    def __len__(self):
        return len(self._element.style_lst)

    def add_style(self, name, style_type, builtin=False):
        """Return a newly added style object of `style_type` and identified by `name`.

        A builtin style can be defined by passing True for the optional `builtin`
        argument.
        """
        style_name = BabelFish.ui2internal(name)
        if style_name in self:
            raise ValueError("document already contains style '%s'" % name)
        style = self._element.add_style_of_type(style_name, style_type, builtin)
        return StyleFactory(style)

    def default(self, style_type: WD_STYLE_TYPE):
        """Return the default style for `style_type` or |None| if no default is defined
        for that type (not common)."""
        style = self._element.default_for(style_type)
        if style is None:
            return None
        return StyleFactory(style)

    def get_by_id(self, style_id: str | None, style_type: WD_STYLE_TYPE):
        """Return the style of `style_type` matching `style_id`.

        Returns the default for `style_type` if `style_id` is not found or is |None|, or
        if the style having `style_id` is not of `style_type`.
        """
        if style_id is None:
            return self.default(style_type)
        return self._get_by_id(style_id, style_type)

    def get_style_id(self, style_or_name, style_type):
        """Return the id of the style corresponding to `style_or_name`, or |None| if
        `style_or_name` is |None|.

        If `style_or_name` is not a style object, the style is looked up using
        `style_or_name` as a style name, raising |ValueError| if no style with that name
        is defined. Raises |ValueError| if the target style is not of `style_type`.
        """
        if style_or_name is None:
            return None
        elif isinstance(style_or_name, BaseStyle):
            return self._get_style_id_from_style(style_or_name, style_type)
        else:
            return self._get_style_id_from_name(style_or_name, style_type)

    @property
    def latent_styles(self):
        """A |LatentStyles| object providing access to the default behaviors for latent
        styles and the collection of |_LatentStyle| objects that define overrides of
        those defaults for a particular named latent style."""
        return LatentStyles(self._element.get_or_add_latentStyles())

    def _get_by_id(self, style_id: str | None, style_type: WD_STYLE_TYPE):
        """Return the style of `style_type` matching `style_id`.

        Returns the default for `style_type` if `style_id` is not found or if the style
        having `style_id` is not of `style_type`.
        """
        style = self._element.get_by_id(style_id) if style_id else None
        if style is None or style.type != style_type:
            return self.default(style_type)
        return StyleFactory(style)

    def _get_style_id_from_name(
        self, style_name: str, style_type: WD_STYLE_TYPE
    ) -> str | None:
        """Return the id of the style of `style_type` corresponding to `style_name`.

        Returns |None| if that style is the default style for `style_type`. Raises
        |ValueError| if the named style is not found in the document or does not match
        `style_type`.
        """
        return self._get_style_id_from_style(self[style_name], style_type)

    def _get_style_id_from_style(
        self, style: BaseStyle, style_type: WD_STYLE_TYPE
    ) -> str | None:
        """Id of `style`, or |None| if it is the default style of `style_type`.

        Raises |ValueError| if style is not of `style_type`.
        """
        if style.type != style_type:
            raise ValueError(
                "assigned style is type %s, need type %s" % (style.type, style_type)
            )
        if style == self.default(style_type):
            return None
        return style.style_id


# __init__.py
"""Sub-package module for docx.styles sub-package."""

from __future__ import annotations

from typing import Dict


class BabelFish:
    """Translates special-case style names from UI name (e.g. Heading 1) to
    internal/styles.xml name (e.g. heading 1) and back."""

    style_aliases = (
        ("Caption", "caption"),
        ("Footer", "footer"),
        ("Header", "header"),
        ("Heading 1", "heading 1"),
        ("Heading 2", "heading 2"),
        ("Heading 3", "heading 3"),
        ("Heading 4", "heading 4"),
        ("Heading 5", "heading 5"),
        ("Heading 6", "heading 6"),
        ("Heading 7", "heading 7"),
        ("Heading 8", "heading 8"),
        ("Heading 9", "heading 9"),
    )

    internal_style_names: Dict[str, str] = dict(style_aliases)
    ui_style_names = {item[1]: item[0] for item in style_aliases}

    @classmethod
    def ui2internal(cls, ui_style_name: str) -> str:
        """Return the internal style name corresponding to `ui_style_name`, such as
        'heading 1' for 'Heading 1'."""
        return cls.internal_style_names.get(ui_style_name, ui_style_name)

    @classmethod
    def internal2ui(cls, internal_style_name: str) -> str:
        """Return the user interface style name corresponding to `internal_style_name`,
        such as 'Heading 1' for 'heading 1'."""
        return cls.ui_style_names.get(internal_style_name, internal_style_name)


# table.py
"""The |Table| object and related proxy classes."""

from __future__ import annotations

from typing import TYPE_CHECKING, Iterator, cast, overload

from typing_extensions import TypeAlias

from docx.blkcntnr import BlockItemContainer
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.simpletypes import ST_Merge
from docx.oxml.table import CT_TblGridCol
from docx.shared import Inches, Parented, StoryChild, lazyproperty

if TYPE_CHECKING:
    import docx.types as t
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
    from docx.oxml.table import CT_Row, CT_Tbl, CT_TblPr, CT_Tc
    from docx.shared import Length
    from docx.styles.style import (
        ParagraphStyle,
        _TableStyle,  # pyright: ignore[reportPrivateUsage]
    )

TableParent: TypeAlias = "Table | _Columns | _Rows"


class Table(StoryChild):
    """Proxy class for a WordprocessingML ``<w:tbl>`` element."""

    def __init__(self, tbl: CT_Tbl, parent: t.ProvidesStoryPart):
        super(Table, self).__init__(parent)
        self._element = tbl
        self._tbl = tbl

    def add_column(self, width: Length):
        """Return a |_Column| object of `width`, newly added rightmost to the table."""
        tblGrid = self._tbl.tblGrid
        gridCol = tblGrid.add_gridCol()
        gridCol.w = width
        for tr in self._tbl.tr_lst:
            tc = tr.add_tc()
            tc.width = width
        return _Column(gridCol, self)

    def add_row(self):
        """Return a |_Row| instance, newly added bottom-most to the table."""
        tbl = self._tbl
        tr = tbl.add_tr()
        for gridCol in tbl.tblGrid.gridCol_lst:
            tc = tr.add_tc()
            if gridCol.w is not None:
                tc.width = gridCol.w
        return _Row(tr, self)

    @property
    def alignment(self) -> WD_TABLE_ALIGNMENT | None:
        """Read/write.

        A member of :ref:`WdRowAlignment` or None, specifying the positioning of this
        table between the page margins. |None| if no setting is specified, causing the
        effective value to be inherited from the style hierarchy.
        """
        return self._tblPr.alignment

    @alignment.setter
    def alignment(self, value: WD_TABLE_ALIGNMENT | None):
        self._tblPr.alignment = value

    @property
    def autofit(self) -> bool:
        """|True| if column widths can be automatically adjusted to improve the fit of
        cell contents.

        |False| if table layout is fixed. Column widths are adjusted in either case if
        total column width exceeds page width. Read/write boolean.
        """
        return self._tblPr.autofit

    @autofit.setter
    def autofit(self, value: bool):
        self._tblPr.autofit = value

    def cell(self, row_idx: int, col_idx: int) -> _Cell:
        """|_Cell| at `row_idx`, `col_idx` intersection.

        (0, 0) is the top, left-most cell.
        """
        cell_idx = col_idx + (row_idx * self._column_count)
        return self._cells[cell_idx]

    def column_cells(self, column_idx: int) -> list[_Cell]:
        """Sequence of cells in the column at `column_idx` in this table."""
        cells = self._cells
        idxs = range(column_idx, len(cells), self._column_count)
        return [cells[idx] for idx in idxs]

    @lazyproperty
    def columns(self):
        """|_Columns| instance representing the sequence of columns in this table."""
        return _Columns(self._tbl, self)

    def row_cells(self, row_idx: int) -> list[_Cell]:
        """DEPRECATED: Use `table.rows[row_idx].cells` instead.

        Sequence of cells in the row at `row_idx` in this table.
        """
        column_count = self._column_count
        start = row_idx * column_count
        end = start + column_count
        return self._cells[start:end]

    @lazyproperty
    def rows(self) -> _Rows:
        """|_Rows| instance containing the sequence of rows in this table."""
        return _Rows(self._tbl, self)

    @property
    def style(self) -> _TableStyle | None:
        """|_TableStyle| object representing the style applied to this table.

        Read/write. The default table style for the document (often `Normal Table`) is
        returned if the table has no directly-applied style. Assigning |None| to this
        property removes any directly-applied table style causing it to inherit the
        default table style of the document.

        Note that the style name of a table style differs slightly from that displayed
        in the user interface; a hyphen, if it appears, must be removed. For example,
        `Light Shading - Accent 1` becomes `Light Shading Accent 1`.
        """
        style_id = self._tbl.tblStyle_val
        return cast("_TableStyle | None", self.part.get_style(style_id, WD_STYLE_TYPE.TABLE))

    @style.setter
    def style(self, style_or_name: _TableStyle | str | None):
        style_id = self.part.get_style_id(style_or_name, WD_STYLE_TYPE.TABLE)
        self._tbl.tblStyle_val = style_id

    @property
    def table(self):
        """Provide child objects with reference to the |Table| object they belong to,
        without them having to know their direct parent is a |Table| object.

        This is the terminus of a series of `parent._table` calls from an arbitrary
        child through its ancestors.
        """
        return self

    @property
    def table_direction(self) -> WD_TABLE_DIRECTION | None:
        """Member of :ref:`WdTableDirection` indicating cell-ordering direction.

        For example: `WD_TABLE_DIRECTION.LTR`. |None| indicates the value is inherited
        from the style hierarchy.
        """
        return cast("WD_TABLE_DIRECTION | None", self._tbl.bidiVisual_val)

    @table_direction.setter
    def table_direction(self, value: WD_TABLE_DIRECTION | None):
        self._element.bidiVisual_val = value

    @property
    def _cells(self) -> list[_Cell]:
        """A sequence of |_Cell| objects, one for each cell of the layout grid.

        If the table contains a span, one or more |_Cell| object references are
        repeated.
        """
        col_count = self._column_count
        cells: list[_Cell] = []
        for tc in self._tbl.iter_tcs():
            for grid_span_idx in range(tc.grid_span):
                if tc.vMerge == ST_Merge.CONTINUE:
                    cells.append(cells[-col_count])
                elif grid_span_idx > 0:
                    cells.append(cells[-1])
                else:
                    cells.append(_Cell(tc, self))
        return cells

    @property
    def _column_count(self):
        """The number of grid columns in this table."""
        return self._tbl.col_count

    @property
    def _tblPr(self) -> CT_TblPr:
        return self._tbl.tblPr


class _Cell(BlockItemContainer):
    """Table cell."""

    def __init__(self, tc: CT_Tc, parent: TableParent):
        super(_Cell, self).__init__(tc, cast("t.ProvidesStoryPart", parent))
        self._parent = parent
        self._tc = self._element = tc

    def add_paragraph(self, text: str = "", style: str | ParagraphStyle | None = None):
        """Return a paragraph newly added to the end of the content in this cell.

        If present, `text` is added to the paragraph in a single run. If specified, the
        paragraph style `style` is applied. If `style` is not specified or is |None|,
        the result is as though the 'Normal' style was applied. Note that the formatting
        of text in a cell can be influenced by the table style. `text` can contain tab
        (``\\t``) characters, which are converted to the appropriate XML form for a tab.
        `text` can also include newline (``\\n``) or carriage return (``\\r``)
        characters, each of which is converted to a line break.
        """
        return super(_Cell, self).add_paragraph(text, style)

    def add_table(  # pyright: ignore[reportIncompatibleMethodOverride]
        self, rows: int, cols: int
    ) -> Table:
        """Return a table newly added to this cell after any existing cell content.

        The new table will have `rows` rows and `cols` columns.

        An empty paragraph is added after the table because Word requires a paragraph
        element as the last element in every cell.
        """
        width = self.width if self.width is not None else Inches(1)
        table = super(_Cell, self).add_table(rows, cols, width)
        self.add_paragraph()
        return table

    @property
    def grid_span(self) -> int:
        """Number of layout-grid cells this cell spans horizontally.

        A "normal" cell has a grid-span of 1. A horizontally merged cell has a grid-span of 2 or
        more.
        """
        return self._tc.grid_span

    def merge(self, other_cell: _Cell):
        """Return a merged cell created by spanning the rectangular region having this
        cell and `other_cell` as diagonal corners.

        Raises |InvalidSpanError| if the cells do not define a rectangular region.
        """
        tc, tc_2 = self._tc, other_cell._tc
        merged_tc = tc.merge(tc_2)
        return _Cell(merged_tc, self._parent)

    @property
    def paragraphs(self):
        """List of paragraphs in the cell.

        A table cell is required to contain at least one block-level element and end
        with a paragraph. By default, a new cell contains a single paragraph. Read-only
        """
        return super(_Cell, self).paragraphs

    @property
    def tables(self):
        """List of tables in the cell, in the order they appear.

        Read-only.
        """
        return super(_Cell, self).tables

    @property
    def text(self) -> str:
        """The entire contents of this cell as a string of text.

        Assigning a string to this property replaces all existing content with a single
        paragraph containing the assigned text in a single run.
        """
        return "\n".join(p.text for p in self.paragraphs)

    @text.setter
    def text(self, text: str):
        """Write-only.

        Set entire contents of cell to the string `text`. Any existing content or
        revisions are replaced.
        """
        tc = self._tc
        tc.clear_content()
        p = tc.add_p()
        r = p.add_r()
        r.text = text

    @property
    def vertical_alignment(self):
        """Member of :ref:`WdCellVerticalAlignment` or None.

        A value of |None| indicates vertical alignment for this cell is inherited.
        Assigning |None| causes any explicitly defined vertical alignment to be removed,
        restoring inheritance.
        """
        tcPr = self._element.tcPr
        if tcPr is None:
            return None
        return tcPr.vAlign_val

    @vertical_alignment.setter
    def vertical_alignment(self, value: WD_CELL_VERTICAL_ALIGNMENT | None):
        tcPr = self._element.get_or_add_tcPr()
        tcPr.vAlign_val = value

    @property
    def width(self):
        """The width of this cell in EMU, or |None| if no explicit width is set."""
        return self._tc.width

    @width.setter
    def width(self, value: Length):
        self._tc.width = value


class _Column(Parented):
    """Table column."""

    def __init__(self, gridCol: CT_TblGridCol, parent: TableParent):
        super(_Column, self).__init__(parent)
        self._parent = parent
        self._gridCol = gridCol

    @property
    def cells(self) -> tuple[_Cell, ...]:
        """Sequence of |_Cell| instances corresponding to cells in this column."""
        return tuple(self.table.column_cells(self._index))

    @property
    def table(self) -> Table:
        """Reference to the |Table| object this column belongs to."""
        return self._parent.table

    @property
    def width(self) -> Length | None:
        """The width of this column in EMU, or |None| if no explicit width is set."""
        return self._gridCol.w

    @width.setter
    def width(self, value: Length | None):
        self._gridCol.w = value

    @property
    def _index(self):
        """Index of this column in its table, starting from zero."""
        return self._gridCol.gridCol_idx


class _Columns(Parented):
    """Sequence of |_Column| instances corresponding to the columns in a table.

    Supports ``len()``, iteration and indexed access.
    """

    def __init__(self, tbl: CT_Tbl, parent: TableParent):
        super(_Columns, self).__init__(parent)
        self._parent = parent
        self._tbl = tbl

    def __getitem__(self, idx: int):
        """Provide indexed access, e.g. 'columns[0]'."""
        try:
            gridCol = self._gridCol_lst[idx]
        except IndexError:
            msg = "column index [%d] is out of range" % idx
            raise IndexError(msg)
        return _Column(gridCol, self)

    def __iter__(self):
        for gridCol in self._gridCol_lst:
            yield _Column(gridCol, self)

    def __len__(self):
        return len(self._gridCol_lst)

    @property
    def table(self) -> Table:
        """Reference to the |Table| object this column collection belongs to."""
        return self._parent.table

    @property
    def _gridCol_lst(self):
        """Sequence containing ``<w:gridCol>`` elements for this table, each
        representing a table column."""
        tblGrid = self._tbl.tblGrid
        return tblGrid.gridCol_lst


class _Row(Parented):
    """Table row."""

    def __init__(self, tr: CT_Row, parent: TableParent):
        super(_Row, self).__init__(parent)
        self._parent = parent
        self._tr = self._element = tr

    @property
    def cells(self) -> tuple[_Cell, ...]:
        """Sequence of |_Cell| instances corresponding to cells in this row.

        Note that Word allows table rows to start later than the first column and end before the
        last column.

        - Only cells actually present are included in the return value.
        - This implies the length of this cell sequence may differ between rows of the same table.
        - If you are reading the cells from each row to form a rectangular "matrix" data structure
          of the table cell values, you will need to account for empty leading and/or trailing
          layout-grid positions using `.grid_cols_before` and `.grid_cols_after`.

        """

        def iter_tc_cells(tc: CT_Tc) -> Iterator[_Cell]:
            """Generate a cell object for each layout-grid cell in `tc`.

            In particular, a `<w:tc>` element with a horizontal "span" with generate the same cell
            multiple times, one for each grid-cell being spanned. This approximates a row in a
            "uniform" table, where each row has a cell for each column in the table.
            """
            # -- a cell comprising the second or later row of a vertical span is indicated by
            # -- tc.vMerge="continue" (the default value of the `w:vMerge` attribute, when it is
            # -- present in the XML). The `w:tc` element at the same grid-offset in the prior row
            # -- is guaranteed to be the same width (gridSpan). So we can delegate content
            # -- discovery to that prior-row `w:tc` element (recursively) until we arrive at the
            # -- "root" cell -- for the vertical span.
            if tc.vMerge == "continue":
                yield from iter_tc_cells(tc._tc_above)  # pyright: ignore[reportPrivateUsage]
                return

            # -- Otherwise, vMerge is either "restart" or None, meaning this `tc` holds the actual
            # -- content of the cell (whether it is vertically merged or not).
            cell = _Cell(tc, self.table)
            for _ in range(tc.grid_span):
                yield cell

        def _iter_row_cells() -> Iterator[_Cell]:
            """Generate `_Cell` instance for each populated layout-grid cell in this row."""
            for tc in self._tr.tc_lst:
                yield from iter_tc_cells(tc)

        return tuple(_iter_row_cells())

    @property
    def grid_cols_after(self) -> int:
        """Count of unpopulated grid-columns after the last cell in this row.

        Word allows a row to "end early", meaning that one or more cells are not present at the
        end of that row.

        Note these are not simply "empty" cells. The renderer reads this value and "skips" this
        many columns after drawing the last cell.

        Note this also implies that not all rows are guaranteed to have the same number of cells,
        e.g. `_Row.cells` could have length `n` for one row and `n - m` for the next row in the same
        table. Visually this appears as a column (at the beginning or end, not in the middle) with
        one or more cells missing.
        """
        return self._tr.grid_after

    @property
    def grid_cols_before(self) -> int:
        """Count of unpopulated grid-columns before the first cell in this row.

        Word allows a row to "start late", meaning that one or more cells are not present at the
        beginning of that row.

        Note these are not simply "empty" cells. The renderer reads this value and skips forward to
        the table layout-grid position of the first cell in this row; the renderer "skips" this many
        columns before drawing the first cell.

        Note this also implies that not all rows are guaranteed to have the same number of cells,
        e.g. `_Row.cells` could have length `n` for one row and `n - m` for the next row in the same
        table.
        """
        return self._tr.grid_before

    @property
    def height(self) -> Length | None:
        """Return a |Length| object representing the height of this cell, or |None| if
        no explicit height is set."""
        return self._tr.trHeight_val

    @height.setter
    def height(self, value: Length | None):
        self._tr.trHeight_val = value

    @property
    def height_rule(self) -> WD_ROW_HEIGHT_RULE | None:
        """Return the height rule of this cell as a member of the :ref:`WdRowHeightRule`.

        This value is |None| if no explicit height_rule is set.
        """
        return self._tr.trHeight_hRule

    @height_rule.setter
    def height_rule(self, value: WD_ROW_HEIGHT_RULE | None):
        self._tr.trHeight_hRule = value

    @property
    def table(self) -> Table:
        """Reference to the |Table| object this row belongs to."""
        return self._parent.table

    @property
    def _index(self) -> int:
        """Index of this row in its table, starting from zero."""
        return self._tr.tr_idx


class _Rows(Parented):
    """Sequence of |_Row| objects corresponding to the rows in a table.

    Supports ``len()``, iteration, indexed access, and slicing.
    """

    def __init__(self, tbl: CT_Tbl, parent: TableParent):
        super(_Rows, self).__init__(parent)
        self._parent = parent
        self._tbl = tbl

    @overload
    def __getitem__(self, idx: int) -> _Row: ...

    @overload
    def __getitem__(self, idx: slice) -> list[_Row]: ...

    def __getitem__(self, idx: int | slice) -> _Row | list[_Row]:
        """Provide indexed access, (e.g. `rows[0]` or `rows[1:3]`)"""
        return list(self)[idx]

    def __iter__(self):
        return (_Row(tr, self) for tr in self._tbl.tr_lst)

    def __len__(self):
        return len(self._tbl.tr_lst)

    @property
    def table(self) -> Table:
        """Reference to the |Table| object this row collection belongs to."""
        return self._parent.table


# font.py
"""Font-related proxy objects."""

from __future__ import annotations

from typing import TYPE_CHECKING, Any

from docx.dml.color import ColorFormat
from docx.enum.text import WD_UNDERLINE
from docx.shared import ElementProxy, Emu

if TYPE_CHECKING:
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.text.run import CT_R
    from docx.shared import Length


class Font(ElementProxy):
    """Proxy object for parent of a `<w:rPr>` element and providing access to
    character properties such as font name, font size, bold, and subscript."""

    def __init__(self, r: CT_R, parent: Any | None = None):
        super().__init__(r, parent)
        self._element = r
        self._r = r

    @property
    def all_caps(self) -> bool | None:
        """Read/write.

        Causes text in this font to appear in capital letters.
        """
        return self._get_bool_prop("caps")

    @all_caps.setter
    def all_caps(self, value: bool | None) -> None:
        self._set_bool_prop("caps", value)

    @property
    def bold(self) -> bool | None:
        """Read/write.

        Causes text in this font to appear in bold.
        """
        return self._get_bool_prop("b")

    @bold.setter
    def bold(self, value: bool | None) -> None:
        self._set_bool_prop("b", value)

    @property
    def color(self):
        """A |ColorFormat| object providing a way to get and set the text color for this
        font."""
        return ColorFormat(self._element)

    @property
    def complex_script(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the characters in the run to be treated as complex script
        regardless of their Unicode values.
        """
        return self._get_bool_prop("cs")

    @complex_script.setter
    def complex_script(self, value: bool | None) -> None:
        self._set_bool_prop("cs", value)

    @property
    def cs_bold(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the complex script characters in the run to be displayed in
        bold typeface.
        """
        return self._get_bool_prop("bCs")

    @cs_bold.setter
    def cs_bold(self, value: bool | None) -> None:
        self._set_bool_prop("bCs", value)

    @property
    def cs_italic(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the complex script characters in the run to be displayed in
        italic typeface.
        """
        return self._get_bool_prop("iCs")

    @cs_italic.setter
    def cs_italic(self, value: bool | None) -> None:
        self._set_bool_prop("iCs", value)

    @property
    def double_strike(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text in the run to appear with double strikethrough.
        """
        return self._get_bool_prop("dstrike")

    @double_strike.setter
    def double_strike(self, value: bool | None) -> None:
        self._set_bool_prop("dstrike", value)

    @property
    def emboss(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text in the run to appear as if raised off the page in
        relief.
        """
        return self._get_bool_prop("emboss")

    @emboss.setter
    def emboss(self, value: bool | None) -> None:
        self._set_bool_prop("emboss", value)

    @property
    def hidden(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text in the run to be hidden from display, unless
        applications settings force hidden text to be shown.
        """
        return self._get_bool_prop("vanish")

    @hidden.setter
    def hidden(self, value: bool | None) -> None:
        self._set_bool_prop("vanish", value)

    @property
    def highlight_color(self) -> WD_COLOR_INDEX | None:
        """Color of highlighing applied or |None| if not highlighted."""
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.highlight_val

    @highlight_color.setter
    def highlight_color(self, value: WD_COLOR_INDEX | None):
        rPr = self._element.get_or_add_rPr()
        rPr.highlight_val = value

    @property
    def italic(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text of the run to appear in italics. |None| indicates
        the effective value is inherited from the style hierarchy.
        """
        return self._get_bool_prop("i")

    @italic.setter
    def italic(self, value: bool | None) -> None:
        self._set_bool_prop("i", value)

    @property
    def imprint(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text in the run to appear as if pressed into the page.
        """
        return self._get_bool_prop("imprint")

    @imprint.setter
    def imprint(self, value: bool | None) -> None:
        self._set_bool_prop("imprint", value)

    @property
    def math(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, specifies this run contains WML that should be handled as though it
        was Office Open XML Math.
        """
        return self._get_bool_prop("oMath")

    @math.setter
    def math(self, value: bool | None) -> None:
        self._set_bool_prop("oMath", value)

    @property
    def name(self) -> str | None:
        """The typeface name for this |Font|.

        Causes the text it controls to appear in the named font, if a matching font is
        found. |None| indicates the typeface is inherited from the style hierarchy.
        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.rFonts_ascii

    @name.setter
    def name(self, value: str | None) -> None:
        rPr = self._element.get_or_add_rPr()
        rPr.rFonts_ascii = value
        rPr.rFonts_hAnsi = value

    @property
    def no_proof(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, specifies that the contents of this run should not report any
        errors when the document is scanned for spelling and grammar.
        """
        return self._get_bool_prop("noProof")

    @no_proof.setter
    def no_proof(self, value: bool | None) -> None:
        self._set_bool_prop("noProof", value)

    @property
    def outline(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the characters in the run to appear as if they have an
        outline, by drawing a one pixel wide border around the inside and outside
        borders of each character glyph.
        """
        return self._get_bool_prop("outline")

    @outline.setter
    def outline(self, value: bool | None) -> None:
        self._set_bool_prop("outline", value)

    @property
    def rtl(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the text in the run to have right-to-left characteristics.
        """
        return self._get_bool_prop("rtl")

    @rtl.setter
    def rtl(self, value: bool | None) -> None:
        self._set_bool_prop("rtl", value)

    @property
    def shadow(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the text in the run to appear as if each character has a
        shadow.
        """
        return self._get_bool_prop("shadow")

    @shadow.setter
    def shadow(self, value: bool | None) -> None:
        self._set_bool_prop("shadow", value)

    @property
    def size(self) -> Length | None:
        """Font height in English Metric Units (EMU).

        |None| indicates the font size should be inherited from the style hierarchy.
        |Length| is a subclass of |int| having properties for convenient conversion into
        points or other length units. The :class:`docx.shared.Pt` class allows
        convenient specification of point values::

            >>> font.size = Pt(24)
            >>> font.size
            304800
            >>> font.size.pt
            24.0

        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.sz_val

    @size.setter
    def size(self, emu: int | Length | None) -> None:
        rPr = self._element.get_or_add_rPr()
        rPr.sz_val = None if emu is None else Emu(emu)

    @property
    def small_caps(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the lowercase characters in the run to appear as capital
        letters two points smaller than the font size specified for the run.
        """
        return self._get_bool_prop("smallCaps")

    @small_caps.setter
    def small_caps(self, value: bool | None) -> None:
        self._set_bool_prop("smallCaps", value)

    @property
    def snap_to_grid(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the run to use the document grid characters per line settings
        defined in the docGrid element when laying out the characters in this run.
        """
        return self._get_bool_prop("snapToGrid")

    @snap_to_grid.setter
    def snap_to_grid(self, value: bool | None) -> None:
        self._set_bool_prop("snapToGrid", value)

    @property
    def spec_vanish(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, specifies that the given run shall always behave as if it is
        hidden, even when hidden text is being displayed in the current document. The
        property has a very narrow, specialized use related to the table of contents.
        Consult the spec (§17.3.2.36) for more details.
        """
        return self._get_bool_prop("specVanish")

    @spec_vanish.setter
    def spec_vanish(self, value: bool | None) -> None:
        self._set_bool_prop("specVanish", value)

    @property
    def strike(self) -> bool | None:
        """Read/write tri-state value.

        When |True| causes the text in the run to appear with a single horizontal line
        through the center of the line.
        """
        return self._get_bool_prop("strike")

    @strike.setter
    def strike(self, value: bool | None) -> None:
        self._set_bool_prop("strike", value)

    @property
    def subscript(self) -> bool | None:
        """Boolean indicating whether the characters in this |Font| appear as subscript.

        |None| indicates the subscript/subscript value is inherited from the style
        hierarchy.
        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.subscript

    @subscript.setter
    def subscript(self, value: bool | None) -> None:
        rPr = self._element.get_or_add_rPr()
        rPr.subscript = value

    @property
    def superscript(self) -> bool | None:
        """Boolean indicating whether the characters in this |Font| appear as
        superscript.

        |None| indicates the subscript/superscript value is inherited from the style
        hierarchy.
        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr.superscript

    @superscript.setter
    def superscript(self, value: bool | None) -> None:
        rPr = self._element.get_or_add_rPr()
        rPr.superscript = value

    @property
    def underline(self) -> bool | WD_UNDERLINE | None:
        """The underline style for this |Font|.

        The value is one of |None|, |True|, |False|, or a member of :ref:`WdUnderline`.

        |None| indicates the font inherits its underline value from the style hierarchy.
        |False| indicates no underline. |True| indicates single underline. The values
        from :ref:`WdUnderline` are used to specify other outline styles such as double,
        wavy, and dotted.
        """
        rPr = self._element.rPr
        if rPr is None:
            return None
        val = rPr.u_val
        return (
            None
            if val == WD_UNDERLINE.INHERITED
            else True
            if val == WD_UNDERLINE.SINGLE
            else False
            if val == WD_UNDERLINE.NONE
            else val
        )

    @underline.setter
    def underline(self, value: bool | WD_UNDERLINE | None) -> None:
        rPr = self._element.get_or_add_rPr()
        # -- works fine without these two mappings, but only because True == 1 and
        # -- False == 0, which happen to match the mapping for WD_UNDERLINE.SINGLE
        # -- and .NONE respectively.
        val = (
            WD_UNDERLINE.SINGLE
            if value is True
            else WD_UNDERLINE.NONE
            if value is False
            else value
        )
        rPr.u_val = val

    @property
    def web_hidden(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, specifies that the contents of this run shall be hidden when the
        document is displayed in web page view.
        """
        return self._get_bool_prop("webHidden")

    @web_hidden.setter
    def web_hidden(self, value: bool | None) -> None:
        self._set_bool_prop("webHidden", value)

    def _get_bool_prop(self, name: str) -> bool | None:
        """Return the value of boolean child of `w:rPr` having `name`."""
        rPr = self._element.rPr
        if rPr is None:
            return None
        return rPr._get_bool_val(name)  # pyright: ignore[reportPrivateUsage]

    def _set_bool_prop(self, name: str, value: bool | None):
        """Assign `value` to the boolean child `name` of `w:rPr`."""
        rPr = self._element.get_or_add_rPr()
        rPr._set_bool_val(name, value)  # pyright: ignore[reportPrivateUsage]


# hyperlink.py
"""Hyperlink-related proxy objects for python-docx, Hyperlink in particular.

A hyperlink occurs in a paragraph, at the same level as a Run, and a hyperlink itself
contains runs, which is where the visible text of the hyperlink is stored. So it's kind
of in-between, less than a paragraph and more than a run. So it gets its own module.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.shared import Parented
from docx.text.run import Run

if TYPE_CHECKING:
    import docx.types as t
    from docx.oxml.text.hyperlink import CT_Hyperlink


class Hyperlink(Parented):
    """Proxy object wrapping a `<w:hyperlink>` element.

    A hyperlink occurs as a child of a paragraph, at the same level as a Run. A
    hyperlink itself contains runs, which is where the visible text of the hyperlink is
    stored.
    """

    def __init__(self, hyperlink: CT_Hyperlink, parent: t.ProvidesStoryPart):
        super().__init__(parent)
        self._parent = parent
        self._hyperlink = self._element = hyperlink

    @property
    def address(self) -> str:
        """The "URL" of the hyperlink (but not necessarily a web link).

        While commonly a web link like "https://google.com" the hyperlink address can
        take a variety of forms including "internal links" to bookmarked locations
        within the document. When this hyperlink is an internal "jump" to for example a
        heading from the table-of-contents (TOC), the address is blank. The bookmark
        reference (like "_Toc147925734") is stored in the `.fragment` property.
        """
        rId = self._hyperlink.rId
        return self._parent.part.rels[rId].target_ref if rId else ""

    @property
    def contains_page_break(self) -> bool:
        """True when the text of this hyperlink is broken across page boundaries.

        This is not uncommon and can happen for example when the hyperlink text is
        multiple words and occurs in the last line of a page. Theoretically, a hyperlink
        can contain more than one page break but that would be extremely uncommon in
        practice. Still, this value should be understood to mean that "one-or-more"
        rendered page breaks are present.
        """
        return bool(self._hyperlink.lastRenderedPageBreaks)

    @property
    def fragment(self) -> str:
        """Reference like `#glossary` at end of URL that refers to a sub-resource.

        Note that this value does not include the fragment-separator character ("#").

        This value is known as a "named anchor" in an HTML context and "anchor" in the
        MS API, but an "anchor" element (`<a>`) represents a full hyperlink in HTML so
        we avoid confusion by using the more precise RFC 3986 naming "URI fragment".

        These are also used to refer to bookmarks within the same document, in which
        case the `.address` value with be blank ("") and this property will hold a
        value like "_Toc147925734".

        To reliably get an entire web URL you will need to concatenate this with the
        `.address` value, separated by "#" when both are present. Consider using the
        `.url` property for that purpose.

        Word sometimes stores a fragment in this property (an XML attribute) and
        sometimes with the address, depending on how the URL is inserted, so don't
        depend on this field being empty to indicate no fragment is present.
        """
        return self._hyperlink.anchor or ""

    @property
    def runs(self) -> list[Run]:
        """List of |Run| instances in this hyperlink.

        Together these define the visible text of the hyperlink. The text of a hyperlink
        is typically contained in a single run will be broken into multiple runs if for
        example part of the hyperlink is bold or the text was changed after the document
        was saved.
        """
        return [Run(r, self._parent) for r in self._hyperlink.r_lst]

    @property
    def text(self) -> str:
        """String formed by concatenating the text of each run in the hyperlink.

        Tabs and line breaks in the XML are mapped to ``\\t`` and ``\\n`` characters
        respectively. Note that rendered page-breaks can occur within a hyperlink but
        they are not reflected in this text.
        """
        return self._hyperlink.text

    @property
    def url(self) -> str:
        """Convenience property to get web URLs from hyperlinks that contain them.

        This value is the empty string ("") when there is no address portion, so its
        boolean value can also be used to distinguish external URIs from internal "jump"
        hyperlinks like those found in a table-of-contents.

        Note that this value may also be a link to a file, so if you only want web-urls
        you'll need to check for a protocol prefix like `https://`.

        When both an address and fragment are present, the return value joins the two
        separated by the fragment-separator hash ("#"). Otherwise this value is the same
        as that of the `.address` property.
        """
        address, fragment = self.address, self.fragment
        if not address:
            return ""
        return f"{address}#{fragment}" if fragment else address


# pagebreak.py
"""Proxy objects related to rendered page-breaks."""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak
from docx.shared import Parented

if TYPE_CHECKING:
    import docx.types as t
    from docx.text.paragraph import Paragraph


class RenderedPageBreak(Parented):
    """A page-break inserted by Word during page-layout for print or display purposes.

    This usually does not correspond to a "hard" page-break inserted by the document
    author, rather just that Word ran out of room on one page and needed to start
    another. The position of these can change depending on the printer and page-size, as
    well as margins, etc. They also will change in response to edits, but not until Word
    loads and saves the document.

    Note these are never inserted by `python-docx` because it has no rendering function.
    These are generally only useful for text-extraction of existing documents when
    `python-docx` is being used solely as a document "reader".

    NOTE: a rendered page-break can occur within a hyperlink; consider a multi-word
    hyperlink like "excellent Wikipedia article on LLMs" that happens to fall close to
    the end of the last line on a page such that the page breaks between "Wikipedia" and
    "article". In such a "page-breaks-in-hyperlink" case, THESE METHODS WILL "MOVE" THE
    PAGE-BREAK to occur after the hyperlink, such that the entire hyperlink appears in
    the paragraph returned by `.preceding_paragraph_fragment`. While this places the
    "tail" text of the hyperlink on the "wrong" page, it avoids having two hyperlinks
    each with a fragment of the actual text and pointing to the same address.
    """

    def __init__(
        self,
        lastRenderedPageBreak: CT_LastRenderedPageBreak,
        parent: t.ProvidesStoryPart,
    ):
        super().__init__(parent)
        self._element = lastRenderedPageBreak
        self._lastRenderedPageBreak = lastRenderedPageBreak

    @property
    def preceding_paragraph_fragment(self) -> Paragraph | None:
        """A "loose" paragraph containing the content preceding this page-break.

        Compare `.following_paragraph_fragment` as these two are intended to be used
        together.

        This value is `None` when no content precedes this page-break. This case is
        common and occurs whenever a page breaks on an even paragraph boundary.
        Returning `None` for this case avoids "inserting" a non-existent paragraph into
        the content stream. Note that content can include DrawingML items like images or
        charts.

        Note the returned paragraph *is divorced from the document body*. Any changes
        made to it will not be reflected in the document. It is intended to provide a
        familiar container (`Paragraph`) to interrogate for the content preceding this
        page-break in the paragraph in which it occured.

        Contains the entire hyperlink when this break occurs within a hyperlink.
        """
        if self._lastRenderedPageBreak.precedes_all_content:
            return None

        from docx.text.paragraph import Paragraph

        return Paragraph(self._lastRenderedPageBreak.preceding_fragment_p, self._parent)

    @property
    def following_paragraph_fragment(self) -> Paragraph | None:
        """A "loose" paragraph containing the content following this page-break.

        HAS POTENTIALLY SURPRISING BEHAVIORS so read carefully to be sure this is what
        you want. This is primarily targeted toward text-extraction use-cases for which
        precisely associating text with the page it occurs on is important.

        Compare `.preceding_paragraph_fragment` as these two are intended to be used
        together.

        This value is `None` when no content follows this page-break. This case is
        unlikely to occur in practice because Word places even-paragraph-boundary
        page-breaks on the paragraph *following* the page-break. Still, it is possible
        and must be checked for. Returning `None` for this case avoids "inserting" an
        extra, non-existent paragraph into the content stream. Note that content can
        include DrawingML items like images or charts, not just text.

        The returned paragraph *is divorced from the document body*. Any changes made to
        it will not be reflected in the document. It is intended to provide a container
        (`Paragraph`) with familiar properties and methods that can be used to
        characterize the paragraph content following a mid-paragraph page-break.

        Contains no portion of the hyperlink when this break occurs within a hyperlink.
        """
        if self._lastRenderedPageBreak.follows_all_content:
            return None

        from docx.text.paragraph import Paragraph

        return Paragraph(self._lastRenderedPageBreak.following_fragment_p, self._parent)


# paragraph.py
"""Paragraph-related proxy types."""

from __future__ import annotations

from typing import TYPE_CHECKING, Iterator, List, cast

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.text.run import CT_R
from docx.shared import StoryChild
from docx.styles.style import ParagraphStyle
from docx.text.hyperlink import Hyperlink
from docx.text.pagebreak import RenderedPageBreak
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Run

if TYPE_CHECKING:
    import docx.types as t
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.text.paragraph import CT_P
    from docx.styles.style import CharacterStyle


class Paragraph(StoryChild):
    """Proxy object wrapping a `<w:p>` element."""

    def __init__(self, p: CT_P, parent: t.ProvidesStoryPart):
        super(Paragraph, self).__init__(parent)
        self._p = self._element = p

    def add_run(self, text: str | None = None, style: str | CharacterStyle | None = None) -> Run:
        """Append run containing `text` and having character-style `style`.

        `text` can contain tab (``\\t``) characters, which are converted to the
        appropriate XML form for a tab. `text` can also include newline (``\\n``) or
        carriage return (``\\r``) characters, each of which is converted to a line
        break. When `text` is `None`, the new run is empty.
        """
        r = self._p.add_r()
        run = Run(r, self)
        if text:
            run.text = text
        if style:
            run.style = style
        return run

    @property
    def alignment(self) -> WD_PARAGRAPH_ALIGNMENT | None:
        """A member of the :ref:`WdParagraphAlignment` enumeration specifying the
        justification setting for this paragraph.

        A value of |None| indicates the paragraph has no directly-applied alignment
        value and will inherit its alignment value from its style hierarchy. Assigning
        |None| to this property removes any directly-applied alignment value.
        """
        return self._p.alignment

    @alignment.setter
    def alignment(self, value: WD_PARAGRAPH_ALIGNMENT):
        self._p.alignment = value

    def clear(self):
        """Return this same paragraph after removing all its content.

        Paragraph-level formatting, such as style, is preserved.
        """
        self._p.clear_content()
        return self

    @property
    def contains_page_break(self) -> bool:
        """`True` when one or more rendered page-breaks occur in this paragraph."""
        return bool(self._p.lastRenderedPageBreaks)

    @property
    def hyperlinks(self) -> List[Hyperlink]:
        """A |Hyperlink| instance for each hyperlink in this paragraph."""
        return [Hyperlink(hyperlink, self) for hyperlink in self._p.hyperlink_lst]

    def insert_paragraph_before(
        self, text: str | None = None, style: str | ParagraphStyle | None = None
    ) -> Paragraph:
        """Return a newly created paragraph, inserted directly before this paragraph.

        If `text` is supplied, the new paragraph contains that text in a single run. If
        `style` is provided, that style is assigned to the new paragraph.
        """
        paragraph = self._insert_paragraph_before()
        if text:
            paragraph.add_run(text)
        if style is not None:
            paragraph.style = style
        return paragraph

    def iter_inner_content(self) -> Iterator[Run | Hyperlink]:
        """Generate the runs and hyperlinks in this paragraph, in the order they appear.

        The content in a paragraph consists of both runs and hyperlinks. This method
        allows accessing each of those separately, in document order, for when the
        precise position of the hyperlink within the paragraph text is important. Note
        that a hyperlink itself contains runs.
        """
        for r_or_hlink in self._p.inner_content_elements:
            yield (
                Run(r_or_hlink, self)
                if isinstance(r_or_hlink, CT_R)
                else Hyperlink(r_or_hlink, self)
            )

    @property
    def paragraph_format(self):
        """The |ParagraphFormat| object providing access to the formatting properties
        for this paragraph, such as line spacing and indentation."""
        return ParagraphFormat(self._element)

    @property
    def rendered_page_breaks(self) -> List[RenderedPageBreak]:
        """All rendered page-breaks in this paragraph.

        Most often an empty list, sometimes contains one page-break, but can contain
        more than one is rare or contrived cases.
        """
        return [RenderedPageBreak(lrpb, self) for lrpb in self._p.lastRenderedPageBreaks]

    @property
    def runs(self) -> List[Run]:
        """Sequence of |Run| instances corresponding to the <w:r> elements in this
        paragraph."""
        return [Run(r, self) for r in self._p.r_lst]

    @property
    def style(self) -> ParagraphStyle | None:
        """Read/Write.

        |_ParagraphStyle| object representing the style assigned to this paragraph. If
        no explicit style is assigned to this paragraph, its value is the default
        paragraph style for the document. A paragraph style name can be assigned in lieu
        of a paragraph style object. Assigning |None| removes any applied style, making
        its effective value the default paragraph style for the document.
        """
        style_id = self._p.style
        style = self.part.get_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        return cast(ParagraphStyle, style)

    @style.setter
    def style(self, style_or_name: str | ParagraphStyle | None):
        style_id = self.part.get_style_id(style_or_name, WD_STYLE_TYPE.PARAGRAPH)
        self._p.style = style_id

    @property
    def text(self) -> str:
        """The textual content of this paragraph.

        The text includes the visible-text portion of any hyperlinks in the paragraph.
        Tabs and line breaks in the XML are mapped to ``\\t`` and ``\\n`` characters
        respectively.

        Assigning text to this property causes all existing paragraph content to be
        replaced with a single run containing the assigned text. A ``\\t`` character in
        the text is mapped to a ``<w:tab/>`` element and each ``\\n`` or ``\\r``
        character is mapped to a line break. Paragraph-level formatting, such as style,
        is preserved. All run-level formatting, such as bold or italic, is removed.
        """
        return self._p.text

    @text.setter
    def text(self, text: str | None):
        self.clear()
        self.add_run(text)

    def _insert_paragraph_before(self):
        """Return a newly created paragraph, inserted directly before this paragraph."""
        p = self._p.add_p_before()
        return Paragraph(p, self._parent)


# parfmt.py
"""Paragraph-related proxy types."""

from docx.enum.text import WD_LINE_SPACING
from docx.shared import ElementProxy, Emu, Length, Pt, Twips, lazyproperty
from docx.text.tabstops import TabStops


class ParagraphFormat(ElementProxy):
    """Provides access to paragraph formatting such as justification, indentation, line
    spacing, space before and after, and widow/orphan control."""

    @property
    def alignment(self):
        """A member of the :ref:`WdParagraphAlignment` enumeration specifying the
        justification setting for this paragraph.

        A value of |None| indicates paragraph alignment is inherited from the style
        hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.jc_val

    @alignment.setter
    def alignment(self, value):
        pPr = self._element.get_or_add_pPr()
        pPr.jc_val = value

    @property
    def first_line_indent(self):
        """|Length| value specifying the relative difference in indentation for the
        first line of the paragraph.

        A positive value causes the first line to be indented. A negative value produces
        a hanging indent. |None| indicates first line indentation is inherited from the
        style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.first_line_indent

    @first_line_indent.setter
    def first_line_indent(self, value):
        pPr = self._element.get_or_add_pPr()
        pPr.first_line_indent = value

    @property
    def keep_together(self):
        """|True| if the paragraph should be kept "in one piece" and not broken across a
        page boundary when the document is rendered.

        |None| indicates its effective value is inherited from the style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.keepLines_val

    @keep_together.setter
    def keep_together(self, value):
        self._element.get_or_add_pPr().keepLines_val = value

    @property
    def keep_with_next(self):
        """|True| if the paragraph should be kept on the same page as the subsequent
        paragraph when the document is rendered.

        For example, this property could be used to keep a section heading on the same
        page as its first paragraph. |None| indicates its effective value is inherited
        from the style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.keepNext_val

    @keep_with_next.setter
    def keep_with_next(self, value):
        self._element.get_or_add_pPr().keepNext_val = value

    @property
    def left_indent(self):
        """|Length| value specifying the space between the left margin and the left side
        of the paragraph.

        |None| indicates the left indent value is inherited from the style hierarchy.
        Use an |Inches| value object as a convenient way to apply indentation in units
        of inches.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.ind_left

    @left_indent.setter
    def left_indent(self, value):
        pPr = self._element.get_or_add_pPr()
        pPr.ind_left = value

    @property
    def line_spacing(self):
        """|float| or |Length| value specifying the space between baselines in
        successive lines of the paragraph.

        A value of |None| indicates line spacing is inherited from the style hierarchy.
        A float value, e.g. ``2.0`` or ``1.75``, indicates spacing is applied in
        multiples of line heights. A |Length| value such as ``Pt(12)`` indicates spacing
        is a fixed height. The |Pt| value class is a convenient way to apply line
        spacing in units of points. Assigning |None| resets line spacing to inherit from
        the style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return self._line_spacing(pPr.spacing_line, pPr.spacing_lineRule)

    @line_spacing.setter
    def line_spacing(self, value):
        pPr = self._element.get_or_add_pPr()
        if value is None:
            pPr.spacing_line = None
            pPr.spacing_lineRule = None
        elif isinstance(value, Length):
            pPr.spacing_line = value
            if pPr.spacing_lineRule != WD_LINE_SPACING.AT_LEAST:
                pPr.spacing_lineRule = WD_LINE_SPACING.EXACTLY
        else:
            pPr.spacing_line = Emu(value * Twips(240))
            pPr.spacing_lineRule = WD_LINE_SPACING.MULTIPLE

    @property
    def line_spacing_rule(self):
        """A member of the :ref:`WdLineSpacing` enumeration indicating how the value of
        :attr:`line_spacing` should be interpreted.

        Assigning any of the :ref:`WdLineSpacing` members :attr:`SINGLE`,
        :attr:`DOUBLE`, or :attr:`ONE_POINT_FIVE` will cause the value of
        :attr:`line_spacing` to be updated to produce the corresponding line spacing.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return self._line_spacing_rule(pPr.spacing_line, pPr.spacing_lineRule)

    @line_spacing_rule.setter
    def line_spacing_rule(self, value):
        pPr = self._element.get_or_add_pPr()
        if value == WD_LINE_SPACING.SINGLE:
            pPr.spacing_line = Twips(240)
            pPr.spacing_lineRule = WD_LINE_SPACING.MULTIPLE
        elif value == WD_LINE_SPACING.ONE_POINT_FIVE:
            pPr.spacing_line = Twips(360)
            pPr.spacing_lineRule = WD_LINE_SPACING.MULTIPLE
        elif value == WD_LINE_SPACING.DOUBLE:
            pPr.spacing_line = Twips(480)
            pPr.spacing_lineRule = WD_LINE_SPACING.MULTIPLE
        else:
            pPr.spacing_lineRule = value

    @property
    def page_break_before(self):
        """|True| if the paragraph should appear at the top of the page following the
        prior paragraph.

        |None| indicates its effective value is inherited from the style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.pageBreakBefore_val

    @page_break_before.setter
    def page_break_before(self, value):
        self._element.get_or_add_pPr().pageBreakBefore_val = value

    @property
    def right_indent(self):
        """|Length| value specifying the space between the right margin and the right
        side of the paragraph.

        |None| indicates the right indent value is inherited from the style hierarchy.
        Use a |Cm| value object as a convenient way to apply indentation in units of
        centimeters.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.ind_right

    @right_indent.setter
    def right_indent(self, value):
        pPr = self._element.get_or_add_pPr()
        pPr.ind_right = value

    @property
    def space_after(self):
        """|Length| value specifying the spacing to appear between this paragraph and
        the subsequent paragraph.

        |None| indicates this value is inherited from the style hierarchy. |Length|
        objects provide convenience properties, such as :attr:`~.Length.pt` and
        :attr:`~.Length.inches`, that allow easy conversion to various length units.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.spacing_after

    @space_after.setter
    def space_after(self, value):
        self._element.get_or_add_pPr().spacing_after = value

    @property
    def space_before(self):
        """|Length| value specifying the spacing to appear between this paragraph and
        the prior paragraph.

        |None| indicates this value is inherited from the style hierarchy. |Length|
        objects provide convenience properties, such as :attr:`~.Length.pt` and
        :attr:`~.Length.cm`, that allow easy conversion to various length units.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.spacing_before

    @space_before.setter
    def space_before(self, value):
        self._element.get_or_add_pPr().spacing_before = value

    @lazyproperty
    def tab_stops(self):
        """|TabStops| object providing access to the tab stops defined for this
        paragraph format."""
        pPr = self._element.get_or_add_pPr()
        return TabStops(pPr)

    @property
    def widow_control(self):
        """|True| if the first and last lines in the paragraph remain on the same page
        as the rest of the paragraph when Word repaginates the document.

        |None| indicates its effective value is inherited from the style hierarchy.
        """
        pPr = self._element.pPr
        if pPr is None:
            return None
        return pPr.widowControl_val

    @widow_control.setter
    def widow_control(self, value):
        self._element.get_or_add_pPr().widowControl_val = value

    @staticmethod
    def _line_spacing(spacing_line, spacing_lineRule):
        """Return the line spacing value calculated from the combination of
        `spacing_line` and `spacing_lineRule`.

        Returns a |float| number of lines when `spacing_lineRule` is
        ``WD_LINE_SPACING.MULTIPLE``, otherwise a |Length| object of absolute line
        height is returned. Returns |None| when `spacing_line` is |None|.
        """
        if spacing_line is None:
            return None
        if spacing_lineRule == WD_LINE_SPACING.MULTIPLE:
            return spacing_line / Pt(12)
        return spacing_line

    @staticmethod
    def _line_spacing_rule(line, lineRule):
        """Return the line spacing rule value calculated from the combination of `line`
        and `lineRule`.

        Returns special members of the :ref:`WdLineSpacing` enumeration when line
        spacing is single, double, or 1.5 lines.
        """
        if lineRule == WD_LINE_SPACING.MULTIPLE:
            if line == Twips(240):
                return WD_LINE_SPACING.SINGLE
            if line == Twips(360):
                return WD_LINE_SPACING.ONE_POINT_FIVE
            if line == Twips(480):
                return WD_LINE_SPACING.DOUBLE
        return lineRule


# run.py
"""Run-related proxy objects for python-docx, Run in particular."""

from __future__ import annotations

from typing import IO, TYPE_CHECKING, Iterator, cast

from docx.drawing import Drawing
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.drawing import CT_Drawing
from docx.oxml.text.pagebreak import CT_LastRenderedPageBreak
from docx.shape import InlineShape
from docx.shared import StoryChild
from docx.styles.style import CharacterStyle
from docx.text.font import Font
from docx.text.pagebreak import RenderedPageBreak

if TYPE_CHECKING:
    import docx.types as t
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.text.run import CT_R, CT_Text
    from docx.shared import Length


class Run(StoryChild):
    """Proxy object wrapping `<w:r>` element.

    Several of the properties on Run take a tri-state value, |True|, |False|, or |None|.
    |True| and |False| correspond to on and off respectively. |None| indicates the
    property is not specified directly on the run and its effective value is taken from
    the style hierarchy.
    """

    def __init__(self, r: CT_R, parent: t.ProvidesStoryPart):
        super().__init__(parent)
        self._r = self._element = self.element = r

    def add_break(self, break_type: WD_BREAK = WD_BREAK.LINE):
        """Add a break element of `break_type` to this run.

        `break_type` can take the values `WD_BREAK.LINE`, `WD_BREAK.PAGE`, and
        `WD_BREAK.COLUMN` where `WD_BREAK` is imported from `docx.enum.text`.
        `break_type` defaults to `WD_BREAK.LINE`.
        """
        type_, clear = {
            WD_BREAK.LINE: (None, None),
            WD_BREAK.PAGE: ("page", None),
            WD_BREAK.COLUMN: ("column", None),
            WD_BREAK.LINE_CLEAR_LEFT: ("textWrapping", "left"),
            WD_BREAK.LINE_CLEAR_RIGHT: ("textWrapping", "right"),
            WD_BREAK.LINE_CLEAR_ALL: ("textWrapping", "all"),
        }[break_type]
        br = self._r.add_br()
        if type_ is not None:
            br.type = type_
        if clear is not None:
            br.clear = clear

    def add_picture(
        self,
        image_path_or_stream: str | IO[bytes],
        width: int | Length | None = None,
        height: int | Length | None = None,
    ) -> InlineShape:
        """Return |InlineShape| containing image identified by `image_path_or_stream`.

        The picture is added to the end of this run.

        `image_path_or_stream` can be a path (a string) or a file-like object containing
        a binary image.

        If neither width nor height is specified, the picture appears at
        its native size. If only one is specified, it is used to compute a scaling
        factor that is then applied to the unspecified dimension, preserving the aspect
        ratio of the image. The native size of the picture is calculated using the dots-
        per-inch (dpi) value specified in the image file, defaulting to 72 dpi if no
        value is specified, as is often the case.
        """
        inline = self.part.new_pic_inline(image_path_or_stream, width, height)
        self._r.add_drawing(inline)
        return InlineShape(inline)

    def add_tab(self) -> None:
        """Add a ``<w:tab/>`` element at the end of the run, which Word interprets as a
        tab character."""
        self._r.add_tab()

    def add_text(self, text: str):
        """Returns a newly appended |_Text| object (corresponding to a new ``<w:t>``
        child element) to the run, containing `text`.

        Compare with the possibly more friendly approach of assigning text to the
        :attr:`Run.text` property.
        """
        t = self._r.add_t(text)
        return _Text(t)

    @property
    def bold(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text of the run to appear in bold face. When |False|,
        the text unconditionally appears non-bold. When |None| the bold setting for this
        run is inherited from the style hierarchy.
        """
        return self.font.bold

    @bold.setter
    def bold(self, value: bool | None):
        self.font.bold = value

    def clear(self):
        """Return reference to this run after removing all its content.

        All run formatting is preserved.
        """
        self._r.clear_content()
        return self

    @property
    def contains_page_break(self) -> bool:
        """`True` when one or more rendered page-breaks occur in this run.

        Note that "hard" page-breaks inserted by the author are not included. A hard
        page-break gives rise to a rendered page-break in the right position so if those
        were included that page-break would be "double-counted".

        It would be very rare for multiple rendered page-breaks to occur in a single
        run, but it is possible.
        """
        return bool(self._r.lastRenderedPageBreaks)

    @property
    def font(self) -> Font:
        """The |Font| object providing access to the character formatting properties for
        this run, such as font name and size."""
        return Font(self._element)

    @property
    def italic(self) -> bool | None:
        """Read/write tri-state value.

        When |True|, causes the text of the run to appear in italics. When |False|, the
        text unconditionally appears non-italic. When |None| the italic setting for this
        run is inherited from the style hierarchy.
        """
        return self.font.italic

    @italic.setter
    def italic(self, value: bool | None):
        self.font.italic = value

    def iter_inner_content(self) -> Iterator[str | Drawing | RenderedPageBreak]:
        """Generate the content-items in this run in the order they appear.

        NOTE: only content-types currently supported by `python-docx` are generated. In
        this version, that is text and rendered page-breaks. Drawing is included but
        currently only provides access to its XML element (CT_Drawing) on its
        `._drawing` attribute. `Drawing` attributes and methods may be expanded in
        future releases.

        There are a number of element-types that can appear inside a run, but most of
        those (w:br, w:cr, w:noBreakHyphen, w:t, w:tab) have a clear plain-text
        equivalent. Any contiguous range of such elements is generated as a single
        `str`. Rendered page-break and drawing elements are generated individually. Any
        other elements are ignored.
        """
        for item in self._r.inner_content_items:
            if isinstance(item, str):
                yield item
            elif isinstance(item, CT_LastRenderedPageBreak):
                yield RenderedPageBreak(item, self)
            elif isinstance(item, CT_Drawing):  # pyright: ignore[reportUnnecessaryIsInstance]
                yield Drawing(item, self)

    @property
    def style(self) -> CharacterStyle:
        """Read/write.

        A |CharacterStyle| object representing the character style applied to this run.
        The default character style for the document (often `Default Character Font`) is
        returned if the run has no directly-applied character style. Setting this
        property to |None| removes any directly-applied character style.
        """
        style_id = self._r.style
        return cast(CharacterStyle, self.part.get_style(style_id, WD_STYLE_TYPE.CHARACTER))

    @style.setter
    def style(self, style_or_name: str | CharacterStyle | None):
        style_id = self.part.get_style_id(style_or_name, WD_STYLE_TYPE.CHARACTER)
        self._r.style = style_id

    @property
    def text(self) -> str:
        """String formed by concatenating the text equivalent of each run.

        Each `<w:t>` element adds the text characters it contains. A `<w:tab/>` element
        adds a `\\t` character. A `<w:cr/>` or `<w:br>` element each add a `\\n`
        character. Note that a `<w:br>` element can indicate a page break or column
        break as well as a line break. Only line-break `<w:br>` elements translate to
        a `\\n` character. Others are ignored. All other content child elements, such as
        `<w:drawing>`, are ignored.

        Assigning text to this property has the reverse effect, translating each `\\t`
        character to a `<w:tab/>` element and each `\\n` or `\\r` character to a
        `<w:cr/>` element. Any existing run content is replaced. Run formatting is
        preserved.
        """
        return self._r.text

    @text.setter
    def text(self, text: str):
        self._r.text = text

    @property
    def underline(self) -> bool | WD_UNDERLINE | None:
        """The underline style for this |Run|.

        Value is one of |None|, |True|, |False|, or a member of :ref:`WdUnderline`.

        A value of |None| indicates the run has no directly-applied underline value and
        so will inherit the underline value of its containing paragraph. Assigning
        |None| to this property removes any directly-applied underline value.

        A value of |False| indicates a directly-applied setting of no underline,
        overriding any inherited value.

        A value of |True| indicates single underline.

        The values from :ref:`WdUnderline` are used to specify other outline styles such
        as double, wavy, and dotted.
        """
        return self.font.underline

    @underline.setter
    def underline(self, value: bool):
        self.font.underline = value


class _Text:
    """Proxy object wrapping `<w:t>` element."""

    def __init__(self, t_elm: CT_Text):
        super(_Text, self).__init__()
        self._t = t_elm


# tabstops.py
"""Tabstop-related proxy types."""

from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import ElementProxy


class TabStops(ElementProxy):
    """A sequence of |TabStop| objects providing access to the tab stops of a paragraph
    or paragraph style.

    Supports iteration, indexed access, del, and len(). It is accesed using the
    :attr:`~.ParagraphFormat.tab_stops` property of ParagraphFormat; it is not intended
    to be constructed directly.
    """

    def __init__(self, element):
        super(TabStops, self).__init__(element, None)
        self._pPr = element

    def __delitem__(self, idx):
        """Remove the tab at offset `idx` in this sequence."""
        tabs = self._pPr.tabs
        try:
            tabs.remove(tabs[idx])
        except (AttributeError, IndexError):
            raise IndexError("tab index out of range")

        if len(tabs) == 0:
            self._pPr.remove(tabs)

    def __getitem__(self, idx):
        """Enables list-style access by index."""
        tabs = self._pPr.tabs
        if tabs is None:
            raise IndexError("TabStops object is empty")
        tab = tabs.tab_lst[idx]
        return TabStop(tab)

    def __iter__(self):
        """Generate a TabStop object for each of the w:tab elements, in XML document
        order."""
        tabs = self._pPr.tabs
        if tabs is not None:
            for tab in tabs.tab_lst:
                yield TabStop(tab)

    def __len__(self):
        tabs = self._pPr.tabs
        if tabs is None:
            return 0
        return len(tabs.tab_lst)

    def add_tab_stop(
        self, position, alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES
    ):
        """Add a new tab stop at `position`, a |Length| object specifying the location
        of the tab stop relative to the paragraph edge.

        A negative `position` value is valid and appears in hanging indentation. Tab
        alignment defaults to left, but may be specified by passing a member of the
        :ref:`WdTabAlignment` enumeration as `alignment`. An optional leader character
        can be specified by passing a member of the :ref:`WdTabLeader` enumeration as
        `leader`.
        """
        tabs = self._pPr.get_or_add_tabs()
        tab = tabs.insert_tab_in_order(position, alignment, leader)
        return TabStop(tab)

    def clear_all(self):
        """Remove all custom tab stops."""
        self._pPr._remove_tabs()


class TabStop(ElementProxy):
    """An individual tab stop applying to a paragraph or style.

    Accessed using list semantics on its containing |TabStops| object.
    """

    def __init__(self, element):
        super(TabStop, self).__init__(element, None)
        self._tab = element

    @property
    def alignment(self):
        """A member of :ref:`WdTabAlignment` specifying the alignment setting for this
        tab stop.

        Read/write.
        """
        return self._tab.val

    @alignment.setter
    def alignment(self, value):
        self._tab.val = value

    @property
    def leader(self):
        """A member of :ref:`WdTabLeader` specifying a repeating character used as a
        "leader", filling in the space spanned by this tab.

        Assigning |None| produces the same result as assigning `WD_TAB_LEADER.SPACES`.
        Read/write.
        """
        return self._tab.leader

    @leader.setter
    def leader(self, value):
        self._tab.leader = value

    @property
    def position(self):
        """A |Length| object representing the distance of this tab stop from the inside
        edge of the paragraph.

        May be positive or negative. Read/write.
        """
        return self._tab.pos

    @position.setter
    def position(self, value):
        tab = self._tab
        tabs = tab.getparent()
        self._tab = tabs.insert_tab_in_order(value, tab.val, tab.leader)
        tabs.remove(tab)


# __init__.py


# types.py
"""Abstract types used by `python-docx`."""

from __future__ import annotations

from typing import TYPE_CHECKING

from typing_extensions import Protocol

if TYPE_CHECKING:
    from docx.opc.part import XmlPart
    from docx.parts.story import StoryPart


class ProvidesStoryPart(Protocol):
    """An object that provides access to the StoryPart.

    This type is for objects that have a story part like document or header as their
    root part.
    """

    @property
    def part(self) -> StoryPart:
        ...


class ProvidesXmlPart(Protocol):
    """An object that provides access to its XmlPart.

    This type is for objects that need access to their part but it either isn't a
    StoryPart or they don't care, possibly because they just need access to the package
    or related parts.
    """

    @property
    def part(self) -> XmlPart:
        ...


# __init__.py
"""Initialize `docx` package.

Export the `Document` constructor function and establish the mapping of part-type to
the part-classe that implements that type.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Type

from docx.api import Document

if TYPE_CHECKING:
    from docx.opc.part import Part

__version__ = "1.1.2"


__all__ = ["Document"]


# -- register custom Part classes with opc package reader --

from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import PartFactory
from docx.opc.parts.coreprops import CorePropertiesPart
from docx.parts.document import DocumentPart
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.parts.image import ImagePart
from docx.parts.numbering import NumberingPart
from docx.parts.settings import SettingsPart
from docx.parts.styles import StylesPart


def part_class_selector(content_type: str, reltype: str) -> Type[Part] | None:
    if reltype == RT.IMAGE:
        return ImagePart
    return None


PartFactory.part_class_selector = part_class_selector
PartFactory.part_type_for[CT.OPC_CORE_PROPERTIES] = CorePropertiesPart
PartFactory.part_type_for[CT.WML_DOCUMENT_MAIN] = DocumentPart
PartFactory.part_type_for[CT.WML_FOOTER] = FooterPart
PartFactory.part_type_for[CT.WML_HEADER] = HeaderPart
PartFactory.part_type_for[CT.WML_NUMBERING] = NumberingPart
PartFactory.part_type_for[CT.WML_SETTINGS] = SettingsPart
PartFactory.part_type_for[CT.WML_STYLES] = StylesPart

del (
    CT,
    CorePropertiesPart,
    DocumentPart,
    FooterPart,
    HeaderPart,
    NumberingPart,
    PartFactory,
    SettingsPart,
    StylesPart,
    part_class_selector,
)


